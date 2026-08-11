
from __future__ import annotations

import asyncio
import ast
import base64
import binascii
import csv
import math
import hashlib
import ipaddress
import hmac
import io
import json
import os
import re
import secrets
import shutil
import sqlite3
import tempfile
import time
import threading
import unicodedata
import zipfile
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timedelta, timezone
from difflib import SequenceMatcher
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP, getcontext
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urlparse
from xml.sax.saxutils import escape as xml_escape

from cryptography.fernet import Fernet, InvalidToken

import httpx
import jwt
from bs4 import BeautifulSoup
from docx import Document
from fastapi import BackgroundTasks, Cookie, Depends, FastAPI, File, Form, Header, HTTPException, Request, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, PlainTextResponse, StreamingResponse
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from pydantic import BaseModel, EmailStr, Field
from pypdf import PdfReader
from rag_engine import cosine_similarity, extract_structured_facts, pack_vector, rerank_hybrid_candidates, unpack_vector
from deep_rag import analyze_query, merge_multiretrieval, semantic_rerank_candidates, build_rule_exception_map, format_rule_map_for_prompt, evidence_confidence
from barsan_cargo import calculate_cargo_fit, format_cargo_result
from barsan_location import extract_location_query, format_location_answer, lookup_neshan, normalize_location_text
from release_info import APP_VERSION, RELEASE_ID, ASSET_VERSION, INGESTION_VERSION, ANSWER_CACHE_NAMESPACE_DEFAULT, SCHEMA_REVISION
from source_quality import source_number_unit_tokens, text_health_score, page_fidelity_metrics, generic_source_quality
from provider_runtime import parse_capability_overrides, infer_provider_capabilities
from runtime_guards import document_source_ids, preserve_live_document_status, document_job_priority
from ops_runtime import acquire_sqlite_replica_lock, release_sqlite_replica_lock, same_origin_allowed, ReplicaLockError

try:
    import fitz  # PyMuPDF: page rendering and embedded-image detection for PDF vision indexing
except ImportError:  # pragma: no cover - validated at runtime when PDF vision is enabled
    fitz = None

APP_NAME = os.getenv("APP_NAME", "Barsan AI Chatbot")
ENVIRONMENT = os.getenv("ENVIRONMENT", "production").lower()
PORT = int(os.getenv("PORT", "8080"))

# Railway runtime awareness. RAILWAY_PUBLIC_DOMAIN and RAILWAY_VOLUME_MOUNT_PATH
# are injected by Railway when public networking / a volume are attached.
_RAILWAY_PUBLIC_DOMAIN = os.getenv("RAILWAY_PUBLIC_DOMAIN", "").strip().strip("/")
_RAILWAY_VOLUME_MOUNT_PATH = os.getenv("RAILWAY_VOLUME_MOUNT_PATH", "").strip()
_DATA_ROOT = Path(_RAILWAY_VOLUME_MOUNT_PATH or os.getenv("DATA_DIR", "/data"))
_configured_public_base = os.getenv("PUBLIC_BASE_URL", "").strip().rstrip("/")

def resolve_public_base_url(configured: str, railway_domain: str) -> str:
    configured=str(configured or '').strip().rstrip('/')
    railway_domain=str(railway_domain or '').strip().strip('/')
    return configured or (f"https://{railway_domain}" if railway_domain else "http://localhost:8080")

PUBLIC_BASE_URL = resolve_public_base_url(_configured_public_base,_RAILWAY_PUBLIC_DOMAIN)

def _railway_path_env(name: str, legacy_default: str, relative_name: str) -> str:
    configured = os.getenv(name, "").strip()
    if _RAILWAY_VOLUME_MOUNT_PATH and configured in {"", legacy_default}:
        return str(_DATA_ROOT / relative_name)
    return configured or str(_DATA_ROOT / relative_name)

_raw_database_url = os.getenv("DATABASE_URL", "").strip()
if _RAILWAY_VOLUME_MOUNT_PATH and _raw_database_url in {"", "sqlite:////data/barsan.db"}:
    DATABASE_URL = f"sqlite:///{_DATA_ROOT / 'barsan.db'}"
else:
    DATABASE_URL = _raw_database_url or f"sqlite:///{_DATA_ROOT / 'barsan.db'}"
UPLOAD_DIR = Path(_railway_path_env("UPLOAD_DIR", "/data/uploads", "uploads"))
UPLOAD_SESSION_DIR = Path(_railway_path_env("UPLOAD_SESSION_DIR", "/data/upload-sessions", "upload-sessions"))
CARGO_IMAGE_TRAINING_DIR = UPLOAD_DIR / 'cargo-image-training'
CARGO_IMAGE_MAX_MB = max(1, min(25, int(os.getenv('CARGO_IMAGE_MAX_MB', '12'))))
MAX_UPLOAD_MB = int(os.getenv("MAX_UPLOAD_MB", "500"))
UPLOAD_CHUNK_MB = max(1, min(16, int(os.getenv("UPLOAD_CHUNK_MB", "8"))))
UPLOAD_CHUNK_BYTES = UPLOAD_CHUNK_MB * 1024 * 1024
UPLOAD_SESSION_TTL_HOURS = max(1, min(168, int(os.getenv("UPLOAD_SESSION_TTL_HOURS", "24"))))
GOOGLE_DOC_MAX_MB = max(1, min(MAX_UPLOAD_MB, int(os.getenv("GOOGLE_DOC_MAX_MB", "100"))))
DOCUMENT_SUFFIXES = {'.pdf','.docx','.xlsx','.csv','.txt','.json','.html','.htm','.md'}
JWT_SECRET = os.getenv("JWT_SECRET", "")
PROVIDER_KEY_ENCRYPTION_SECRET = os.getenv("PROVIDER_KEY_ENCRYPTION_SECRET", JWT_SECRET).strip()
JWT_EXPIRE_MINUTES = int(os.getenv("JWT_EXPIRE_MINUTES", "720"))
COOKIE_SECURE = os.getenv("COOKIE_SECURE", "true").lower() in {"1", "true", "yes", "on"}
INITIAL_ADMIN_USERNAME = os.getenv("INITIAL_ADMIN_USERNAME", "admin")
INITIAL_ADMIN_EMAIL = os.getenv("INITIAL_ADMIN_EMAIL", "admin@barsanai.com")
INITIAL_ADMIN_PASSWORD = os.getenv("INITIAL_ADMIN_PASSWORD", "")
INITIAL_ADMIN_NAME = os.getenv("INITIAL_ADMIN_NAME", "مدیر بارسان")
MASTER_ADMIN_USERNAME = os.getenv("MASTER_ADMIN_USERNAME", INITIAL_ADMIN_USERNAME)
MASTER_ADMIN_EMAIL = os.getenv("MASTER_ADMIN_EMAIL", INITIAL_ADMIN_EMAIL)
SYNC_MASTER_ADMIN_CREDENTIALS = os.getenv("SYNC_MASTER_ADMIN_CREDENTIALS", "false").strip().lower() in {"1","true","yes","on"}
AI_PROVIDER = os.getenv("AI_PROVIDER", "openai_compatible").strip().lower()

# Generic OpenAI-compatible provider. Switch API vendors later by changing only
# AI_API_KEY, AI_BASE_URL and AI_MODEL in Railway Variables.
AI_API_KEY = os.getenv("AI_API_KEY", "") or os.getenv("GAPGPT_API_KEY", "") or os.getenv("OPENAI_API_KEY", "")
AI_BASE_URL = os.getenv("AI_BASE_URL", "https://api.gapgpt.app/v1").rstrip("/")
AI_MODEL = os.getenv("AI_MODEL", "gpt-4o-mini").strip()
AI_PROVIDER_LABEL = os.getenv("AI_PROVIDER_LABEL", "GapGPT").strip() or "AI provider"
AI_CHAT_COMPLETIONS_PATH = os.getenv("AI_CHAT_COMPLETIONS_PATH", "/chat/completions").strip() or "/chat/completions"
AI_TIMEOUT_SECONDS = float(os.getenv("AI_TIMEOUT_SECONDS", "180"))
AI_MAX_RETRIES = max(0, int(os.getenv("AI_MAX_RETRIES", "3")))
AI_MAX_COMPLETION_TOKENS = max(256, int(os.getenv("AI_MAX_COMPLETION_TOKENS", "2400")))
AI_DEFAULT_MAX_COMPLETION_TOKENS = max(128, min(AI_MAX_COMPLETION_TOKENS, int(os.getenv("AI_DEFAULT_MAX_COMPLETION_TOKENS", "450"))))
AI_DETAILED_MAX_COMPLETION_TOKENS = max(AI_DEFAULT_MAX_COMPLETION_TOKENS, min(AI_MAX_COMPLETION_TOKENS, int(os.getenv("AI_DETAILED_MAX_COMPLETION_TOKENS", "2400"))))
AI_DEFAULT_MAX_ANSWER_CHARS = max(220, int(os.getenv("AI_DEFAULT_MAX_ANSWER_CHARS", "700")))
AI_DETAILED_MAX_ANSWER_CHARS = max(1200, int(os.getenv("AI_DETAILED_MAX_ANSWER_CHARS", "12000")))
AI_TEMPERATURE = min(1.0, max(0.0, float(os.getenv("AI_TEMPERATURE", "0.12"))))
AI_TOKEN_PARAMETER = os.getenv("AI_TOKEN_PARAMETER", "auto").strip().lower()
AI_SEND_TEMPERATURE = os.getenv("AI_SEND_TEMPERATURE", "true").lower() in {"1", "true", "yes", "on"}
AI_AUTH_HEADER = os.getenv("AI_AUTH_HEADER", "Authorization").strip() or "Authorization"
AI_AUTH_SCHEME = os.getenv("AI_AUTH_SCHEME", "Bearer").strip()
AI_EXTRA_HEADERS_JSON = os.getenv("AI_EXTRA_HEADERS_JSON", "{}").strip() or "{}"
AI_CONTINUATION_ROUNDS = max(0, min(3, int(os.getenv("AI_CONTINUATION_ROUNDS", "1"))))
AI_ENFORCE_DETAILED_ANSWER = os.getenv("AI_ENFORCE_DETAILED_ANSWER", "false").lower() in {"1", "true", "yes", "on"}
AI_MIN_RESPONSE_CHARS = max(60, int(os.getenv("AI_MIN_RESPONSE_CHARS", "120")))
DIRECT_TRAINING_ANSWER_ENABLED = os.getenv("DIRECT_TRAINING_ANSWER_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
ANSWER_CACHE_ENABLED = os.getenv("ANSWER_CACHE_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
ANSWER_CACHE_FUZZY_THRESHOLD = min(1.0, max(0.90, float(os.getenv("ANSWER_CACHE_FUZZY_THRESHOLD", "0.97"))))
ANSWER_CACHE_MAX_AGE_DAYS = max(1, int(os.getenv("ANSWER_CACHE_MAX_AGE_DAYS", "90")))
ANSWER_CACHE_MAX_ROWS = max(100, int(os.getenv("ANSWER_CACHE_MAX_ROWS", "5000")))
ANSWER_CACHE_FUZZY_CANDIDATES = max(20, min(500, int(os.getenv("ANSWER_CACHE_FUZZY_CANDIDATES", "120"))))
ANSWER_CACHE_NAMESPACE = os.getenv("ANSWER_CACHE_NAMESPACE", ANSWER_CACHE_NAMESPACE_DEFAULT).strip() or ANSWER_CACHE_NAMESPACE_DEFAULT

# R31 dynamic OpenAI-compatible API pool. Slots 1-6 remain bootstrapped from Railway
# Variables for backwards compatibility; the owner/admin can add slots 5-20 at runtime.
MAX_AI_API_SLOTS = max(4, min(20, int(os.getenv("MAX_AI_API_SLOTS", "20"))))
AI_API_KEY_1 = os.getenv("AI_API_KEY_1", "").strip()
AI_API_KEY_2 = os.getenv("AI_API_KEY_2", "").strip()
AI_API_KEY_3 = os.getenv("AI_API_KEY_3", "").strip()
AI_API_KEY_4 = os.getenv("AI_API_KEY_4", "").strip()
AI_API_KEY_5 = os.getenv("AI_API_KEY_5", "").strip()
AI_API_KEY_6 = os.getenv("AI_API_KEY_6", "").strip()
AI_BASE_URL_1 = os.getenv("AI_BASE_URL_1", "").strip().rstrip("/")
AI_BASE_URL_2 = os.getenv("AI_BASE_URL_2", "").strip().rstrip("/")
AI_BASE_URL_3 = os.getenv("AI_BASE_URL_3", "").strip().rstrip("/")
AI_BASE_URL_4 = os.getenv("AI_BASE_URL_4", "").strip().rstrip("/")
AI_BASE_URL_5 = os.getenv("AI_BASE_URL_5", "").strip().rstrip("/")
AI_BASE_URL_6 = os.getenv("AI_BASE_URL_6", "").strip().rstrip("/")
AI_MODEL_1 = os.getenv("AI_MODEL_1", "").strip()
AI_MODEL_2 = os.getenv("AI_MODEL_2", "").strip()
AI_MODEL_3 = os.getenv("AI_MODEL_3", "").strip()
AI_MODEL_4 = os.getenv("AI_MODEL_4", "").strip()
AI_MODEL_5 = os.getenv("AI_MODEL_5", "").strip()
AI_MODEL_6 = os.getenv("AI_MODEL_6", "").strip()
AI_PROVIDER_LABEL_1 = os.getenv("AI_PROVIDER_LABEL_1", "").strip()
AI_PROVIDER_LABEL_2 = os.getenv("AI_PROVIDER_LABEL_2", "").strip()
AI_PROVIDER_LABEL_3 = os.getenv("AI_PROVIDER_LABEL_3", "").strip()
AI_PROVIDER_LABEL_4 = os.getenv("AI_PROVIDER_LABEL_4", "").strip()
AI_PROVIDER_LABEL_5 = os.getenv("AI_PROVIDER_LABEL_5", "").strip()
AI_PROVIDER_LABEL_6 = os.getenv("AI_PROVIDER_LABEL_6", "").strip()
AI_KEY_RATE_LIMIT_COOLDOWN_SECONDS = max(30, int(os.getenv("AI_KEY_RATE_LIMIT_COOLDOWN_SECONDS", "1800")))
AI_KEY_QUOTA_COOLDOWN_SECONDS = max(300, int(os.getenv("AI_KEY_QUOTA_COOLDOWN_SECONDS", "86400")))
AI_KEY_TRANSIENT_COOLDOWN_SECONDS = max(15, int(os.getenv("AI_KEY_TRANSIENT_COOLDOWN_SECONDS", "120")))
AI_ROTATE_ON_TRANSIENT_ERRORS = os.getenv("AI_ROTATE_ON_TRANSIENT_ERRORS", "true").lower() in {"1", "true", "yes", "on"}

# Model routing. Direct training is authoritative; FAQ/cache are fallback-only zero-token routes;
# economy, standard and advanced routes select a model per API slot.
MODEL_ROUTING_ENABLED = os.getenv("MODEL_ROUTING_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
MODEL_ROUTE_SIMPLE_MAX_WORDS = max(3, int(os.getenv("MODEL_ROUTE_SIMPLE_MAX_WORDS", "24")))
MODEL_ROUTE_MULTI_SOURCE_THRESHOLD = max(2, int(os.getenv("MODEL_ROUTE_MULTI_SOURCE_THRESHOLD", "3")))
MODEL_ROUTE_COMPLEX_KEYWORDS = tuple(x.strip() for x in os.getenv("MODEL_ROUTE_COMPLEX_KEYWORDS", "مقایسه,تحلیل,مرحله به مرحله,جزئیات,کامل,دقیق,علت,سناریو,استثنا").split(',') if x.strip())


# R34 deep-analysis pipeline. All stages are source-first and can be tuned independently.
DEEP_QUERY_ANALYSIS_ENABLED = os.getenv("DEEP_QUERY_ANALYSIS_ENABLED", "true").lower() in {"1","true","yes","on"}
DEEP_MULTI_RETRIEVAL_ENABLED = os.getenv("DEEP_MULTI_RETRIEVAL_ENABLED", "true").lower() in {"1","true","yes","on"}
DEEP_MULTI_RETRIEVAL_MAX_QUERIES = max(2, min(8, int(os.getenv("DEEP_MULTI_RETRIEVAL_MAX_QUERIES", "6"))))
DEEP_SEMANTIC_RERANK_ENABLED = os.getenv("DEEP_SEMANTIC_RERANK_ENABLED", "true").lower() in {"1","true","yes","on"}
MODEL_SEMANTIC_RERANK_ENABLED = os.getenv("MODEL_SEMANTIC_RERANK_ENABLED", "true").lower() in {"1","true","yes","on"}
MODEL_SEMANTIC_RERANK_MIN_COMPLEXITY = max(0.0,min(1.0,float(os.getenv("MODEL_SEMANTIC_RERANK_MIN_COMPLEXITY","0.55"))))
MODEL_SEMANTIC_RERANK_MAX_ITEMS = max(4,min(18,int(os.getenv("MODEL_SEMANTIC_RERANK_MAX_ITEMS","12"))))
MODEL_SEMANTIC_RERANK_MAX_TOKENS = max(120,min(700,int(os.getenv("MODEL_SEMANTIC_RERANK_MAX_TOKENS","360"))))
DEEP_RULE_ENGINE_ENABLED = os.getenv("DEEP_RULE_ENGINE_ENABLED", "true").lower() in {"1","true","yes","on"}
AI_REASONING_ENABLED = os.getenv("AI_REASONING_ENABLED", "true").lower() in {"1","true","yes","on"}
AI_REASONING_EFFORT_STANDARD = os.getenv("AI_REASONING_EFFORT_STANDARD", "medium").strip().lower() or "medium"
AI_REASONING_EFFORT_ADVANCED = os.getenv("AI_REASONING_EFFORT_ADVANCED", "high").strip().lower() or "high"
AI_REASONING_PARAMETER = os.getenv("AI_REASONING_PARAMETER", "auto").strip().lower() or "auto"
AI_REASONING_SEND_TEMPERATURE = os.getenv("AI_REASONING_SEND_TEMPERATURE", "false").lower() in {"1","true","yes","on"}
AI_PROVIDER_CAPABILITIES_JSON = os.getenv("AI_PROVIDER_CAPABILITIES_JSON", "{}").strip() or "{}"
DEEP_ANSWER_VERIFICATION_ENABLED = os.getenv("DEEP_ANSWER_VERIFICATION_ENABLED", "true").lower() in {"1","true","yes","on"}
DEEP_CONFIDENCE_GATE_ENABLED = os.getenv("DEEP_CONFIDENCE_GATE_ENABLED", "true").lower() in {"1","true","yes","on"}
DEEP_CONFIDENCE_MIN = min(0.95, max(0.20, float(os.getenv("DEEP_CONFIDENCE_MIN", "0.48"))))
DEEP_CONFIDENCE_COMPLEX_MIN = min(0.98, max(DEEP_CONFIDENCE_MIN, float(os.getenv("DEEP_CONFIDENCE_COMPLEX_MIN", "0.56"))))

# R29.7 operational engines: deterministic cargo/calculation and Neshan routing.
NESHAN_API_KEY = os.getenv("NESHAN_API_KEY", "").strip()
NESHAN_GEOCODING_URL = os.getenv("NESHAN_GEOCODING_URL", "https://api.neshan.org/geocoding/v1").strip().rstrip("/")
NESHAN_SEARCH_URL = os.getenv("NESHAN_SEARCH_URL", "https://api.neshan.org/v3/search").strip().rstrip("/")
NESHAN_USE_PLUS = os.getenv("NESHAN_USE_PLUS", "true").lower() in {"1","true","yes","on"}
NESHAN_SEARCH_ENRICHMENT = os.getenv("NESHAN_SEARCH_ENRICHMENT", "true").lower() in {"1","true","yes","on"}
NESHAN_TIMEOUT_SECONDS = max(3.0, min(30.0, float(os.getenv("NESHAN_TIMEOUT_SECONDS", "10"))))
NESHAN_MAX_RESULTS = max(1, min(10, int(os.getenv("NESHAN_MAX_RESULTS", "5"))))
NESHAN_CACHE_TTL_DAYS = max(1, min(365, int(os.getenv("NESHAN_CACHE_TTL_DAYS", "30"))))
CARGO_VEHICLE_PROFILES_JSON = os.getenv("CARGO_VEHICLE_PROFILES_JSON", "{}").strip() or "{}"

# Estimated provider credit/cost. These values are configured by the owner because
# OpenAI-compatible vendors do not expose one common balance endpoint.
AI_INPUT_COST_PER_1M = [float(os.getenv(f"AI_INPUT_COST_PER_1M_{i}", "0")) for i in range(1,7)]
AI_OUTPUT_COST_PER_1M = [float(os.getenv(f"AI_OUTPUT_COST_PER_1M_{i}", "0")) for i in range(1,7)]
AI_CREDIT_AMOUNT = [float(os.getenv(f"AI_CREDIT_AMOUNT_{i}", "0")) for i in range(1,7)]
AI_CREDIT_CURRENCY = [os.getenv(f"AI_CREDIT_CURRENCY_{i}", "IRR").strip() or "IRR" for i in range(1,7)]

# Persistent abuse protection and account quotas.
REQUESTS_PER_MINUTE = max(1, int(os.getenv("REQUESTS_PER_MINUTE", "30")))
GUEST_REQUESTS_PER_MINUTE = max(1, int(os.getenv("GUEST_REQUESTS_PER_MINUTE", str(REQUESTS_PER_MINUTE))))
GUEST_DAILY_LIMIT = max(0, int(os.getenv("GUEST_DAILY_LIMIT", "100")))
GUEST_MONTHLY_LIMIT = max(0, int(os.getenv("GUEST_MONTHLY_LIMIT", "1000")))
INTEGRATION_DAILY_LIMIT = max(0, int(os.getenv("INTEGRATION_DAILY_LIMIT", "5000")))
RATE_LIMIT_RETENTION_DAYS = max(2, int(os.getenv("RATE_LIMIT_RETENTION_DAYS", "45")))
SQLITE_CACHE_KB = max(8000, min(131072, int(os.getenv("SQLITE_CACHE_KB", "64000"))))
SQLITE_MMAP_MB = max(0, min(512, int(os.getenv("SQLITE_MMAP_MB", "128"))))

# Approved cache has no time expiry and is invalidated only by knowledge/access changes.
# Self-analysis is kept in a separate temporary cache for at most 24 hours.
TEMPORARY_CACHE_HOURS = max(1, min(168, int(os.getenv("TEMPORARY_CACHE_HOURS", "24"))))
APPROVED_CACHE_MAX_ROWS = max(100, int(os.getenv("APPROVED_CACHE_MAX_ROWS", "10000")))
TEMPORARY_CACHE_MAX_ROWS = max(50, int(os.getenv("TEMPORARY_CACHE_MAX_ROWS", "2000")))

# Local hybrid semantic index. It adds zero-token semantic fingerprints to FTS/BM25.
SEMANTIC_SEARCH_ENABLED = os.getenv("SEMANTIC_SEARCH_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
SEMANTIC_BUCKET_LIMIT = max(8, min(64, int(os.getenv("SEMANTIC_BUCKET_LIMIT", "28"))))
SEMANTIC_CANDIDATE_LIMIT = max(100, int(os.getenv("SEMANTIC_CANDIDATE_LIMIT", "800")))
RETRIEVAL_HIGH_CONFIDENCE = min(1.5, max(0.1, float(os.getenv("RETRIEVAL_HIGH_CONFIDENCE", "0.62"))))
RETRIEVAL_MEDIUM_CONFIDENCE = min(RETRIEVAL_HIGH_CONFIDENCE, max(0.05, float(os.getenv("RETRIEVAL_MEDIUM_CONFIDENCE", "0.30"))))

# R29 real semantic retrieval. OpenRouter slot 3 is the default because it exposes
# an OpenAI-compatible /embeddings endpoint. If unavailable, the local semantic
# index remains active so chat never depends on embeddings being online.
REMOTE_EMBEDDING_ENABLED = os.getenv("REMOTE_EMBEDDING_ENABLED", "true").lower() in {"1","true","yes","on"}
EMBEDDING_API_SLOT = max(1, min(MAX_AI_API_SLOTS, int(os.getenv("EMBEDDING_API_SLOT", "3"))))
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "openai/text-embedding-3-small").strip()
EMBEDDING_DIMENSIONS = max(128, min(1536, int(os.getenv("EMBEDDING_DIMENSIONS", "384"))))
EMBEDDING_BATCH_SIZE = max(1, min(96, int(os.getenv("EMBEDDING_BATCH_SIZE", "32"))))
EMBEDDING_MAX_CHARS = max(500, min(12000, int(os.getenv("EMBEDDING_MAX_CHARS", "6000"))))
EMBEDDING_SCAN_LIMIT = max(1000, min(50000, int(os.getenv("EMBEDDING_SCAN_LIMIT", "20000"))))
EMBEDDING_TIMEOUT_SECONDS = max(15.0, float(os.getenv("EMBEDDING_TIMEOUT_SECONDS", "75")))
EMBEDDING_QUERY_SKIP_CONFIDENCE = max(0.35, min(1.2, float(os.getenv("EMBEDDING_QUERY_SKIP_CONFIDENCE", "0.52"))))

# R30 latency engine: exact-answer fast path, authoritative retrieval memoization,
# and adaptive provider ordering. Knowledge-version is part of every cache key, so
# manager training remains authoritative and any knowledge mutation invalidates hits.
FAST_EXACT_CACHE_ENABLED = os.getenv("FAST_EXACT_CACHE_ENABLED", "true").lower() in {"1","true","yes","on"}
RETRIEVAL_STAGE_CACHE_TTL_SECONDS = max(15, min(3600, int(os.getenv("RETRIEVAL_STAGE_CACHE_TTL_SECONDS", "300"))))
RETRIEVAL_STAGE_CACHE_MAX_ENTRIES = max(64, min(5000, int(os.getenv("RETRIEVAL_STAGE_CACHE_MAX_ENTRIES", "800"))))
PROVIDER_SPEED_ROUTING_ENABLED = os.getenv("PROVIDER_SPEED_ROUTING_ENABLED", "true").lower() in {"1","true","yes","on"}
PROVIDER_SPEED_CACHE_SECONDS = max(5, min(300, int(os.getenv("PROVIDER_SPEED_CACHE_SECONDS", "30"))))
PROVIDER_SPEED_MIN_SAMPLES = max(1, min(20, int(os.getenv("PROVIDER_SPEED_MIN_SAMPLES", "2"))))
PDF_VISION_BALANCE_PROVIDERS = os.getenv("PDF_VISION_BALANCE_PROVIDERS", "true").lower() in {"1","true","yes","on"}
PDF_VISION_DELTA_MODE = os.getenv("PDF_VISION_DELTA_MODE", "true").lower() in {"1","true","yes","on"}
PDF_VISION_DELTA_TEXT_THRESHOLD = max(250, min(5000, int(os.getenv("PDF_VISION_DELTA_TEXT_THRESHOLD", "700"))))
SOURCE_VERIFICATION_FAST_GUARD = os.getenv("SOURCE_VERIFICATION_FAST_GUARD", "true").lower() in {"1","true","yes","on"}

HYBRID_RERANK_ENABLED = os.getenv("HYBRID_RERANK_ENABLED", "true").lower() in {"1","true","yes","on"}
HYBRID_RERANK_TOP_N = max(max(1,int(os.getenv("RETRIEVAL_TOP_K","12"))), min(80, int(os.getenv("HYBRID_RERANK_TOP_N", "32"))))
SOURCE_FACT_EXTRACTION_ENABLED = os.getenv("SOURCE_FACT_EXTRACTION_ENABLED", "true").lower() in {"1","true","yes","on"}
FACT_RETRIEVAL_LIMIT = max(10, min(200, int(os.getenv("FACT_RETRIEVAL_LIMIT", "60"))))
RETRIEVAL_DEBUG_MAX_ITEMS = max(5, min(100, int(os.getenv("RETRIEVAL_DEBUG_MAX_ITEMS", "30"))))

# Controlled conversational memory; used only for follow-up wording.
CONVERSATION_MEMORY_ENABLED = os.getenv("CONVERSATION_MEMORY_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
CONVERSATION_MEMORY_MESSAGES = max(0, min(8, int(os.getenv("CONVERSATION_MEMORY_MESSAGES", "4"))))
CONVERSATION_MEMORY_MAX_CHARS = max(300, min(4000, int(os.getenv("CONVERSATION_MEMORY_MAX_CHARS", "1400"))))

# Health monitoring and backup limits.
HEALTH_MONITOR_ENABLED = os.getenv("HEALTH_MONITOR_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
HEALTH_MONITOR_INTERVAL_SECONDS = max(60, int(os.getenv("HEALTH_MONITOR_INTERVAL_SECONDS", "300")))
HEALTH_DISK_WARNING_PERCENT = min(99, max(50, int(os.getenv("HEALTH_DISK_WARNING_PERCENT", "85"))))
HEALTH_ERROR_RATE_WARNING_PERCENT = min(100, max(1, int(os.getenv("HEALTH_ERROR_RATE_WARNING_PERCENT", "25"))))
BACKUP_MAX_MB = max(20, int(os.getenv("BACKUP_MAX_MB", str(max(2048, MAX_UPLOAD_MB * 2)))))
BACKUP_DIR = Path(_railway_path_env("BACKUP_DIR", "/data/backups", "backups"))
AUTO_BACKUP_ENABLED = os.getenv("AUTO_BACKUP_ENABLED", "true").lower() in {"1","true","yes","on"}
AUTO_BACKUP_INTERVAL_HOURS = max(1, int(os.getenv("AUTO_BACKUP_INTERVAL_HOURS", "24")))
AUTO_BACKUP_RETENTION = max(1, min(30, int(os.getenv("AUTO_BACKUP_RETENTION", "7"))))
BACKUP_OVERDUE_WARNING_HOURS = max(AUTO_BACKUP_INTERVAL_HOURS, int(os.getenv("BACKUP_OVERDUE_WARNING_HOURS", "36")))
BACKGROUND_DOCUMENT_PROCESSING = os.getenv("BACKGROUND_DOCUMENT_PROCESSING", "true").lower() in {"1","true","yes","on"}
DOCUMENT_JOB_WORKER_ENABLED = os.getenv("DOCUMENT_JOB_WORKER_ENABLED", "true").lower() in {"1","true","yes","on"}
DOCUMENT_JOB_POLL_SECONDS = max(0.5, min(10.0, float(os.getenv("DOCUMENT_JOB_POLL_SECONDS", "1.5"))))
DOCUMENT_JOB_MAX_ATTEMPTS = max(1, min(8, int(os.getenv("DOCUMENT_JOB_MAX_ATTEMPTS", "3"))))
DOCUMENT_JOB_STALE_SECONDS = max(60, int(os.getenv("DOCUMENT_JOB_STALE_SECONDS", "900")))
GOLDEN_EVAL_ENABLED = os.getenv("GOLDEN_EVAL_ENABLED", "true").lower() in {"1","true","yes","on"}
GOLDEN_EVAL_MIN_PASS_PERCENT = max(0.0, min(100.0, float(os.getenv("GOLDEN_EVAL_MIN_PASS_PERCENT", "90"))))
GOLDEN_MIN_TOKEN_RECALL = max(0.1, min(1.0, float(os.getenv("GOLDEN_MIN_TOKEN_RECALL", "0.55"))))
GOLDEN_MAX_CASES = max(10, min(2000, int(os.getenv("GOLDEN_MAX_CASES", "500"))))
CONFIDENCE_CALIBRATION_ENABLED = os.getenv("CONFIDENCE_CALIBRATION_ENABLED","true").lower() in {"1","true","yes","on"}
CONFIDENCE_CALIBRATION_MIN_CASES = max(8,min(100,int(os.getenv("CONFIDENCE_CALIBRATION_MIN_CASES","12"))))
CONFIDENCE_CALIBRATION_MAX_SHIFT = max(0.02,min(0.20,float(os.getenv("CONFIDENCE_CALIBRATION_MAX_SHIFT","0.10"))))
SQLITE_SINGLE_REPLICA_MODE = os.getenv("SQLITE_SINGLE_REPLICA_MODE", "true").lower() in {"1","true","yes","on"}
SQLITE_REPLICA_LOCK_ENABLED = os.getenv("SQLITE_REPLICA_LOCK_ENABLED", "true").lower() in {"1","true","yes","on"}
ORIGIN_GUARD_ENABLED = os.getenv("ORIGIN_GUARD_ENABLED", "true").lower() in {"1","true","yes","on"}
CONTENT_SECURITY_POLICY_ENABLED = os.getenv("CONTENT_SECURITY_POLICY_ENABLED", "true").lower() in {"1","true","yes","on"}

# FAQ matching is always local and consumes zero model tokens.
FAQ_FUZZY_THRESHOLD = min(0.99, max(0.65, float(os.getenv("FAQ_FUZZY_THRESHOLD", "0.84"))))
FAQ_MAX_ROWS = max(100, int(os.getenv("FAQ_MAX_ROWS", "5000")))
TRAINING_STAGE_MIN_SCORE = min(1.5, max(0.05, float(os.getenv("TRAINING_STAGE_MIN_SCORE", "0.22"))))
TRAINING_RESCUE_SCAN_LIMIT = max(500, min(20000, int(os.getenv("TRAINING_RESCUE_SCAN_LIMIT", "5000"))))
SELF_ANALYSIS_ENABLED = os.getenv("SELF_ANALYSIS_ENABLED", "true").lower() in {"1", "true", "yes", "on"}

# Direct Gemini compatibility remains supported. The primary path is the OpenAI-compatible API pool.
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
GEMINI_TIMEOUT_SECONDS = float(os.getenv("GEMINI_TIMEOUT_SECONDS", "120"))
GEMINI_MAX_RETRIES = max(0, int(os.getenv("GEMINI_MAX_RETRIES", "3")))
PUBLIC_REGISTRATION_ENABLED = os.getenv("PUBLIC_REGISTRATION_ENABLED", "false").lower() in {"1", "true", "yes", "on"}
GUEST_CHAT_ENABLED = os.getenv("GUEST_CHAT_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
RETRIEVAL_TOP_K = max(1, int(os.getenv("RETRIEVAL_TOP_K", "12")))
RETRIEVAL_TOTAL_ITEMS = max(RETRIEVAL_TOP_K, int(os.getenv("RETRIEVAL_TOTAL_ITEMS", "16")))
RETRIEVAL_MIN_SCORE = float(os.getenv("RETRIEVAL_MIN_SCORE", "0.08"))
RETRIEVAL_NEIGHBOR_CHUNKS = max(0, min(4, int(os.getenv("RETRIEVAL_NEIGHBOR_CHUNKS", "3"))))
RETRIEVAL_MAX_CONTEXT_CHARS = max(8000, int(os.getenv("RETRIEVAL_MAX_CONTEXT_CHARS", "36000")))
RETRIEVAL_SCAN_LIMIT = max(1000, int(os.getenv("RETRIEVAL_SCAN_LIMIT", "8000")))
RETRIEVAL_CHUNK_SIZE = max(900, int(os.getenv("RETRIEVAL_CHUNK_SIZE", "1700")))
RETRIEVAL_CHUNK_OVERLAP = max(120, min(RETRIEVAL_CHUNK_SIZE // 2, int(os.getenv("RETRIEVAL_CHUNK_OVERLAP", "320"))))
MIN_EXTRACTED_TEXT_CHARS = max(10, int(os.getenv("MIN_EXTRACTED_TEXT_CHARS", "40")))

# PDF visual indexing. Pages containing images (or too little extractable text) are
# rendered and sent to a vision-capable API slot once during indexing. The extracted
# text is then stored beside the normal PDF text and used by the same local RAG index.
PDF_VISION_ENABLED = os.getenv("PDF_VISION_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
PDF_VISION_MAX_PAGES = max(1, min(500, int(os.getenv("PDF_VISION_MAX_PAGES", "500"))))
PDF_VISION_CONCURRENCY = max(1, min(8, int(os.getenv("PDF_VISION_CONCURRENCY", "4"))))
PDF_VISION_SCAN_ALL_PAGES = os.getenv("PDF_VISION_SCAN_ALL_PAGES", "true").lower() in {"1", "true", "yes", "on"}
PDF_VISION_SCAN_IMAGE_PAGES = os.getenv("PDF_VISION_SCAN_IMAGE_PAGES", "true").lower() in {"1","true","yes","on"}
PDF_VISION_SCAN_TABLE_LIKE_PAGES = os.getenv("PDF_VISION_SCAN_TABLE_LIKE_PAGES", "true").lower() in {"1","true","yes","on"}
SOURCE_INGESTION_V3 = os.getenv("SOURCE_INGESTION_V3", os.getenv("SOURCE_INGESTION_V2", "true")).lower() in {"1", "true", "yes", "on"}
SOURCE_INGESTION_V2 = SOURCE_INGESTION_V3  # backward-compatible alias for older Railway Variables
SOURCE_INCLUDE_PARTIAL_DOCUMENTS = os.getenv("SOURCE_INCLUDE_PARTIAL_DOCUMENTS", "false").lower() in {"1", "true", "yes", "on"}
SOURCE_MIN_QUALITY_PCT = max(0.0, min(100.0, float(os.getenv("SOURCE_MIN_QUALITY_PCT", "92"))))
SOURCE_MIN_PAGE_FIDELITY = max(0.0, min(1.0, float(os.getenv("SOURCE_MIN_PAGE_FIDELITY", "0.82"))))
SOURCE_MIN_NUMERIC_AGREEMENT = max(0.0, min(1.0, float(os.getenv("SOURCE_MIN_NUMERIC_AGREEMENT", "0.72"))))
AUTO_REINDEX_LEGACY_SOURCES = os.getenv('AUTO_REINDEX_LEGACY_SOURCES','true').lower() in {'1','true','yes','on'}
AUTO_REINDEX_LEGACY_LIMIT = max(1,min(500,int(os.getenv('AUTO_REINDEX_LEGACY_LIMIT','100'))))
SOURCE_PAGE_STORAGE_ENABLED = os.getenv("SOURCE_PAGE_STORAGE_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
SOURCE_STRUCTURE_AWARE_INDEXING = os.getenv("SOURCE_STRUCTURE_AWARE_INDEXING", "true").lower() in {"1", "true", "yes", "on"}
SOURCE_PAGE_BY_PAGE_STRICT = os.getenv("SOURCE_PAGE_BY_PAGE_STRICT", "true").lower() in {"1", "true", "yes", "on"}
SOURCE_NUMERIC_FACT_BOOST = max(0.0, min(0.5, float(os.getenv("SOURCE_NUMERIC_FACT_BOOST", "0.14"))))
PDF_VISION_DPI = max(96, min(240, int(os.getenv("PDF_VISION_DPI", "180"))))
PDF_VISION_MIN_TEXT_CHARS = max(0, int(os.getenv("PDF_VISION_MIN_TEXT_CHARS", "100")))
PDF_VISION_MAX_RENDER_MB = max(1, int(os.getenv("PDF_VISION_MAX_RENDER_MB", "5")))
PDF_VISION_MAX_TOKENS = max(500, int(os.getenv("PDF_VISION_MAX_TOKENS", "2400")))
SOURCE_FIRST_STRICT = os.getenv("SOURCE_FIRST_STRICT", "true").lower() in {"1", "true", "yes", "on"}
SOURCE_ANSWER_VERIFICATION = os.getenv("SOURCE_ANSWER_VERIFICATION", "true").lower() in {"1", "true", "yes", "on"}
SOURCE_VERIFICATION_MAX_TOKENS = max(250, int(os.getenv("SOURCE_VERIFICATION_MAX_TOKENS", "700")))

# R32 built-in Barsan booklets: immediate local pre-index + page-by-page Vision enrichment.
BUILTIN_SOURCE_DIR = Path(__file__).resolve().parent / 'builtin_sources'
BUILTIN_SOURCE_PREINDEX = BUILTIN_SOURCE_DIR / 'preindex.json'
BUILTIN_SOURCE_MANIFEST = BUILTIN_SOURCE_DIR / 'source_manifest.json'
BUILTIN_SOURCE_AUTO_INSTALL = os.getenv('BUILTIN_SOURCE_AUTO_INSTALL','true').lower() in {'1','true','yes','on'}
BUILTIN_SOURCE_AUTO_ENRICH = os.getenv('BUILTIN_SOURCE_AUTO_ENRICH','false').lower() in {'1','true','yes','on'}
BUILTIN_SOURCE_GLOBAL_DEFAULT = os.getenv('BUILTIN_SOURCE_GLOBAL_DEFAULT','true').lower() in {'1','true','yes','on'}

# Unified chat: the chat panel answers from all organizational knowledge directly.
CHAT_UNIFIED_KNOWLEDGE_MODE = os.getenv("CHAT_UNIFIED_KNOWLEDGE_MODE", "true").lower() in {"1", "true", "yes", "on"}
CHAT_DISABLE_SECTION_REDIRECTS = os.getenv("CHAT_DISABLE_SECTION_REDIRECTS", "true").lower() in {"1", "true", "yes", "on"}
CHAT_CORPUS_RESCUE_ENABLED = os.getenv("CHAT_CORPUS_RESCUE_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
CHAT_CORPUS_RESCUE_LIMIT = max(500, min(20000, int(os.getenv("CHAT_CORPUS_RESCUE_LIMIT", "8000"))))
CHAT_REDIRECT_RETRY_MAX_TOKENS = max(250, int(os.getenv("CHAT_REDIRECT_RETRY_MAX_TOKENS", "650")))

# Voice-to-text for FAQ and training. Browser speech recognition is used first;
# uploaded audio can fall back to any API slot that has a transcription model.
TRANSCRIPTION_MAX_MB = max(1, min(50, int(os.getenv("TRANSCRIPTION_MAX_MB", "20"))))
AI_TRANSCRIPTION_MODEL_1 = os.getenv("AI_TRANSCRIPTION_MODEL_1", "").strip()
AI_TRANSCRIPTION_MODEL_2 = os.getenv("AI_TRANSCRIPTION_MODEL_2", "").strip()
AI_TRANSCRIPTION_MODEL_3 = os.getenv("AI_TRANSCRIPTION_MODEL_3", "").strip()
AI_TRANSCRIPTION_MODEL_4 = os.getenv("AI_TRANSCRIPTION_MODEL_4", "").strip()
AI_TRANSCRIPTION_MODEL_5 = os.getenv("AI_TRANSCRIPTION_MODEL_5", "").strip()
AI_TRANSCRIPTION_MODEL_6 = os.getenv("AI_TRANSCRIPTION_MODEL_6", "").strip()
AI_VISION_MODEL_1 = os.getenv("AI_VISION_MODEL_1", "").strip()
AI_VISION_MODEL_2 = os.getenv("AI_VISION_MODEL_2", "").strip()
AI_VISION_MODEL_3 = os.getenv("AI_VISION_MODEL_3", "").strip()
AI_VISION_MODEL_4 = os.getenv("AI_VISION_MODEL_4", "").strip()
AI_VISION_MODEL_5 = os.getenv("AI_VISION_MODEL_5", "").strip()
AI_VISION_MODEL_6 = os.getenv("AI_VISION_MODEL_6", "").strip()
_configured_origins = [x.strip().rstrip("/") for x in os.getenv("ALLOWED_ORIGINS", "").split(",") if x.strip()]
ALLOWED_ORIGINS = list(dict.fromkeys([*(_configured_origins or [PUBLIC_BASE_URL]), PUBLIC_BASE_URL]))

def _normalized_origin(value: str) -> str:
    raw=str(value or '').strip().rstrip('/')
    parsed=urlparse(raw)
    if parsed.scheme not in {'http','https'} or not parsed.netloc:
        return ''
    return f"{parsed.scheme}://{parsed.netloc}"

_widget_origin_values=[x.strip() for x in os.getenv('WIDGET_ALLOWED_ORIGINS','').split(',') if x.strip()]
WIDGET_ALLOWED_ORIGINS=list(dict.fromkeys(filter(None,(_normalized_origin(x) for x in (_widget_origin_values or ALLOWED_ORIGINS)))))
INTEGRATION_API_KEYS_RAW = os.getenv("INTEGRATION_API_KEYS", "{}")
LOGIN_MAX_ATTEMPTS = int(os.getenv("LOGIN_MAX_ATTEMPTS", "5"))
LOGIN_LOCK_MINUTES = int(os.getenv("LOGIN_LOCK_MINUTES", "15"))
LOGIN_WINDOW_MINUTES = int(os.getenv("LOGIN_WINDOW_MINUTES", "15"))
LOCAL_TIMEZONE_OFFSET_MINUTES = int(os.getenv("LOCAL_TIMEZONE_OFFSET_MINUTES", "210"))
EXPORT_ROW_LIMIT = int(os.getenv("EXPORT_ROW_LIMIT", "100000"))
MAX_ARCHIVE_EXPANDED_MB = int(os.getenv("MAX_ARCHIVE_EXPANDED_MB", "100"))
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "500"))

# Compatibility-only memory bucket. Enforcement is persisted in SQLite; clearing this
# object must not reset limits (covered by the release test).
RATE_BUCKETS: dict[str, list[float]] = {}
_HTTP_CLIENT: httpx.AsyncClient | None = None
_HEALTH_TASK: asyncio.Task[Any] | None = None
_DOCUMENT_JOB_TASK: asyncio.Task[Any] | None = None
_REPLICA_LOCK_HANDLE: Any | None = None
_QUERY_EMBED_CACHE: dict[str, tuple[float,list[float]]] = {}
_RETRIEVAL_STAGE_CACHE: dict[str, tuple[float,str,list[dict[str,Any]]]] = {}
_RETRIEVAL_STAGE_CACHE_LOCK = threading.RLock()
_PROVIDER_SPEED_CACHE: dict[str, tuple[float,dict[int,tuple[int,float]]]] = {}
_PROVIDER_SPEED_CACHE_LOCK = threading.RLock()
_BACKUP_LOCK = threading.RLock()
_MAINTENANCE_LOCK = threading.RLock()
_RESTORE_IN_PROGRESS = threading.Event()
_UPLOAD_SESSION_LOCKS: dict[str,threading.Lock] = {}
_CONFIDENCE_CALIBRATION_CACHE: tuple[float,float|None,dict[str,Any]] = (0.0,None,{})
_CONFIDENCE_CALIBRATION_LOCK = threading.RLock()


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def db_path() -> Path:
    if not DATABASE_URL.startswith("sqlite:///"):
        raise RuntimeError("این نسخه فقط SQLite را پشتیبانی می‌کند.")
    return Path(DATABASE_URL.removeprefix("sqlite:///"))

DB_PATH = db_path()


OPENAI_COMPATIBLE_PROVIDERS = {"openai_compatible", "gapgpt", "openai", "custom"}

_PROVIDER_KEY_PREFIX='fernet:v1:'

def _provider_fernet() -> Fernet:
    secret=PROVIDER_KEY_ENCRYPTION_SECRET or JWT_SECRET
    if not secret:
        raise RuntimeError('PROVIDER_KEY_ENCRYPTION_SECRET یا JWT_SECRET برای محافظت از کلیدهای API لازم است.')
    key=base64.urlsafe_b64encode(hashlib.sha256(secret.encode('utf-8')).digest())
    return Fernet(key)


def _protect_api_key(value: str) -> str:
    raw=str(value or '').strip()
    if not raw or raw.startswith(_PROVIDER_KEY_PREFIX):
        return raw
    token=_provider_fernet().encrypt(raw.encode('utf-8')).decode('ascii')
    return _PROVIDER_KEY_PREFIX+token


def _unprotect_api_key(value: str) -> str:
    stored=str(value or '').strip()
    if not stored:
        return ''
    if not stored.startswith(_PROVIDER_KEY_PREFIX):
        # Backward-compatible plaintext read; ensure_schema migrates it in-place.
        return stored
    token=stored[len(_PROVIDER_KEY_PREFIX):]
    try:
        return _provider_fernet().decrypt(token.encode('ascii')).decode('utf-8')
    except (InvalidToken,UnicodeDecodeError,binascii.Error) as exc:
        raise RuntimeError('کلید API ذخیره‌شده قابل رمزگشایی نیست؛ PROVIDER_KEY_ENCRYPTION_SECRET را بررسی کنید.') from exc


def _slot_model_variants(slot: int, fallback: str) -> dict[str, str]:
    return {
        'economy': os.getenv(f'AI_MODEL_ECONOMY_{slot}', fallback).strip() or fallback,
        'standard': os.getenv(f'AI_MODEL_STANDARD_{slot}', fallback).strip() or fallback,
        'advanced': os.getenv(f'AI_MODEL_ADVANCED_{slot}', fallback).strip() or fallback,
    }


def _static_ai_slots() -> list[dict[str, Any]]:
    keys = [AI_API_KEY_1 or AI_API_KEY, AI_API_KEY_2, AI_API_KEY_3, AI_API_KEY_4, AI_API_KEY_5, AI_API_KEY_6]
    bases = [AI_BASE_URL_1 or AI_BASE_URL, AI_BASE_URL_2 or AI_BASE_URL, AI_BASE_URL_3 or AI_BASE_URL, AI_BASE_URL_4 or AI_BASE_URL, AI_BASE_URL_5 or AI_BASE_URL, AI_BASE_URL_6 or AI_BASE_URL]
    models = [AI_MODEL_1 or AI_MODEL, AI_MODEL_2 or AI_MODEL, AI_MODEL_3 or AI_MODEL, AI_MODEL_4 or AI_MODEL, AI_MODEL_5 or AI_MODEL, AI_MODEL_6 or AI_MODEL]
    labels = [AI_PROVIDER_LABEL_1 or AI_PROVIDER_LABEL, AI_PROVIDER_LABEL_2 or AI_PROVIDER_LABEL, AI_PROVIDER_LABEL_3 or AI_PROVIDER_LABEL, AI_PROVIDER_LABEL_4 or AI_PROVIDER_LABEL, AI_PROVIDER_LABEL_5 or AI_PROVIDER_LABEL, AI_PROVIDER_LABEL_6 or AI_PROVIDER_LABEL]
    vision_models=[AI_VISION_MODEL_1,AI_VISION_MODEL_2,AI_VISION_MODEL_3,AI_VISION_MODEL_4,AI_VISION_MODEL_5,AI_VISION_MODEL_6]
    voice_models=[AI_TRANSCRIPTION_MODEL_1,AI_TRANSCRIPTION_MODEL_2,AI_TRANSCRIPTION_MODEL_3,AI_TRANSCRIPTION_MODEL_4,AI_TRANSCRIPTION_MODEL_5,AI_TRANSCRIPTION_MODEL_6]
    slots: list[dict[str, Any]] = []
    for index, key in enumerate(keys, 1):
        if not key:
            continue
        base_model = models[index-1]
        slots.append({
            "slot": index, "api_key": key, "base_url": bases[index-1].rstrip('/'),
            "model": base_model, "models": _slot_model_variants(index, base_model),
            "label": labels[index-1] or f"API {index}",
            "vision_model":vision_models[index-1],"transcription_model":voice_models[index-1],
            "embedding_model": EMBEDDING_MODEL if index==EMBEDDING_API_SLOT else '',
            "input_cost_per_1m": AI_INPUT_COST_PER_1M[index-1],
            "output_cost_per_1m": AI_OUTPUT_COST_PER_1M[index-1],
            "credit_amount": AI_CREDIT_AMOUNT[index-1],
            "credit_currency": AI_CREDIT_CURRENCY[index-1],
            "managed_by":"variables","enabled":True,
        })
    return slots

def _dynamic_ai_slots() -> list[dict[str, Any]]:
    if not DB_PATH.exists():
        return []
    try:
        with get_db() as db:
            exists=db.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='ai_provider_configs'").fetchone()
            if not exists:
                return []
            rows=db.execute("SELECT * FROM ai_provider_configs WHERE enabled=1 ORDER BY slot").fetchall()
    except sqlite3.Error:
        return []
    result=[]
    for row in rows:
        item=dict(row);slot=int(item['slot'])
        if slot<1 or slot>MAX_AI_API_SLOTS or not item.get('api_key'):
            continue
        try:
            decrypted_key=_unprotect_api_key(str(item.get('api_key') or ''))
        except RuntimeError:
            continue
        if not decrypted_key:
            continue
        model=str(item.get('model') or '').strip()
        result.append({
            'slot':slot,'api_key':decrypted_key,
            'base_url':str(item.get('base_url') or '').strip().rstrip('/'),'model':model,
            'models':{
                'economy':str(item.get('model_economy') or model).strip() or model,
                'standard':str(item.get('model_standard') or model).strip() or model,
                'advanced':str(item.get('model_advanced') or model).strip() or model,
            },
            'label':str(item.get('label') or f'API {slot}').strip(),
            'vision_model':str(item.get('vision_model') or '').strip(),
            'transcription_model':str(item.get('transcription_model') or '').strip(),
            'embedding_model':str(item.get('embedding_model') or '').strip(),
            'input_cost_per_1m':float(item.get('input_cost_per_1m') or 0),
            'output_cost_per_1m':float(item.get('output_cost_per_1m') or 0),
            'credit_amount':float(item.get('credit_amount') or 0),
            'credit_currency':str(item.get('credit_currency') or 'USD').strip() or 'USD',
            'managed_by':'admin','enabled':True,
        })
    return result

def configured_ai_slots() -> list[dict[str, Any]]:
    unique={int(item['slot']):item for item in (_dynamic_ai_slots()+_static_ai_slots())}
    return [unique[k] for k in sorted(unique)]

def _provider_config_for_slot(slot: int) -> dict[str,Any] | None:
    return next((item for item in configured_ai_slots() if int(item['slot'])==int(slot)),None)


def _provider_capability_overrides() -> dict[str,Any]:
    return parse_capability_overrides(AI_PROVIDER_CAPABILITIES_JSON)


def provider_capabilities(slot: dict[str,Any], model_name: str | None = None) -> dict[str,Any]:
    """Return explicit/inferred provider capabilities from the isolated R35 registry."""
    return infer_provider_capabilities(slot, model_name, overrides=_provider_capability_overrides())


def active_model_name() -> str:
    if AI_PROVIDER in OPENAI_COMPATIBLE_PROVIDERS:
        slots = configured_ai_slots()
        if not slots:
            return AI_MODEL
        try:
            active = _active_api_slot_number() if '_active_api_slot_number' in globals() else 1
            return next((item["model"] for item in slots if int(item["slot"]) == active), slots[0]["model"])
        except Exception:
            return slots[0]["model"]
    if AI_PROVIDER == "gemini":
        return GEMINI_MODEL
    return "local-fallback"


def active_max_completion_tokens(detailed: bool = False) -> int:
    requested = AI_DETAILED_MAX_COMPLETION_TOKENS if detailed else AI_DEFAULT_MAX_COMPLETION_TOKENS
    if AI_PROVIDER in OPENAI_COMPATIBLE_PROVIDERS:
        return min(AI_MAX_COMPLETION_TOKENS, requested)
    return requested


def active_temperature() -> float:
    if AI_PROVIDER in OPENAI_COMPATIBLE_PROVIDERS:
        return AI_TEMPERATURE
    return AI_TEMPERATURE


def validate_environment() -> None:
    allowed = OPENAI_COMPATIBLE_PROVIDERS | {"gemini", "local"}
    if AI_PROVIDER not in allowed:
        raise RuntimeError("AI_PROVIDER باید openai_compatible، gapgpt، gemini یا local باشد.")
    if AI_TOKEN_PARAMETER not in {"auto", "max_tokens", "max_completion_tokens"}:
        raise RuntimeError("AI_TOKEN_PARAMETER باید auto، max_tokens یا max_completion_tokens باشد.")
    try:
        extra_headers = json.loads(AI_EXTRA_HEADERS_JSON)
        if not isinstance(extra_headers, dict):
            raise ValueError
    except (json.JSONDecodeError, ValueError) as exc:
        raise RuntimeError("AI_EXTRA_HEADERS_JSON باید یک شیء JSON معتبر باشد.") from exc
    if ENVIRONMENT == "production":
        if len(JWT_SECRET) < 32:
            raise RuntimeError("JWT_SECRET باید حداقل 32 کاراکتر باشد.")
        if len(PROVIDER_KEY_ENCRYPTION_SECRET or JWT_SECRET) < 32:
            raise RuntimeError("PROVIDER_KEY_ENCRYPTION_SECRET باید حداقل 32 کاراکتر باشد.")
        if len(INITIAL_ADMIN_PASSWORD) < 8:
            raise RuntimeError("INITIAL_ADMIN_PASSWORD باید حداقل 8 کاراکتر باشد.")
        if AI_PROVIDER in OPENAI_COMPATIBLE_PROVIDERS:
            slots = configured_ai_slots()
            if not slots:
                raise RuntimeError("برای راه‌اندازی اولیه حداقل یک API در Railway Variables لازم است؛ پس از اجرا می‌توانید از پنل ادمین تا ۲۰ API اضافه کنید.")
            for slot in slots:
                if not slot["base_url"].startswith(("https://", "http://")):
                    raise RuntimeError(f"AI_BASE_URL اسلات {slot['slot']} معتبر نیست.")
                if not slot["model"]:
                    raise RuntimeError(f"AI_MODEL اسلات {slot['slot']} الزامی است.")
        if AI_PROVIDER == "gemini" and not GEMINI_API_KEY:
            raise RuntimeError("GEMINI_API_KEY برای AI_PROVIDER=gemini الزامی است.")
        raw_integration=str(INTEGRATION_API_KEYS_RAW or '').strip()
        if raw_integration and raw_integration!='{}':
            parsed_integration=parse_integration_keys()
            if not parsed_integration:
                raise RuntimeError('INTEGRATION_API_KEYS معتبر نیست؛ از JSON یا قالب name:key استفاده کنید.')
            for name,value in parsed_integration.items():
                lowered=value.lower()
                if len(value)<24 or 'change-me' in lowered or 'example' in lowered or 'placeholder' in lowered:
                    raise RuntimeError(f'کلید Integration برای {name} باید تصادفی، واقعی و حداقل 24 کاراکتر باشد.')


class ClosingSQLiteConnection(sqlite3.Connection):
    """SQLite connection that closes when used as a context manager.

    sqlite3.Connection.__exit__ commits/rolls back but does not close; Barsan uses
    `with get_db()` extensively, so production connections must be released there.
    """
    def __exit__(self, exc_type, exc_value, traceback):
        try:
            return super().__exit__(exc_type, exc_value, traceback)
        finally:
            self.close()


def get_db() -> sqlite3.Connection:
    con = sqlite3.connect(DB_PATH, timeout=30, check_same_thread=False, factory=ClosingSQLiteConnection)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys = ON")
    con.execute("PRAGMA journal_mode = WAL")
    con.execute("PRAGMA synchronous = NORMAL")
    con.execute("PRAGMA busy_timeout = 30000")
    con.execute("PRAGMA wal_autocheckpoint = 1000")
    con.execute("PRAGMA temp_store = MEMORY")
    con.execute(f"PRAGMA cache_size = -{SQLITE_CACHE_KB}")
    con.execute(f"PRAGMA mmap_size = {SQLITE_MMAP_MB * 1024 * 1024}")
    return con

def init_storage() -> None:
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_SESSION_DIR.mkdir(parents=True, exist_ok=True)
    CARGO_IMAGE_TRAINING_DIR.mkdir(parents=True, exist_ok=True)
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)


def column_names(db: sqlite3.Connection, table: str) -> set[str]:
    return {row[1] for row in db.execute(f"PRAGMA table_info({table})").fetchall()}


def _ensure_base_schema() -> None:
    with get_db() as db:
        db.executescript(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE,
                email TEXT UNIQUE,
                name TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                salt TEXT NOT NULL,
                role TEXT NOT NULL DEFAULT 'user' CHECK(role IN ('user','manager','admin')),
                is_active INTEGER NOT NULL DEFAULT 1,
                created_by INTEGER,
                created_at TEXT NOT NULL,
                FOREIGN KEY(created_by) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                stored_path TEXT NOT NULL,
                mime_type TEXT,
                visibility TEXT NOT NULL DEFAULT 'public' CHECK(visibility IN ('public','authenticated','internal')),
                status TEXT NOT NULL DEFAULT 'ready',
                character_count INTEGER NOT NULL DEFAULT 0,
                chunk_count INTEGER NOT NULL DEFAULT 0,
                version INTEGER NOT NULL DEFAULT 1,
                created_by INTEGER,
                created_at TEXT NOT NULL,
                FOREIGN KEY(created_by) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_id TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                content TEXT NOT NULL,
                FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE,
                UNIQUE(document_id, chunk_index)
            );

            CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(
                content,
                file_name,
                visibility,
                document_id UNINDEXED,
                chunk_index UNINDEXED,
                tokenize = 'unicode61'
            );

            CREATE TABLE IF NOT EXISTS conversations (
                id TEXT PRIMARY KEY,
                user_id INTEGER,
                external_user_id TEXT,
                title TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY(user_id) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS messages (
                id TEXT PRIMARY KEY,
                conversation_id TEXT NOT NULL,
                role TEXT NOT NULL CHECK(role IN ('user','assistant')),
                content TEXT NOT NULL,
                sources_json TEXT NOT NULL DEFAULT '[]',
                model TEXT,
                response_ms INTEGER,
                status TEXT NOT NULL DEFAULT 'answered',
                created_at TEXT NOT NULL,
                FOREIGN KEY(conversation_id) REFERENCES conversations(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS audit_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                action TEXT NOT NULL,
                details TEXT NOT NULL DEFAULT '{}',
                created_at TEXT NOT NULL,
                FOREIGN KEY(user_id) REFERENCES users(id)
            );

            CREATE INDEX IF NOT EXISTS idx_messages_conversation ON messages(conversation_id);
            CREATE INDEX IF NOT EXISTS idx_conversations_user ON conversations(user_id);
            CREATE INDEX IF NOT EXISTS idx_documents_created_at ON documents(created_at DESC);
            CREATE INDEX IF NOT EXISTS idx_users_username ON users(username);
            CREATE INDEX IF NOT EXISTS idx_users_email ON users(email);
            """
        )

        cols = column_names(db, 'users')
        if 'username' not in cols:
            db.execute("ALTER TABLE users ADD COLUMN username TEXT")
        if 'is_active' not in cols:
            db.execute("ALTER TABLE users ADD COLUMN is_active INTEGER NOT NULL DEFAULT 1")
        if 'created_by' not in cols:
            db.execute("ALTER TABLE users ADD COLUMN created_by INTEGER")
        if 'question_limit' not in cols:
            db.execute("ALTER TABLE users ADD COLUMN question_limit INTEGER")
        if 'questions_used' not in cols:
            db.execute("ALTER TABLE users ADD COLUMN questions_used INTEGER NOT NULL DEFAULT 0")

        dcols = column_names(db, 'documents')
        if 'status' not in dcols:
            db.execute("ALTER TABLE documents ADD COLUMN status TEXT NOT NULL DEFAULT 'ready'")
        if 'chunk_count' not in dcols:
            db.execute("ALTER TABLE documents ADD COLUMN chunk_count INTEGER NOT NULL DEFAULT 0")
        if 'version' not in dcols:
            db.execute("ALTER TABLE documents ADD COLUMN version INTEGER NOT NULL DEFAULT 1")
        for name,definition in (
            ('page_count','INTEGER NOT NULL DEFAULT 0'),
            ('vision_candidate_pages','INTEGER NOT NULL DEFAULT 0'),
            ('vision_success_pages','INTEGER NOT NULL DEFAULT 0'),
            ('vision_failed_pages','INTEGER NOT NULL DEFAULT 0'),
            ('ingestion_quality_pct','REAL NOT NULL DEFAULT 100'),
            ('page_fidelity_pct','REAL NOT NULL DEFAULT 100'),
            ('numeric_agreement_pct','REAL NOT NULL DEFAULT 100'),
            ('quality_gate_reason',"TEXT NOT NULL DEFAULT ''"),
            ('ingestion_warnings_json',"TEXT NOT NULL DEFAULT '[]'"),
            ('last_indexed_at','TEXT'),
            ('ingestion_version','INTEGER NOT NULL DEFAULT 1'),
            ('is_builtin','INTEGER NOT NULL DEFAULT 0'),
            ('is_enabled','INTEGER NOT NULL DEFAULT 1'),
            ('source_key','TEXT'),
            ('reindex_status',"TEXT NOT NULL DEFAULT 'idle'"),
            ('reindex_error',"TEXT NOT NULL DEFAULT ''"),
            ('reindex_started_at','TEXT'),
            ('reindex_completed_at','TEXT'),
        ):
            if name not in dcols:
                db.execute(f"ALTER TABLE documents ADD COLUMN {name} {definition}")

        db.executescript("""
            CREATE TABLE IF NOT EXISTS document_pages (
                document_id TEXT NOT NULL,
                page_number INTEGER NOT NULL,
                base_text TEXT NOT NULL DEFAULT '',
                vision_text TEXT NOT NULL DEFAULT '',
                combined_text TEXT NOT NULL DEFAULT '',
                image_count INTEGER NOT NULL DEFAULT 0,
                status TEXT NOT NULL DEFAULT 'text',
                vision_error TEXT,
                character_count INTEGER NOT NULL DEFAULT 0,
                page_fidelity REAL NOT NULL DEFAULT 1.0,
                numeric_agreement REAL,
                PRIMARY KEY(document_id,page_number),
                FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE
            );
            CREATE INDEX IF NOT EXISTS idx_document_pages_status ON document_pages(document_id,status,page_number);
        """)
        pcols = column_names(db, 'document_pages')
        if 'page_fidelity' not in pcols:
            db.execute("ALTER TABLE document_pages ADD COLUMN page_fidelity REAL NOT NULL DEFAULT 1.0")
        if 'numeric_agreement' not in pcols:
            db.execute("ALTER TABLE document_pages ADD COLUMN numeric_agreement REAL")

        # R35: migrate legacy dynamic provider secrets from plaintext to Fernet.
        try:
            legacy_keys=db.execute("SELECT slot,api_key FROM ai_provider_configs WHERE api_key IS NOT NULL AND api_key!=''").fetchall()
            for key_row in legacy_keys:
                stored=str(key_row['api_key'] or '')
                if stored and not stored.startswith(_PROVIDER_KEY_PREFIX):
                    db.execute("UPDATE ai_provider_configs SET api_key=?,updated_at=? WHERE slot=?",(_protect_api_key(stored),now_iso(),int(key_row['slot'])))
        except sqlite3.Error:
            pass

        ccols = column_names(db, 'chunks')
        for name,definition in (
            ('page_start','INTEGER'),('page_end','INTEGER'),('section_title','TEXT'),
            ('chunk_type',"TEXT NOT NULL DEFAULT 'text'"),('search_aliases',"TEXT NOT NULL DEFAULT ''")
        ):
            if name not in ccols:
                db.execute(f"ALTER TABLE chunks ADD COLUMN {name} {definition}")

        mcols = column_names(db, 'messages')
        if 'response_ms' not in mcols:
            db.execute("ALTER TABLE messages ADD COLUMN response_ms INTEGER")
        if 'status' not in mcols:
            db.execute("ALTER TABLE messages ADD COLUMN status TEXT NOT NULL DEFAULT 'answered'")

        db.execute("UPDATE users SET username = lower(substr(email,1,instr(email,'@')-1)) WHERE (username IS NULL OR username = '') AND email LIKE '%@%'")
        backfill_fts(db)


def backfill_fts(db: sqlite3.Connection) -> None:
    current = db.execute("SELECT COUNT(*) FROM chunks_fts").fetchone()[0]
    source = db.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
    if current == source:
        return
    db.execute("DELETE FROM chunks_fts")
    rows = db.execute(
        """
        SELECT c.document_id, c.chunk_index, c.content, COALESCE(c.search_aliases,'') AS search_aliases, d.filename, d.visibility
        FROM chunks c JOIN documents d ON d.id = c.document_id
        ORDER BY c.id ASC
        """
    ).fetchall()
    db.executemany(
        "INSERT INTO chunks_fts(content, file_name, visibility, document_id, chunk_index) VALUES (?, ?, ?, ?, ?)",
        [((row['content']+'\nکلیدواژه‌های جست‌وجو: '+row['search_aliases']).strip() if row['search_aliases'] else row['content'], row['filename'], row['visibility'], row['document_id'], row['chunk_index']) for row in rows],
    )


def hash_password(password: str, salt_hex: str | None = None) -> tuple[str, str]:
    salt = bytes.fromhex(salt_hex) if salt_hex else secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac('sha256', password.encode(), salt, 250_000)
    return digest.hex(), salt.hex()


def verify_password(password: str, expected: str, salt_hex: str) -> bool:
    actual, _ = hash_password(password, salt_hex)
    return hmac.compare_digest(actual, expected)


def seed_admin() -> None:
    with get_db() as db:
        row = db.execute(
            "SELECT id FROM users WHERE lower(username)=lower(?) OR lower(email)=lower(?) ORDER BY CASE WHEN lower(email)=lower(?) THEN 0 ELSE 1 END, id LIMIT 1",
            (MASTER_ADMIN_USERNAME, MASTER_ADMIN_EMAIL, MASTER_ADMIN_EMAIL),
        ).fetchone()
        if row:
            db.execute("UPDATE users SET is_owner=0 WHERE is_owner=1 AND id != ?", (row['id'],))
            if SYNC_MASTER_ADMIN_CREDENTIALS:
                password_hash, salt = hash_password(INITIAL_ADMIN_PASSWORD)
                db.execute(
                    "UPDATE users SET username=?, email=?, password_hash=?, salt=?, role='admin', is_active=1, is_owner=1 WHERE id=?",
                    (MASTER_ADMIN_USERNAME, MASTER_ADMIN_EMAIL, password_hash, salt, row['id']),
                )
            else:
                db.execute(
                    "UPDATE users SET username=?, email=?, role='admin', is_active=1, is_owner=1 WHERE id=?",
                    (MASTER_ADMIN_USERNAME, MASTER_ADMIN_EMAIL, row['id']),
                )
            return
        password_hash, salt = hash_password(INITIAL_ADMIN_PASSWORD)
        db.execute(
            """
            INSERT INTO users(username,email,name,password_hash,salt,role,is_active,is_owner,created_at)
            VALUES (?,?,?,?,?,'admin',1,1,?)
            """,
            (MASTER_ADMIN_USERNAME, MASTER_ADMIN_EMAIL, INITIAL_ADMIN_NAME, password_hash, salt, now_iso()),
        )


def create_token(user_id: int, role: str) -> str:
    exp = datetime.now(timezone.utc) + timedelta(minutes=JWT_EXPIRE_MINUTES)
    return jwt.encode({'sub': str(user_id), 'role': role, 'exp': exp}, JWT_SECRET, algorithm='HS256')


def decode_token(token: str) -> dict[str, Any]:
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
    except jwt.PyJWTError as exc:
        raise HTTPException(status_code=401, detail='توکن نامعتبر یا منقضی است.') from exc


def current_user_optional(
    authorization: str | None = Header(default=None),
    barsan_token: str | None = Cookie(default=None),
) -> dict[str, Any] | None:
    token = barsan_token
    if authorization and authorization.lower().startswith('bearer '):
        token = authorization.split(' ', 1)[1].strip()
    if not token:
        return None
    payload = decode_token(token)
    with get_db() as db:
        row = db.execute(
            "SELECT id,username,email,name,role,is_active,is_owner,department,question_limit,questions_used,daily_question_limit,daily_questions_used,daily_quota_date,monthly_question_limit,monthly_questions_used,monthly_quota_month,created_at,last_login_at FROM users WHERE id=?",
            (int(payload['sub']),),
        ).fetchone()
    if not row or not row['is_active']:
        raise HTTPException(status_code=401, detail='حساب کاربری غیرفعال یا نامعتبر است.')
    return dict(row)


def current_user(user: dict[str, Any] | None = Depends(current_user_optional)) -> dict[str, Any]:
    if not user:
        raise HTTPException(status_code=401, detail='ابتدا وارد حساب شوید.')
    return user


def require_roles(*roles: str):
    def dep(user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
        if user['role'] not in roles:
            raise HTTPException(status_code=403, detail='دسترسی کافی ندارید.')
        return user
    return dep


def require_owner(user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    if user['role'] != 'admin' or not bool(user.get('is_owner')):
        raise HTTPException(status_code=403, detail='این عملیات فقط در اختیار ادمین اصلی است.')
    return user


def audit(user_id: int | None, action: str, details: dict[str, Any] | None = None) -> None:
    with get_db() as db:
        db.execute("INSERT INTO audit_logs(user_id, action, details, created_at) VALUES (?, ?, ?, ?)", (user_id, action, json.dumps(details or {}, ensure_ascii=False), now_iso()))


def _local_now() -> datetime:
    return datetime.now(timezone.utc) + timedelta(minutes=LOCAL_TIMEZONE_OFFSET_MINUTES)


def _quota_period_keys() -> tuple[str, str]:
    local = _local_now()
    return local.strftime('%Y-%m-%d'), local.strftime('%Y-%m')


def _refresh_all_quota_periods(db: sqlite3.Connection) -> None:
    day_key, month_key = _quota_period_keys()
    db.execute(
        """UPDATE users SET daily_questions_used=0,daily_quota_date=?
           WHERE daily_quota_date IS NULL OR daily_quota_date<>?""",
        (day_key, day_key),
    )
    db.execute(
        """UPDATE users SET monthly_questions_used=0,monthly_quota_month=?
           WHERE monthly_quota_month IS NULL OR monthly_quota_month<>?""",
        (month_key, month_key),
    )


def _persistent_counter(identity_key: str, window_type: str, window_key: str, limit: int) -> int:
    if limit <= 0:
        return 0
    ts = now_iso()
    with get_db() as db:
        db.execute('BEGIN IMMEDIATE')
        row = db.execute(
            "SELECT count FROM request_counters WHERE identity_key=? AND window_type=? AND window_key=?",
            (identity_key, window_type, window_key),
        ).fetchone()
        count = int(row['count'] or 0) if row else 0
        if count >= limit:
            db.rollback()
            raise HTTPException(status_code=429, detail='سقف درخواست این بازه زمانی تمام شده است.')
        db.execute(
            """INSERT INTO request_counters(identity_key,window_type,window_key,count,updated_at)
               VALUES(?,?,?,?,?) ON CONFLICT(identity_key,window_type,window_key)
               DO UPDATE SET count=request_counters.count+1,updated_at=excluded.updated_at""",
            (identity_key, window_type, window_key, 1, ts),
        )
        # Bounded cleanup keeps the limiter stable across redeploys without unbounded growth.
        if secrets.randbelow(100) == 0:
            cutoff = (datetime.now(timezone.utc)-timedelta(days=RATE_LIMIT_RETENTION_DAYS)).isoformat()
            db.execute("DELETE FROM request_counters WHERE updated_at<?", (cutoff,))
        db.commit()
    return count + 1


def enforce_rate_limit(
    key: str,
    *,
    per_minute: int | None = None,
    daily_limit: int = 0,
    monthly_limit: int = 0,
) -> None:
    """Atomically enforce persistent minute/day/month windows.

    All counters are checked before any is incremented. A request rejected by the
    daily or monthly ceiling therefore does not consume another minute counter.
    """
    current = _local_now()
    windows = [
        ('minute', current.strftime('%Y-%m-%dT%H:%M'), per_minute or REQUESTS_PER_MINUTE),
        ('day', current.strftime('%Y-%m-%d'), daily_limit),
        ('month', current.strftime('%Y-%m'), monthly_limit),
    ]
    windows = [(kind, period, int(limit)) for kind, period, limit in windows if int(limit or 0) > 0]
    if not windows:
        return
    ts = now_iso()
    with get_db() as db:
        db.execute('BEGIN IMMEDIATE')
        try:
            for kind, period, limit in windows:
                row = db.execute(
                    "SELECT count FROM request_counters WHERE identity_key=? AND window_type=? AND window_key=?",
                    (key, kind, period),
                ).fetchone()
                if row and int(row['count'] or 0) >= limit:
                    label = {'minute':'دقیقه‌ای','day':'روزانه','month':'ماهانه'}[kind]
                    raise HTTPException(status_code=429, detail=f'سقف درخواست {label} تمام شده است.')
            for kind, period, _limit in windows:
                db.execute(
                    """INSERT INTO request_counters(identity_key,window_type,window_key,count,updated_at)
                       VALUES(?,?,?,?,?) ON CONFLICT(identity_key,window_type,window_key)
                       DO UPDATE SET count=request_counters.count+1,updated_at=excluded.updated_at""",
                    (key, kind, period, 1, ts),
                )
            if secrets.randbelow(100) == 0:
                cutoff = (datetime.now(timezone.utc)-timedelta(days=RATE_LIMIT_RETENTION_DAYS)).isoformat()
                db.execute("DELETE FROM request_counters WHERE updated_at<?", (cutoff,))
            db.commit()
        except Exception:
            db.rollback()
            raise


_PERSIAN_STOPWORDS = {
    "از", "به", "در", "با", "برای", "و", "یا", "که", "را", "این", "آن", "یک", "چی", "چه",
    "چطور", "چگونه", "کدام", "است", "هست", "هستم", "هستند", "بود", "شد", "شود", "می", "من",
    "شما", "ما", "لطفا", "لطفاً", "درباره", "مورد", "بگو", "کن", "کرد", "آیا", "روی", "تا",
}
_SEARCH_ALIASES = {
    "لغو": "کنسلی", "لغوی": "کنسلی", "جریمه": "کنسلی", "بازپرداخت": "استرداد",
    "قیمت": "هزینه", "مبلغ": "هزینه", "بها": "هزینه", "راهنما": "آموزش",
    "ثبتنام": "ثبت نام", "نامنویسی": "ثبت نام", "پشتیبانی": "تیکت", "درخواست": "تیکت",
    # واژه‌های محاوره‌ای رایج در عملیات حمل بار
    "روبار": "روباری", "روباره": "روباری", "روبارى": "روباری", "خوابار": "خاور",
    "نیسانی": "نیسان", "پیکانی": "پیکان", "وانتی": "وانت",
    "روباری": "روبار باربند بار روی سقف", "باربندی": "باربند روبار", "باربند": "روبار بار روی سقف",
    "ظرفیت": "وزن مجاز حداکثر بار", "تناژ": "وزن ظرفیت بار", "کیلو": "کیلوگرم kg",
    "کیلویی": "کیلوگرم kg", "تن": "تن وزن", "محدودیت": "استثنا شرط", "تبصره": "استثنا شرط",
}
_PERSIAN_SUFFIXES = (
    "هایشان", "هایتان", "هایمان", "های", "هایی", "ها", "ترین", "تر", "شان", "تان", "مان",
    "یم", "ید", "ند", "ام", "ات", "اش", "ای", "ی",
)


def reserve_question_quota(user: dict[str, Any] | None) -> bool:
    """Atomically reserve lifetime, daily and monthly quota for a registered account."""
    if not user:
        return False
    day_key, month_key = _quota_period_keys()
    with get_db() as db:
        db.execute('BEGIN IMMEDIATE')
        row = db.execute(
            """SELECT question_limit,questions_used,daily_question_limit,daily_questions_used,
                      daily_quota_date,monthly_question_limit,monthly_questions_used,monthly_quota_month
               FROM users WHERE id=? AND is_active=1""",
            (user['id'],),
        ).fetchone()
        if not row:
            db.rollback()
            raise HTTPException(status_code=401, detail='حساب کاربری پیدا نشد یا غیرفعال است.')
        lifetime_used = int(row['questions_used'] or 0)
        daily_used = int(row['daily_questions_used'] or 0) if row['daily_quota_date'] == day_key else 0
        monthly_used = int(row['monthly_questions_used'] or 0) if row['monthly_quota_month'] == month_key else 0
        checks = [
            ('کل', row['question_limit'], lifetime_used),
            ('روزانه', row['daily_question_limit'], daily_used),
            ('ماهانه', row['monthly_question_limit'], monthly_used),
        ]
        for label, limit_value, used in checks:
            if limit_value is not None and used >= int(limit_value):
                db.rollback()
                raise HTTPException(
                    status_code=429,
                    detail=f"سهمیه سؤال {label} این حساب تمام شده است ({used} از {int(limit_value)}). مدیر باید سهمیه را افزایش دهد یا دوره جدید آغاز شود.",
                )
        db.execute(
            """UPDATE users SET questions_used=?,daily_questions_used=?,monthly_questions_used=?,
                      daily_quota_date=?,monthly_quota_month=? WHERE id=?""",
            (lifetime_used+1, daily_used+1, monthly_used+1, day_key, month_key, user['id']),
        )
        db.commit()
    return True


def refund_question_quota(user: dict[str, Any] | None, reserved: bool) -> None:
    if not user or not reserved:
        return
    day_key, month_key = _quota_period_keys()
    with get_db() as db:
        db.execute(
            """UPDATE users SET questions_used=MAX(0,questions_used-1),
               daily_questions_used=CASE WHEN daily_quota_date=? THEN MAX(0,daily_questions_used-1) ELSE daily_questions_used END,
               monthly_questions_used=CASE WHEN monthly_quota_month=? THEN MAX(0,monthly_questions_used-1) ELSE monthly_questions_used END
               WHERE id=?""",
            (day_key, month_key, user['id']),
        )


def question_quota_snapshot(user_id: int | None) -> dict[str, Any] | None:
    if user_id is None:
        return None
    day_key, month_key = _quota_period_keys()
    with get_db() as db:
        row = db.execute(
            """SELECT question_limit,questions_used,daily_question_limit,daily_questions_used,daily_quota_date,
                      monthly_question_limit,monthly_questions_used,monthly_quota_month
               FROM users WHERE id=?""",
            (user_id,),
        ).fetchone()
    if not row:
        return None
    lifetime_limit = row['question_limit']
    lifetime_used = int(row['questions_used'] or 0)
    daily_used = int(row['daily_questions_used'] or 0) if row['daily_quota_date'] == day_key else 0
    monthly_used = int(row['monthly_questions_used'] or 0) if row['monthly_quota_month'] == month_key else 0
    def section(limit_value: Any, used: int) -> dict[str, int | None]:
        limit_int = int(limit_value) if limit_value is not None else None
        return {'limit': limit_int, 'used': used, 'remaining': max(0, limit_int-used) if limit_int is not None else None}
    return {
        **section(lifetime_limit, lifetime_used),
        'lifetime': section(lifetime_limit, lifetime_used),
        'daily': section(row['daily_question_limit'], daily_used),
        'monthly': section(row['monthly_question_limit'], monthly_used),
        'daily_period': day_key,
        'monthly_period': month_key,
    }


def normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFKC", str(value or "")).lower()
    value = value.translate(str.maketrans({
        "ي": "ی", "ى": "ی", "ك": "ک", "ة": "ه", "ۀ": "ه", "ؤ": "و", "إ": "ا", "أ": "ا",
        "۰": "0", "۱": "1", "۲": "2", "۳": "3", "۴": "4", "۵": "5", "۶": "6", "۷": "7", "۸": "8", "۹": "9",
    }))
    value = value.replace("\u200c", " ").replace("\u200f", " ").replace("\u200e", " ")
    value = re.sub(r"[\u064b-\u065f\u0670\u06d6-\u06ed]", "", value)
    value = re.sub(r"[^\w\u0600-\u06ff]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _stem_search_token(token: str) -> str:
    mapped=_SEARCH_ALIASES.get(token)
    if mapped and ' ' not in mapped:
        token=mapped
    if token.startswith("نمی") and len(token) > 6:
        token = token[3:]
    elif token.startswith("می") and len(token) > 5:
        token = token[2:]
    for suffix in _PERSIAN_SUFFIXES:
        if token.endswith(suffix) and len(token) - len(suffix) >= 3:
            token = token[:-len(suffix)]
            break
    mapped=_SEARCH_ALIASES.get(token)
    return mapped if mapped and ' ' not in mapped else token


def search_tokens(value: str) -> list[str]:
    tokens: list[str] = []
    for raw in normalize_text(value).split():
        if len(raw) < 2 or raw in _PERSIAN_STOPWORDS:
            continue
        alias = _SEARCH_ALIASES.get(raw, raw)
        for candidate in alias.split():
            stem = _stem_search_token(candidate)
            if len(stem) >= 2 and stem not in _PERSIAN_STOPWORDS and stem not in tokens:
                tokens.append(stem)
    return tokens


def build_fts_query(question: str) -> str:
    terms = search_tokens(question)
    if not terms:
        terms = [token for token in normalize_text(question).split() if len(token) >= 2]
    if not terms:
        return 'barsan'
    unique: list[str] = []
    for term in terms[:16]:
        if term not in unique:
            unique.append(term)
    return ' OR '.join(f'{term}*' if len(term) >= 3 else f'"{term}"' for term in unique)


def _fuzzy_token_score(query_tokens: list[str], haystack_tokens: list[str]) -> float:
    if not query_tokens or not haystack_tokens:
        return 0.0
    haystack_set = set(haystack_tokens)
    total = 0.0
    for query_token in query_tokens:
        if query_token in haystack_set:
            total += 1.0
            continue
        if any(
            len(query_token) >= 3 and (candidate.startswith(query_token) or query_token.startswith(candidate))
            for candidate in haystack_set
        ):
            total += 0.86
            continue
        best = 0.0
        for candidate in haystack_tokens:
            if abs(len(candidate) - len(query_token)) > 4:
                continue
            best = max(best, SequenceMatcher(None, query_token, candidate).ratio())
            if best >= 0.92:
                break
        if best >= 0.72:
            total += best * 0.72
    return total / len(query_tokens)


def _retrieval_score(question: str, content: str, file_name: str = "") -> float:
    q_norm = normalize_text(question)
    c_norm = normalize_text(content)
    q_tokens = search_tokens(question)
    c_tokens = search_tokens(content)
    if not q_tokens:
        return 1.0 if q_norm and q_norm in c_norm else 0.0
    coverage = _fuzzy_token_score(q_tokens, c_tokens)
    exact = sum(1 for token in q_tokens if token in c_tokens) / len(q_tokens)
    phrase_bonus = 0.18 if len(q_norm) >= 6 and q_norm in c_norm else 0.0
    filename_bonus = 0.10 * _fuzzy_token_score(q_tokens, search_tokens(file_name)) if file_name else 0.0
    source_bonus=0.0
    fn=normalize_text(file_name)
    if any(x in q_norm for x in ('سلب','کنسلی','لغو')) and ('04_barsan' in fn or 'فرایند سلب' in fn):
        source_bonus=0.06 if coverage>=0.35 else 0.0
    elif any(x in q_norm for x in ('بارگیری','ظرفیت','اضافه بار','اضافه‌بار','باسکول','پیکان','نیسان','خاور','باربند')) and ('01_barsan' in fn or 'شناخت ناوگان' in fn):
        source_bonus=0.035 if coverage>=0.35 else 0.0
    elif any(x in q_norm for x in ('صف رانندگان','غیرفعال','جابجایی','جابه جایی','مرخصی','موقعیت قدیمی')) and ('02_barsan' in fn or 'عملیات رانندگان' in fn):
        source_bonus=0.035 if coverage>=0.35 else 0.0
    elif any(x in q_norm for x in ('مراحل سرویس','توقف','بازنگری','فاکتور','مالی','گزارشات','اتمام سرویس')) and ('03_barsan' in fn or 'پشتیبانی چرخه' in fn):
        source_bonus=0.035 if coverage>=0.35 else 0.0
    return min(1.5, coverage * 0.72 + exact * 0.28 + phrase_bonus + filename_bonus + source_bonus)

def allowed_visibilities(user: dict[str, Any] | None, integration: bool = False) -> tuple[str, ...]:
    if integration or user is None:
        return ('public',)
    if user['role'] in {'admin', 'manager'}:
        return ('public', 'authenticated', 'internal')
    return ('public', 'authenticated')


def _json_list(value: Any) -> list[Any]:
    if isinstance(value, list):
        return value
    try:
        parsed = json.loads(value or '[]')
        return parsed if isinstance(parsed, list) else []
    except (TypeError, ValueError, json.JSONDecodeError):
        return []


def content_accessible(row: Any, user: dict[str, Any] | None, integration: bool = False) -> bool:
    item = dict(row)
    visibility = str(item.get('visibility') or 'public')
    if visibility not in allowed_visibilities(user, integration):
        return False
    roles = {str(x) for x in _json_list(item.get('allowed_roles_json')) if str(x)}
    user_ids = {int(x) for x in _json_list(item.get('allowed_user_ids_json')) if str(x).isdigit()}
    department = str(item.get('department') or '').strip().lower()
    if roles and (not user or user.get('role') not in roles):
        return False
    if user_ids and (not user or int(user.get('id') or 0) not in user_ids):
        return False
    if department and (not user or str(user.get('department') or '').strip().lower() != department):
        return False
    return True


def _validate_archive_safety(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos=archive.infolist()
            if len(infos)>5000:
                raise HTTPException(status_code=400,detail='تعداد فایل‌های داخلی سند بیش از حد مجاز است.')
            expanded=sum(max(0,i.file_size) for i in infos)
            if expanded>MAX_ARCHIVE_EXPANDED_MB*1024*1024:
                raise HTTPException(status_code=400,detail='حجم بازشده سند بیش از حد مجاز است.')
            compressed=max(1,sum(max(0,i.compress_size) for i in infos))
            if expanded/compressed>150:
                raise HTTPException(status_code=400,detail='نسبت فشرده‌سازی سند غیرعادی است.')
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=400,detail='ساختار فایل Office معتبر نیست.') from exc



def _run_coro_sync(coro: Any) -> Any:
    """Run async provider work safely from synchronous FastAPI worker threads."""
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro)
    result: dict[str, Any] = {}
    errors: list[BaseException] = []
    def runner() -> None:
        try:
            result['value'] = asyncio.run(coro)
        except BaseException as exc:
            errors.append(exc)
    thread=threading.Thread(target=runner,daemon=True)
    thread.start();thread.join()
    if errors: raise errors[0]
    return result.get('value')


def _normalize_extracted_text(text: str) -> str:
    text = unicodedata.normalize('NFKC', str(text or ''))
    text = re.sub(r'[\u200b-\u200f\u202a-\u202e\u2066-\u2069\ufeff\ufffd]', '', text)
    text = text.translate(str.maketrans({'ي':'ی','ى':'ی','ك':'ک','ة':'ه'}))
    text = re.sub(r'\r\n?', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{4,}', '\n\n\n', text).strip()
    return text



def _vision_model_for_slot(slot: int) -> str:
    config=_provider_config_for_slot(slot)
    return str(config.get('vision_model') or '') if config else '' 


async def _generate_vision_text(messages: list[dict[str, Any]], preferred_slot: int | None = None, max_tokens: int | None = None) -> str:
    """Vision requests share the same provider health/cooldown state as chat."""
    errors=[]
    slots=ordered_available_ai_slots_for_route('vision')
    if preferred_slot is not None:
        slots=sorted(slots,key=lambda slot: (0 if int(slot['slot'])==int(preferred_slot) else 1,))
    if not any(_vision_model_for_slot(int(slot['slot'])) for slot in slots):
        raise RuntimeError('هیچ مدل تصویری سالم و در دسترسی در مدل تصویری یکی از APIهای فعال تنظیم نشده است.')
    for attempt_index,slot in enumerate(slots,1):
        model=_vision_model_for_slot(int(slot['slot']))
        if not model:
            continue
        started=time.perf_counter()
        headers={'Content-Type':'application/json',AI_AUTH_HEADER:f"{AI_AUTH_SCHEME} {slot['api_key']}".strip()}
        token_parameters=['max_tokens','max_completion_tokens'] if AI_TOKEN_PARAMETER=='auto' else [AI_TOKEN_PARAMETER]
        compatibility_error=None
        try:
            data=None
            for token_parameter in token_parameters:
                payload={'model':model,'messages':messages,token_parameter:int(max_tokens or PDF_VISION_MAX_TOKENS),'stream':False}
                if AI_SEND_TEMPERATURE:
                    payload['temperature']=0.0
                try:
                    data=await _request_json_with_retries(
                        url=f"{slot['base_url']}{AI_CHAT_COMPLETIONS_PATH}",payload=payload,params=None,headers=headers,
                        timeout_seconds=AI_TIMEOUT_SECONDS,max_retries=AI_MAX_RETRIES,provider_label=f"{slot['label']} Vision (API {slot['slot']})",rotate_on_quota=True,
                    )
                    break
                except HTTPException as exc:
                    if any(code in str(exc.detail) for code in ('HTTP 400','HTTP 404','HTTP 422')):
                        compatibility_error=exc
                        continue
                    raise
            if data is None:
                raise compatibility_error or RuntimeError('مدل Vision پاسخ سازگار نداد.')
        except ApiSlotLimitError as exc:
            elapsed=int((time.perf_counter()-started)*1000)
            record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=model,model_route='vision',status='limited',http_status=exc.status_code,error_class='quota_or_rate_limit',error_detail=exc.detail,response_ms=elapsed,attempt_index=attempt_index,was_failover=attempt_index>1)
            cooldown=exc.retry_after_seconds or (AI_KEY_RATE_LIMIT_COOLDOWN_SECONDS if exc.status_code==429 else AI_KEY_QUOTA_COOLDOWN_SECONDS)
            mark_api_slot_blocked(int(slot['slot']),exc.detail,cooldown,'limited');errors.append(f"API {slot['slot']}: {exc.detail}");continue
        except Exception as exc:
            detail=sanitize_answer_text(str(getattr(exc,'detail',exc)))[:300]
            elapsed=int((time.perf_counter()-started)*1000)
            record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=model,model_route='vision',status='transient_error',http_status=None,error_class='vision_error',error_detail=detail,response_ms=elapsed,attempt_index=attempt_index,was_failover=attempt_index>1)
            mark_api_slot_blocked(int(slot['slot']),detail,AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'transient_error');errors.append(f"API {slot['slot']}: {detail}");continue
        choices=data.get('choices') or []
        raw=((choices[0].get('message') or {}).get('content') if choices else None)
        if isinstance(raw,list):
            text='\n'.join(str(x.get('text') or x.get('content') or '') for x in raw if isinstance(x,dict)).strip()
        else:
            text=str(raw or (choices[0].get('text') if choices else '') or '').strip()
        if not text:
            errors.append(f"API {slot['slot']}: پاسخ تصویری خالی")
            mark_api_slot_blocked(int(slot['slot']),'پاسخ Vision خالی',AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'invalid_response')
            continue
        mark_api_slot_success(int(slot['slot']))
        elapsed=int((time.perf_counter()-started)*1000)
        record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=model,model_route='vision',status='success',http_status=200,error_class=None,error_detail=None,response_ms=elapsed,attempt_index=attempt_index,was_failover=attempt_index>1)
        return text
    raise RuntimeError('پردازش تصویر با APIهای تصویری ناموفق بود. '+(' | '.join(errors[-4:])))


def _source_number_unit_tokens(text: str) -> set[str]:
    return source_number_unit_tokens(text, normalizer=_normalize_extracted_text)


def _text_health_score(text: str) -> float:
    return text_health_score(text)


def _page_fidelity_metrics(base_text: str, vision_text: str, *, vision_full: bool, status: str) -> tuple[float,float|None]:
    return page_fidelity_metrics(base_text, vision_text, vision_full=vision_full, status=status, normalizer=_normalize_extracted_text)


def _generic_source_quality(text: str, kind: str) -> tuple[float,list[str]]:
    return generic_source_quality(text, kind, min_extracted_text_chars=MIN_EXTRACTED_TEXT_CHARS, normalizer=_normalize_extracted_text)


def _looks_table_like_page(text: str) -> bool:
    clean=_normalize_extracted_text(text)
    if not clean:
        return False
    lines=[line.strip() for line in clean.splitlines() if line.strip()]
    numeric_lines=sum(1 for line in lines if len(re.findall(r'\d|[۰-۹]',line))>=2)
    separators=sum(1 for line in lines if '|' in line or '\t' in line or re.search(r'\s{3,}',line))
    return numeric_lines>=3 or separators>=2


async def _vision_read_pdf_page(page_number: int, image_bytes: bytes, existing_text: str, image_count: int, preferred_slot: int | None = None) -> str:
    encoded=base64.b64encode(image_bytes).decode('ascii')
    delta_mode=bool(PDF_VISION_DELTA_MODE and len(existing_text)>=PDF_VISION_DELTA_TEXT_THRESHOLD and image_count==0 and not _looks_table_like_page(existing_text))
    if delta_mode:
        prompt=(
            f"این تصویر صفحه {page_number} یک سند سازمانی فارسی است. متن پایه این صفحه قبلاً با استخراج مستقیم دریافت شده است. "
            "تصویر را دقیق کنترل کن و فقط مواردی را برگردان که در متن پایه جا افتاده، اشتباه خوانده شده یا به‌صورت پاورقی/عدد/واحد/تبصره/شرط/استثنا قابل مشاهده است. "
            "اگر متن پایه کامل است فقط بنویس «تأیید تصویری: متن پایه کامل است.»؛ متن پایه را دوباره رونویسی نکن. هیچ حدس یا دانش بیرونی اضافه نکن."
        )
    else:
        prompt=(
            f"این تصویر صفحه {page_number} یک سند سازمانی فارسی است. تعداد تصاویر داخلی صفحه: {image_count}. "
            "تمام متن قابل مشاهده را خط‌به‌خط استخراج کن؛ هیچ سطر، سلول جدول، عدد، واحد، عنوان، پاورقی، تبصره، شرط یا استثنایی را حذف نکن. "
            "جدول‌ها را با ردیف و ستون‌های قابل فهم بازنویسی کن و ارتباط هر عدد با وسیله/عنوان همان ردیف را حفظ کن. "
            "قاعده عمومی را صریحاً با برچسب «قاعده اصلی» و هر محدودیت را با برچسب «شرط/استثنا» جدا کن. "
            "اگر تصویر یا نمودار مفهوم عملی دارد، فقط همان مفهوم قابل مشاهده را توصیف کن؛ هیچ دانش بیرونی یا حدسی اضافه نکن. "
            "خروجی باید کامل، وفادار به همین صفحه و مناسب جست‌وجوی فارسی باشد؛ خلاصه‌سازی نکن. "
            "اگر بخشی ناخواناست دقیقاً با برچسب [ناخوانا] مشخص کن و آن را حدس نزن. هیچ اطلاعاتی از صفحه قبلی یا بعدی اختراع نکن."
        )
    if existing_text:
        prompt += f"\nمتن استخراج‌شده فعلی صفحه برای تطبیق و رفع جاافتادگی:\n{existing_text[:5000]}"
    messages=[{'role':'user','content':[
        {'type':'text','text':prompt},
        {'type':'image_url','image_url':{'url':f'data:image/png;base64,{encoded}','detail':'high'}},
    ]}]
    answer=await _generate_vision_text(messages,preferred_slot=preferred_slot,max_tokens=(700 if delta_mode else PDF_VISION_MAX_TOKENS))
    return _normalize_extracted_text(answer)


async def _extract_pdf_source_result_async(
    source: bytes | Path,
    progress_callback: Callable[[int, str], None] | None = None,
) -> dict[str, Any]:
    """Source Ingestion 3.0 for PDF: smart page-aware text/vision with explicit quality state."""
    is_path=isinstance(source,Path)
    reader=PdfReader(str(source)) if is_path else PdfReader(io.BytesIO(source))
    page_count=len(reader.pages)
    if page_count>MAX_PDF_PAGES:
        raise HTTPException(status_code=400,detail='تعداد صفحات PDF بیش از حد مجاز است.')
    pdf_doc=(fitz.open(str(source)) if is_path else fitz.open(stream=source,filetype='pdf')) if fitz is not None else None
    pages: list[dict[str,Any]]=[]
    candidates: list[dict[str,Any]]=[]
    warnings: list[str]=[]
    configured_vision_slot_ids=[int(slot['slot']) for slot in configured_ai_slots() if _vision_model_for_slot(int(slot['slot']))]
    vision_configured=bool(configured_vision_slot_ids)
    if PDF_VISION_ENABLED and PDF_VISION_SCAN_ALL_PAGES and not vision_configured:
        warnings.append('مدل Vision تنظیم نشده است؛ صفحات متنی سالم با کنترل سلامت متن پذیرفته می‌شوند اما صفحات تصویری/جدولی/کم‌متن تا تنظیم Vision مسدود می‌مانند.')
    for index,page in enumerate(reader.pages):
        page_number=index+1
        try:
            base_text=_normalize_extracted_text((page.extract_text() or '').strip())
        except Exception as exc:
            base_text=''
            warnings.append(f'استخراج متن صفحه {page_number} ناموفق بود: {sanitize_answer_text(str(exc))[:120]}')
        image_count=0
        page_obj=None
        if pdf_doc is not None and index<len(pdf_doc):
            page_obj=pdf_doc[index]
            try:
                image_count=len(page_obj.get_images(full=True))
            except Exception:
                image_count=0
        essential_vision=bool(
            page_obj is not None and (
                (PDF_VISION_SCAN_IMAGE_PAGES and image_count>0)
                or len(normalize_text(base_text))<PDF_VISION_MIN_TEXT_CHARS
                or (PDF_VISION_SCAN_TABLE_LIKE_PAGES and _looks_table_like_page(base_text))
            )
        )
        audit_vision=bool(page_obj is not None and PDF_VISION_SCAN_ALL_PAGES and vision_configured)
        needs_vision=bool(PDF_VISION_ENABLED and (essential_vision or audit_vision))
        unavailable_required=bool(PDF_VISION_ENABLED and essential_vision and not vision_configured)
        initial_status=(
            'vision_unavailable' if unavailable_required else
            'vision_pending' if needs_vision else
            ('text' if base_text else 'blank' if image_count==0 else 'empty')
        )
        row={
            'page_number':page_number,'base_text':base_text,'vision_text':'','combined_text':base_text,
            'image_count':image_count,'status':initial_status,
            'vision_error':'مدل Vision تنظیم نشده است' if unavailable_required else None,
            'character_count':len(base_text),'needs_vision':needs_vision,'vision_full':False,'page_fidelity':1.0,'numeric_agreement':None,
        }
        pages.append(row)
        if needs_vision and not unavailable_required:
            candidates.append({'page_number':page_number,'page_obj':page_obj,'page_text':base_text,'image_count':image_count})
    if progress_callback:
        progress_callback(28,'شناسایی متن، تصاویر، جداول و صفحات اسکن‌شده PDF')

    limited=candidates[:PDF_VISION_MAX_PAGES]
    skipped=max(0,len(candidates)-len(limited))
    if skipped:
        warnings.append(f'{skipped} صفحه نیازمند Vision به‌دلیل سقف PDF_VISION_MAX_PAGES بررسی نشد.')
        for item in candidates[len(limited):]:
            row=pages[item['page_number']-1]
            row['status']='vision_skipped'
            row['vision_error']='سقف پردازش تصویری'

    semaphore=asyncio.Semaphore(PDF_VISION_CONCURRENCY)
    available_vision_slot_ids=[int(slot['slot']) for slot in ordered_available_ai_slots_for_route('vision') if _vision_model_for_slot(int(slot['slot']))]
    vision_slot_ids=available_vision_slot_ids or configured_vision_slot_ids
    completed=0
    progress_lock=asyncio.Lock()
    async def process_candidate(item: dict[str,Any]) -> tuple[int,str,str|None,bool]:
        nonlocal completed
        page_number=int(item['page_number'])
        delta_mode=bool(PDF_VISION_DELTA_MODE and len(item['page_text'])>=PDF_VISION_DELTA_TEXT_THRESHOLD and int(item['image_count'])==0 and not _looks_table_like_page(item['page_text']))
        vision_full=not delta_mode
        try:
            async with semaphore:
                zoom=PDF_VISION_DPI/72.0
                pix=item['page_obj'].get_pixmap(matrix=fitz.Matrix(zoom,zoom),alpha=False)
                image_bytes=pix.tobytes('png')
                if len(image_bytes)>PDF_VISION_MAX_RENDER_MB*1024*1024:
                    reduced=max(1.0,zoom*0.72)
                    pix=item['page_obj'].get_pixmap(matrix=fitz.Matrix(reduced,reduced),alpha=False)
                    image_bytes=pix.tobytes('png')
                preferred_slot=(vision_slot_ids[(page_number-1)%len(vision_slot_ids)] if PDF_VISION_BALANCE_PROVIDERS and vision_slot_ids else None)
                try:
                    visual_text=await _vision_read_pdf_page(page_number,image_bytes,item['page_text'],int(item['image_count']),preferred_slot=preferred_slot)
                except TypeError as exc:
                    # Backward-compatible hook for tests/custom deployments that monkeypatch
                    # the historical four-argument page reader.
                    if 'preferred_slot' not in str(exc):
                        raise
                    visual_text=await _vision_read_pdf_page(page_number,image_bytes,item['page_text'],int(item['image_count']))
                if not visual_text:
                    raise RuntimeError('خروجی Vision خالی بود.')
            return page_number,visual_text,None,vision_full
        except Exception as exc:
            return page_number,'',sanitize_answer_text(str(exc))[:300],vision_full
        finally:
            async with progress_lock:
                completed+=1
                if progress_callback:
                    pct=28+int(25*completed/max(1,len(limited)))
                    progress_callback(min(53,pct),f'تحلیل تصویری {completed} از {len(limited)} صفحه')

    if limited:
        results=await asyncio.gather(*(process_candidate(item) for item in limited))
        for page_number,visual_text,error,vision_full in results:
            row=pages[page_number-1]
            row['vision_full']=bool(vision_full)
            if visual_text:
                row['vision_text']=visual_text
                row['combined_text']=_normalize_extracted_text('\n\n'.join(x for x in (row['base_text'],'جزئیات استخراج‌شده از تصویر، جدول یا اسکن صفحه:\n'+visual_text) if x))
                row['status']='vision_ok'
                row['vision_error']=None
            else:
                row['status']='vision_error'
                row['vision_error']=error or 'خطای نامشخص Vision'
                warnings.append(f'Vision صفحه {page_number} ناموفق بود: {row["vision_error"]}')
            row['character_count']=len(row['combined_text'])
    if pdf_doc is not None:
        pdf_doc.close()
    if PDF_VISION_ENABLED and fitz is None:
        warnings.append('PyMuPDF در دسترس نیست؛ پردازش تصویری PDF انجام نشد.')
        for row in pages:
            if len(normalize_text(row['base_text']))<PDF_VISION_MIN_TEXT_CHARS:
                row['status']='vision_unavailable'
                row['vision_error']='PyMuPDF در دسترس نیست'

    fidelity_scores=[];numeric_agreements=[]
    for row in pages:
        fidelity,numeric_agreement=_page_fidelity_metrics(str(row.get('base_text') or ''),str(row.get('vision_text') or ''),vision_full=bool(row.get('vision_full')),status=str(row.get('status') or ''))
        row['page_fidelity']=round(fidelity,4);row['numeric_agreement']=None if numeric_agreement is None else round(numeric_agreement,4)
        fidelity_scores.append(fidelity)
        if numeric_agreement is not None:
            numeric_agreements.append(numeric_agreement)
            if numeric_agreement<SOURCE_MIN_NUMERIC_AGREEMENT:
                warnings.append(f"صفحه {row['page_number']}: تطابق عدد/واحد متن و Vision فقط {numeric_agreement*100:.0f}٪ است.")
    average_fidelity=(sum(fidelity_scores)/len(fidelity_scores)) if fidelity_scores else 0.0
    average_numeric_agreement=(sum(numeric_agreements)/len(numeric_agreements)) if numeric_agreements else 1.0

    page_outputs=[]
    nonempty_pages=0
    for row in pages:
        text=_normalize_extracted_text(row['combined_text'])
        row['combined_text']=text
        row['character_count']=len(text)
        if text:
            nonempty_pages+=1
            page_outputs.append(f"--- صفحه {row['page_number']} ---\n{text}")
    vision_success=sum(1 for row in pages if row['status']=='vision_ok')
    vision_failed=sum(1 for row in pages if row['status'] in {'vision_error','vision_skipped','vision_unavailable'})
    required_unavailable=sum(1 for row in pages if row['status']=='vision_unavailable')
    candidate_count=len(candidates)+required_unavailable
    meaningful_pages=sum(1 for row in pages if row.get('base_text') or int(row.get('image_count') or 0)>0 or row.get('needs_vision'))
    covered_meaningful=sum(1 for row in pages if (row.get('combined_text') or '').strip() and (row.get('base_text') or int(row.get('image_count') or 0)>0 or row.get('needs_vision')))
    text_coverage=(covered_meaningful/meaningful_pages) if meaningful_pages else 1.0
    vision_coverage=(vision_success/candidate_count) if candidate_count else 1.0
    quality=round(max(0.0,min(100.0,(text_coverage*0.35+vision_coverage*0.20+average_fidelity*0.35+average_numeric_agreement*0.10)*100.0)),1)
    document_status='ready'
    if vision_failed or covered_meaningful<meaningful_pages or quality<SOURCE_MIN_QUALITY_PCT or average_fidelity<SOURCE_MIN_PAGE_FIDELITY or average_numeric_agreement<SOURCE_MIN_NUMERIC_AGREEMENT or (SOURCE_PAGE_BY_PAGE_STRICT and candidate_count and vision_success<candidate_count):
        document_status='partial' if nonempty_pages else 'error'
    text='\n\n'.join(page_outputs).strip()
    if len(normalize_text(text))<MIN_EXTRACTED_TEXT_CHARS:
        document_status='error'
    if progress_callback:
        progress_callback(54,f'کنترل کیفیت منبع: {quality:.0f}٪ — Vision موفق {vision_success} از {candidate_count}')
    return {
        'text':text,'kind':'pdf','pages':pages,'warnings':warnings,'status':document_status,
        'stats':{
            'page_count':page_count,'nonempty_pages':nonempty_pages,'meaningful_pages':meaningful_pages,'vision_candidate_pages':candidate_count,
            'vision_success_pages':vision_success,'vision_failed_pages':vision_failed,'ingestion_quality_pct':quality,'average_page_fidelity':round(average_fidelity,4),'average_numeric_agreement':round(average_numeric_agreement,4),
        },
    }


async def _extract_pdf_with_vision_async(source: bytes | Path, progress_callback: Callable[[int, str], None] | None = None) -> str:
    return str((await _extract_pdf_source_result_async(source,progress_callback))['text'])


def _generic_source_result(text: str, kind: str='text') -> dict[str,Any]:
    clean=_normalize_extracted_text(text)
    quality,warnings=_generic_source_quality(clean,kind)
    status='ready' if clean and quality>=SOURCE_MIN_QUALITY_PCT else ('partial' if clean else 'error')
    return {
        'text':clean,'kind':kind,'pages':[], 'warnings':warnings, 'status':status,
        'stats':{'page_count':0,'nonempty_pages':1 if clean else 0,'vision_candidate_pages':0,'vision_success_pages':0,'vision_failed_pages':0,'ingestion_quality_pct':quality,'average_page_fidelity':round(_text_health_score(clean),4),'average_numeric_agreement':1.0},
    }

def extract_text(filename: str, data: bytes, progress_callback: Callable[[int, str], None] | None = None) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {'.txt', '.md', '.csv'}:
        try:
            text = data.decode('utf-8-sig')
        except UnicodeDecodeError:
            text = data.decode('cp1256', errors='replace')
    elif suffix == '.json':
        try:
            decoded = data.decode('utf-8-sig')
        except UnicodeDecodeError:
            decoded = data.decode('cp1256', errors='replace')
        try:
            text = json.dumps(json.loads(decoded), ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            text = decoded
    elif suffix in {'.html', '.htm'}:
        text = BeautifulSoup(data, 'html.parser').get_text('\n')
    elif suffix == '.pdf':
        text = _run_coro_sync(_extract_pdf_with_vision_async(data,progress_callback))
    elif suffix == '.docx':
        _validate_archive_safety(data)
        doc = Document(io.BytesIO(data))
        out = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table_index, table in enumerate(doc.tables, 1):
            out.append(f'--- جدول {table_index} ---')
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    out.append(' | '.join(values))
        text = '\n'.join(out)
    elif suffix == '.xlsx':
        _validate_archive_safety(data)
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f'--- شیت: {ws.title} ---')
            for row in ws.iter_rows(values_only=True):
                values = ['' if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    out.append(' | '.join(values))
        text = '\n'.join(out)
    else:
        raise HTTPException(status_code=400, detail='فرمت مجاز: PDF, DOCX, XLSX, CSV, TXT, JSON, HTML, MD')
    text = unicodedata.normalize('NFKC', text)
    text = re.sub(r'[\u200b-\u200f\u202a-\u202e\u2066-\u2069\ufeff\ufffd]', '', text)
    text = text.translate(str.maketrans({'ي':'ی','ى':'ی','ك':'ک','ة':'ه'}))
    text = re.sub(r'\r\n?', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{4,}', '\n\n\n', text).strip()
    if len(text) > MAX_ARCHIVE_EXPANDED_MB * 1024 * 1024:
        raise HTTPException(status_code=400, detail='متن استخراج‌شده بیش از حد مجاز است.')
    if len(normalize_text(text)) < MIN_EXTRACTED_TEXT_CHARS:
        raise HTTPException(
            status_code=400,
            detail='متن قابل جستجوی کافی از فایل استخراج نشد. برای PDF تصویری، مدل تصویری را در مدل تصویری یکی از APIهای فعال تنظیم و فایل را دوباره ایندکس کنید.',
        )
    return text


def extract_text_from_path(filename: str, path: Path, progress_callback: Callable[[int,str],None] | None = None) -> str:
    if Path(filename).suffix.lower()=='.pdf':
        text=_run_coro_sync(_extract_pdf_with_vision_async(path,progress_callback))
        text=unicodedata.normalize('NFKC',text)
        text=re.sub(r'[\u200b-\u200f\u202a-\u202e\u2066-\u2069\ufeff\ufffd]','',text)
        text=text.translate(str.maketrans({'ي':'ی','ى':'ی','ك':'ک','ة':'ه'}))
        text=re.sub(r'\r\n?','\n',text);text=re.sub(r'[ \t]+',' ',text);text=re.sub(r'\n{4,}','\n\n\n',text).strip()
        if len(normalize_text(text))<MIN_EXTRACTED_TEXT_CHARS:
            raise HTTPException(status_code=400,detail='متن قابل جستجوی کافی از فایل استخراج نشد.')
        return text
    return extract_text(filename,path.read_bytes(),progress_callback)



def extract_document_result_from_path(filename: str, path: Path, progress_callback: Callable[[int,str],None] | None = None) -> dict[str,Any]:
    """Return normalized source text plus page-level ingestion diagnostics."""
    suffix=Path(filename).suffix.lower()
    if suffix=='.pdf':
        result=_run_coro_sync(_extract_pdf_source_result_async(path,progress_callback))
        text=_normalize_extracted_text(str(result.get('text') or ''))
        if len(normalize_text(text))<MIN_EXTRACTED_TEXT_CHARS:
            raise HTTPException(status_code=400,detail='متن قابل جستجوی کافی از فایل استخراج نشد.')
        result['text']=text
        return result
    text=extract_text(filename,path.read_bytes(),progress_callback)
    return _generic_source_result(text,suffix.lstrip('.') or 'text')


def _infer_section_title(text: str) -> str:
    for line in str(text or '').splitlines():
        clean=line.strip()
        if not clean:
            continue
        if _is_heading_line(clean):
            return clean[:180]
        if len(clean)<=90 and not re.search(r'[.!؟؛]$',clean):
            return clean[:180]
        break
    return ''


def _persian_digit_text(value: str) -> str:
    return str(value or '').translate(str.maketrans('0123456789','۰۱۲۳۴۵۶۷۸۹'))


def _chunk_search_aliases(content: str) -> str:
    """Create deterministic search expansions without inventing new business facts."""
    norm=normalize_text(content)
    aliases: list[str]=[]
    synonym_pairs=(
        ('روبار','روباری باربند بار روی سقف بار سقفی'),
        ('باربند','روبار روباری بار روی سقف'),
        ('ظرفیت','وزن مجاز حداکثر بار میزان بار تناژ'),
        ('وزن','ظرفیت وزن مجاز حداکثر بار'),
        ('نیسان','نیسان وانت وانت نیسان'),
        ('پیکان','پیکان وانت وانت پیکان'),
        ('خاور','کامیون خاور'),
        ('تبصره','شرط استثنا محدودیت'),
        ('استثنا','تبصره شرط محدودیت'),
        ('شمال تهران','مناطق شمال تهران'),
    )
    for needle,expanded in synonym_pairs:
        if normalize_text(needle) in norm:
            aliases.append(expanded)
    numbers=re.findall(r'(?<!\w)\d+(?:[.,]\d+)?(?!\w)',norm)
    for number in numbers[:24]:
        aliases.extend((number,_persian_digit_text(number)))
    if any(unit in norm for unit in ('کیلو','کیلوگرم','kg')):
        aliases.append('کیلو کیلوگرم kg')
    if 'تن' in norm:
        aliases.append('تن وزن تناژ')
    return ' '.join(dict.fromkeys(x.strip() for x in aliases if x.strip()))[:1600]


def chunk_source_result(result: dict[str,Any]) -> list[dict[str,Any]]:
    """Page/section-aware chunks. Rules and exceptions on the same page stay close."""
    pages=list(result.get('pages') or [])
    output: list[dict[str,Any]]=[]
    if SOURCE_STRUCTURE_AWARE_INDEXING and pages:
        for page in pages:
            page_number=int(page.get('page_number') or 0)
            page_text=_normalize_extracted_text(page.get('combined_text') or '')
            if not page_text:
                continue
            section=_infer_section_title(page_text)
            local_parts=chunk_text(page_text,size=RETRIEVAL_CHUNK_SIZE,overlap=RETRIEVAL_CHUNK_OVERLAP)
            for part in local_parts:
                visual=bool(page.get('vision_text'))
                prefix=f'صفحه {page_number}' + (f' — {section}' if section else '')
                content=f'{prefix}\n{part}'.strip()
                output.append({
                    'content':content,'page_start':page_number,'page_end':page_number,'section_title':section,
                    'chunk_type':'vision+text' if visual else 'text','search_aliases':_chunk_search_aliases(content),
                })
    else:
        for part in chunk_text(str(result.get('text') or '')):
            output.append({'content':part,'page_start':None,'page_end':None,'section_title':_infer_section_title(part),'chunk_type':'text','search_aliases':_chunk_search_aliases(part)})
    if pages and len(pages)>1:
        bridges=[]
        for left,right in zip(pages,pages[1:]):
            left_text=_normalize_extracted_text(left.get('combined_text') or '')
            right_text=_normalize_extracted_text(right.get('combined_text') or '')
            if not left_text or not right_text:
                continue
            tail=left_text[-650:].strip(); head=right_text[:850].strip()
            bridge=f"صفحات {left.get('page_number')}–{right.get('page_number')} — ادامه مطلب\n{tail}\n\n{head}".strip()
            bridges.append({'content':bridge,'page_start':int(left.get('page_number') or 0),'page_end':int(right.get('page_number') or 0),'section_title':'ادامه بین صفحات','chunk_type':'bridge','search_aliases':_chunk_search_aliases(bridge)})
        output.extend(bridges)
    seen=set();dedup=[]
    for row in output:
        key=hashlib.blake2b(normalize_text(row['content']).encode('utf-8'),digest_size=12).hexdigest()
        if key not in seen:
            seen.add(key);dedup.append(row)
    return dedup


def persist_document_ingestion(db: sqlite3.Connection, document_id: str, doc: Any, result: dict[str,Any], chunks: list[dict[str,Any]], embedding_rows: list[tuple[int,str,int,bytes]] | None = None) -> None:
    """Atomically replace page/chunk indexes and diagnostics for a document."""
    db.execute('DELETE FROM document_pages WHERE document_id=?',(document_id,))
    pages=list(result.get('pages') or [])
    if SOURCE_PAGE_STORAGE_ENABLED and pages:
        db.executemany(
            '''INSERT INTO document_pages(document_id,page_number,base_text,vision_text,combined_text,image_count,status,vision_error,character_count,page_fidelity,numeric_agreement)
               VALUES(?,?,?,?,?,?,?,?,?,?,?)''',
            [(document_id,int(r.get('page_number') or 0),str(r.get('base_text') or ''),str(r.get('vision_text') or ''),str(r.get('combined_text') or ''),int(r.get('image_count') or 0),str(r.get('status') or 'text'),r.get('vision_error'),int(r.get('character_count') or 0),float(r.get('page_fidelity') if r.get('page_fidelity') is not None else 1.0),None if r.get('numeric_agreement') is None else float(r.get('numeric_agreement'))) for r in pages],
        )
    db.execute('DELETE FROM chunks WHERE document_id=?',(document_id,))
    db.execute('DELETE FROM chunks_fts WHERE document_id=?',(document_id,))
    rows=[];fts=[]
    for idx,row in enumerate(chunks):
        content=str(row['content'])
        aliases=str(row.get('search_aliases') or '')
        rows.append((document_id,idx,content,row.get('page_start'),row.get('page_end'),row.get('section_title'),row.get('chunk_type') or 'text',aliases))
        searchable=(content+'\nکلیدواژه‌های جست‌وجو: '+aliases).strip() if aliases else content
        fts.append((searchable,doc['filename'],doc['visibility'],document_id,idx))
    db.executemany('''INSERT INTO chunks(document_id,chunk_index,content,page_start,page_end,section_title,chunk_type,search_aliases)
                      VALUES(?,?,?,?,?,?,?,?)''',rows)
    db.executemany('INSERT INTO chunks_fts(content,file_name,visibility,document_id,chunk_index) VALUES(?,?,?,?,?)',fts)
    index_semantic_chunks(db,document_id,[(row['content']+' '+str(row.get('search_aliases') or '')).strip() for row in chunks])
    db.execute('DELETE FROM chunk_embeddings WHERE document_id=?',(document_id,))
    if embedding_rows:
        db.executemany('INSERT INTO chunk_embeddings(document_id,chunk_index,model,dimensions,vector_blob,updated_at) VALUES(?,?,?,?,?,?)',
            [(document_id,idx,model,dimensions,blob,now_iso()) for idx,model,dimensions,blob in embedding_rows])
    db.execute('DELETE FROM knowledge_facts WHERE document_id=?',(document_id,))
    db.execute('DELETE FROM knowledge_facts_fts WHERE document_id=?',(document_id,))
    if SOURCE_FACT_EXTRACTION_ENABLED:
        fact_rows=[];fact_fts=[]
        for idx,row in enumerate(chunks):
            for fact in extract_structured_facts(str(row.get('content') or ''),page_start=row.get('page_start'),page_end=row.get('page_end'),section_title=row.get('section_title')):
                fact_rows.append((document_id,idx,fact['subject'],fact['fact_type'],fact['value_text'],fact['condition_text'],fact['fact_text'],fact.get('page_start'),fact.get('page_end'),fact.get('section_title') or '',now_iso()))
                fact_fts.append((fact['subject'],fact['fact_text'],fact['condition_text'],fact.get('section_title') or '',document_id,idx))
        if fact_rows:
            db.executemany('''INSERT INTO knowledge_facts(document_id,chunk_index,subject,fact_type,value_text,condition_text,fact_text,page_start,page_end,section_title,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)''',fact_rows)
            db.executemany('INSERT INTO knowledge_facts_fts(subject,fact_text,condition_text,section_title,document_id,chunk_index) VALUES(?,?,?,?,?,?)',fact_fts)
    stats=result.get('stats') or {};warnings=list(result.get('warnings') or [])
    quality=float(stats.get('ingestion_quality_pct') or 0.0)
    page_fidelity=max(0.0,min(1.0,float(stats.get('average_page_fidelity') if stats.get('average_page_fidelity') is not None else quality/100.0)))
    numeric_agreement=max(0.0,min(1.0,float(stats.get('average_numeric_agreement') if stats.get('average_numeric_agreement') is not None else 1.0)))
    source_status=str(result.get('status') or 'ready')
    gate_reasons=[]
    if quality<SOURCE_MIN_QUALITY_PCT:
        gate_reasons.append(f'کیفیت کل {quality:.1f}٪ کمتر از حد {SOURCE_MIN_QUALITY_PCT:.1f}٪ است')
    if page_fidelity<SOURCE_MIN_PAGE_FIDELITY:
        gate_reasons.append(f'وفاداری صفحه {page_fidelity*100:.1f}٪ کمتر از حد {SOURCE_MIN_PAGE_FIDELITY*100:.1f}٪ است')
    if numeric_agreement<SOURCE_MIN_NUMERIC_AGREEMENT:
        gate_reasons.append(f'تطابق عدد/واحد {numeric_agreement*100:.1f}٪ کمتر از حد {SOURCE_MIN_NUMERIC_AGREEMENT*100:.1f}٪ است')
    if source_status!='ready' and not gate_reasons:
        gate_reasons.append('پردازش منبع کامل نشده است')
    status=source_status
    if status=='partial' and not SOURCE_INCLUDE_PARTIAL_DOCUMENTS:
        status='error'
    quality_gate_reason='؛ '.join(gate_reasons) if gate_reasons else 'passed'
    db.execute('''UPDATE documents SET status=?,character_count=?,chunk_count=?,page_count=?,vision_candidate_pages=?,vision_success_pages=?,
                  vision_failed_pages=?,ingestion_quality_pct=?,page_fidelity_pct=?,numeric_agreement_pct=?,quality_gate_reason=?,ingestion_warnings_json=?,last_indexed_at=?,ingestion_version=? WHERE id=?''',
        (status,len(str(result.get('text') or '')),len(chunks),int(stats.get('page_count') or 0),int(stats.get('vision_candidate_pages') or 0),int(stats.get('vision_success_pages') or 0),int(stats.get('vision_failed_pages') or 0),quality,round(page_fidelity*100,1),round(numeric_agreement*100,1),quality_gate_reason,json.dumps(warnings,ensure_ascii=False),now_iso(),INGESTION_VERSION,document_id))
    bump_knowledge_version(db)


def _is_heading_line(line: str) -> bool:
    clean = line.strip()
    if not clean or len(clean) > 140:
        return False
    return bool(
        clean.startswith(('--- ', '#', 'فصل ', 'بخش ', 'ماده ', 'تبصره ', 'سؤال ', 'سوال '))
        or clean.endswith((':', '：'))
        or re.match(r'^(?:\d+|[۰-۹]+)[\).\-:]\s+', clean)
    )


def _split_long_block(block: str, size: int) -> list[str]:
    sentences = [x.strip() for x in re.split(r'(?<=[\.؟!؛])\s+|\n+', block) if x.strip()]
    if not sentences:
        return [block[i:i+size] for i in range(0, len(block), size)]
    out: list[str] = []
    current = ''
    for sentence in sentences:
        if current and len(current) + len(sentence) + 1 > size:
            out.append(current.strip())
            current = ''
        if len(sentence) > size:
            if current:
                out.append(current.strip()); current = ''
            for offset in range(0, len(sentence), size):
                out.append(sentence[offset:offset+size].strip())
        else:
            current = f'{current} {sentence}'.strip()
    if current:
        out.append(current.strip())
    return [x for x in out if x]


def chunk_text(text: str, size: int | None = None, overlap: int | None = None) -> list[str]:
    """Structure-aware chunking that keeps headings, paragraphs, pages and tables together."""
    size = size or RETRIEVAL_CHUNK_SIZE
    overlap = RETRIEVAL_CHUNK_OVERLAP if overlap is None else overlap
    txt = re.sub(r'\r\n?', '\n', text).strip()
    if not txt:
        return []
    raw_blocks = [b.strip() for b in re.split(r'\n\s*\n+', txt) if b.strip()]
    blocks: list[str] = []
    current_heading = ''
    for raw in raw_blocks:
        lines = [line.strip() for line in raw.splitlines() if line.strip()]
        if not lines:
            continue
        if _is_heading_line(lines[0]):
            current_heading = lines[0]
        body = '\n'.join(lines)
        if current_heading and not body.startswith(current_heading):
            body = f'{current_heading}\n{body}'
        blocks.extend(_split_long_block(body, size))
    chunks: list[str] = []
    current = ''
    for block in blocks:
        candidate = f'{current}\n\n{block}'.strip() if current else block
        if current and len(candidate) > size:
            chunks.append(current.strip())
            tail = current[-overlap:].strip() if overlap else ''
            current = f'{tail}\n\n{block}'.strip() if tail else block
            if len(current) > size * 1.35:
                parts = _split_long_block(current, size)
                chunks.extend(parts[:-1])
                current = parts[-1] if parts else ''
        else:
            current = candidate
    if current:
        chunks.append(current.strip())
    # Deduplicate overlap-only duplicates while preserving order.
    result: list[str] = []
    seen: set[str] = set()
    for chunk in chunks:
        key = hashlib.sha1(normalize_text(chunk).encode('utf-8'), usedforsecurity=False).hexdigest()
        if chunk and key not in seen:
            result.append(chunk); seen.add(key)
    return result


_SEMANTIC_GROUPS = [
    {'لغو','کنسلی','فسخ','خاتمه','پایان','انصراف'},
    {'قیمت','هزینه','مبلغ','تعرفه','بها','پرداخت'},
    {'بازپرداخت','استرداد','عودت','برگشت'},
    {'ثبت نام','ثبتنام','نامنویسی','عضویت'},
    {'پشتیبانی','تیکت','درخواست','پیگیری','کارشناس'},
    {'رمز','گذرواژه','پسورد','ورود','لاگین'},
    {'مرخصی','تعطیلی','غیبت','استراحت'},
    {'قرارداد','همکاری','توافق','تعهد'},
    {'ارسال','تحویل','بارگذاری','آپلود'},
    {'روبار','روباری','باربند','بار روی سقف','بار سقفی'},
    {'ظرفیت','وزن','وزن مجاز','حداکثر بار','تناژ','کیلوگرم','تن'},
    {'نیسان','نیسان وانت','وانت نیسان'},
    {'پیکان','پیکان وانت','وانت پیکان'},
    {'خاور','کامیون خاور'},
    {'شرط','استثنا','تبصره','محدودیت','در صورتی','به جز'},
]


def _semantic_profile(text: str) -> tuple[dict[str, float], set[str]]:
    normalized = normalize_text(text)
    tokens = search_tokens(text)
    expanded = set(tokens)
    normalized_tokens = set(normalized.split())
    for group in _SEMANTIC_GROUPS:
        if group & (expanded | normalized_tokens):
            expanded.update(_stem_search_token(x.replace(' ', '')) for x in group)
    weights: dict[str, float] = {token: 1.0 for token in expanded if len(token) >= 2}
    ordered = [t for t in tokens if len(t) >= 2]
    for i in range(len(ordered)-1):
        weights[f'{ordered[i]}_{ordered[i+1]}'] = 1.35
    compact = normalized.replace(' ', '')
    ngrams: set[str] = set()
    if len(compact) >= 4:
        for i in range(min(len(compact)-2, 240)):
            gram = compact[i:i+3]
            if len(gram) == 3:
                ngrams.add(gram)
    return weights, ngrams


def _semantic_bucket_values(text: str) -> list[str]:
    weights, ngrams = _semantic_profile(text)
    values = list(weights.keys()) + sorted(ngrams)[:SEMANTIC_BUCKET_LIMIT]
    buckets: list[str] = []
    for value in values:
        bucket = hashlib.blake2b(value.encode('utf-8'), digest_size=6).hexdigest()
        if bucket not in buckets:
            buckets.append(bucket)
        if len(buckets) >= SEMANTIC_BUCKET_LIMIT:
            break
    return buckets


def _semantic_similarity(question: str, terms_json: str | None, ngrams_json: str | None, fallback_text: str = '') -> float:
    q_weights, q_ngrams = _semantic_profile(question)
    try:
        d_weights = {str(k): float(v) for k,v in json.loads(terms_json or '{}').items()}
        d_ngrams = set(json.loads(ngrams_json or '[]'))
    except (ValueError, TypeError, json.JSONDecodeError):
        d_weights, d_ngrams = _semantic_profile(fallback_text)
    if not q_weights or not d_weights:
        term_score = 0.0
    else:
        dot = sum(weight * d_weights.get(term, 0.0) for term, weight in q_weights.items())
        q_norm = math.sqrt(sum(v*v for v in q_weights.values()))
        d_norm = math.sqrt(sum(v*v for v in d_weights.values()))
        term_score = dot / (q_norm*d_norm) if q_norm and d_norm else 0.0
    ngram_score = len(q_ngrams & d_ngrams) / len(q_ngrams | d_ngrams) if q_ngrams and d_ngrams else 0.0
    return min(1.0, term_score * 0.78 + ngram_score * 0.22)



async def _generate_remote_embeddings(texts: list[str]) -> list[list[float]]:
    if not REMOTE_EMBEDDING_ENABLED or not EMBEDDING_MODEL or not texts:
        return []
    available=ordered_available_ai_slots_for_route('embedding')
    desired = next((slot for slot in available if int(slot['slot']) == EMBEDDING_API_SLOT), None)
    if not desired:
        desired=next((slot for slot in available if str(slot.get('embedding_model') or '').strip()),None)
    if not desired:
        raise RuntimeError('هیچ API فعال با مدل Embedding در دسترس نیست.')
    embedding_model=str(desired.get('embedding_model') or EMBEDDING_MODEL).strip() or EMBEDDING_MODEL
    started=time.perf_counter()
    headers={'Content-Type':'application/json',AI_AUTH_HEADER:f"{AI_AUTH_SCHEME} {desired['api_key']}".strip()}
    try:
        headers.update({str(k):str(v) for k,v in json.loads(AI_EXTRA_HEADERS_JSON).items()})
    except (json.JSONDecodeError,AttributeError):
        pass
    payload: dict[str,Any]={'model':embedding_model,'input':[str(x)[:EMBEDDING_MAX_CHARS] for x in texts]}
    if embedding_model.startswith('openai/text-embedding-3'):
        payload['dimensions']=EMBEDDING_DIMENSIONS
    client=_HTTP_CLIENT
    owns_client=False
    if client is None:
        client=httpx.AsyncClient(timeout=EMBEDDING_TIMEOUT_SECONDS,follow_redirects=True);owns_client=True
    try:
        response=await client.post(f"{desired['base_url']}/embeddings",headers=headers,json=payload,timeout=EMBEDDING_TIMEOUT_SECONDS)
        if response.status_code>=400:
            detail=response.text[:350]
            elapsed=int((time.perf_counter()-started)*1000)
            record_api_call_event(slot=int(desired['slot']),provider_label=desired['label'],model=embedding_model,model_route='embedding',status='error',http_status=response.status_code,error_class='embedding_http_error',error_detail=detail,response_ms=elapsed,attempt_index=1,was_failover=False)
            if response.status_code in {401,402,403,429}:
                mark_api_slot_blocked(int(desired['slot']),detail,AI_KEY_RATE_LIMIT_COOLDOWN_SECONDS if response.status_code==429 else AI_KEY_QUOTA_COOLDOWN_SECONDS,'limited')
            elif response.status_code in {408,425,500,502,503,504,529}:
                mark_api_slot_blocked(int(desired['slot']),detail,AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'transient_error')
            raise RuntimeError(f'Embedding API HTTP {response.status_code}: {detail}')
        data=response.json();items=data.get('data') or []
        ordered=sorted(items,key=lambda row:int(row.get('index') or 0))
        vectors=[list(map(float,row.get('embedding') or [])) for row in ordered]
        if len(vectors)!=len(texts) or any(not vector for vector in vectors):
            raise RuntimeError('Embedding API تعداد بردار کامل برنگرداند.')
        mark_api_slot_success(int(desired['slot']))
        elapsed=int((time.perf_counter()-started)*1000)
        record_api_call_event(slot=int(desired['slot']),provider_label=desired['label'],model=embedding_model,model_route='embedding',status='success',http_status=200,error_class=None,error_detail=None,response_ms=elapsed,attempt_index=1,was_failover=False)
        return vectors
    finally:
        if owns_client:
            await client.aclose()


def build_chunk_embedding_rows(chunks: list[dict[str,Any]]) -> tuple[list[tuple[int,str,int,bytes]], str | None]:
    if not REMOTE_EMBEDDING_ENABLED or not chunks:
        return [],None
    rows: list[tuple[int,str,int,bytes]]=[]
    try:
        for offset in range(0,len(chunks),EMBEDDING_BATCH_SIZE):
            batch=chunks[offset:offset+EMBEDDING_BATCH_SIZE]
            texts=[(str(row.get('content') or '')+'\n'+str(row.get('search_aliases') or '')).strip()[:EMBEDDING_MAX_CHARS] for row in batch]
            vectors=_run_coro_sync(_generate_remote_embeddings(texts))
            for local_index,vector in enumerate(vectors):
                rows.append((offset+local_index,EMBEDDING_MODEL,len(vector),pack_vector(vector)))
        return rows,None
    except Exception as exc:
        return [],sanitize_answer_text(str(exc))[:500]


def _query_embedding(question: str) -> list[float]:
    if not REMOTE_EMBEDDING_ENABLED:
        return []
    key=hashlib.sha256((EMBEDDING_MODEL+'|'+normalize_text(question)).encode('utf-8')).hexdigest()
    cached=_QUERY_EMBED_CACHE.get(key)
    if cached and time.time()-cached[0] < 3600:
        return cached[1]
    try:
        vectors=_run_coro_sync(_generate_remote_embeddings([question[:EMBEDDING_MAX_CHARS]]))
        vector=vectors[0] if vectors else []
        if vector:
            if len(_QUERY_EMBED_CACHE)>300:
                oldest=min(_QUERY_EMBED_CACHE.items(),key=lambda item:item[1][0])[0]
                _QUERY_EMBED_CACHE.pop(oldest,None)
            _QUERY_EMBED_CACHE[key]=(time.time(),vector)
        return vector
    except Exception:
        return []


def _embedding_candidate_scores(question: str, vis: list[str]) -> dict[tuple[str,int],float]:
    query_vector=_query_embedding(question)
    if not query_vector:
        return {}
    placeholders=','.join('?' for _ in vis)
    with get_db() as db:
        rows=db.execute(
            f'''SELECT e.document_id,e.chunk_index,e.dimensions,e.vector_blob
                FROM chunk_embeddings e JOIN documents d ON d.id=e.document_id
                WHERE e.model=? AND d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                ORDER BY e.document_id,e.chunk_index LIMIT ?''',
            (EMBEDDING_MODEL,*vis,EMBEDDING_SCAN_LIMIT),
        ).fetchall()
    scores: dict[tuple[str,int],float]={}
    for row in rows:
        if int(row['dimensions'])!=len(query_vector):
            continue
        sim=cosine_similarity(query_vector,unpack_vector(row['vector_blob']))
        score=max(0.0,min(1.0,(sim+0.15)/1.15))
        if score>=0.18:
            scores[(row['document_id'],int(row['chunk_index']))]=score
    return scores


def _fact_candidates(question: str, vis: list[str]) -> dict[tuple[str,int],float]:
    if not SOURCE_FACT_EXTRACTION_ENABLED:
        return {}
    match=build_fts_query(question)
    placeholders=','.join('?' for _ in vis)
    scores: dict[tuple[str,int],float]={}
    with get_db() as db:
        try:
            rows=db.execute(
                f'''SELECT f.document_id,f.chunk_index,bm25(knowledge_facts_fts) AS rank
                    FROM knowledge_facts_fts f JOIN documents d ON d.id=f.document_id
                    WHERE knowledge_facts_fts MATCH ? AND d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                    ORDER BY rank ASC LIMIT ?''',
                (match,*vis,FACT_RETRIEVAL_LIMIT),
            ).fetchall()
        except sqlite3.OperationalError:
            rows=[]
    for pos,row in enumerate(rows):
        scores[(row['document_id'],int(row['chunk_index']))]=max(0.25,0.88-pos*0.012)
    return scores

def index_semantic_chunks(db: sqlite3.Connection, document_id: str, parts: list[str]) -> None:
    db.execute('DELETE FROM chunk_semantic_buckets WHERE document_id=?', (document_id,))
    db.execute('DELETE FROM chunk_semantic WHERE document_id=?', (document_id,))
    semantic_rows=[]; bucket_rows=[]; ts=now_iso()
    for idx, content in enumerate(parts):
        weights, ngrams = _semantic_profile(content)
        semantic_rows.append((document_id, idx, json.dumps(weights, ensure_ascii=False, separators=(',',':')), json.dumps(sorted(ngrams), ensure_ascii=False), ts))
        bucket_rows.extend((bucket, document_id, idx) for bucket in _semantic_bucket_values(content))
    if semantic_rows:
        db.executemany('INSERT INTO chunk_semantic(document_id,chunk_index,terms_json,ngrams_json,updated_at) VALUES(?,?,?,?,?)', semantic_rows)
    if bucket_rows:
        db.executemany('INSERT OR IGNORE INTO chunk_semantic_buckets(bucket,document_id,chunk_index) VALUES(?,?,?)', bucket_rows)


def backfill_semantic_index(db: sqlite3.Connection) -> None:
    if not SEMANTIC_SEARCH_ENABLED:
        return
    source_count = int(db.execute('SELECT COUNT(*) FROM chunks').fetchone()[0])
    indexed_count = int(db.execute('SELECT COUNT(*) FROM chunk_semantic').fetchone()[0])
    if source_count == indexed_count:
        return
    document_ids = [row['document_id'] for row in db.execute('SELECT DISTINCT document_id FROM chunks ORDER BY document_id').fetchall()]
    for document_id in document_ids:
        parts = [(row['content']+' '+str(row['search_aliases'] or '')).strip() for row in db.execute('SELECT content,search_aliases FROM chunks WHERE document_id=? ORDER BY chunk_index', (document_id,)).fetchall()]
        index_semantic_chunks(db, document_id, parts)


def _fast_fts_probe_score(question: str, vis: list[str]) -> float:
    """Cheap local confidence probe used to avoid a remote embedding call for obvious matches."""
    match=build_fts_query(question);placeholders=','.join('?' for _ in vis)
    with get_db() as db:
        try:
            rows=db.execute(
                f"""SELECT c.content,COALESCE(c.search_aliases,'') AS search_aliases,d.filename
                    FROM chunks_fts f JOIN chunks c ON c.document_id=f.document_id AND c.chunk_index=f.chunk_index
                    JOIN documents d ON d.id=f.document_id
                    WHERE chunks_fts MATCH ? AND d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                    ORDER BY bm25(chunks_fts) ASC LIMIT 12""",(match,*vis)).fetchall()
        except sqlite3.OperationalError:
            rows=[]
    return max([_retrieval_score(question,(row['content']+' '+str(row['search_aliases'] or '')).strip(),row['filename']) for row in rows] or [0.0])


def _numeric_alignment_score(question: str, content: str) -> float:
    qnums=set(re.findall(r'(?<!\w)\d+(?:[.,]\d+)?(?!\w)',normalize_text(question)))
    if not qnums:
        return 0.0
    cnums=set(re.findall(r'(?<!\w)\d+(?:[.,]\d+)?(?!\w)',normalize_text(content)))
    if not cnums:
        return 0.0
    return len(qnums & cnums)/max(1,len(qnums))


def retrieve(question: str, user: dict[str, Any] | None, integration: bool = False) -> list[dict[str, Any]]:
    """R29 hybrid retrieval: FTS + local semantic + real embeddings + structured facts + rerank."""
    vis=allowed_visibilities(user,integration)
    placeholders=','.join('?' for _ in vis)
    match_expr=build_fts_query(question)
    candidate_limit=max(HYBRID_RERANK_TOP_N*3,RETRIEVAL_TOP_K*8,96)
    query_buckets=_semantic_bucket_values(question) if SEMANTIC_SEARCH_ENABLED else []
    fact_scores=_fact_candidates(question,vis)
    # Remote query embeddings are most valuable for paraphrases and weak lexical
    # matches. Clear FTS/fact hits stay fully local for lower latency.
    local_probe=max(_fast_fts_probe_score(question,vis),max(fact_scores.values(),default=0.0))
    embedding_scores=_embedding_candidate_scores(question,vis) if REMOTE_EMBEDDING_ENABLED and local_probe<EMBEDDING_QUERY_SKIP_CONFIDENCE else {}
    # Do not flood SQLite/Python with weak embedding matches; keep the strongest
    # semantic candidates and let the second-stage reranker decide the final order.
    embedding_top=dict(sorted(embedding_scores.items(),key=lambda x:-x[1])[:HYBRID_RERANK_TOP_N*5])
    fact_top=dict(sorted(fact_scores.items(),key=lambda x:-x[1])[:HYBRID_RERANK_TOP_N*3])

    with get_db() as db:
        candidates: dict[tuple[str,int],dict[str,Any]]={}
        try:
            rows=db.execute(
                f"""SELECT f.document_id,f.chunk_index,c.content,COALESCE(c.search_aliases,'') AS search_aliases,c.page_start,c.page_end,c.section_title,c.chunk_type,
                           d.filename AS file_name,d.visibility,bm25(chunks_fts) AS rank
                    FROM chunks_fts f JOIN chunks c ON c.document_id=f.document_id AND c.chunk_index=f.chunk_index
                    JOIN documents d ON d.id=f.document_id
                    WHERE chunks_fts MATCH ? AND d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                    ORDER BY rank ASC LIMIT ?""",
                (match_expr,*vis,candidate_limit),
            ).fetchall()
        except sqlite3.OperationalError:
            rows=[]
        for row in rows:
            candidates[(row['document_id'],int(row['chunk_index']))]=dict(row)

        # Diversity rescue: broad OR queries over large corpora can let a few long
        # booklets dominate SQLite BM25. Probe several distinctive query terms
        # separately so a small uploaded document with an exact phrase is not buried.
        distinct_terms=sorted(set(search_tokens(question)),key=lambda x:(-len(x),x))[:5]
        for term in distinct_terms:
            try:
                term_expr=(f'{term}*' if len(term)>=3 else f'"{term}"')
                term_rows=db.execute(
                    f"""SELECT f.document_id,f.chunk_index,c.content,COALESCE(c.search_aliases,'') AS search_aliases,c.page_start,c.page_end,c.section_title,c.chunk_type,
                               d.filename AS file_name,d.visibility,bm25(chunks_fts) AS rank
                        FROM chunks_fts f JOIN chunks c ON c.document_id=f.document_id AND c.chunk_index=f.chunk_index
                        JOIN documents d ON d.id=f.document_id
                        WHERE chunks_fts MATCH ? AND d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                        ORDER BY rank ASC LIMIT 10""",
                    (term_expr,*vis),
                ).fetchall()
            except sqlite3.OperationalError:
                term_rows=[]
            for row in term_rows:
                candidates.setdefault((row['document_id'],int(row['chunk_index'])),dict(row))

        if query_buckets:
            bucket_placeholders=','.join('?' for _ in query_buckets)
            semantic_rows=db.execute(
                f"""SELECT b.document_id,b.chunk_index,c.content,COALESCE(c.search_aliases,'') AS search_aliases,c.page_start,c.page_end,c.section_title,c.chunk_type,d.filename AS file_name,d.visibility,
                           0.0 AS rank,COUNT(*) AS bucket_hits
                    FROM chunk_semantic_buckets b
                    JOIN chunks c ON c.document_id=b.document_id AND c.chunk_index=b.chunk_index
                    JOIN documents d ON d.id=b.document_id
                    WHERE b.bucket IN ({bucket_placeholders}) AND d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                    GROUP BY b.document_id,b.chunk_index,c.content,c.search_aliases,c.page_start,c.page_end,c.section_title,c.chunk_type,d.filename,d.visibility
                    ORDER BY bucket_hits DESC LIMIT ?""",
                (*query_buckets,*vis,SEMANTIC_CANDIDATE_LIMIT),
            ).fetchall()
            for row in semantic_rows:
                key=(row['document_id'],int(row['chunk_index']))
                if key in candidates:
                    candidates[key]['bucket_hits']=int(row['bucket_hits'])
                else:
                    candidates[key]=dict(row)

        # Materialize candidates found only by real embeddings or structured facts.
        extra_keys=[key for key in set(embedding_top)|set(fact_top) if key not in candidates]
        for offset in range(0,len(extra_keys),200):
            batch=extra_keys[offset:offset+200]
            if not batch:
                continue
            where=' OR '.join('(c.document_id=? AND c.chunk_index=?)' for _ in batch)
            params=[value for key in batch for value in key]
            extra_rows=db.execute(
                f"""SELECT c.document_id,c.chunk_index,c.content,COALESCE(c.search_aliases,'') AS search_aliases,c.page_start,c.page_end,c.section_title,c.chunk_type,
                           d.filename AS file_name,d.visibility,0.0 AS rank
                    FROM chunks c JOIN documents d ON d.id=c.document_id
                    WHERE ({where}) AND d.is_enabled=1 AND d.status='ready'""",
                params,
            ).fetchall()
            for row in extra_rows:
                candidates[(row['document_id'],int(row['chunk_index']))]=dict(row)

        if len(candidates)<RETRIEVAL_TOP_K:
            seed=int(hashlib.blake2b(normalize_text(question).encode('utf-8'),digest_size=4).hexdigest(),16)
            fallback_rows=db.execute(
                f"""SELECT c.document_id,c.chunk_index,c.content,COALESCE(c.search_aliases,'') AS search_aliases,c.page_start,c.page_end,c.section_title,c.chunk_type,d.filename AS file_name,d.visibility,0.0 AS rank
                    FROM chunks c JOIN documents d ON d.id=c.document_id
                    WHERE d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                    ORDER BY abs(((c.id * 1103515245) + ?) % 2147483647) LIMIT ?""",
                (*vis,seed,min(RETRIEVAL_SCAN_LIMIT,420)),
            ).fetchall()
            for row in fallback_rows:
                candidates.setdefault((row['document_id'],int(row['chunk_index'])),dict(row))

        semantic_map: dict[tuple[str,int],sqlite3.Row]={}
        keys=list(candidates.keys())
        if SEMANTIC_SEARCH_ENABLED and keys:
            for offset in range(0,len(keys),250):
                batch=keys[offset:offset+250]
                where=' OR '.join('(document_id=? AND chunk_index=?)' for _ in batch)
                params=[value for key in batch for value in key]
                for row in db.execute(f'SELECT document_id,chunk_index,terms_json,ngrams_json FROM chunk_semantic WHERE {where}',params).fetchall():
                    semantic_map[(row['document_id'],int(row['chunk_index']))]=row

        scored=[]
        for key,row in candidates.items():
            semantic=semantic_map.get(key)
            searchable=(row['content']+' '+str(row.get('search_aliases') or '')).strip()
            lexical=_retrieval_score(question,searchable,row['file_name'])
            semantic_score=_semantic_similarity(question,semantic['terms_json'] if semantic else None,semantic['ngrams_json'] if semantic else None,searchable) if SEMANTIC_SEARCH_ENABLED else 0.0
            numeric_score=_numeric_alignment_score(question,searchable)
            embedding_score=float(embedding_top.get(key) or 0.0)
            fact_score=float(fact_top.get(key) or 0.0)
            base_score=min(1.5,lexical*0.58+semantic_score*0.40+embedding_score*0.50+fact_score*0.16+min(0.10,int(row.get('bucket_hits') or 0)*0.008)+numeric_score*SOURCE_NUMERIC_FACT_BOOST)
            if base_score<RETRIEVAL_MIN_SCORE and embedding_score<0.24 and fact_score<0.30:
                continue
            confidence='high' if base_score>=RETRIEVAL_HIGH_CONFIDENCE else 'medium' if base_score>=RETRIEVAL_MEDIUM_CONFIDENCE else 'low'
            scored.append({'source_type':'document','document_id':row['document_id'],'chunk_index':int(row['chunk_index']),
                'content':row['content'],'file_name':row['file_name'],'score':round(base_score,4),'lexical_score':round(lexical,4),
                'semantic_score':round(semantic_score,4),'embedding_score':round(embedding_score,4),'fact_score':round(fact_score,4),
                'numeric_score':round(numeric_score,4),'bucket_hits':int(row.get('bucket_hits') or 0),'confidence':confidence,'rank':float(row.get('rank') or 0.0),
                'page_start':row.get('page_start'),'page_end':row.get('page_end'),'section_title':row.get('section_title'),'chunk_type':row.get('chunk_type')})

        if HYBRID_RERANK_ENABLED:
            scored=rerank_hybrid_candidates(question,scored,top_n=HYBRID_RERANK_TOP_N)
        else:
            scored.sort(key=lambda item:(-item['score'],item['rank'],item['document_id'],item['chunk_index']))
        primary_hits=scored[:RETRIEVAL_TOP_K]
        results=[];covered_chunks={};used_chars=0
        for hit in primary_hits:
            if hit['chunk_index'] in covered_chunks.get(hit['document_id'],set()):
                continue
            start_index=max(0,hit['chunk_index']-RETRIEVAL_NEIGHBOR_CHUNKS)
            end_index=hit['chunk_index']+RETRIEVAL_NEIGHBOR_CHUNKS
            neighbor_rows=db.execute('SELECT chunk_index,content FROM chunks WHERE document_id=? AND chunk_index BETWEEN ? AND ? ORDER BY chunk_index',(hit['document_id'],start_index,end_index)).fetchall()
            covered_chunks.setdefault(hit['document_id'],set()).update(int(row['chunk_index']) for row in neighbor_rows)
            expanded='\n\n'.join(row['content'].strip() for row in neighbor_rows if row['content'].strip()) or hit['content']
            remaining=RETRIEVAL_MAX_CONTEXT_CHARS-used_chars
            if remaining<=0:
                break
            if len(expanded)>remaining:
                expanded=expanded[:remaining].rsplit(' ',1)[0].rstrip() or expanded[:remaining]
            used_chars+=len(expanded)
            results.append({**hit,'content':expanded,'excerpt':expanded[:500]})
            if used_chars>=RETRIEVAL_MAX_CONTEXT_CHARS:
                break

    # The built-in Barsan corpus is only 103 pages. Compare normal chunk retrieval
    # with a deterministic full-page pass so OCR noise or FTS ranking can never
    # make an existing baseline rule disappear from chat.
    builtin_pages=_builtin_page_rescue(question,vis,limit=RETRIEVAL_TOP_K)
    if builtin_pages:
        merged=results+builtin_pages
        merged.sort(key=lambda item:-float(item.get('score') or 0))
        dedup=[];seen_pages=set();seen_chunks=set()
        for item in merged:
            page_key=(item.get('document_id'),item.get('page_start'),item.get('page_end'))
            chunk_key=(item.get('document_id'),item.get('chunk_index'))
            if item.get('chunk_type')=='builtin-page-rescue':
                if page_key in seen_pages:
                    continue
                seen_pages.add(page_key)
            elif chunk_key in seen_chunks:
                continue
            seen_chunks.add(chunk_key)
            dedup.append(item)
            if len(dedup)>=RETRIEVAL_TOTAL_ITEMS:
                break
        results=dedup

    best_score=max([float(item.get('score') or 0) for item in results] or [0.0])
    if CHAT_UNIFIED_KNOWLEDGE_MODE and CHAT_CORPUS_RESCUE_ENABLED and best_score < RETRIEVAL_MEDIUM_CONFIDENCE:
        existing={(item['document_id'],int(item['chunk_index'])) for item in results}
        rescue=[]
        with get_db() as rescue_db:
            rows=rescue_db.execute(
                f"""SELECT c.document_id,c.chunk_index,c.content,COALESCE(c.search_aliases,'') AS search_aliases,c.page_start,c.page_end,c.section_title,c.chunk_type,d.filename AS file_name,d.visibility
                    FROM chunks c JOIN documents d ON d.id=c.document_id
                    WHERE d.visibility IN ({placeholders}) AND d.is_enabled=1 AND d.status='ready'
                    ORDER BY d.is_builtin DESC,d.id,c.chunk_index LIMIT ?""",
                (*vis,CHAT_CORPUS_RESCUE_LIMIT),
            ).fetchall()
        for raw_row in rows:
            row=dict(raw_row);key=(row['document_id'],int(row['chunk_index']))
            if key in existing:
                continue
            searchable=(row['content']+' '+str(row.get('search_aliases') or '')).strip()
            lexical=_retrieval_score(question,searchable,row['file_name'])
            numeric_score=_numeric_alignment_score(question,searchable)
            embedding_score=float(embedding_scores.get(key) or 0.0)
            fact_score=float(fact_scores.get(key) or 0.0)
            rescue_score=min(1.5,lexical*0.72+embedding_score*0.62+fact_score*0.18+numeric_score*SOURCE_NUMERIC_FACT_BOOST)
            if rescue_score < RETRIEVAL_MIN_SCORE:
                continue
            rescue.append({'source_type':'document','document_id':row['document_id'],'chunk_index':int(row['chunk_index']),
                'content':row['content'],'file_name':row['file_name'],'score':round(rescue_score,4),'lexical_score':round(lexical,4),
                'semantic_score':0.0,'embedding_score':round(embedding_score,4),'fact_score':round(fact_score,4),'numeric_score':round(numeric_score,4),
                'page_start':row.get('page_start'),'page_end':row.get('page_end'),'section_title':row.get('section_title'),'chunk_type':row.get('chunk_type'),
                'confidence':'high' if rescue_score>=RETRIEVAL_HIGH_CONFIDENCE else 'medium' if rescue_score>=RETRIEVAL_MEDIUM_CONFIDENCE else 'low','rank':0.0,'excerpt':row['content'][:500]})
        if rescue:
            rescue=rerank_hybrid_candidates(question,rescue,top_n=RETRIEVAL_TOP_K) if HYBRID_RERANK_ENABLED else sorted(rescue,key=lambda item:-item['score'])[:RETRIEVAL_TOP_K]
            merged=results+rescue
            merged.sort(key=lambda item:-float(item.get('score') or 0))
            return merged[:RETRIEVAL_TOTAL_ITEMS]
    return results


def _builtin_page_rescue(question: str, vis: tuple[str, ...] | list[str], limit: int | None = None) -> list[dict[str, Any]]:
    """Exhaustive fallback over the four reviewed built-in Barsan booklets.

    Manager training is resolved before document retrieval, so this can never
    override the user's own training. It exists only to prevent an OCR/FTS miss
    from turning an answer that is visibly present in a booklet into a gap.
    """
    q_norm=normalize_text(question)
    q_tokens=search_tokens(question)
    if not q_norm or not q_tokens:
        return []
    placeholders=','.join('?' for _ in vis)
    with get_db() as db:
        rows=db.execute(
            f"""SELECT p.document_id,p.page_number,p.combined_text,p.base_text,p.vision_text,
                       d.filename,d.visibility,d.source_key
                FROM document_pages p JOIN documents d ON d.id=p.document_id
                WHERE d.is_builtin=1 AND d.is_enabled=1
                  AND d.visibility IN ({placeholders})
                  AND d.status='ready'
                ORDER BY d.source_key,p.page_number""",
            tuple(vis),
        ).fetchall()
    high_signal={'کنسلی','لغو','سلب','توقف','بارگیری','ظرفیت','وزن','اضافه','باسکول','روبار','روباری',
                 'نیسان','پیکان','خاور','کفی','مسقف','روباز','صف','راننده','غیرفعال','جابجایی','موقعیت',
                 'مبدا','مقصد','بازنگری','فاکتور','رزرو','کمیسیون','خرابی','تصادف','طرح','مرخصی'}
    q_set=set(q_tokens);scored=[]
    for row in rows:
        text=str(row['combined_text'] or row['base_text'] or row['vision_text'] or '').strip()
        if not text:
            continue
        c_set=set(search_tokens(text));shared=q_set & c_set
        if not shared:
            continue
        weighted_total=sum(2.0 if t in high_signal else 1.0 for t in q_set)
        weighted_shared=sum(2.0 if t in high_signal else 1.0 for t in shared)
        coverage=weighted_shared/max(1.0,weighted_total)
        lexical=_retrieval_score(question,text,str(row['filename']))
        # A rule is often one short line inside a long operational page. Score the
        # strongest local sentence as well so the exact rule outranks a long page
        # that merely contains the same words in unrelated paragraphs.
        fragments=[x.strip() for x in re.split(r'[\n؟!؛]+|(?<=[.])\s+',text) if len(x.strip())>=12]
        sentence_score=max([_retrieval_score(question,fragment,str(row['filename'])) for fragment in fragments[:180]] or [0.0])
        numeric=_numeric_alignment_score(question,text)
        page_norm=normalize_text(text)
        ordered=[t for t in q_tokens if len(t)>=2]
        pair_hits=sum(1 for a,b in zip(ordered,ordered[1:]) if a in page_norm and b in page_norm)
        phrase_bonus=min(0.18,0.06*pair_hits)
        source_bonus=0.0
        fn=normalize_text(str(row['filename']))
        if any(x in q_norm for x in ('سلب','کنسلی','لغو')) and ('فرایند سلب' in fn or '04_barsan' in fn):
            source_bonus=0.14
        elif any(x in q_norm for x in ('توقف','بازنگری','مراحل سرویس','فاکتور','رزرو')) and ('پشتیبانی چرخه' in fn or '03_barsan' in fn):
            source_bonus=0.12
        elif any(x in q_norm for x in ('صف','راننده','غیرفعال','جابجایی','مرخصی','موقعیت')) and ('عملیات رانندگان' in fn or '02_barsan' in fn):
            source_bonus=0.12
        elif any(x in q_norm for x in ('بارگیری','ظرفیت','وزن','باسکول','روبار','نیسان','پیکان','خاور')) and ('شناخت ناوگان' in fn or '01_barsan' in fn):
            source_bonus=0.10
        local_rule_bonus=min(0.24,max(0.0,sentence_score-0.52)*0.55)
        score=min(1.5,max(lexical,sentence_score,coverage*0.82+numeric*0.10)+phrase_bonus+source_bonus+local_rule_bonus)
        if len(q_set)>=3 and len(shared)<2 and score<0.62:
            continue
        if score<0.24:
            continue
        scored.append({
            'source_type':'document','document_id':str(row['document_id']),
            'chunk_index':-int(row['page_number'] or 0),'content':text,
            'file_name':str(row['filename']),'score':round(score,4),'lexical_score':round(lexical,4),
            'sentence_score':round(sentence_score,4),
            'semantic_score':0.0,'embedding_score':0.0,'fact_score':0.0,'numeric_score':round(numeric,4),
            'page_start':int(row['page_number'] or 0),'page_end':int(row['page_number'] or 0),
            'section_title':_infer_section_title(text),'chunk_type':'builtin-page-rescue',
            'confidence':'high' if score>=RETRIEVAL_HIGH_CONFIDENCE else 'medium' if score>=RETRIEVAL_MEDIUM_CONFIDENCE else 'low',
            'rank':0.0,'excerpt':text[:500],
        })
    scored.sort(key=lambda item:(-float(item['score']),item['file_name'],int(item['page_start'] or 0)))
    out=[];seen=set()
    for item in scored:
        key=(item['document_id'],item['page_start'])
        if key in seen:
            continue
        seen.add(key);out.append(item)
        if len(out)>=int(limit or RETRIEVAL_TOP_K):
            break
    return out


def create_conversation(user_id: int | None, external_user_id: str | None, title: str) -> str:
    cid = str(uuid.uuid4())
    ts = now_iso()
    with get_db() as db:
        db.execute("INSERT INTO conversations(id, user_id, external_user_id, title, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?)", (cid, user_id, external_user_id, title[:120], ts, ts))
    return cid


def ensure_conversation_access(conversation_id: str, user: dict[str, Any] | None, external_user_id: str | None = None) -> sqlite3.Row:
    with get_db() as db:
        row = db.execute("SELECT * FROM conversations WHERE id = ?", (conversation_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail='گفت‌وگو پیدا نشد.')
    if user and user['role'] in {'admin', 'manager'}:
        return row
    if user and row['user_id'] == user['id']:
        return row
    if external_user_id and row['external_user_id'] == external_user_id:
        return row
    raise HTTPException(status_code=403, detail='به این گفت‌وگو دسترسی ندارید.')


def parse_integration_keys() -> dict[str, str]:
    raw=str(INTEGRATION_API_KEYS_RAW or '').strip()
    if not raw or raw=='{}':
        return {}
    try:
        data=json.loads(raw)
        if isinstance(data,dict):
            return {str(k).strip():str(v).strip() for k,v in data.items() if str(k).strip() and str(v).strip()}
    except json.JSONDecodeError:
        pass
    # Backward compatibility for the historical env sample: site:key,other:key2
    parsed: dict[str,str]={}
    for item in re.split(r'[\n,;]+',raw):
        name,sep,value=item.partition(':')
        if sep and name.strip() and value.strip():
            parsed[name.strip()]=value.strip()
    return parsed


def validate_integration_key(value: str | None) -> str:
    if not value:
        raise HTTPException(status_code=401, detail='X-API-Key الزامی است.')
    for name, expected in parse_integration_keys().items():
        if expected and hmac.compare_digest(value, expected):
            return name
    raise HTTPException(status_code=401, detail='کلید اتصال نامعتبر است.')


class LoginInput(BaseModel):
    identifier: str = Field(min_length=1, max_length=200)
    password: str = Field(min_length=1, max_length=200)


class ChatInput(BaseModel):
    message: str = Field(min_length=1, max_length=6000)
    conversation_id: str | None = None


class GuestChatInput(ChatInput):
    external_user_id: str = Field(min_length=8, max_length=200)


class IntegrationChatInput(ChatInput):
    external_user_id: str = Field(min_length=3, max_length=200)


class CargoToolInput(BaseModel):
    vehicle: str = Field(min_length=1, max_length=100)
    weight_kg: Decimal | None = Field(default=None, ge=0, le=1000000)
    description: str | None = Field(default=None, max_length=4000)
    route: str | None = Field(default=None, max_length=1500)
    notes: str | None = Field(default=None, max_length=3000)


class RouteToolInput(BaseModel):
    origin: str = Field(min_length=1, max_length=500)
    destination: str = Field(min_length=1, max_length=500)
    vehicle: str = Field(min_length=1, max_length=100)
    weight_kg: Decimal | None = Field(default=None, ge=0, le=1000000)
    load_notes: str | None = Field(default=None, max_length=4000)


class CalculationToolInput(BaseModel):
    query: str = Field(min_length=1, max_length=6000)


class LocationSearchInput(BaseModel):
    query: str = Field(min_length=3, max_length=1500)
    city: str | None = Field(default=None, max_length=100)
    province: str | None = Field(default=None, max_length=100)


class CargoItemInput(BaseModel):
    name: str = Field(default="بار", min_length=1, max_length=100)
    count: int = Field(default=1, ge=1, le=100000)
    length_cm: float = Field(gt=0, le=100000)
    width_cm: float = Field(gt=0, le=100000)
    height_cm: float = Field(gt=0, le=100000)
    weight_kg: float | None = Field(default=None, gt=0, le=1000000)
    rotatable: bool = True


class CargoVehicleProfileInput(BaseModel):
    length_cm: float = Field(gt=0, le=100000)
    width_cm: float = Field(gt=0, le=100000)
    height_cm: float = Field(gt=0, le=100000)
    max_weight_kg: float | None = Field(default=None, gt=0, le=1000000)


class CargoCheckV2Input(BaseModel):
    vehicle: str = Field(min_length=2, max_length=40)
    items: list[CargoItemInput] = Field(min_length=1, max_length=100)
    notes: str | None = Field(default=None, max_length=3000)


class CancellationCalculationInput(BaseModel):
    vehicle: str = Field(min_length=2, max_length=40)
    origin_wait_minutes: Decimal | None = Field(default=None, ge=0, le=100000)
    billable_wait_minutes: Decimal | None = Field(default=None, ge=0, le=100000)


class WaitingClockIntervalInput(BaseModel):
    start_time: str = Field(min_length=4, max_length=5)
    end_time: str = Field(min_length=4, max_length=5)
    destination_number: int | None = Field(default=None, ge=1, le=100)


class WaitingCalculationInput(BaseModel):
    vehicle: str = Field(min_length=2, max_length=40)
    calculation_mode: str = Field(default='minutes', min_length=3, max_length=20)
    origin_wait_minutes: Decimal = Field(default=Decimal('0'), ge=0, le=100000)
    destination_wait_minutes: Decimal = Field(default=Decimal('0'), ge=0, le=100000)
    origin_start_time: str | None = Field(default=None, min_length=4, max_length=5)
    origin_end_time: str | None = Field(default=None, min_length=4, max_length=5)
    destination_time_ranges: list[WaitingClockIntervalInput] = Field(default_factory=list, max_length=100)
    destination_count: int = Field(default=1, ge=1, le=100)


class DeviationCalculationInput(BaseModel):
    vehicle: str = Field(min_length=2, max_length=40)
    mode: str | None = Field(default=None, max_length=20)
    distance_km: Decimal | None = Field(default=None, gt=0, le=100000)
    wait_minutes: Decimal | None = Field(default=None, ge=0, le=100000)
    service_amount_toman: Decimal | None = Field(default=None, ge=0, le=1000000000000)


class CalculationSettingsInput(BaseModel):
    cancellation_base_toman: int = Field(gt=0, le=1000000000000)
    waiting_hourly_toman: int = Field(gt=0, le=1000000000000)
    deviation_per_km_toman: int = Field(ge=0, le=1000000000000)
    deviation_time_unit_minutes: int = Field(default=0, ge=0, le=100000)
    deviation_time_unit_toman: int = Field(default=0, ge=0, le=1000000000000)
    deviation_per_minute_toman: int = Field(default=0, ge=0, le=1000000000000)
    free_wait_minutes: int = Field(ge=0, le=100000)
    extra_destination_free_minutes: int = Field(default=15, ge=0, le=10000)
    deviation_use_distance: bool = True
    deviation_use_time: bool = True


class DynamicApiProviderInput(BaseModel):
    label: str = Field(min_length=1,max_length=120)
    base_url: str = Field(min_length=8,max_length=500)
    api_key: str = Field(min_length=8,max_length=1000)
    model: str = Field(min_length=1,max_length=240)
    model_economy: str | None = Field(default=None,max_length=240)
    model_standard: str | None = Field(default=None,max_length=240)
    model_advanced: str | None = Field(default=None,max_length=240)
    vision_model: str | None = Field(default=None,max_length=240)
    transcription_model: str | None = Field(default=None,max_length=240)
    embedding_model: str | None = Field(default=None,max_length=240)
    input_cost_per_1m: float = Field(default=0,ge=0)
    output_cost_per_1m: float = Field(default=0,ge=0)
    credit_amount: float = Field(default=0,ge=0)
    credit_currency: str = Field(default='USD',min_length=1,max_length=12)

class DynamicApiProviderUpdate(BaseModel):
    label: str | None = Field(default=None,min_length=1,max_length=120)
    base_url: str | None = Field(default=None,min_length=8,max_length=500)
    api_key: str | None = Field(default=None,min_length=8,max_length=1000)
    model: str | None = Field(default=None,min_length=1,max_length=240)
    model_economy: str | None = Field(default=None,max_length=240)
    model_standard: str | None = Field(default=None,max_length=240)
    model_advanced: str | None = Field(default=None,max_length=240)
    vision_model: str | None = Field(default=None,max_length=240)
    transcription_model: str | None = Field(default=None,max_length=240)
    embedding_model: str | None = Field(default=None,max_length=240)
    input_cost_per_1m: float | None = Field(default=None,ge=0)
    output_cost_per_1m: float | None = Field(default=None,ge=0)
    credit_amount: float | None = Field(default=None,ge=0)
    credit_currency: str | None = Field(default=None,min_length=1,max_length=12)
    enabled: bool | None = None

class CreateUserInput(BaseModel):
    username: str = Field(min_length=3, max_length=50)
    password: str = Field(min_length=8, max_length=200)
    name: str = Field(min_length=2, max_length=100)
    email: EmailStr | None = None
    role: str = 'user'
    department: str | None = Field(default=None, max_length=100)
    question_limit: int | None = Field(default=None, ge=0, le=1000000)
    daily_question_limit: int | None = Field(default=None, ge=0, le=1000000)
    monthly_question_limit: int | None = Field(default=None, ge=0, le=10000000)


class UpdateUserInput(BaseModel):
    role: str | None = None
    is_active: bool | None = None
    password: str | None = Field(default=None, min_length=8, max_length=200)
    department: str | None = Field(default=None, max_length=100)
    set_department: bool = False
    question_limit: int | None = Field(default=None, ge=0, le=1000000)
    daily_question_limit: int | None = Field(default=None, ge=0, le=1000000)
    monthly_question_limit: int | None = Field(default=None, ge=0, le=10000000)
    set_question_limit: bool = False
    set_daily_question_limit: bool = False
    set_monthly_question_limit: bool = False
    reset_questions_used: bool = False
    reset_daily_questions_used: bool = False
    reset_monthly_questions_used: bool = False


class FaqInput(BaseModel):
    question: str = Field(min_length=3, max_length=2000)
    answer: str = Field(min_length=1, max_length=20000)
    aliases: list[str] = Field(default_factory=list)
    is_active: bool = True


class FaqUpdateInput(BaseModel):
    question: str | None = Field(default=None, min_length=3, max_length=2000)
    answer: str | None = Field(default=None, min_length=1, max_length=20000)
    aliases: list[str] | None = None
    is_active: bool | None = None


class GoldenCaseInput(BaseModel):
    question: str = Field(min_length=2, max_length=2000)
    expected_answer: str = Field(min_length=1, max_length=20000)
    expected_source: str | None = Field(default=None, max_length=500)
    is_active: bool = True


class GoldenCaseUpdateInput(BaseModel):
    question: str | None = Field(default=None, min_length=2, max_length=2000)
    expected_answer: str | None = Field(default=None, min_length=1, max_length=20000)
    expected_source: str | None = Field(default=None, max_length=500)
    is_active: bool | None = None


class FeedbackInput(BaseModel):
    rating: str
    note: str | None = Field(default=None, max_length=1000)


class RoleInput(BaseModel):
    role: str


class UploadSessionStartInput(BaseModel):
    filename: str = Field(min_length=1, max_length=255)
    size_bytes: int = Field(gt=0)
    visibility: str = 'public'
    content_type: str | None = Field(default=None, max_length=200)


class GoogleDocImportInput(BaseModel):
    url: str = Field(min_length=10, max_length=2000)
    visibility: str = 'public'


def _load_builtin_source_payloads() -> tuple[dict[str,Any],dict[str,Any]]:
    if not BUILTIN_SOURCE_MANIFEST.is_file() or not BUILTIN_SOURCE_PREINDEX.is_file():
        return {},{}
    try:
        manifest=json.loads(BUILTIN_SOURCE_MANIFEST.read_text(encoding='utf-8'))
        preindex=json.loads(BUILTIN_SOURCE_PREINDEX.read_text(encoding='utf-8'))
        return manifest,preindex
    except Exception:
        return {},{}


def ensure_builtin_sources() -> dict[str,Any]:
    """Install the four reviewed Barsan booklets as immutable baseline sources.

    Base text/page metadata is indexed immediately without network calls. A durable
    background reindex then enriches every page with Vision. New manager training is
    always authoritative and therefore overrides these booklets.
    """
    manifest,preindex=_load_builtin_source_payloads()
    entries=list(manifest.get('sources') or []) if isinstance(manifest,dict) else []
    global_enabled=system_setting_bool('builtin_sources_enabled',BUILTIN_SOURCE_GLOBAL_DEFAULT)
    installed=[];queued=[]
    with get_db() as db:
        owner=db.execute("SELECT id FROM users WHERE is_owner=1 ORDER BY id LIMIT 1").fetchone()
        actor_id=int(owner['id']) if owner else None
    for item in entries:
        key=str(item.get('source_key') or '').strip(); filename=str(item.get('filename') or '').strip()
        if not key or not filename:
            continue
        path=(BUILTIN_SOURCE_DIR/filename).resolve()
        if not path.is_file():
            continue
        sha=str(item.get('sha256') or hashlib.sha256(path.read_bytes()).hexdigest())
        doc_id=f'builtin-barsan-{key}'
        new_or_changed=False
        with get_db() as db:
            row=db.execute('SELECT * FROM documents WHERE source_key=? OR id=?',(key,doc_id)).fetchone()
            if row is None:
                db.execute("""INSERT INTO documents(id,filename,stored_path,mime_type,visibility,status,character_count,chunk_count,version,created_by,created_at,content_sha256,file_size_bytes,is_builtin,is_enabled,source_key)
                              VALUES(?,?,?,?,?,'processing',0,0,1,?,?,?,?,1,?,?)""",
                           (doc_id,filename,str(path),'application/pdf','public',actor_id,now_iso(),sha,path.stat().st_size,int(global_enabled),key))
                new_or_changed=True
            else:
                doc_id=str(row['id'])
                changed=str(row['content_sha256'] or '')!=sha or int(row['chunk_count'] or 0)==0
                repaired_status=str(row['status'] or '')
                if int(row['chunk_count'] or 0)>0 and repaired_status not in {'ready','partial'}:
                    repaired_status='partial'
                db.execute("UPDATE documents SET filename=?,stored_path=?,mime_type=?,visibility='public',content_sha256=?,file_size_bytes=?,is_builtin=1,is_enabled=?,source_key=?,status=? WHERE id=?",
                           (filename,str(path),'application/pdf',sha,path.stat().st_size,int(global_enabled),key,repaired_status or 'processing',doc_id))
                new_or_changed=changed
        if new_or_changed:
            payload=((preindex.get('sources') or {}).get(key) or {}).get('result') if isinstance(preindex,dict) else None
            if isinstance(payload,dict):
                result=json.loads(json.dumps(payload,ensure_ascii=False))
                chunks=chunk_source_result(result)
                with get_db() as db:
                    doc=db.execute('SELECT * FROM documents WHERE id=?',(doc_id,)).fetchone()
                    persist_document_ingestion(db,doc_id,doc,result,chunks,None)
                    db.execute('UPDATE documents SET is_builtin=1,is_enabled=?,source_key=?,content_sha256=?,file_size_bytes=? WHERE id=?',(int(global_enabled),key,sha,path.stat().st_size,doc_id))
        installed.append(doc_id)
        if BUILTIN_SOURCE_AUTO_ENRICH:
            with get_db() as db:
                row=db.execute('SELECT page_count,vision_success_pages,status FROM documents WHERE id=?',(doc_id,)).fetchone()
            if row and (int(row['vision_success_pages'] or 0)<int(row['page_count'] or 0) or str(row['status'])!='ready'):
                job=_enqueue_reindex_job(doc_id,actor_id or 0,'غنی‌سازی صفحه‌به‌صفحه جزوه پایه با Vision')
                queued.append(job)
    return {'installed':len(installed),'queued':len(set(queued)),'enabled':global_enabled}


@asynccontextmanager
async def lifespan(_: FastAPI):
    global _HTTP_CLIENT, _HEALTH_TASK, _DOCUMENT_JOB_TASK, _REPLICA_LOCK_HANDLE
    validate_environment()
    init_storage()
    if SQLITE_SINGLE_REPLICA_MODE and SQLITE_REPLICA_LOCK_ENABLED and DATABASE_URL.startswith('sqlite:'):
        try:
            _REPLICA_LOCK_HANDLE=acquire_sqlite_replica_lock(_DATA_ROOT,enabled=True)
        except ReplicaLockError as exc:
            raise RuntimeError('اجرای هم‌زمان چند Replica روی SQLite مسدود شد. فقط یک Replica برای بارسان فعال کنید.') from exc
    ensure_schema()
    seed_admin()
    if BUILTIN_SOURCE_AUTO_INSTALL:
        await asyncio.to_thread(ensure_builtin_sources)
    if AUTO_REINDEX_LEGACY_SOURCES and DOCUMENT_JOB_WORKER_ENABLED:
        await asyncio.to_thread(_enqueue_legacy_source_reindex_jobs)
    _HTTP_CLIENT=httpx.AsyncClient(
        timeout=httpx.Timeout(AI_TIMEOUT_SECONDS,connect=min(20.0,AI_TIMEOUT_SECONDS)),
        limits=httpx.Limits(max_connections=80,max_keepalive_connections=40,keepalive_expiry=45.0),
        follow_redirects=True,
    )
    if HEALTH_MONITOR_ENABLED:
        _HEALTH_TASK=asyncio.create_task(_health_monitor_loop(),name='barsan-health-monitor')
    if DOCUMENT_JOB_WORKER_ENABLED:
        await asyncio.to_thread(_recover_stale_document_jobs)
        _DOCUMENT_JOB_TASK=asyncio.create_task(_document_job_worker_loop(),name='barsan-document-worker')
    try:
        yield
    finally:
        if _HEALTH_TASK:
            _HEALTH_TASK.cancel()
            try:
                await _HEALTH_TASK
            except asyncio.CancelledError:
                pass
            _HEALTH_TASK=None
        if _DOCUMENT_JOB_TASK:
            _DOCUMENT_JOB_TASK.cancel()
            try:
                await _DOCUMENT_JOB_TASK
            except asyncio.CancelledError:
                pass
            _DOCUMENT_JOB_TASK=None
        if _HTTP_CLIENT:
            await _HTTP_CLIENT.aclose()
            _HTTP_CLIENT=None
        release_sqlite_replica_lock(_REPLICA_LOCK_HANDLE)
        _REPLICA_LOCK_HANDLE=None


app = FastAPI(title=APP_NAME, version=APP_VERSION, description='Barsan R35.2: operational reliability, safe restore, embed/integration compatibility and production workflow hardening', lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=ALLOWED_ORIGINS, allow_credentials=True, allow_methods=['GET','POST','PUT','PATCH','DELETE','OPTIONS'], allow_headers=['Authorization','Content-Type','X-API-Key','X-Integration-Key','X-Chunk-SHA256'])


def _record_request_metric(request_id: str, method: str, path: str, status_code: int, duration_ms: int) -> None:
    if not path.startswith('/api/') and path not in {'/healthz','/readyz'}:
        return
    try:
        with get_db() as db:
            db.execute('INSERT INTO request_metrics(request_id,method,path,status_code,duration_ms,created_at) VALUES(?,?,?,?,?,?)',
                (request_id,method[:12],path[:240],int(status_code),max(0,int(duration_ms)),now_iso()))
            if secrets.randbelow(200)==0:
                cutoff=(datetime.now(timezone.utc)-timedelta(days=14)).isoformat()
                db.execute('DELETE FROM request_metrics WHERE created_at<?',(cutoff,))
    except sqlite3.Error:
        pass


def _content_security_policy_for_path(path: str) -> str:
    frame_ancestors="'self'"
    if path=='/widget':
        extras=' '.join(x for x in WIDGET_ALLOWED_ORIGINS if x and x!=_normalized_origin(PUBLIC_BASE_URL))
        if extras:
            frame_ancestors=f"'self' {extras}"
    return ("default-src 'self'; script-src 'self' 'unsafe-inline'; style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
            "font-src 'self' https://fonts.gstatic.com data:; img-src 'self' data: blob:; media-src 'self' blob:; connect-src 'self'; "
            f"object-src 'none'; base-uri 'self'; frame-ancestors {frame_ancestors}; form-action 'self'")


@app.middleware('http')
async def request_observability(request: Request, call_next: Callable[...,Any]) -> Response:
    incoming=str(request.headers.get('x-request-id') or '').strip()
    request_id=incoming if re.fullmatch(r'[A-Za-z0-9._-]{8,80}',incoming) else uuid.uuid4().hex
    request.state.request_id=request_id
    if _RESTORE_IN_PROGRESS.is_set() and request.url.path.startswith('/api/'):
        return JSONResponse({'detail':'سیستم در حال بازیابی امن پشتیبان است؛ چند لحظه بعد دوباره تلاش کنید.','request_id':request_id},status_code=503,headers={'X-Request-ID':request_id,'Retry-After':'5'})
    # Cookie-authenticated browser mutations must be same-origin. Bearer/integration
    # requests are unaffected so API clients and internal integrations remain compatible.
    if ORIGIN_GUARD_ENABLED and request.method.upper() in {'POST','PUT','PATCH','DELETE'} and request.cookies.get('barsan_token') and not str(request.headers.get('authorization') or '').lower().startswith('bearer '):
        allowed={PUBLIC_BASE_URL.rstrip('/')} | {str(x).rstrip('/') for x in ALLOWED_ORIGINS if str(x).strip()}
        request_origin=str(request.headers.get('origin') or '') or None
        request_referer=str(request.headers.get('referer') or '') or None
        if not same_origin_allowed(request_origin,request_referer,allowed):
            return JSONResponse({'detail':'درخواست مرورگر از مبدأ نامعتبر رد شد.','request_id':request_id},status_code=403,headers={'X-Request-ID':request_id})
    started=time.perf_counter();status=500
    try:
        response=await call_next(request)
        status=int(response.status_code)
        response.headers['X-Request-ID']=request_id
        response.headers['X-Content-Type-Options']='nosniff'
        response.headers['Referrer-Policy']='same-origin'
        response.headers['Permissions-Policy']='camera=(), geolocation=(), payment=()'
        if CONTENT_SECURITY_POLICY_ENABLED:
            response.headers['Content-Security-Policy']=_content_security_policy_for_path(request.url.path)
        if COOKIE_SECURE and str(PUBLIC_BASE_URL).startswith('https://'):
            response.headers['Strict-Transport-Security']='max-age=15552000; includeSubDomains'
        return response
    finally:
        duration_ms=int((time.perf_counter()-started)*1000)
        await asyncio.to_thread(_record_request_metric,request_id,request.method,request.url.path,status,duration_ms)


@app.get('/healthz')
def healthz() -> dict[str, Any]:
    return {'status': 'ok', 'app': APP_NAME, 'version': APP_VERSION, 'port': PORT, 'railway': bool(os.getenv('RAILWAY_SERVICE_ID'))}


@app.get('/readyz')
def readyz() -> JSONResponse:
    if _RESTORE_IN_PROGRESS.is_set():
        return JSONResponse({'status':'maintenance','release':RELEASE_ID,'error':'restore_in_progress'},status_code=503,headers={'Retry-After':'5'})
    db_ok = False
    storage_ok = False
    try:
        with get_db() as db:
            db.execute('SELECT 1').fetchone()
        db_ok = True
        probe = UPLOAD_DIR / '.probe'
        probe.write_text('ok', encoding='utf-8')
        probe.unlink(missing_ok=True)
        storage_ok = True
    except Exception:
        # Public readiness must not leak database paths, SQL details, or filesystem internals.
        pass
    code = 200 if db_ok and storage_ok else 503
    error_code=None if code==200 else ('database_unavailable' if not db_ok else 'storage_unavailable')
    return JSONResponse({'status': 'ready' if code == 200 else 'not_ready', 'database': db_ok, 'storage': storage_ok, 'error': error_code, 'release': RELEASE_ID}, status_code=code)


@app.get('/api/v1/system/info')
def system_info() -> dict[str, Any]:
    return {
        'app': APP_NAME,
        'version': APP_VERSION,
        'release': RELEASE_ID,
        'registration_enabled': PUBLIC_REGISTRATION_ENABLED,
    }


@app.post('/api/v1/auth/register')
def register_disabled() -> JSONResponse:
    raise HTTPException(status_code=403, detail='ثبت‌نام عمومی غیرفعال است. حساب کاربری فقط توسط مدیر ایجاد می‌شود.')


def _request_ip(request: Request) -> str:
    # Railway documents X-Real-IP as the client remote IP added by its public
    # edge. Only use it for a request carrying Railway's own request marker;
    # elsewhere, use the socket peer Uvicorn has already validated. Never trust
    # arbitrary X-Forwarded-For values directly.
    if os.getenv('RAILWAY_SERVICE_ID') and request.headers.get('x-railway-request-id'):
        candidate=str(request.headers.get('x-real-ip') or '').strip()
        try:
            if candidate:
                return str(ipaddress.ip_address(candidate))
        except ValueError:
            pass
    return request.client.host if request.client else 'unknown'


def _login_lock(identifier: str, ip_address: str) -> datetime | None:
    with get_db() as db:
        row=db.execute("SELECT locked_until FROM login_attempts WHERE identifier=? AND ip_address=?",(identifier,ip_address)).fetchone()
    if not row or not row['locked_until']:
        return None
    try:
        locked=datetime.fromisoformat(row['locked_until'])
        if locked.tzinfo is None: locked=locked.replace(tzinfo=timezone.utc)
        return locked.astimezone(timezone.utc)
    except ValueError:
        return None


def _record_login_failure(identifier: str, ip_address: str) -> int:
    now=datetime.now(timezone.utc)
    window_start=now-timedelta(minutes=LOGIN_WINDOW_MINUTES)
    with get_db() as db:
        row=db.execute("SELECT * FROM login_attempts WHERE identifier=? AND ip_address=?",(identifier,ip_address)).fetchone()
        if not row or not row['first_failed_at'] or datetime.fromisoformat(row['first_failed_at']).astimezone(timezone.utc)<window_start:
            count=1; first=now.isoformat()
        else:
            count=int(row['failed_count'])+1; first=row['first_failed_at']
        locked_until=(now+timedelta(minutes=LOGIN_LOCK_MINUTES)).isoformat() if count>=LOGIN_MAX_ATTEMPTS else None
        db.execute(
            """INSERT INTO login_attempts(identifier,ip_address,failed_count,first_failed_at,last_failed_at,locked_until)
            VALUES(?,?,?,?,?,?) ON CONFLICT(identifier,ip_address) DO UPDATE SET
            failed_count=excluded.failed_count,first_failed_at=excluded.first_failed_at,
            last_failed_at=excluded.last_failed_at,locked_until=excluded.locked_until""",
            (identifier,ip_address,count,first,now.isoformat(),locked_until),
        )
    audit(None,'login.failed',{'identifier':identifier,'ip':ip_address,'failed_count':count})
    return count


def _clear_login_failures(identifier: str, ip_address: str) -> None:
    with get_db() as db:
        db.execute("DELETE FROM login_attempts WHERE identifier=? AND ip_address=?",(identifier,ip_address))


@app.post('/api/v1/auth/login')
def login(data: LoginInput, response: Response, request: Request) -> dict[str, Any]:
    identifier=data.identifier.strip().lower()
    ip_address=_request_ip(request)
    locked=_login_lock(identifier,ip_address)
    now=datetime.now(timezone.utc)
    if locked and locked>now:
        seconds=max(1,int((locked-now).total_seconds()))
        raise HTTPException(status_code=429,detail=f'ورود موقتاً قفل شده است. {seconds} ثانیه دیگر تلاش کنید.')
    with get_db() as db:
        row=db.execute("SELECT * FROM users WHERE lower(username)=? OR lower(email)=?",(identifier,identifier)).fetchone()
    if not row or not row['is_active'] or not verify_password(data.password,row['password_hash'],row['salt']):
        count=_record_login_failure(identifier,ip_address)
        if count>=LOGIN_MAX_ATTEMPTS:
            raise HTTPException(status_code=429,detail=f'تعداد تلاش ناموفق زیاد است. حساب ورود برای {LOGIN_LOCK_MINUTES} دقیقه قفل شد.')
        raise HTTPException(status_code=401,detail='نام کاربری یا رمز عبور نادرست است.')
    _clear_login_failures(identifier,ip_address)
    token=create_token(row['id'],row['role'])
    response.set_cookie('barsan_token',token,httponly=True,secure=COOKIE_SECURE,samesite='strict',max_age=JWT_EXPIRE_MINUTES*60,path='/')
    with get_db() as db:
        db.execute("UPDATE users SET last_login_at=? WHERE id=?",(now_iso(),row['id']))
    audit(row['id'],'user.login',{'username':row['username'],'ip':ip_address})
    return {'user':{'id':row['id'],'username':row['username'],'email':row['email'],'name':row['name'],'role':row['role'],'is_owner':bool(row['is_owner']),'department':row['department'],'question_limit':row['question_limit'],'questions_used':row['questions_used'],'daily_question_limit':row['daily_question_limit'],'daily_questions_used':row['daily_questions_used'],'monthly_question_limit':row['monthly_question_limit'],'monthly_questions_used':row['monthly_questions_used']}}


@app.post('/api/v1/auth/logout')
def logout(response: Response) -> dict[str,bool]:
    response.delete_cookie('barsan_token',path='/')
    return {'ok':True}


@app.get('/api/v1/auth/me')
def me(user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    return user




def _anonymous_rate_identity(prefix: str, request: Request, external_id: str | None = None) -> str:
    """Bind anonymous limits to network/browser signals so rotating a client ID cannot bypass limits."""
    raw = "|".join((prefix, _request_ip(request), request.headers.get("user-agent", "")[:240], str(external_id or "")[:240]))
    key = JWT_SECRET.encode("utf-8") if JWT_SECRET else b"barsan-rate-key"
    digest = hmac.new(key, raw.encode("utf-8"), hashlib.sha256).hexdigest()
    return f"{prefix}:{digest}"


@app.post('/api/v1/chat')
async def chat(data: ChatInput, user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    return await process_chat(message=data.message, conversation_id=data.conversation_id, user=user, external_user_id=None, integration=False, rate_identity=f"user:{user['id']}")


@app.post('/api/v1/tools/cargo')
async def tool_cargo(data: CargoToolInput, user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    if data.weight_kg is None and not (data.description or '').strip():
        raise HTTPException(status_code=400,detail='شرح بار یا وزن بار را وارد کنید.')
    parts=[f"بررسی بار. خودرو: {data.vehicle}."]
    if data.weight_kg is not None: parts.append(f"وزن بار: {_format_decimal_result(Decimal(data.weight_kg))} کیلوگرم.")
    if (data.description or '').strip(): parts.append(f"شرح بار: {sanitize_answer_text(data.description or '')}.")
    if (data.route or '').strip(): parts.append(f"مسیر یا شرایط منطقه‌ای: {sanitize_answer_text(data.route or '')}.")
    if (data.notes or '').strip(): parts.append(f"توضیحات تکمیلی: {sanitize_answer_text(data.notes or '')}.")
    parts.append('ظرفیت پایه، مجاز یا غیرمجاز بودن این بار، محدودیت‌ها، شرط‌ها، تبصره‌ها و استثناهای مرتبط را مستقیم بررسی کن. آموزش مدیر مرجع اول و قطعی است؛ فقط در نبود آموزش مرتبط از منابع استفاده کن. اگر مقایسه عددی لازم است آن را قطعی محاسبه کن.')
    return await process_operational_module(module='cargo',question=' '.join(parts),user=user)


@app.post('/api/v1/tools/route')
async def tool_route(data: RouteToolInput, user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    parts=[f"بررسی مستقل مسیریابی باربری. مبدأ: {data.origin}. مقصد: {data.destination}. خودرو: {data.vehicle}."]
    if data.weight_kg is not None: parts.append(f"وزن بار: {_format_decimal_result(Decimal(data.weight_kg))} کیلوگرم.")
    if (data.load_notes or '').strip(): parts.append(f"نوع بار و توضیحات: {sanitize_answer_text(data.load_notes or '')}.")
    parts.append('محدودیت‌های مسیر، مناطق خاص، سربالایی، ممنوعیت‌ها، ظرفیت مجاز و همه شرط‌ها و تبصره‌های مرتبط را مستقیم اعلام کن. آموزش مدیر مرجع اول و قطعی است؛ فقط در نبود آموزش مرتبط از منابع استفاده کن. اگر بخشی از مسیر در دانش سازمانی موجود نیست فقط همان بخش را نامشخص اعلام کن.')
    return await process_operational_module(module='route',question=' '.join(parts),user=user)


@app.post('/api/v1/tools/calculate')
async def tool_calculate(data: CalculationToolInput, user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    query=data.query.strip()
    prompt=query if _pure_math_expression(query) else (
        'محاسبه دقیق و مستقل. '+query+' '
        'اگر هر عدد، ظرفیت، نرخ، قانون یا شرط باید از دانش سازمانی گرفته شود، ابتدا آموزش‌های مدیر و فقط در نبود آموزش مرتبط منابع را بررسی کن. '
        'قاعده اصلی و همه استثناهای مرتبط را حفظ کن و هیچ نتیجه عددی را حدس نزن؛ محاسبات را با موتور قطعی سامانه انجام بده.'
    )
    return await process_operational_module(module='calc',question=prompt,user=user)


# =============================================================================
# R29.7 — CARGO COUNT VERTICAL + PER-MINUTE DEVIATION
# Cargo geometry, Neshan routing, cancellation/waiting/deviation calculations.
# These modules are independent from chat history. Manager training remains the
# authoritative knowledge layer whenever contextual rules are needed.
# =============================================================================


def _cargo_profile_public(row: dict[str,Any] | sqlite3.Row) -> dict[str,Any]:
    d=dict(row)
    return {
        'vehicle':d['vehicle'],'vehicle_label':d['vehicle_label'],
        'length_cm':d.get('length_cm'),'width_cm':d.get('width_cm'),'height_cm':d.get('height_cm'),
        'max_weight_kg':d.get('max_weight_kg'),'configured':bool(d.get('configured')),
        'updated_at':d.get('updated_at'),
    }


@app.get('/api/v1/cargo/vehicles')
def cargo_vehicle_profiles(user:dict[str,Any]=Depends(current_user)) -> list[dict[str,Any]]:
    with get_db() as db:
        rows=db.execute("SELECT * FROM cargo_vehicle_profiles WHERE vehicle IN ('peykan_flatbed','peykan_no_flatbed','nissan_flatbed','nissan_no_flatbed','khavar_covered','khavar_open') ORDER BY CASE vehicle WHEN 'peykan_flatbed' THEN 1 WHEN 'peykan_no_flatbed' THEN 2 WHEN 'nissan_flatbed' THEN 3 WHEN 'nissan_no_flatbed' THEN 4 WHEN 'khavar_covered' THEN 5 WHEN 'khavar_open' THEN 6 ELSE 99 END").fetchall()
    return [_cargo_profile_public(r) for r in rows]


@app.put('/api/v1/cargo/vehicles/{vehicle}')
def update_cargo_vehicle_profile(vehicle:str,data:CargoVehicleProfileInput,user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    key=_normalize_cargo_vehicle_key(vehicle)
    with get_db() as db:
        row=db.execute("SELECT vehicle_label FROM cargo_vehicle_profiles WHERE vehicle=?",(key,)).fetchone()
        if not row:
            raise HTTPException(status_code=404,detail='پروفایل خودرو پیدا نشد.')
        db.execute(
            """UPDATE cargo_vehicle_profiles SET length_cm=?,width_cm=?,height_cm=?,max_weight_kg=?,configured=1,updated_by=?,updated_at=? WHERE vehicle=?""",
            (data.length_cm,data.width_cm,data.height_cm,data.max_weight_kg,user['id'],now_iso(),key),
        )
        updated=db.execute("SELECT * FROM cargo_vehicle_profiles WHERE vehicle=?",(key,)).fetchone()
    audit(user['id'],'cargo.vehicle_profile.updated',{'vehicle':key,'length_cm':data.length_cm,'width_cm':data.width_cm,'height_cm':data.height_cm,'max_weight_kg':data.max_weight_kg})
    return _cargo_profile_public(updated)


def _cargo_profile_for_request(data:CargoCheckV2Input) -> dict[str,Any]:
    key=_normalize_cargo_vehicle_key(data.vehicle)
    with get_db() as db:
        row=db.execute("SELECT * FROM cargo_vehicle_profiles WHERE vehicle=?",(key,)).fetchone()
    profile=dict(row) if row else {'vehicle':key,'vehicle_label':CARGO_VEHICLES[key],'configured':0}
    if not profile.get('configured') or any(profile.get(k) in (None,'') for k in ('length_cm','width_cm','height_cm')):
        raise HTTPException(status_code=409,detail='ابعاد ثابت این خودرو هنوز یک‌بار توسط مدیر ثبت نشده است. مدیر از بخش «پروفایل خودروها» ابعاد مفید را ثبت کند؛ بعد از آن اپراتور فقط خودرو و بار را وارد می‌کند.')
    return {
        'vehicle':key,'name':profile['vehicle_label'],
        'length_cm':float(profile['length_cm']),'width_cm':float(profile['width_cm']),'height_cm':float(profile['height_cm']),
        'max_weight_kg':float(profile['max_weight_kg']) if profile.get('max_weight_kg') not in (None,'') else None,
        'profile_source':'manager_profile',
    }


async def _cargo_context_note(question:str,user:dict[str,Any]) -> tuple[str,list[dict[str,Any]]]:
    """Return the authoritative rule note without letting it alter geometry math."""
    try:
        stage,sources=await retrieve_priority_stage_async(question,user,integration=False)
    except Exception:
        return '',[]
    if not sources:
        return '',[]
    public=public_source_items(sources)
    if stage=='training':
        raw=str(sources[0].get('answer') or sources[0].get('content') or '').strip()
        return (format_answer_for_mode(raw,False) if raw else ''),public
    # Document content can contain the legal exception that geometry alone cannot
    # infer (for example a lower payload limit in a specific district).
    try:
        prompt=(
            'فقط قوانین، محدودیت‌ها، شرط‌ها و استثناهای مرتبط با همین بررسی بار را از متن منبع خلاصه کن. '
            'نتیجه هندسی را دوباره محاسبه نکن و عددی خارج از منبع نساز. پاسخ کوتاه و دقیق باشد.\n'+question
        )
        answer,_usage=await ask_ai(prompt,sources,detailed=False,route='standard',memory='')
        return answer,public
    except Exception:
        raw='\n'.join(str(x.get('content') or x.get('answer') or '') for x in sources[:2])
        return format_answer_for_mode(raw,False),public


@app.post('/api/v1/cargo/check')
async def cargo_check_v2(data:CargoCheckV2Input,user:dict[str,Any]=Depends(current_user)) -> dict[str,Any]:
    enforce_rate_limit(f"cargo-fit:{user['id']}")
    started=time.perf_counter()
    profile=_cargo_profile_for_request(data)
    items=[x.model_dump() for x in data.items]
    try:
        result=calculate_cargo_fit(profile,items)
    except (TypeError,ValueError,KeyError) as exc:
        raise HTTPException(status_code=400,detail=str(exc)) from exc
    note_parts=[f"بررسی بار {profile['name']}."]
    if data.notes:
        note_parts.append(sanitize_answer_text(data.notes))
    for item in data.items[:8]:
        note_parts.append(f"{item.count} عدد {item.name} با ابعاد {item.length_cm} در {item.width_cm} در {item.height_cm} سانتی‌متر")
    knowledge_note,knowledge_sources=await _cargo_context_note(' '.join(note_parts),user)
    answer=format_cargo_result(result)
    if knowledge_note:
        answer += '\n\nقانون/آموزش مرتبط:\n' + knowledge_note
    audit(user['id'],'cargo.calculated',{'vehicle':profile['vehicle'],'item_count':len(items),'fits':bool(result['fits']),'profile_source':profile['profile_source']})
    return {
        'answer':answer,'calculation':result,'vehicle_profile':profile,
        'knowledge_note':knowledge_note,'sources':knowledge_sources,
        'response_ms':int((time.perf_counter()-started)*1000),'zero_token':not bool(knowledge_sources),
    }




def _normalize_cargo_image_vehicle(value: str) -> str:
    raw=normalize_text(value)
    compact=raw.replace(' ','')
    key=CARGO_IMAGE_TRAINING_ALIASES.get(raw) or CARGO_IMAGE_TRAINING_ALIASES.get(compact) or CARGO_IMAGE_TRAINING_ALIASES.get(str(value).strip().lower())
    if key not in CARGO_IMAGE_TRAINING_VEHICLES:
        raise HTTPException(status_code=400,detail='نوع خودرو برای آموزش تصویری معتبر نیست.')
    return key


def _detect_safe_image(payload: bytes) -> tuple[str,str]:
    if payload.startswith(b'\xff\xd8\xff'):
        return '.jpg','image/jpeg'
    if payload.startswith(b'\x89PNG\r\n\x1a\n'):
        return '.png','image/png'
    if len(payload)>=12 and payload[:4]==b'RIFF' and payload[8:12]==b'WEBP':
        return '.webp','image/webp'
    raise HTTPException(status_code=400,detail='فقط تصویر واقعی JPG، PNG یا WEBP پذیرفته می‌شود.')


def _cargo_image_training_public(row: dict[str,Any] | sqlite3.Row) -> dict[str,Any]:
    d=dict(row)
    return {
        'id':d['id'],'vehicle':d['vehicle'],'vehicle_label':d['vehicle_label'],'cargo_name':d['cargo_name'],
        'original_filename':d.get('original_filename'),'file_size_bytes':int(d.get('file_size_bytes') or 0),
        'created_at':d.get('created_at'),'image_url':f"/api/v1/cargo/image-training/{d['id']}/image",
    }


def _cargo_image_training_score(query: str, row: dict[str,Any] | sqlite3.Row) -> float:
    d=dict(row);q=normalize_text(query);hay=normalize_text(f"{d.get('vehicle_label','')} {d.get('cargo_name','')}")
    q_tokens=search_tokens(q);h_tokens=search_tokens(hay)
    token_score=_fuzzy_token_score(q_tokens,h_tokens) if q_tokens else 0.0
    sequence=SequenceMatcher(None,q,hay).ratio() if q and hay else 0.0
    cargo_norm=normalize_text(str(d.get('cargo_name') or ''))
    exact_bonus=0.34 if cargo_norm and cargo_norm in q else 0.0
    vehicle_norm=normalize_text(str(d.get('vehicle_label') or ''))
    vehicle_bonus=0.24 if vehicle_norm and vehicle_norm in q else 0.0
    return min(1.5,token_score*0.58+sequence*0.22+exact_bonus+vehicle_bonus)


@app.post('/api/v1/cargo/image-training')
async def create_cargo_image_training(
    vehicle: str = Form(...),
    cargo_name: str = Form(...),
    image: UploadFile = File(...),
    user: dict[str,Any] = Depends(require_roles('manager','admin')),
) -> dict[str,Any]:
    key=_normalize_cargo_image_vehicle(vehicle)
    clean_name=' '.join(sanitize_answer_text(cargo_name).split()).strip()
    if len(clean_name)<2 or len(clean_name)>180:
        raise HTTPException(status_code=400,detail='نام بار باید بین ۲ تا ۱۸۰ نویسه باشد.')
    max_bytes=CARGO_IMAGE_MAX_MB*1024*1024
    payload=await image.read(max_bytes+1)
    if not payload:
        raise HTTPException(status_code=400,detail='فایل تصویر خالی است.')
    if len(payload)>max_bytes:
        raise HTTPException(status_code=413,detail=f'حجم عکس باید حداکثر {CARGO_IMAGE_MAX_MB} مگابایت باشد.')
    suffix,mime=_detect_safe_image(payload)
    item_id=str(uuid.uuid4())
    path=CARGO_IMAGE_TRAINING_DIR/f'{item_id}{suffix}'
    path.parent.mkdir(parents=True,exist_ok=True)
    with path.open('wb') as out:
        out.write(payload);out.flush();os.fsync(out.fileno())
    try:
        with get_db() as db:
            db.execute("""INSERT INTO cargo_image_training(id,vehicle,vehicle_label,cargo_name,normalized_cargo_name,image_path,mime_type,original_filename,file_size_bytes,created_by,created_at)
                          VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                       (item_id,key,CARGO_IMAGE_TRAINING_VEHICLES[key],clean_name,normalize_text(clean_name),str(path.relative_to(UPLOAD_DIR)),mime,Path(image.filename or f'image{suffix}').name,len(payload),user['id'],now_iso()))
            row=db.execute('SELECT * FROM cargo_image_training WHERE id=?',(item_id,)).fetchone()
    except Exception:
        path.unlink(missing_ok=True);raise
    audit(user['id'],'cargo.image_training.created',{'id':item_id,'vehicle':key,'cargo_name':clean_name,'size_bytes':len(payload)})
    return _cargo_image_training_public(row)


@app.get('/api/v1/cargo/image-training')
def list_cargo_image_training(user: dict[str,Any]=Depends(current_user)) -> list[dict[str,Any]]:
    with get_db() as db:
        rows=db.execute('SELECT * FROM cargo_image_training ORDER BY created_at DESC LIMIT 300').fetchall()
    return [_cargo_image_training_public(row) for row in rows]


@app.get('/api/v1/cargo/image-training/search')
def search_cargo_image_training(q: str, user: dict[str,Any]=Depends(current_user)) -> dict[str,Any]:
    query=' '.join(sanitize_answer_text(q).split()).strip()
    if len(query)<2:
        raise HTTPException(status_code=400,detail='برای بررسی عکس، نام خودرو و/یا نوع بار را وارد کنید.')
    with get_db() as db:
        rows=db.execute('SELECT * FROM cargo_image_training ORDER BY created_at DESC LIMIT 800').fetchall()
    scored=[]
    for row in rows:
        score=_cargo_image_training_score(query,row)
        if score>=0.24:
            item=_cargo_image_training_public(row);item['score']=round(score,4);scored.append(item)
    scored.sort(key=lambda x:(-float(x['score']),x.get('created_at') or ''))
    return {'query':query,'items':scored[:8],'matched':bool(scored),'total_candidates':len(rows)}


def _resolve_cargo_training_image_path(stored_path: str) -> Path:
    raw=Path(str(stored_path or '').strip())
    path=raw if raw.is_absolute() else (UPLOAD_DIR/raw)
    resolved=path.resolve()
    try:
        resolved.relative_to(CARGO_IMAGE_TRAINING_DIR.resolve())
    except ValueError as exc:
        raise HTTPException(status_code=403,detail='مسیر تصویر نامعتبر است.') from exc
    return resolved


@app.get('/api/v1/cargo/image-training/{training_id}/image')
def cargo_image_training_image(training_id: str, user: dict[str,Any]=Depends(current_user)) -> FileResponse:
    with get_db() as db:
        row=db.execute('SELECT image_path,mime_type,original_filename FROM cargo_image_training WHERE id=?',(training_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404,detail='آموزش تصویری پیدا نشد.')
    path=_resolve_cargo_training_image_path(row['image_path'])
    if not path.is_file():
        raise HTTPException(status_code=404,detail='فایل تصویر آموزش پیدا نشد.')
    return FileResponse(path,media_type=row['mime_type'],headers={'Cache-Control':'private, max-age=3600'})


@app.delete('/api/v1/cargo/image-training/{training_id}')
def delete_cargo_image_training(training_id: str, user: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    with get_db() as db:
        row=db.execute('SELECT * FROM cargo_image_training WHERE id=?',(training_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404,detail='آموزش تصویری پیدا نشد.')
        db.execute('DELETE FROM cargo_image_training WHERE id=?',(training_id,))
    try:
        _resolve_cargo_training_image_path(row['image_path']).unlink(missing_ok=True)
    except (HTTPException,OSError):
        pass
    audit(user['id'],'cargo.image_training.deleted',{'id':training_id,'vehicle':row['vehicle'],'cargo_name':row['cargo_name']})
    return {'ok':True}


def _location_cache_key_v2(query:str,city:str='',province:str='') -> str:
    raw='|'.join((normalize_location_text(query),normalize_location_text(city),normalize_location_text(province)))
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()


def _get_location_cache_v2(query:str,city:str='',province:str='') -> dict[str,Any] | None:
    key=_location_cache_key_v2(query,city,province);ts=now_iso()
    with get_db() as db:
        row=db.execute("SELECT result_json FROM location_cache WHERE query_key=? AND expires_at>?",(key,ts)).fetchone()
        if not row:
            return None
        db.execute("UPDATE location_cache SET hit_count=hit_count+1,last_used_at=? WHERE query_key=?",(ts,key))
    try:
        data=json.loads(row['result_json'])
    except (TypeError,json.JSONDecodeError):
        return None
    if isinstance(data,dict):
        data['cached']=True;data['provider_calls']=0
        return data
    return None


def _put_location_cache_v2(query:str,city:str,province:str,result:dict[str,Any]) -> None:
    key=_location_cache_key_v2(query,city,province);ts=now_iso();expires=(datetime.now(timezone.utc)+timedelta(days=NESHAN_CACHE_TTL_DAYS)).isoformat()
    stored={**result,'cached':False}
    with get_db() as db:
        db.execute(
            """INSERT INTO location_cache(query_key,normalized_query,result_json,created_at,expires_at,last_used_at)
               VALUES(?,?,?,?,?,?) ON CONFLICT(query_key) DO UPDATE SET result_json=excluded.result_json,created_at=excluded.created_at,expires_at=excluded.expires_at,last_used_at=excluded.last_used_at""",
            (key,normalize_location_text(query),json.dumps(stored,ensure_ascii=False),ts,expires,ts),
        )
        db.execute("DELETE FROM location_cache WHERE expires_at<=?",(ts,))


def _location_sources_v2(items:list[dict[str,Any]]) -> list[dict[str,Any]]:
    out=[]
    for i,item in enumerate(items):
        lat=float(item['latitude']);lon=float(item['longitude'])
        out.append({
            'source_type':'location','document_id':f'neshan:{lat:.6f},{lon:.6f}',
            'file_name':f"نشان — {item.get('title') or 'موقعیت'}",'chunk_index':i,
            'score':float(item.get('confidence') or 0),'excerpt':str(item.get('address') or '')[:500],
            'title':item.get('title'),'address':item.get('address'),'latitude':lat,'longitude':lon,
            'map_url':item.get('map_url'),'navigation_url':item.get('navigation_url'),
            'google_maps_url':item.get('google_maps_url'),'balad_url':item.get('balad_url'),'provider':item.get('provider'),
        })
    return out


@app.post('/api/v1/routing/search')
async def routing_search_v2(data:LocationSearchInput,user:dict[str,Any]=Depends(current_user)) -> dict[str,Any]:
    enforce_rate_limit(f"routing:{user['id']}")
    started=time.perf_counter();address=extract_location_query(data.query)
    if len(address)<3:
        raise HTTPException(status_code=400,detail='آدرس کامل‌تری وارد کنید.')
    cached=_get_location_cache_v2(address,data.city or '',data.province or '')
    if cached:
        return {**cached,'response_ms':int((time.perf_counter()-started)*1000)}
    if not NESHAN_API_KEY:
        raise HTTPException(status_code=503,detail='کلید NESHAN_API_KEY در Railway تنظیم نشده است. کلید سرویس نشان را در Variables قرار دهید.')
    client=_HTTP_CLIENT;owns=False
    if client is None:
        client=httpx.AsyncClient(follow_redirects=True);owns=True
    try:
        result=await lookup_neshan(
            client,api_key=NESHAN_API_KEY,query=address,city=data.city or '',province=data.province or '',
            geocoding_url=NESHAN_GEOCODING_URL,search_url=NESHAN_SEARCH_URL,use_plus=NESHAN_USE_PLUS,
            search_enrichment=NESHAN_SEARCH_ENRICHMENT,timeout_seconds=NESHAN_TIMEOUT_SECONDS,max_results=NESHAN_MAX_RESULTS,
        )
    except Exception as exc:
        raise HTTPException(status_code=502,detail=f'سرویس نشان پاسخ قابل استفاده نداد: {str(exc)[:250]}') from exc
    finally:
        if owns:
            await client.aclose()
    if not result.items:
        raise HTTPException(status_code=404,detail='برای این آدرس نتیجه قابل استفاده‌ای در نشان پیدا نشد.')
    sources=_location_sources_v2(result.items)
    payload={
        'query':address,'items':result.items,'sources':sources,'answer':format_location_answer(address,result.items),
        'cached':False,'provider_calls':result.provider_calls,'used_plus':result.used_plus,'used_search':result.used_search,
    }
    _put_location_cache_v2(address,data.city or '',data.province or '',payload)
    audit(user['id'],'location.searched',{'query':address[:300],'result_count':len(result.items),'provider_calls':result.provider_calls})
    return {**payload,'response_ms':int((time.perf_counter()-started)*1000)}


def _calculation_setting(vehicle:str) -> dict[str,Any]:
    with get_db() as db:
        row=db.execute('SELECT * FROM calculation_settings WHERE vehicle=?',(vehicle,)).fetchone()
    if not row:
        raise HTTPException(status_code=404,detail='تنظیمات محاسبات این خودرو پیدا نشد.')
    out=dict(row);out['configured']=out.get('updated_by') is not None
    return out


def _require_calculation_setting(vehicle:str) -> dict[str,Any]:
    row=_calculation_setting(vehicle)
    if not row.get('configured'):
        raise HTTPException(status_code=409,detail=f"آموزش محاسبات {row['vehicle_label']} هنوز ثبت نشده است. مدیر یک‌بار نرخ‌ها را در «آموزش محاسبات» ثبت کند.")
    return row


def _round_toman(value:Decimal) -> int:
    return int(value.quantize(Decimal('1'),rounding=ROUND_HALF_UP))


def _decimal_number(value:Decimal) -> int | float:
    n=value.normalize()
    return int(n) if n==n.to_integral() else float(n)


_CLOCK_DIGITS=str.maketrans('۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩','01234567890123456789')


def _clock_minutes(value:str|None,label:str) -> int:
    raw=str(value or '').strip().translate(_CLOCK_DIGITS)
    if not re.fullmatch(r'(?:[01]?\d|2[0-3]):[0-5]\d',raw):
        raise HTTPException(status_code=400,detail=f'{label} باید مثل 08:30 وارد شود.')
    h,m=raw.split(':',1);return int(h)*60+int(m)


def _clock_interval(start:str|None,end:str|None,label:str) -> Decimal:
    a=_clock_minutes(start,f'ساعت شروع {label}');b=_clock_minutes(end,f'ساعت پایان {label}')
    if b<a:b+=1440
    return Decimal(b-a)


def _setting_public(row:dict[str,Any]) -> dict[str,Any]:
    return {
        'vehicle':row['vehicle'],'vehicle_label':row['vehicle_label'],
        'cancellation_base_toman':int(row['cancellation_base_toman']),'waiting_hourly_toman':int(row['waiting_hourly_toman']),
        'deviation_per_km_toman':int(row['deviation_per_km_toman']),
        'deviation_time_unit_minutes':int(row.get('deviation_time_unit_minutes') or 0),
        'deviation_time_unit_toman':int(row.get('deviation_time_unit_toman') or 0),
        'deviation_per_minute_toman':_round_toman(Decimal(int(row.get('deviation_time_unit_toman') or 0))/Decimal(max(1,int(row.get('deviation_time_unit_minutes') or 1)))) if int(row.get('deviation_time_unit_toman') or 0)>0 else 0,
        'free_wait_minutes':int(row['free_wait_minutes']),
        'extra_destination_free_minutes':int(row['extra_destination_free_minutes']),
        'deviation_use_distance':bool(row.get('deviation_use_distance',1)),'deviation_use_time':bool(row.get('deviation_use_time',1)),
        'configured':bool(row.get('updated_by')),'updated_at':row.get('updated_at'),
    }


def _record_calculation(user_id:int,calc_type:str,vehicle:str,input_data:dict[str,Any],setting:dict[str,Any],result:dict[str,Any]) -> str:
    calc_id=str(uuid.uuid4())
    keys=('vehicle_label','cancellation_base_toman','waiting_hourly_toman','deviation_per_km_toman','deviation_time_unit_minutes','deviation_time_unit_toman','free_wait_minutes','extra_destination_free_minutes','deviation_use_distance','deviation_use_time')
    snapshot={k:setting[k] for k in keys}
    with get_db() as db:
        db.execute('INSERT INTO calculation_history(id,user_id,calculation_type,vehicle,input_json,rate_snapshot_json,result_json,created_at) VALUES(?,?,?,?,?,?,?,?)',
                   (calc_id,user_id,calc_type,vehicle,json.dumps(input_data,ensure_ascii=False),json.dumps(snapshot,ensure_ascii=False),json.dumps(result,ensure_ascii=False),now_iso()))
    audit(user_id,'calculation.completed',{'calculation_id':calc_id,'type':calc_type,'vehicle':vehicle,'amount_toman':result.get('final_amount_toman',result.get('amount_toman'))})
    return calc_id


@app.get('/api/v1/calculations/settings')
def calculation_settings_v2(user:dict[str,Any]=Depends(current_user)) -> list[dict[str,Any]]:
    with get_db() as db:
        rows=[dict(r) for r in db.execute("SELECT * FROM calculation_settings ORDER BY CASE vehicle WHEN 'nissan' THEN 1 WHEN 'peykan' THEN 2 ELSE 3 END").fetchall()]
    return [_setting_public(r) for r in rows]


@app.put('/api/v1/calculations/settings/{vehicle}')
def update_calculation_settings_v2(vehicle:str,data:CalculationSettingsInput,user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    key=_normalize_calculation_vehicle(vehicle)
    per_minute_toman=int(data.deviation_per_minute_toman or 0)
    if per_minute_toman>0:
        time_unit_minutes=1;time_unit_toman=per_minute_toman
    else:
        time_unit_minutes=int(data.deviation_time_unit_minutes or 0)
        time_unit_toman=int(data.deviation_time_unit_toman or 0)
        # Legacy settings are accepted, but normalized to a one-minute rate when possible.
        if time_unit_minutes>0 and time_unit_toman>0:
            time_unit_toman=_round_toman(Decimal(time_unit_toman)/Decimal(time_unit_minutes));time_unit_minutes=1
    if data.deviation_use_time and time_unit_minutes<=0 and time_unit_toman<=0 and data.waiting_hourly_toman>0:
        # Backward compatibility: derive a per-minute rate from the old hourly waiting rate.
        time_unit_minutes=1;time_unit_toman=_round_toman(Decimal(data.waiting_hourly_toman)/Decimal(60))
    if data.deviation_per_km_toman<=0 and (time_unit_minutes<=0 or time_unit_toman<=0):
        raise HTTPException(status_code=400,detail='حداقل یکی از نرخ‌های انحراف کیلومتری یا انحراف زمانی باید ثبت شود.')
    if (time_unit_minutes>0) != (time_unit_toman>0):
        raise HTTPException(status_code=400,detail='برای انحراف زمانی، مبلغ «انحراف مسیر به ازای هر یک دقیقه» را وارد کنید.')
    use_distance=data.deviation_per_km_toman>0
    use_time=time_unit_minutes>0 and time_unit_toman>0
    with get_db() as db:
        db.execute(
            """UPDATE calculation_settings SET cancellation_base_toman=?,waiting_hourly_toman=?,deviation_per_km_toman=?,deviation_time_unit_minutes=?,deviation_time_unit_toman=?,free_wait_minutes=?,extra_destination_free_minutes=?,deviation_use_distance=?,deviation_use_time=?,updated_by=?,updated_at=? WHERE vehicle=?""",
            (data.cancellation_base_toman,data.waiting_hourly_toman,data.deviation_per_km_toman,time_unit_minutes,time_unit_toman,data.free_wait_minutes,data.extra_destination_free_minutes,int(use_distance),int(use_time),user['id'],now_iso(),key),
        )
        row=dict(db.execute('SELECT * FROM calculation_settings WHERE vehicle=?',(key,)).fetchone())
    audit(user['id'],'calculation.settings.updated',{'vehicle':key})
    return _setting_public(row)


@app.post('/api/v1/calculations/cancellation')
def calculate_cancellation_v2(data:CancellationCalculationInput,user:dict[str,Any]=Depends(current_user)) -> dict[str,Any]:
    vehicle=_normalize_calculation_vehicle(data.vehicle);setting=_require_calculation_setting(vehicle)
    free=Decimal(15)
    if data.origin_wait_minutes is not None:
        origin=Decimal(data.origin_wait_minutes);billable=max(Decimal(0),origin-free);mode='origin_total_minus_15'
    elif data.billable_wait_minutes is not None:
        billable=Decimal(data.billable_wait_minutes);origin=billable+free;mode='legacy_billable_input'
    else:
        origin=Decimal(0);billable=Decimal(0);mode='origin_total_minus_15'
    wait_amount=Decimal(setting['waiting_hourly_toman'])*billable/Decimal(60)
    final=Decimal(setting['cancellation_base_toman'])+wait_amount
    result={'type':'cancellation','vehicle':vehicle,'vehicle_label':setting['vehicle_label'],'origin_wait_minutes':_decimal_number(origin),'free_wait_minutes':15,'billable_wait_minutes':_decimal_number(billable),'cancellation_base_toman':int(setting['cancellation_base_toman']),'waiting_hourly_toman':int(setting['waiting_hourly_toman']),'waiting_amount_toman':_round_toman(wait_amount),'final_amount_toman':_round_toman(final),'calculation_mode':mode}
    result['calculation_id']=_record_calculation(user['id'],'cancellation',vehicle,{'origin_wait_minutes':_decimal_number(origin),'billable_wait_minutes':_decimal_number(billable)},setting,result)
    return result


@app.post('/api/v1/calculations/waiting')
def calculate_waiting_v2(data:WaitingCalculationInput,user:dict[str,Any]=Depends(current_user)) -> dict[str,Any]:
    vehicle=_normalize_calculation_vehicle(data.vehicle);setting=_require_calculation_setting(vehicle);mode=str(data.calculation_mode or 'minutes').lower()
    breakdown=[];clock=None
    if mode=='minutes':
        origin=Decimal(data.origin_wait_minutes);destination=Decimal(data.destination_wait_minutes)
    elif mode=='clock':
        if not data.origin_start_time or not data.origin_end_time:
            raise HTTPException(status_code=400,detail='ساعت ورود و حرکت راننده در مبدا را کامل وارد کنید.')
        if len(data.destination_time_ranges)!=data.destination_count:
            raise HTTPException(status_code=400,detail='برای همه مقصدها ساعت رسیدن و پایان توقف را وارد کنید.')
        origin=_clock_interval(data.origin_start_time,data.origin_end_time,'توقف مبدا')
        breakdown=[_clock_interval(x.start_time,x.end_time,f'توقف مقصد {i}') for i,x in enumerate(data.destination_time_ranges,start=1)]
        destination=sum(breakdown,Decimal(0))
        clock={'origin':{'start_time':data.origin_start_time,'end_time':data.origin_end_time,'minutes':_decimal_number(origin)},'destinations':[{'destination_number':i,'start_time':x.start_time,'end_time':x.end_time,'minutes':_decimal_number(breakdown[i-1])} for i,x in enumerate(data.destination_time_ranges,start=1)]}
    else:
        raise HTTPException(status_code=400,detail='روش محاسبه توقف نامعتبر است.')
    total=origin+destination
    free=Decimal(setting['free_wait_minutes'])+Decimal(setting['extra_destination_free_minutes'])*Decimal(max(0,data.destination_count-1))
    billable=max(Decimal(0),total-free)
    amount=Decimal(setting['waiting_hourly_toman'])*billable/Decimal(60)
    result={'type':'waiting','vehicle':vehicle,'vehicle_label':setting['vehicle_label'],'calculation_mode':mode,'origin_wait_minutes':_decimal_number(origin),'destination_wait_minutes':_decimal_number(destination),'destination_breakdown_minutes':[_decimal_number(x) for x in breakdown],'destination_count':data.destination_count,'total_wait_minutes':_decimal_number(total),'base_free_wait_minutes':int(setting['free_wait_minutes']),'extra_destination_free_minutes':int(setting['extra_destination_free_minutes']),'free_wait_minutes':_decimal_number(free),'billable_wait_minutes':_decimal_number(billable),'waiting_hourly_toman':int(setting['waiting_hourly_toman']),'amount_toman':_round_toman(amount),'final_amount_toman':_round_toman(amount)}
    if clock:result['clock_times']=clock
    result['calculation_id']=_record_calculation(user['id'],'waiting',vehicle,{'calculation_mode':mode,'origin_wait_minutes':_decimal_number(origin),'destination_wait_minutes':_decimal_number(destination),'destination_count':data.destination_count},setting,result)
    return result


@app.post('/api/v1/calculations/deviation')
def calculate_deviation_v2(data:DeviationCalculationInput,user:dict[str,Any]=Depends(current_user)) -> dict[str,Any]:
    vehicle=_normalize_calculation_vehicle(data.vehicle);setting=_require_calculation_setting(vehicle)
    requested_mode=str(data.mode or '').strip().lower()
    if requested_mode not in {'distance','time'}:
        if data.distance_km is not None and data.wait_minutes is None:
            requested_mode='distance'
        elif data.wait_minutes is not None and data.distance_km is None:
            requested_mode='time'
        else:
            requested_mode='distance' if bool(setting.get('deviation_use_distance',1)) else 'time'
    distance=Decimal(data.distance_km) if data.distance_km is not None else None
    minutes=Decimal(data.wait_minutes) if data.wait_minutes is not None else None
    if requested_mode=='distance':
        rate=Decimal(int(setting.get('deviation_per_km_toman') or 0))
        if rate<=0:
            raise HTTPException(status_code=409,detail='نرخ انحراف کیلومتری این خودرو هنوز در «آموزش محاسبات» ثبت نشده است.')
        if distance is None or distance<=0:
            raise HTTPException(status_code=400,detail='میزان انحراف مسیر را به کیلومتر وارد کنید.')
        surcharge=rate*distance;mode='distance'
    else:
        unit_minutes=Decimal(int(setting.get('deviation_time_unit_minutes') or 0))
        unit_toman=Decimal(int(setting.get('deviation_time_unit_toman') or 0))
        if unit_minutes<=0 or unit_toman<=0:
            raise HTTPException(status_code=409,detail='نرخ انحراف زمانی این خودرو هنوز در «آموزش محاسبات» ثبت نشده است.')
        if minutes is None or minutes<=0:
            raise HTTPException(status_code=400,detail='مدت معطلی ناشی از انحراف مسیر را به دقیقه وارد کنید.')
        surcharge=unit_toman*minutes/unit_minutes;mode='time'
    service=Decimal(data.service_amount_toman) if data.service_amount_toman is not None else None
    final=service+surcharge if service is not None else surcharge
    result={
        'type':'deviation','vehicle':vehicle,'vehicle_label':setting['vehicle_label'],'deviation_mode':mode,
        'deviation_use_distance':bool(setting.get('deviation_use_distance',1)),'deviation_use_time':bool(setting.get('deviation_use_time',1)),
        'distance_km':_decimal_number(distance) if distance is not None else None,'wait_minutes':_decimal_number(minutes) if minutes is not None else None,
        'deviation_per_km_toman':int(setting.get('deviation_per_km_toman') or 0),
        'deviation_time_unit_minutes':int(setting.get('deviation_time_unit_minutes') or 0),
        'deviation_time_unit_toman':int(setting.get('deviation_time_unit_toman') or 0),
        'deviation_per_minute_toman':_round_toman(Decimal(int(setting.get('deviation_time_unit_toman') or 0))/Decimal(max(1,int(setting.get('deviation_time_unit_minutes') or 1)))) if int(setting.get('deviation_time_unit_toman') or 0)>0 else 0,
        'deviation_amount_toman':_round_toman(surcharge),'service_amount_toman':_round_toman(service) if service is not None else None,
        'amount_toman':_round_toman(surcharge),'final_amount_toman':_round_toman(final),
    }
    result['calculation_id']=_record_calculation(user['id'],'deviation',vehicle,{'distance_km':result['distance_km'],'wait_minutes':result['wait_minutes'],'service_amount_toman':result['service_amount_toman'],'deviation_mode':mode},setting,result)
    return result


@app.get('/api/v1/calculations/history')
def calculation_history_v2(limit:int=30,user:dict[str,Any]=Depends(current_user)) -> list[dict[str,Any]]:
    limit=max(1,min(int(limit),100))
    with get_db() as db:
        rows=db.execute('SELECT * FROM calculation_history WHERE user_id=? ORDER BY created_at DESC LIMIT ?',(user['id'],limit)).fetchall()
    labels={'cancellation':'کنسلی','waiting':'توقف','deviation':'انحراف مسیر'}
    out=[]
    for row in rows:
        try:result=json.loads(row['result_json'] or '{}')
        except json.JSONDecodeError:result={}
        out.append({'id':row['id'],'calculation_type':row['calculation_type'],'calculation_label':labels.get(row['calculation_type'],row['calculation_type']),'vehicle':row['vehicle'],'vehicle_label':CALCULATION_VEHICLES.get(row['vehicle'],row['vehicle']),'result':result,'created_at':row['created_at']})
    return out


@app.post('/api/v1/guest/chat')
async def guest_chat(data: GuestChatInput, request: Request) -> dict[str, Any]:
    if not GUEST_CHAT_ENABLED:
        raise HTTPException(status_code=403, detail='چت مهمان غیرفعال است.')
    identity=_anonymous_rate_identity('guest',request,None)
    return await process_chat(message=data.message, conversation_id=data.conversation_id, user=None, external_user_id=data.external_user_id, integration=True, rate_identity=identity, rate_per_minute=GUEST_REQUESTS_PER_MINUTE, rate_daily=GUEST_DAILY_LIMIT, rate_monthly=GUEST_MONTHLY_LIMIT)


@app.post('/api/v1/integration/chat')
async def integration_chat(
    data: IntegrationChatInput,
    request: Request,
    x_api_key: str | None = Header(default=None,alias='X-API-Key'),
    x_integration_key: str | None = Header(default=None,alias='X-Integration-Key'),
) -> dict[str, Any]:
    client_name = validate_integration_key(x_api_key or x_integration_key)
    external=f'{client_name}:{data.external_user_id}'
    identity=_anonymous_rate_identity(f'integration:{client_name}',request,None)
    result = await process_chat(message=data.message, conversation_id=data.conversation_id, user=None, external_user_id=external, integration=True, rate_identity=identity, rate_per_minute=REQUESTS_PER_MINUTE, rate_daily=INTEGRATION_DAILY_LIMIT, rate_monthly=0)
    result['client'] = client_name
    return result


@app.get('/api/v1/conversations')
def my_conversations(user: dict[str, Any] = Depends(current_user)) -> list[dict[str, Any]]:
    with get_db() as db:
        rows = db.execute("SELECT id, title, created_at, updated_at FROM conversations WHERE user_id = ? ORDER BY updated_at DESC LIMIT 60", (user['id'],)).fetchall()
    return [dict(r) for r in rows]


@app.get('/api/v1/conversations/{conversation_id}/messages')
def my_messages(conversation_id: str, user: dict[str, Any] = Depends(current_user)) -> list[dict[str, Any]]:
    ensure_conversation_access(conversation_id, user)
    with get_db() as db:
        rows = db.execute("SELECT id, role, content, sources_json, model, response_ms, status, prompt_tokens, output_tokens, total_tokens, api_slot, model_route, estimated_cost, confidence_score, created_at FROM messages WHERE conversation_id = ? ORDER BY created_at ASC", (conversation_id,)).fetchall()
    result=[]
    for row in rows:
        item=dict(row)
        try:
            sources=json.loads(item.get('sources_json') or '[]')
        except (TypeError, json.JSONDecodeError):
            sources=[]
        item['sources']=sources
        if item.get('role')=='assistant':
            item.update(response_section_details(item.get('status'),item.get('model'),sources,item.get('total_tokens') or 0))
        result.append(item)
    return result


def _question_for_assistant_message(message_id: str) -> str | None:
    with get_db() as db:
        row=db.execute(
            """SELECT q.content FROM messages a JOIN messages q
                 ON q.conversation_id=a.conversation_id AND q.role='user'
               WHERE a.id=? AND q.created_at<=a.created_at
               ORDER BY q.created_at DESC,q.id DESC LIMIT 1""",
            (message_id,),
        ).fetchone()
    return str(row['content']) if row else None


def _review_cache_for_message(message_id: str, rating: str, reviewer: dict[str,Any]) -> dict[str,int]:
    question=_question_for_assistant_message(message_id)
    if not question:
        return {'deleted':0,'promoted':0}
    normalized=canonical_question_for_cache(question)
    if not normalized:
        return {'deleted':0,'promoted':0}
    deleted=promoted=0
    with get_db() as db:
        if rating in {'wrong','wrong_source','incomplete'}:
            cur=db.execute("DELETE FROM answer_cache WHERE normalized_question=?",(normalized,))
            deleted=int(cur.rowcount or 0)
        elif rating=='correct' and reviewer.get('role') in {'manager','admin'}:
            cur=db.execute(
                """UPDATE answer_cache SET cache_tier='approved',expires_at=NULL,approved_by=?
                   WHERE normalized_question=? AND cache_tier='temporary'""",
                (reviewer['id'],normalized),
            )
            promoted=int(cur.rowcount or 0)
    return {'deleted':deleted,'promoted':promoted}


@app.post('/api/v1/messages/{message_id}/feedback')
def submit_answer_feedback(message_id: str, data: FeedbackInput, user: dict[str,Any]=Depends(current_user)) -> dict[str,Any]:
    allowed={'correct','incomplete','wrong','wrong_source','more_detail'}
    if data.rating not in allowed:
        raise HTTPException(status_code=400,detail='نوع بازخورد معتبر نیست.')
    with get_db() as db:
        row=db.execute("""SELECT m.id,m.role,c.user_id FROM messages m JOIN conversations c ON c.id=m.conversation_id WHERE m.id=?""",(message_id,)).fetchone()
        if not row or row['role']!='assistant':
            raise HTTPException(status_code=404,detail='پاسخ پیدا نشد.')
        if user['role']=='user' and int(row['user_id'] or 0)!=int(user['id']):
            raise HTTPException(status_code=403,detail='به این پاسخ دسترسی ندارید.')
        ts=now_iso()
        db.execute("""INSERT INTO answer_feedback(message_id,user_id,rating,note,created_at,updated_at)
            VALUES(?,?,?,?,?,?) ON CONFLICT(message_id,user_id) DO UPDATE SET rating=excluded.rating,note=excluded.note,updated_at=excluded.updated_at""",
            (message_id,user['id'],data.rating,(data.note or '').strip() or None,ts,ts))
    cache_action=_review_cache_for_message(message_id,data.rating,user)
    audit(user['id'],'answer.feedback',{'message_id':message_id,'rating':data.rating,'cache_action':cache_action})
    if data.rating in {'wrong','wrong_source','incomplete'}:
        notify_admin('answer_quality','یک پاسخ نیازمند بررسی مدیر است',{'message_id':message_id,'rating':data.rating,'note':data.note,'cache_invalidated':cache_action['deleted']},'warning')
    return {'ok':True,'rating':data.rating,'cache_action':cache_action}


@app.post('/api/v1/admin/users')
def create_user(data: CreateUserInput, admin: dict[str,Any]=Depends(require_roles('admin'))):
    if data.role not in {'user','manager','admin'}:
        raise HTTPException(status_code=400,detail='نقش نامعتبر است.')
    if data.role=='admin' and not admin.get('is_owner'):
        raise HTTPException(status_code=403,detail='فقط ادمین اصلی می‌تواند ادمین جدید ایجاد کند.')
    username=data.username.strip().lower()
    if not re.fullmatch(r'[a-z0-9_.-]{3,50}',username):
        raise HTTPException(status_code=400,detail='نام کاربری فقط شامل حروف انگلیسی، عدد و _-. باشد.')
    password_hash,salt=hash_password(data.password)
    try:
        with get_db() as db:
            cur=db.execute(
                """INSERT INTO users(username,email,name,password_hash,salt,role,is_active,is_owner,department,
                    question_limit,questions_used,daily_question_limit,daily_questions_used,daily_quota_date,
                    monthly_question_limit,monthly_questions_used,monthly_quota_month,created_by,created_at)
                    VALUES(?,?,?,?,?,?,1,0,?,?,0,?,0,?,?,0,?,?,?)""",
                (username,str(data.email).lower() if data.email else None,data.name.strip(),password_hash,salt,data.role,
                 (data.department or '').strip() or None,data.question_limit,data.daily_question_limit,_quota_period_keys()[0],
                 data.monthly_question_limit,_quota_period_keys()[1],admin['id'],now_iso()),
            )
            user_id=cur.lastrowid
    except sqlite3.IntegrityError as exc:
        raise HTTPException(status_code=409,detail='نام کاربری یا ایمیل تکراری است.') from exc
    audit(admin['id'],'user.created',{'user_id':user_id,'username':username,'role':data.role,'question_limit':data.question_limit,'daily_question_limit':data.daily_question_limit,'monthly_question_limit':data.monthly_question_limit,'department':data.department})
    notify_admin('user_created','حساب سازمانی جدید ایجاد شد',{'user_id':user_id,'username':username,'role':data.role,'question_limit':data.question_limit,'daily_question_limit':data.daily_question_limit,'monthly_question_limit':data.monthly_question_limit})
    return {'id':user_id,'username':username,'email':str(data.email).lower() if data.email else None,'name':data.name.strip(),'role':data.role,'department':(data.department or '').strip() or None,'is_active':True,'is_owner':False,'question_limit':data.question_limit,'questions_used':0,'daily_question_limit':data.daily_question_limit,'daily_questions_used':0,'monthly_question_limit':data.monthly_question_limit,'monthly_questions_used':0}


@app.get('/api/v1/admin/users')
def list_users(admin: dict[str,Any]=Depends(require_roles('admin'))):
    with get_db() as db:
        rows=db.execute("SELECT id,username,email,name,role,is_active,is_owner,department,question_limit,questions_used,daily_question_limit,daily_questions_used,daily_quota_date,monthly_question_limit,monthly_questions_used,monthly_quota_month,created_at,last_login_at FROM users ORDER BY is_owner DESC,created_at DESC").fetchall()
    return [dict(r) for r in rows]


@app.patch('/api/v1/admin/users/{user_id}')
def update_user(user_id:int,data:UpdateUserInput,admin:dict[str,Any]=Depends(require_roles('admin'))):
    with get_db() as db:
        row=db.execute("SELECT * FROM users WHERE id=?",(user_id,)).fetchone()
        if not row: raise HTTPException(status_code=404,detail='کاربر پیدا نشد.')
        target=dict(row)
        if target.get('is_owner'):
            if admin['id']!=user_id:
                raise HTTPException(status_code=403,detail='ادمین اصلی قابل تنزل، غیرفعال‌سازی یا ویرایش توسط دیگران نیست.')
            if data.role not in (None,'admin') or data.is_active is False:
                raise HTTPException(status_code=400,detail='ادمین اصلی باید فعال و دارای نقش admin باقی بماند.')
        elif target['role']=='admin' and not admin.get('is_owner'):
            raise HTTPException(status_code=403,detail='ویرایش حساب ادمین فقط توسط ادمین اصلی مجاز است.')
        role=data.role if data.role is not None else target['role']
        if role not in {'user','manager','admin'}: raise HTTPException(status_code=400,detail='نقش نامعتبر است.')
        if role=='admin' and not admin.get('is_owner'): raise HTTPException(status_code=403,detail='فقط ادمین اصلی می‌تواند نقش admin بدهد.')
        if user_id==admin['id'] and (role!='admin' or data.is_active is False):
            raise HTTPException(status_code=400,detail='نمی‌توانید حساب فعال خودتان را تنزل یا غیرفعال کنید.')
        is_active=int(data.is_active) if data.is_active is not None else target['is_active']
        password_hash,salt=target['password_hash'],target['salt']
        if data.password: password_hash,salt=hash_password(data.password)
        question_limit = data.question_limit if data.set_question_limit else target.get('question_limit')
        daily_limit = data.daily_question_limit if data.set_daily_question_limit else target.get('daily_question_limit')
        monthly_limit = data.monthly_question_limit if data.set_monthly_question_limit else target.get('monthly_question_limit')
        department = ((data.department or '').strip() or None) if data.set_department else target.get('department')
        questions_used = 0 if data.reset_questions_used else int(target.get('questions_used') or 0)
        daily_used = 0 if data.reset_daily_questions_used else int(target.get('daily_questions_used') or 0)
        monthly_used = 0 if data.reset_monthly_questions_used else int(target.get('monthly_questions_used') or 0)
        day_key,month_key=_quota_period_keys()
        db.execute("""UPDATE users SET role=?,is_active=?,password_hash=?,salt=?,department=?,question_limit=?,questions_used=?,
            daily_question_limit=?,daily_questions_used=?,daily_quota_date=?,monthly_question_limit=?,monthly_questions_used=?,monthly_quota_month=? WHERE id=?""",
            (role,is_active,password_hash,salt,department,question_limit,questions_used,daily_limit,daily_used,day_key,monthly_limit,monthly_used,month_key,user_id))
    audit(admin['id'],'user.updated',{'user_id':user_id,'role':role,'is_active':bool(is_active),'password_reset':bool(data.password),'department':department,'question_limit':question_limit,'daily_question_limit':daily_limit,'monthly_question_limit':monthly_limit,'questions_used_reset':data.reset_questions_used})
    return {'ok':True,'department':department,'question_limit':question_limit,'questions_used':questions_used,'daily_question_limit':daily_limit,'daily_questions_used':daily_used,'monthly_question_limit':monthly_limit,'monthly_questions_used':monthly_used}


def _file_sha256(path: Path) -> str:
    digest=hashlib.sha256()
    with path.open('rb') as handle:
        for chunk in iter(lambda:handle.read(1024*1024),b''):
            digest.update(chunk)
    return digest.hexdigest()


def _create_document_record(
    *,
    filename: str,
    stored_path: Path,
    content_type: str | None,
    visibility: str,
    admin_id: int,
) -> tuple[str, str]:
    document_id=str(uuid.uuid4());job_id=str(uuid.uuid4())
    suffix=stored_path.suffix.lower()
    final_path=UPLOAD_DIR/f'{document_id}{suffix}'
    if stored_path.resolve()!=final_path.resolve():
        final_path.parent.mkdir(parents=True,exist_ok=True)
        shutil.move(str(stored_path),str(final_path))
    content_sha256=_file_sha256(final_path)
    file_size_bytes=final_path.stat().st_size
    with get_db() as db:
        db.execute("""INSERT INTO documents(id,filename,stored_path,mime_type,visibility,status,character_count,chunk_count,version,created_by,created_at,content_sha256,file_size_bytes)
            VALUES(?,?,?,?,?,'processing',0,0,1,?,?,?,?)""",(document_id,filename,str(final_path),content_type,visibility,admin_id,now_iso(),content_sha256,file_size_bytes))
        ts=now_iso();priority=document_job_priority('ingest','فایل دریافت شد')
        db.execute("INSERT INTO document_jobs(id,document_id,status,progress,phase,created_by,created_at,updated_at,job_type,priority) VALUES(?,?,'queued',5,'فایل دریافت شد',?,?,?,'ingest',?)",(job_id,document_id,admin_id,ts,ts,priority))
        db.execute("UPDATE documents SET reindex_status='queued',reindex_started_at=? WHERE id=?",(ts,document_id))
    return document_id,job_id


async def _stream_upload_to_path(file: UploadFile, destination: Path, max_bytes: int) -> int:
    total=0
    destination.parent.mkdir(parents=True,exist_ok=True)
    try:
        with destination.open('wb') as out:
            while True:
                chunk=await file.read(1024*1024)
                if not chunk: break
                total+=len(chunk)
                if total>max_bytes:
                    raise HTTPException(status_code=413,detail=f'حجم فایل باید حداکثر {MAX_UPLOAD_MB} مگابایت باشد.')
                out.write(chunk)
        if total<=0:
            raise HTTPException(status_code=400,detail='فایل خالی است.')
        return total
    except Exception:
        destination.unlink(missing_ok=True)
        raise


def _upload_session_paths(upload_id: str) -> tuple[Path,Path,Path]:
    if not re.fullmatch(r'[a-f0-9]{32}',upload_id):
        raise HTTPException(status_code=400,detail='شناسه آپلود معتبر نیست.')
    root=UPLOAD_SESSION_DIR/upload_id
    return root,root/'meta.json',root/'payload.part'


def _write_upload_meta(path: Path, meta: dict[str,Any]) -> None:
    tmp=path.with_suffix('.tmp')
    tmp.write_text(json.dumps(meta,ensure_ascii=False),encoding='utf-8')
    tmp.replace(path)


def _read_upload_meta(upload_id: str, admin_id: int) -> tuple[dict[str,Any],Path,Path,Path]:
    root,meta_path,part_path=_upload_session_paths(upload_id)
    if not meta_path.exists():
        raise HTTPException(status_code=404,detail='جلسه آپلود پیدا نشد یا منقضی شده است.')
    try: meta=json.loads(meta_path.read_text(encoding='utf-8'))
    except Exception as exc: raise HTTPException(status_code=500,detail='اطلاعات جلسه آپلود خراب است.') from exc
    if int(meta.get('admin_id') or 0)!=int(admin_id):
        raise HTTPException(status_code=403,detail='به این جلسه آپلود دسترسی ندارید.')
    if part_path.exists():
        actual=part_path.stat().st_size;recorded=int(meta.get('received_bytes') or 0);chunk_bytes=int(meta.get('chunk_bytes') or UPLOAD_CHUNK_BYTES);total=int(meta.get('size_bytes') or 0)
        if actual!=recorded:
            if 0<=actual<=total and (actual==total or actual%chunk_bytes==0):
                meta['received_bytes']=actual;meta['next_index']=math.ceil(actual/chunk_bytes) if actual else 0;meta['updated_at']=now_iso();_write_upload_meta(meta_path,meta)
            else:
                safe=max(0,min(recorded,total));
                with part_path.open('r+b') as handle: handle.truncate(safe)
    return meta,root,meta_path,part_path


def _cleanup_expired_upload_sessions() -> None:
    cutoff=time.time()-UPLOAD_SESSION_TTL_HOURS*3600
    if not UPLOAD_SESSION_DIR.exists(): return
    for root in UPLOAD_SESSION_DIR.iterdir():
        try:
            if root.stat().st_mtime<cutoff:
                if root.is_dir(): shutil.rmtree(root,ignore_errors=True)
                else: root.unlink(missing_ok=True)
        except OSError: pass


def _extract_google_doc_id(raw_url: str) -> str:
    parsed=urlparse(raw_url.strip())
    host=(parsed.hostname or '').lower()
    if host not in {'docs.google.com','www.docs.google.com'}:
        raise HTTPException(status_code=400,detail='فقط لینک معتبر Google Docs پذیرفته می‌شود.')
    match=re.search(r'/document/(?:u/\d+/)?d/([a-zA-Z0-9_-]+)',parsed.path)
    if not match:
        raise HTTPException(status_code=400,detail='شناسه سند از لینک Google Docs پیدا نشد.')
    return match.group(1)



def _recover_stale_document_jobs() -> int:
    cutoff=(datetime.now(timezone.utc)-timedelta(seconds=DOCUMENT_JOB_STALE_SECONDS)).isoformat()
    with get_db() as db:
        cur=db.execute(
            """UPDATE document_jobs SET status='queued',worker_id=NULL,phase='بازیابی پس از راه‌اندازی مجدد',next_run_at=?,updated_at=?
               WHERE status='processing' AND updated_at<?""",
            (now_iso(),now_iso(),cutoff),
        )
        return int(cur.rowcount or 0)


def _claim_next_document_job(worker_id: str) -> dict[str,Any] | None:
    now=now_iso()
    con=get_db()
    try:
        con.execute('BEGIN IMMEDIATE')
        row=con.execute(
            """SELECT * FROM document_jobs WHERE status='queued' AND attempts<max_attempts
               AND (next_run_at IS NULL OR next_run_at<=?) ORDER BY priority DESC,created_at,id LIMIT 1""",
            (now,),
        ).fetchone()
        if not row:
            con.commit();return None
        con.execute(
            """UPDATE document_jobs SET status='processing',attempts=attempts+1,worker_id=?,phase=COALESCE(phase,'در صف پردازش'),updated_at=? WHERE id=?""",
            (worker_id,now,row['id']),
        )
        con.commit()
        return dict(con.execute('SELECT * FROM document_jobs WHERE id=?',(row['id'],)).fetchone())
    finally:
        con.close()


def _run_one_document_job_guarded(worker_id: str) -> bool:
    if _RESTORE_IN_PROGRESS.is_set():
        return False
    with _MAINTENANCE_LOCK:
        if _RESTORE_IN_PROGRESS.is_set():
            return False
        job=_claim_next_document_job(worker_id)
        if not job:
            return False
        _process_document_job(str(job['id']),str(job['document_id']))
        return True


async def _document_job_worker_loop() -> None:
    worker_id=f"worker-{uuid.uuid4().hex[:10]}"
    while True:
        try:
            worked=await asyncio.to_thread(_run_one_document_job_guarded,worker_id)
            if not worked:
                await asyncio.sleep(DOCUMENT_JOB_POLL_SECONDS)
        except asyncio.CancelledError:
            raise
        except Exception as exc:
            try:
                notify_admin_once('document_worker_error','خطای Worker پردازش منابع',{'error':sanitize_answer_text(str(exc))[:350]},'danger',3)
            except Exception:
                pass
            await asyncio.sleep(DOCUMENT_JOB_POLL_SECONDS)


def _process_document_job_guarded(job_id: str, document_id: str) -> None:
    with _MAINTENANCE_LOCK:
        if _RESTORE_IN_PROGRESS.is_set():
            raise HTTPException(status_code=503,detail='بازیابی پشتیبان در حال انجام است؛ پردازش منبع موقتاً متوقف شده است.')
        _process_document_job(job_id,document_id)


def _enqueue_reindex_job(document_id: str, actor_id: int, phase: str='در صف بازسازی منبع') -> str:
    job_id=str(uuid.uuid4())
    with get_db() as db:
        existing=db.execute("SELECT id FROM document_jobs WHERE document_id=? AND status IN ('queued','processing') ORDER BY created_at DESC LIMIT 1",(document_id,)).fetchone()
        if existing:
            return str(existing['id'])
        priority=document_job_priority('reindex',phase)
        ts=now_iso()
        db.execute(
            """INSERT INTO document_jobs(id,document_id,status,progress,phase,created_by,created_at,updated_at,job_type,payload_json,attempts,max_attempts,next_run_at,priority)
               VALUES(?,?,'queued',5,?,?,?,?, 'reindex','{}',0,?,?,?)""",
            (job_id,document_id,phase,actor_id,ts,ts,DOCUMENT_JOB_MAX_ATTEMPTS,ts,priority),
        )
        db.execute("UPDATE documents SET reindex_status='queued',reindex_error='',reindex_started_at=COALESCE(reindex_started_at,?) WHERE id=?",(ts,document_id))
        # R35.2 staging contract: the active document status/chunks are never changed
        # merely because a replacement reindex has been queued.
    return job_id

def _enqueue_legacy_source_reindex_jobs() -> dict[str,int]:
    """Queue R34/older source indexes for R35 quality enrichment without downtime."""
    with get_db() as db:
        owner=db.execute("SELECT id FROM users WHERE is_owner=1 ORDER BY id LIMIT 1").fetchone()
        actor_id=int(owner['id']) if owner else 0
        rows=db.execute("SELECT id,stored_path,ingestion_version FROM documents WHERE COALESCE(ingestion_version,1)<? ORDER BY created_at ASC LIMIT ?",(INGESTION_VERSION,AUTO_REINDEX_LEGACY_LIMIT)).fetchall()
    queued=0;missing=0
    for row in rows:
        if not Path(str(row['stored_path'] or '')).is_file():
            missing+=1;continue
        before=None
        with get_db() as db:
            before=db.execute("SELECT id FROM document_jobs WHERE document_id=? AND status IN ('queued','processing') LIMIT 1",(row['id'],)).fetchone()
        _enqueue_reindex_job(str(row['id']),actor_id,'مهاجرت خودکار منبع به Quality Gate نسخه R35')
        if not before: queued+=1
    return {'candidates':len(rows),'queued':queued,'missing_files':missing}


def _process_document_job(job_id: str, document_id: str) -> None:
    try:
        with get_db() as db:
            doc=db.execute("SELECT * FROM documents WHERE id=?",(document_id,)).fetchone()
            if not doc: raise RuntimeError('سند پیدا نشد.')
            ts=now_iso()
            db.execute("UPDATE document_jobs SET status='processing',progress=20,phase='استخراج متن',updated_at=? WHERE id=?",(ts,job_id))
            db.execute("UPDATE documents SET reindex_status='processing',reindex_error='',reindex_started_at=COALESCE(reindex_started_at,?) WHERE id=?",(ts,document_id))
        path=Path(doc['stored_path'])
        def update_extract_progress(progress: int, phase: str) -> None:
            with get_db() as progress_db:
                progress_db.execute("UPDATE document_jobs SET progress=?,phase=?,updated_at=? WHERE id=?",(max(20,min(50,int(progress))),phase,now_iso(),job_id))
        result=extract_document_result_from_path(doc['filename'],path,update_extract_progress)
        # R35.2 quality gate: uploaded manager resources are never allowed to enter
        # retrieval with missing/failed pages. A partial extraction is retried by
        # the document worker; an existing live index stays untouched during retry.
        if str(result.get('status') or 'ready')=='partial' and not int(doc['is_builtin'] or 0):
            stats=result.get('stats') or {}
            warnings=' | '.join(result.get('warnings') or [])[:420]
            raise RuntimeError(
                f"کنترل کیفیت فایل کامل نشد؛ کیفیت {float(stats.get('ingestion_quality_pct') or 0):.0f}٪، "
                f"صفحات Vision ناموفق {int(stats.get('vision_failed_pages') or 0)}. "
                f"فایل تا تکمیل تحلیل وارد پاسخ‌گویی نمی‌شود. {warnings}"
            )
        text=str(result.get('text') or '').strip()
        parts=chunk_source_result(result)
        if not parts: raise RuntimeError('قطعه دانشی ساخته نشد.')
        with get_db() as db:
            db.execute("UPDATE document_jobs SET progress=58,phase='ساخت Embedding و حقایق ساختاری',updated_at=? WHERE id=?",(now_iso(),job_id))
        embedding_rows,embedding_error=build_chunk_embedding_rows(parts)
        if embedding_error:
            result.setdefault('warnings',[]).append('Embedding واقعی در این نوبت ساخته نشد؛ جست‌وجوی محلی همچنان فعال است: '+embedding_error)
        with get_db() as db:
            db.execute("UPDATE document_jobs SET progress=72,phase='ساخت ایندکس Hybrid و صفحه‌ای',updated_at=? WHERE id=?",(now_iso(),job_id))
            live=db.execute("SELECT status,chunk_count FROM documents WHERE id=?",(document_id,)).fetchone()
            has_ready_active=bool(live and str(live['status'] or '')=='ready' and int(live['chunk_count'] or 0)>0)
            proposed_status=str(result.get('status') or 'ready')
            if has_ready_active and proposed_status!='ready':
                raise RuntimeError('نسخه Staging از Quality Gate عبور نکرد؛ ایندکس Ready قبلی بدون تغییر فعال ماند.')
            persist_document_ingestion(db,document_id,doc,result,parts,embedding_rows)
            db.execute("UPDATE document_jobs SET progress=88,phase='کنترل کیفیت و آماده‌سازی جست‌وجو',updated_at=? WHERE id=?",(now_iso(),job_id))
            final_status=str(result.get('status') or 'ready')
            if final_status=='partial' and not SOURCE_INCLUDE_PARTIAL_DOCUMENTS:
                final_status='error'
            phase='کامل' if final_status=='ready' else 'کامل با هشدار' if final_status=='partial' else 'خطا'
            error_text=None if final_status=='ready' else (' | '.join(result.get('warnings') or [])[:700] or 'بخشی از منبع کامل پردازش نشد.')
            done_ts=now_iso()
            db.execute("UPDATE document_jobs SET status=?,progress=100,phase=?,error=?,updated_at=? WHERE id=?",(final_status,phase,error_text,done_ts,job_id))
            db.execute("UPDATE documents SET reindex_status=?,reindex_error=?,reindex_completed_at=? WHERE id=?",('idle' if final_status=='ready' else final_status,error_text or '',done_ts,document_id))
        stats=result.get('stats') or {}
        severity='success' if final_status=='ready' else 'warning' if final_status=='partial' else 'danger'
        title='منبع جدید آماده جست‌وجو شد' if final_status=='ready' else 'منبع با هشدار آماده جست‌وجو شد' if final_status=='partial' else 'پردازش منبع ناقص ماند'
        notify_admin('document_ready' if final_status!='error' else 'document_error',title,{'document_id':document_id,'filename':doc['filename'],'chunks':len(parts),'quality_pct':stats.get('ingestion_quality_pct'),'vision_failed_pages':stats.get('vision_failed_pages')},severity)
    except Exception as exc:
        error=sanitize_answer_text(str(exc))[:700]
        with get_db() as db:
            job=db.execute('SELECT attempts,max_attempts FROM document_jobs WHERE id=?',(job_id,)).fetchone()
            attempts=int(job['attempts'] or 0) if job else DOCUMENT_JOB_MAX_ATTEMPTS
            max_attempts=int(job['max_attempts'] or DOCUMENT_JOB_MAX_ATTEMPTS) if job else DOCUMENT_JOB_MAX_ATTEMPTS
            existing_doc=db.execute('SELECT chunk_count,status,is_builtin FROM documents WHERE id=?',(document_id,)).fetchone()
            has_live_index=bool(existing_doc and int(existing_doc['chunk_count'] or 0)>0)
            fallback_status=preserve_live_document_status(existing_doc['status'] if existing_doc else None,existing_doc['chunk_count'] if existing_doc else 0)
            if attempts < max_attempts:
                delay=min(300,20*(2**max(0,attempts-1)))
                next_run=(datetime.now(timezone.utc)+timedelta(seconds=delay)).isoformat()
                # Staging failure cannot demote an existing Ready index.
                if existing_doc and str(existing_doc['status'] or '')!='ready':
                    db.execute('UPDATE documents SET status=? WHERE id=?',(fallback_status,document_id))
                db.execute("UPDATE documents SET reindex_status='queued',reindex_error=? WHERE id=?",(error,document_id))
                db.execute("UPDATE document_jobs SET status='queued',progress=5,phase='تلاش مجدد خودکار',error=?,worker_id=NULL,next_run_at=?,updated_at=? WHERE id=?",(error,next_run,now_iso(),job_id))
                retrying=True
            else:
                # Exhausted staging retries still leave the last Ready active index online.
                final_doc_status='ready' if existing_doc and str(existing_doc['status'] or '')=='ready' and has_live_index else ('partial' if has_live_index else 'error')
                if not (existing_doc and str(existing_doc['status'] or '')=='ready' and has_live_index):
                    db.execute('UPDATE documents SET status=? WHERE id=?',(final_doc_status,document_id))
                done_ts=now_iso()
                db.execute("UPDATE documents SET reindex_status='error',reindex_error=?,reindex_completed_at=? WHERE id=?",(error,done_ts,document_id))
                db.execute("UPDATE document_jobs SET status='error',progress=100,phase='خطا پس از چند تلاش',error=?,worker_id=NULL,updated_at=? WHERE id=?",(error,done_ts,job_id))
                retrying=False
        notify_admin('document_retry' if retrying else 'document_error','پردازش منبع دوباره تلاش می‌شود' if retrying else 'پردازش منبع ناموفق بود',{'document_id':document_id,'error':error,'attempts':attempts,'max_attempts':max_attempts},'warning' if retrying else 'danger')


@app.post('/api/v1/admin/documents')
async def upload_document(background_tasks: BackgroundTasks, file: UploadFile = File(...), visibility: str = Form('public'), admin: dict[str, Any] = Depends(require_roles('manager','admin'))):
    if visibility not in {'public','authenticated','internal'}:
        raise HTTPException(status_code=400,detail='سطح دسترسی نامعتبر است.')
    filename=Path(file.filename or 'resource.txt').name
    suffix=Path(filename).suffix.lower()
    if suffix not in DOCUMENT_SUFFIXES:
        raise HTTPException(status_code=400,detail='فرمت فایل پشتیبانی نمی‌شود.')
    temp_path=UPLOAD_SESSION_DIR/f'direct-{uuid.uuid4().hex}{suffix}'
    size_bytes=await _stream_upload_to_path(file,temp_path,MAX_UPLOAD_MB*1024*1024)
    try:
        document_id,job_id=_create_document_record(filename=filename,stored_path=temp_path,content_type=file.content_type,visibility=visibility,admin_id=admin['id'])
    except Exception:
        temp_path.unlink(missing_ok=True);raise
    audit(admin['id'],'document.uploaded',{'document_id':document_id,'job_id':job_id,'name':filename,'status':'processing','access_scope':visibility,'size_bytes':size_bytes,'mode':'stream'})
    if BACKGROUND_DOCUMENT_PROCESSING:
        return {'id':document_id,'job_id':job_id,'filename':filename,'visibility':visibility,'size_kb':round(size_bytes/1024,1),'status':'processing','progress':5}
    _process_document_job_guarded(job_id,document_id)
    with get_db() as db: job=dict(db.execute('SELECT * FROM document_jobs WHERE id=?',(job_id,)).fetchone())
    if job['status']=='error': raise HTTPException(status_code=400,detail=job['error'])
    with get_db() as db: doc=dict(db.execute('SELECT * FROM documents WHERE id=?',(document_id,)).fetchone())
    return {'id':document_id,'job_id':job_id,'filename':filename,'visibility':visibility,'chunks':doc['chunk_count'],'characters':doc['character_count'],'size_kb':round(size_bytes/1024,1),'status':'ready','progress':100}


@app.get('/api/v1/admin/upload-config')
def upload_config(admin: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    return {'max_upload_mb':MAX_UPLOAD_MB,'chunk_mb':UPLOAD_CHUNK_MB,'session_ttl_hours':UPLOAD_SESSION_TTL_HOURS,'google_doc_max_mb':GOOGLE_DOC_MAX_MB,'source_ingestion_version':INGESTION_VERSION,'pdf_vision_scan_all_pages':PDF_VISION_SCAN_ALL_PAGES,'pdf_vision_max_pages':PDF_VISION_MAX_PAGES,'pdf_vision_concurrency':PDF_VISION_CONCURRENCY,'source_min_quality_pct':SOURCE_MIN_QUALITY_PCT,'source_min_page_fidelity_pct':round(SOURCE_MIN_PAGE_FIDELITY*100,1),'source_min_numeric_agreement_pct':round(SOURCE_MIN_NUMERIC_AGREEMENT*100,1)}


@app.post('/api/v1/admin/upload-sessions')
def start_upload_session(data: UploadSessionStartInput, admin: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    _cleanup_expired_upload_sessions()
    filename=Path(data.filename).name
    suffix=Path(filename).suffix.lower()
    if suffix not in DOCUMENT_SUFFIXES:
        raise HTTPException(status_code=400,detail='فرمت فایل پشتیبانی نمی‌شود.')
    if data.visibility not in {'public','authenticated','internal'}:
        raise HTTPException(status_code=400,detail='سطح دسترسی نامعتبر است.')
    if data.size_bytes>MAX_UPLOAD_MB*1024*1024:
        raise HTTPException(status_code=413,detail=f'حجم فایل باید حداکثر {MAX_UPLOAD_MB} مگابایت باشد.')
    upload_id=uuid.uuid4().hex
    root,meta_path,part_path=_upload_session_paths(upload_id)
    root.mkdir(parents=True,exist_ok=False)
    meta={'upload_id':upload_id,'filename':filename,'suffix':suffix,'size_bytes':int(data.size_bytes),'received_bytes':0,'next_index':0,
          'chunk_bytes':UPLOAD_CHUNK_BYTES,'visibility':data.visibility,'content_type':data.content_type or 'application/octet-stream',
          'admin_id':admin['id'],'created_at':now_iso(),'updated_at':now_iso(),'status':'uploading','chunk_hashes':{}}
    _write_upload_meta(meta_path,meta);part_path.touch()
    return {'upload_id':upload_id,'chunk_bytes':UPLOAD_CHUNK_BYTES,'next_index':0,'received_bytes':0,'size_bytes':data.size_bytes,'status':'uploading'}


@app.get('/api/v1/admin/upload-sessions/{upload_id}')
def upload_session_status(upload_id: str, admin: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    meta,_,_,_=_read_upload_meta(upload_id,admin['id'])
    return {k:meta.get(k) for k in ('upload_id','filename','size_bytes','received_bytes','next_index','chunk_bytes','status','updated_at')}


@app.post('/api/v1/admin/upload-sessions/{upload_id}/chunks/{chunk_index}')
async def upload_session_chunk(upload_id: str, chunk_index: int, file: UploadFile=File(...), x_chunk_sha256: str | None=Header(default=None,alias='X-Chunk-SHA256'), admin: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    payload=await file.read(UPLOAD_CHUNK_BYTES+1024)
    if not payload: raise HTTPException(status_code=400,detail='قطعه آپلود خالی است.')
    if len(payload)>UPLOAD_CHUNK_BYTES:
        raise HTTPException(status_code=413,detail=f'هر قطعه باید حداکثر {UPLOAD_CHUNK_MB} مگابایت باشد.')
    actual_chunk_hash=hashlib.sha256(payload).hexdigest()
    if x_chunk_sha256 and not hmac.compare_digest(actual_chunk_hash,x_chunk_sha256.strip().lower()):
        raise HTTPException(status_code=422,detail='هش قطعه آپلود با فایل محلی مطابقت ندارد؛ قطعه دوباره ارسال شود.')
    lock=_UPLOAD_SESSION_LOCKS.setdefault(upload_id,threading.Lock())
    with lock:
        meta,_,meta_path,part_path=_read_upload_meta(upload_id,admin['id'])
        expected=int(meta.get('next_index') or 0)
        if chunk_index<expected:
            return {'ok':True,'duplicate':True,'next_index':expected,'received_bytes':int(meta.get('received_bytes') or 0)}
        if chunk_index!=expected:
            raise HTTPException(status_code=409,detail=f'ترتیب قطعه صحیح نیست؛ قطعه مورد انتظار {expected} است.')
        received=int(meta.get('received_bytes') or 0)
        if received+len(payload)>int(meta['size_bytes']):
            raise HTTPException(status_code=400,detail='حجم قطعات از حجم فایل اصلی بیشتر شد.')
        with part_path.open('ab') as out:
            out.write(payload);out.flush();os.fsync(out.fileno())
        meta['received_bytes']=received+len(payload);meta['next_index']=expected+1;meta['updated_at']=now_iso()
        hashes=meta.setdefault('chunk_hashes',{});hashes[str(chunk_index)]=actual_chunk_hash
        _write_upload_meta(meta_path,meta)
        return {'ok':True,'next_index':meta['next_index'],'received_bytes':meta['received_bytes'],'size_bytes':meta['size_bytes']}


@app.post('/api/v1/admin/upload-sessions/{upload_id}/complete')
def complete_upload_session(upload_id: str, background_tasks: BackgroundTasks, admin: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    lock=_UPLOAD_SESSION_LOCKS.setdefault(upload_id,threading.Lock())
    with lock:
        meta,root,meta_path,part_path=_read_upload_meta(upload_id,admin['id'])
        expected=int(meta['size_bytes']);actual=part_path.stat().st_size if part_path.exists() else 0
        if actual!=expected or int(meta.get('received_bytes') or 0)!=expected:
            raise HTTPException(status_code=409,detail=f'آپلود هنوز کامل نیست؛ {actual} از {expected} بایت دریافت شده است.')
        meta['status']='finalizing';meta['updated_at']=now_iso();_write_upload_meta(meta_path,meta)
        temp_final=root/f"complete{meta['suffix']}";part_path.replace(temp_final)
        try:
            document_id,job_id=_create_document_record(filename=meta['filename'],stored_path=temp_final,content_type=meta.get('content_type'),visibility=meta['visibility'],admin_id=admin['id'])
        except Exception:
            if temp_final.exists(): temp_final.replace(part_path)
            raise
        shutil.rmtree(root,ignore_errors=True)
    _UPLOAD_SESSION_LOCKS.pop(upload_id,None)
    audit(admin['id'],'document.uploaded',{'document_id':document_id,'job_id':job_id,'name':meta['filename'],'status':'processing','access_scope':meta['visibility'],'size_bytes':expected,'mode':'chunked'})
    if not BACKGROUND_DOCUMENT_PROCESSING: _process_document_job_guarded(job_id,document_id)
    return {'id':document_id,'job_id':job_id,'filename':meta['filename'],'visibility':meta['visibility'],'size_kb':round(expected/1024,1),'status':'processing' if BACKGROUND_DOCUMENT_PROCESSING else 'ready','progress':5 if BACKGROUND_DOCUMENT_PROCESSING else 100}


@app.post('/api/v1/admin/documents/google-doc')
async def import_google_doc(data: GoogleDocImportInput, background_tasks: BackgroundTasks, admin: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    if data.visibility not in {'public','authenticated','internal'}:
        raise HTTPException(status_code=400,detail='سطح دسترسی نامعتبر است.')
    doc_id=_extract_google_doc_id(data.url)
    export_url=f'https://docs.google.com/document/d/{doc_id}/export?format=docx'
    temp_path=UPLOAD_SESSION_DIR/f'gdoc-{uuid.uuid4().hex}.docx'
    client=_HTTP_CLIENT or httpx.AsyncClient(follow_redirects=True)
    owns_client=_HTTP_CLIENT is None
    total=0
    try:
        async with client.stream('GET',export_url,timeout=httpx.Timeout(120.0,connect=20.0)) as response:
            if response.status_code>=400:
                raise HTTPException(status_code=400,detail='Google Docs قابل دریافت نیست؛ دسترسی سند را روی «هرکس لینک را دارد» قرار دهید.')
            temp_path.parent.mkdir(parents=True,exist_ok=True)
            with temp_path.open('wb') as out:
                async for chunk in response.aiter_bytes(1024*1024):
                    total+=len(chunk)
                    if total>GOOGLE_DOC_MAX_MB*1024*1024:
                        raise HTTPException(status_code=413,detail=f'حجم Google Docs باید حداکثر {GOOGLE_DOC_MAX_MB} مگابایت باشد.')
                    out.write(chunk)
        with temp_path.open('rb') as probe:
            magic=probe.read(2)
        if total<4 or magic!=b'PK':
            raise HTTPException(status_code=400,detail='خروجی معتبر DOCX دریافت نشد؛ سند باید عمومی یا قابل مشاهده با لینک باشد.')
        filename=f'Google-Docs-{doc_id[:12]}.docx'
        document_id,job_id=_create_document_record(filename=filename,stored_path=temp_path,content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',visibility=data.visibility,admin_id=admin['id'])
        audit(admin['id'],'document.google_doc_imported',{'document_id':document_id,'job_id':job_id,'google_doc_id':doc_id,'size_bytes':total,'access_scope':data.visibility})
        if not BACKGROUND_DOCUMENT_PROCESSING: _process_document_job_guarded(job_id,document_id)
        return {'id':document_id,'job_id':job_id,'filename':filename,'visibility':data.visibility,'size_kb':round(total/1024,1),'status':'processing' if BACKGROUND_DOCUMENT_PROCESSING else 'ready','progress':5 if BACKGROUND_DOCUMENT_PROCESSING else 100}
    finally:
        if temp_path.exists(): temp_path.unlink(missing_ok=True)
        if owns_client: await client.aclose()


@app.get('/api/v1/admin/document-jobs/{job_id}')
def document_job_status(job_id: str, admin: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        row=db.execute("""SELECT j.*,d.filename,d.character_count,d.chunk_count,d.visibility
            FROM document_jobs j JOIN documents d ON d.id=j.document_id WHERE j.id=?""",(job_id,)).fetchone()
    if not row: raise HTTPException(status_code=404,detail='وظیفه پردازش پیدا نشد.')
    return dict(row)


@app.get('/api/v1/admin/builtin-sources/access')
def builtin_sources_access(user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    enabled=system_setting_bool('builtin_sources_enabled',BUILTIN_SOURCE_GLOBAL_DEFAULT)
    with get_db() as db:
        rows=db.execute("SELECT id,filename,source_key,is_enabled,status,page_count,vision_success_pages,vision_failed_pages,ingestion_quality_pct,page_fidelity_pct,numeric_agreement_pct,quality_gate_reason FROM documents WHERE is_builtin=1 ORDER BY source_key").fetchall()
    sources=[]
    for row in rows:
        item=dict(row);item['answer_eligible']=bool(item.get('is_enabled')) and item.get('status')=='ready'
        sources.append(item)
    return {'enabled':enabled,'priority_order':['training','sources','faq','cache'],'sources':sources}


class BuiltinSourceAccessInput(BaseModel):
    enabled: bool


@app.put('/api/v1/admin/builtin-sources/access')
def set_builtin_sources_access(data: BuiltinSourceAccessInput, user: dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    update_system_setting('builtin_sources_enabled','true' if data.enabled else 'false',user['id'])
    with get_db() as db:
        db.execute('UPDATE documents SET is_enabled=? WHERE is_builtin=1',(int(data.enabled),))
        bump_knowledge_version(db)
    audit(user['id'],'builtin_sources.access_changed',{'enabled':bool(data.enabled)})
    return {'ok':True,'enabled':bool(data.enabled),'message':'دسترسی چت به چهار جزوه پایه فعال شد.' if data.enabled else 'دسترسی چت به چهار جزوه پایه قطع شد؛ آموزش‌های مدیر همچنان فعال هستند.'}


@app.get('/api/v1/admin/documents')
def list_documents(user: dict[str, Any] = Depends(require_roles('admin','manager'))):
    with get_db() as db:
        rows = db.execute("SELECT id,filename,visibility,version,status,chunk_count,character_count,page_count,vision_candidate_pages,vision_success_pages,vision_failed_pages,ingestion_quality_pct,page_fidelity_pct,numeric_agreement_pct,quality_gate_reason,ingestion_warnings_json,last_indexed_at,ingestion_version,reindex_status,reindex_error,reindex_started_at,reindex_completed_at,created_at,is_builtin,is_enabled,source_key,round(character_count/1024.0,1) AS size_kb FROM documents ORDER BY created_at DESC").fetchall()
    output=[]
    for row in rows:
        item=dict(row);item['answer_eligible']=bool(item.get('is_enabled')) and item.get('status')=='ready'
        output.append(item)
    return output


@app.get('/api/v1/admin/sources/health')
def source_health(user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        rows=db.execute("SELECT id,filename,status,is_enabled,is_builtin,ingestion_quality_pct,page_fidelity_pct,numeric_agreement_pct,quality_gate_reason,last_indexed_at,reindex_status,reindex_error,vision_candidate_pages,vision_success_pages,vision_failed_pages FROM documents ORDER BY created_at DESC").fetchall()
    items=[]
    for row in rows:
        item=dict(row)
        item['answer_eligible']=bool(item.get('is_enabled')) and item.get('status')=='ready'
        items.append(item)
    total=len(items);eligible=sum(1 for x in items if x['answer_eligible'])
    blocked=[x for x in items if not x['answer_eligible']]
    avg_quality=round(sum(float(x.get('ingestion_quality_pct') or 0) for x in items)/max(1,total),1)
    vision_slots=[int(slot['slot']) for slot in configured_ai_slots() if _vision_model_for_slot(int(slot['slot']))]
    vision_needed=sum(1 for x in items if int(x.get('vision_candidate_pages') or 0)>int(x.get('vision_success_pages') or 0))
    return {
        'release':RELEASE_ID,'schema_revision':SCHEMA_REVISION,'total_sources':total,'answer_eligible':eligible,'blocked_sources':len(blocked),
        'average_quality_pct':avg_quality,'thresholds':{'quality_pct':SOURCE_MIN_QUALITY_PCT,'page_fidelity_pct':SOURCE_MIN_PAGE_FIDELITY*100,'numeric_agreement_pct':SOURCE_MIN_NUMERIC_AGREEMENT*100},
        'vision':{'configured':bool(vision_slots),'configured_slots':vision_slots,'sources_needing_vision':vision_needed,'ready':bool(vision_slots) or vision_needed==0},
        'blocked':blocked[:100],
    }


@app.get('/api/v1/admin/retrieval-debug')
def retrieval_debug(q: str, user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    question=sanitize_answer_text(q)[:2000]
    if len(question)<2:
        raise HTTPException(status_code=400,detail='سؤال برای بررسی بازیابی وارد نشده است.')
    training=_filter_navigation_only_items(retrieve_training(question,user,False))
    documents=_credible_document_items(question,_retrieve_document_chunks(question,user,integration=False))
    faq=find_faq_answer(question,user,False)
    selected_stage='training' if training else ('document' if documents else ('faq' if faq else 'none'))
    def item_view(item: dict[str,Any]) -> dict[str,Any]:
        return {
            'source_type':item.get('source_type'),'document_id':item.get('document_id'),'chunk_index':item.get('chunk_index'),'file_name':item.get('file_name'),
            'score':round(float(item.get('score') or 0),4),'rerank_score':round(float(item.get('rerank_score') or 0),4),
            'lexical_score':round(float(item.get('lexical_score') or 0),4),'semantic_score':round(float(item.get('semantic_score') or 0),4),
            'embedding_score':round(float(item.get('embedding_score') or 0),4),'fact_score':round(float(item.get('fact_score') or 0),4),
            'page_start':item.get('page_start'),'page_end':item.get('page_end'),'section_title':item.get('section_title'),
            'chunk_type':item.get('chunk_type'),'excerpt':sanitize_answer_text(str(item.get('excerpt') or item.get('content') or item.get('answer') or ''))[:900],
        }
    return {
        'question':question,'selected_stage':selected_stage,
        'priority_order':['training','document','faq','cache'],
        'training':[item_view(x) for x in training[:RETRIEVAL_DEBUG_MAX_ITEMS]],
        'documents':[item_view(x) for x in documents[:RETRIEVAL_DEBUG_MAX_ITEMS]],
        'faq':item_view(faq['source']) if faq else None,
        'knowledge_version':knowledge_version(),
    }


@app.get('/api/v1/admin/golden-cases')
def list_golden_cases(user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> list[dict[str,Any]]:
    with get_db() as db:
        rows=db.execute('SELECT * FROM golden_cases ORDER BY updated_at DESC').fetchall()
    return [dict(row) for row in rows]


@app.post('/api/v1/admin/golden-cases')
def create_golden_case(data: GoldenCaseInput, user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    case_id=str(uuid.uuid4());ts=now_iso()
    with get_db() as db:
        db.execute('INSERT INTO golden_cases(id,question,expected_answer,expected_source,is_active,created_by,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)',
                   (case_id,sanitize_answer_text(data.question),sanitize_answer_text(data.expected_answer),sanitize_answer_text(data.expected_source or '') or None,int(data.is_active),user['id'],ts,ts))
    audit(user['id'],'golden_case.created',{'case_id':case_id})
    return {'id':case_id,'ok':True}


@app.patch('/api/v1/admin/golden-cases/{case_id}')
def update_golden_case(case_id: str, data: GoldenCaseUpdateInput, user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        row=db.execute('SELECT * FROM golden_cases WHERE id=?',(case_id,)).fetchone()
        if not row: raise HTTPException(status_code=404,detail='سؤال آزمون پیدا نشد.')
        current=dict(row)
        db.execute('UPDATE golden_cases SET question=?,expected_answer=?,expected_source=?,is_active=?,updated_at=? WHERE id=?',(
            sanitize_answer_text(data.question) if data.question is not None else current['question'],
            sanitize_answer_text(data.expected_answer) if data.expected_answer is not None else current['expected_answer'],
            (sanitize_answer_text(data.expected_source) or None) if data.expected_source is not None else current['expected_source'],
            int(data.is_active) if data.is_active is not None else int(current['is_active']),now_iso(),case_id))
    return {'ok':True}


@app.delete('/api/v1/admin/golden-cases/{case_id}')
def delete_golden_case(case_id: str, user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        cur=db.execute('DELETE FROM golden_cases WHERE id=?',(case_id,))
    if cur.rowcount!=1: raise HTTPException(status_code=404,detail='سؤال آزمون پیدا نشد.')
    return {'ok':True}


def _golden_case_retrieval_result(case: dict[str,Any], user: dict[str,Any]) -> dict[str,Any]:
    question=str(case['question']);expected=str(case['expected_answer'])
    stage,items,deep_meta=retrieve_deep_priority_stage(question,user,False)
    if stage=='none':
        faq=find_faq_answer(question,user,False)
        if faq: stage='faq';items=[faq['source'] | {'content':faq['answer']}]
    evidence='\n'.join(str(x.get('content') or x.get('answer') or x.get('excerpt') or '') for x in items[:12])
    expected_tokens=_significant_tokens(expected);evidence_tokens=_significant_tokens(evidence)
    token_recall=(len(expected_tokens & evidence_tokens)/len(expected_tokens)) if expected_tokens else (1.0 if normalize_text(expected) in normalize_text(evidence) else 0.0)
    expected_numbers=_numeric_claims(expected);evidence_numbers=_numeric_claims(evidence)
    numeric_ok=expected_numbers.issubset(evidence_numbers)
    source_expected=normalize_text(str(case.get('expected_source') or ''))
    source_blob=normalize_text(' '.join(str(x.get('file_name') or '') for x in items))
    source_ok=(not source_expected) or source_expected in source_blob
    passed=bool(items) and token_recall>=GOLDEN_MIN_TOKEN_RECALL and numeric_ok and source_ok
    plan=analyze_query(question,max_subqueries=DEEP_MULTI_RETRIEVAL_MAX_QUERIES)
    rule_map=build_rule_exception_map(items) if items else {}
    confidence,parts=evidence_confidence(plan,items,rule_map,verification_status='deterministic_verified')
    return {'case_id':case['id'],'question':question,'stage':stage,'passed':passed,'token_recall':round(token_recall,4),'numeric_ok':numeric_ok,'source_ok':source_ok,
            'evidence_confidence':confidence,'confidence_parts':parts,'deep_meta':deep_meta,
            'top_sources':[{'file_name':x.get('file_name'),'page_start':x.get('page_start'),'score':round(float(x.get('score') or 0),4)} for x in items[:4]]}


@app.post('/api/v1/admin/golden-runs')
def run_golden_cases(user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        cases=[dict(row) for row in db.execute('SELECT * FROM golden_cases WHERE is_active=1 ORDER BY updated_at DESC LIMIT ?',(GOLDEN_MAX_CASES,)).fetchall()]
    details=[_golden_case_retrieval_result(case,user) for case in cases]
    total=len(details);passed=sum(1 for x in details if x['passed']);percent=round((passed/total*100) if total else 0.0,1)
    run_id=str(uuid.uuid4())
    with get_db() as db:
        db.execute('INSERT INTO golden_runs(id,total,passed,pass_percent,details_json,created_by,created_at) VALUES(?,?,?,?,?,?,?)',
                   (run_id,total,passed,percent,json.dumps(details,ensure_ascii=False),user['id'],now_iso()))
    audit(user['id'],'golden_run.completed',{'run_id':run_id,'total':total,'passed':passed,'pass_percent':percent})
    return {'id':run_id,'total':total,'passed':passed,'pass_percent':percent,'details':details}


@app.get('/api/v1/admin/golden-runs/latest')
def latest_golden_run(user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        row=db.execute('SELECT * FROM golden_runs ORDER BY created_at DESC LIMIT 1').fetchone()
    if not row:return {'total':0,'passed':0,'pass_percent':0,'details':[]}
    item=dict(row)
    try:item['details']=json.loads(item.pop('details_json') or '[]')
    except json.JSONDecodeError:item['details']=[]
    return item


@app.get('/api/v1/admin/documents/{document_id}/pages/{page_number}/preview')
def document_page_preview(document_id: str, page_number: int, user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> StreamingResponse:
    if fitz is None:
        raise HTTPException(status_code=503,detail='نمایش صفحه PDF در این سرور در دسترس نیست.')
    with get_db() as db:
        row=db.execute('SELECT filename,stored_path FROM documents WHERE id=?',(document_id,)).fetchone()
    if not row: raise HTTPException(status_code=404,detail='منبع پیدا نشد.')
    if Path(row['filename']).suffix.lower()!='.pdf': raise HTTPException(status_code=400,detail='پیش‌نمایش صفحه فقط برای PDF است.')
    path=Path(row['stored_path'])
    if not path.is_file(): raise HTTPException(status_code=404,detail='فایل اصلی منبع روی Volume موجود نیست.')
    pdf=fitz.open(str(path))
    try:
        if page_number<1 or page_number>len(pdf): raise HTTPException(status_code=404,detail='شماره صفحه معتبر نیست.')
        page=pdf[page_number-1];zoom=110/72.0
        pix=page.get_pixmap(matrix=fitz.Matrix(zoom,zoom),alpha=False)
        payload=pix.tobytes('png')
    finally:
        pdf.close()
    return StreamingResponse(io.BytesIO(payload),media_type='image/png',headers={'Cache-Control':'private, max-age=300'})


@app.get('/api/v1/admin/documents/{document_id}/diagnostics')
def document_diagnostics(document_id: str, user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        doc=db.execute("SELECT * FROM documents WHERE id=?",(document_id,)).fetchone()
        if not doc:
            raise HTTPException(status_code=404,detail='فایل پیدا نشد.')
        pages=db.execute("SELECT page_number,image_count,status,vision_error,character_count,page_fidelity,numeric_agreement FROM document_pages WHERE document_id=? ORDER BY page_number",(document_id,)).fetchall()
    item=dict(doc)
    try:
        item['warnings']=json.loads(item.get('ingestion_warnings_json') or '[]')
    except (TypeError,ValueError,json.JSONDecodeError):
        item['warnings']=[]
    item['answer_eligible']=bool(item.get('is_enabled',1)) and item.get('status')=='ready'
    item['pages']=[dict(row) for row in pages]
    item['failed_pages']=[dict(row) for row in pages if row['status'] in {'vision_error','vision_skipped','vision_unavailable','empty'}]
    item['vision_pages']=[dict(row) for row in pages if row['status']=='vision_ok']
    return item


def _retry_failed_pdf_pages(document_id: str) -> dict[str,Any]:
    if fitz is None:
        raise HTTPException(status_code=503,detail='PyMuPDF برای ترمیم صفحات PDF در دسترس نیست.')
    with get_db() as db:
        doc=db.execute('SELECT * FROM documents WHERE id=?',(document_id,)).fetchone()
        if not doc:
            raise HTTPException(status_code=404,detail='فایل پیدا نشد.')
        if Path(doc['filename']).suffix.lower()!='.pdf':
            raise HTTPException(status_code=400,detail='ترمیم صفحه‌ای فقط برای PDF قابل استفاده است.')
        page_rows=[dict(row) for row in db.execute('SELECT * FROM document_pages WHERE document_id=? ORDER BY page_number',(document_id,)).fetchall()]
    path=Path(doc['stored_path'])
    if not path.exists():
        raise HTTPException(status_code=404,detail='فایل ذخیره‌شده پیدا نشد.')
    failed=[row for row in page_rows if row.get('status') in {'vision_error','vision_skipped','vision_unavailable','empty'}]
    if not failed:
        return {'ok':True,'status':doc['status'],'retried':0,'remaining_failed':0,'quality_pct':float(doc['ingestion_quality_pct'] or 100),'chunks':int(doc['chunk_count'] or 0)}

    async def rerun() -> list[tuple[int,str,str|None]]:
        pdf=fitz.open(str(path));sem=asyncio.Semaphore(PDF_VISION_CONCURRENCY)
        async def one(row: dict[str,Any]) -> tuple[int,str,str|None]:
            n=int(row['page_number'])
            try:
                async with sem:
                    page=pdf[n-1]
                    zoom=PDF_VISION_DPI/72.0
                    pix=page.get_pixmap(matrix=fitz.Matrix(zoom,zoom),alpha=False)
                    image_bytes=pix.tobytes('png')
                    if len(image_bytes)>PDF_VISION_MAX_RENDER_MB*1024*1024:
                        pix=page.get_pixmap(matrix=fitz.Matrix(max(1.0,zoom*0.72),max(1.0,zoom*0.72)),alpha=False)
                        image_bytes=pix.tobytes('png')
                    text=await _vision_read_pdf_page(n,image_bytes,str(row.get('base_text') or ''),int(row.get('image_count') or 0))
                    if not text: raise RuntimeError('خروجی Vision خالی بود.')
                    return n,text,None
            except Exception as exc:
                return n,'',sanitize_answer_text(str(exc))[:300]
        try:
            return await asyncio.gather(*(one(row) for row in failed))
        finally:
            pdf.close()
    results=_run_coro_sync(rerun())
    by_number={int(row['page_number']):row for row in page_rows}
    for number,vision,error in results:
        row=by_number[number]
        if vision:
            row['vision_text']=vision
            row['combined_text']=_normalize_extracted_text('\n\n'.join(x for x in (str(row.get('base_text') or ''),'جزئیات استخراج‌شده از تصویر، جدول یا اسکن صفحه:\n'+vision) if x))
            row['status']='vision_ok';row['vision_error']=None
            row['vision_full']=not bool(PDF_VISION_DELTA_MODE and len(str(row.get('base_text') or ''))>=PDF_VISION_DELTA_TEXT_THRESHOLD and int(row.get('image_count') or 0)==0 and not _looks_table_like_page(str(row.get('base_text') or '')))
        else:
            row['status']='vision_error';row['vision_error']=error;row['vision_full']=False
        row['character_count']=len(str(row.get('combined_text') or ''))
        fidelity,agreement=_page_fidelity_metrics(str(row.get('base_text') or ''),str(row.get('vision_text') or ''),vision_full=bool(row.get('vision_full')),status=str(row.get('status') or ''))
        row['page_fidelity']=round(fidelity,4);row['numeric_agreement']=None if agreement is None else round(agreement,4)
    pages=[by_number[k] for k in sorted(by_number)]
    nonempty=sum(1 for r in pages if normalize_text(str(r.get('combined_text') or '')))
    meaningful=max(1,sum(1 for r in pages if str(r.get('base_text') or '').strip() or int(r.get('image_count') or 0)>0))
    candidate=max(int(doc['vision_candidate_pages'] or 0),sum(1 for r in pages if r.get('status') in {'vision_ok','vision_error','vision_skipped','vision_unavailable'}))
    success=sum(1 for r in pages if r.get('status')=='vision_ok')
    remaining=sum(1 for r in pages if r.get('status') in {'vision_error','vision_skipped','vision_unavailable','empty'})
    fidelity_scores=[float(r.get('page_fidelity') or 0) for r in pages if str(r.get('base_text') or '').strip() or str(r.get('vision_text') or '').strip() or int(r.get('image_count') or 0)>0]
    agreement_scores=[float(r['numeric_agreement']) for r in pages if r.get('numeric_agreement') is not None]
    average_fidelity=sum(fidelity_scores)/max(1,len(fidelity_scores))
    average_numeric_agreement=sum(agreement_scores)/len(agreement_scores) if agreement_scores else 1.0
    text_coverage=nonempty/max(1,meaningful);vision_coverage=success/max(1,candidate) if candidate else 1.0
    quality=round(max(0.0,min(100.0,(text_coverage*0.35+vision_coverage*0.20+average_fidelity*0.35+average_numeric_agreement*0.10)*100)),1)
    status='ready' if remaining==0 and quality>=SOURCE_MIN_QUALITY_PCT and average_fidelity>=SOURCE_MIN_PAGE_FIDELITY and average_numeric_agreement>=SOURCE_MIN_NUMERIC_AGREEMENT else 'partial'
    text='\n\n'.join(f"--- صفحه {r['page_number']} ---\n{r.get('combined_text') or ''}" for r in pages if str(r.get('combined_text') or '').strip())
    warnings=[f"Vision صفحه {r['page_number']} هنوز ناموفق است: {r.get('vision_error') or 'خطای نامشخص'}" for r in pages if r.get('status') in {'vision_error','vision_skipped','vision_unavailable'}]
    for r in pages:
        if float(r.get('page_fidelity') or 0)<SOURCE_MIN_PAGE_FIDELITY:
            warnings.append(f"صفحه {r['page_number']}: وفاداری استخراج {float(r.get('page_fidelity') or 0)*100:.0f}٪ است.")
        if r.get('numeric_agreement') is not None and float(r.get('numeric_agreement') or 0)<SOURCE_MIN_NUMERIC_AGREEMENT:
            warnings.append(f"صفحه {r['page_number']}: تطابق عدد/واحد {float(r.get('numeric_agreement') or 0)*100:.0f}٪ است.")
    result={'text':text,'kind':'pdf','pages':pages,'warnings':warnings,'status':status,'stats':{'page_count':len(pages),'nonempty_pages':nonempty,'vision_candidate_pages':candidate,'vision_success_pages':success,'vision_failed_pages':remaining,'ingestion_quality_pct':quality,'average_page_fidelity':round(average_fidelity,4),'average_numeric_agreement':round(average_numeric_agreement,4)}}
    chunks=chunk_source_result(result)
    with get_db() as db:
        fresh_doc=db.execute('SELECT * FROM documents WHERE id=?',(document_id,)).fetchone()
        active_ready=bool(fresh_doc and str(fresh_doc['status'] or '')=='ready' and int(fresh_doc['chunk_count'] or 0)>0)
        if active_ready and status!='ready':
            # Page repair is a staging operation too: never replace a known-good
            # active index with a partial repair result. Preserve the live index
            # and expose the repair failure separately for the manager.
            db.execute("UPDATE documents SET reindex_status='error',reindex_error=?,reindex_completed_at=? WHERE id=?",
                       (f'ترمیم صفحه‌ای کامل نشد؛ نسخه فعال قبلی حفظ شد. کیفیت {quality}٪، صفحات ناموفق {remaining}.',now_iso(),document_id))
            return {'ok':False,'status':'ready','staging_status':status,'preserved_active':True,'retried':len(failed),'remaining_failed':remaining,'quality_pct':quality,'chunks':int(fresh_doc['chunk_count'] or 0)}
        persist_document_ingestion(db,document_id,fresh_doc,result,chunks)
        db.execute("UPDATE documents SET version=version+1,reindex_status=?,reindex_error='',reindex_completed_at=? WHERE id=?",('idle' if status=='ready' else status,now_iso(),document_id))
    return {'ok':True,'status':status,'retried':len(failed),'remaining_failed':remaining,'quality_pct':quality,'chunks':len(chunks)}


@app.post('/api/v1/admin/documents/{document_id}/retry-failed-pages')
def retry_failed_document_pages(document_id: str, admin: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    result=_retry_failed_pdf_pages(document_id)
    audit(admin['id'],'document.failed_pages_retried',{'document_id':document_id,**result})
    return result


@app.get('/api/v1/admin/documents/{document_id}/pages/{page_number}')
def document_page_content(document_id: str, page_number: int, user: dict[str,Any]=Depends(require_roles('admin','manager'))) -> dict[str,Any]:
    with get_db() as db:
        row=db.execute('SELECT * FROM document_pages WHERE document_id=? AND page_number=?',(document_id,page_number)).fetchone()
    if not row:
        raise HTTPException(status_code=404,detail='صفحه استخراج‌شده پیدا نشد.')
    return dict(row)


@app.post('/api/v1/admin/documents/{document_id}/reindex')
def reindex_document(document_id: str, admin: dict[str, Any] = Depends(require_roles('manager','admin'))):
    with get_db() as db:
        doc=db.execute('SELECT id,stored_path FROM documents WHERE id=?',(document_id,)).fetchone()
    if not doc:
        raise HTTPException(status_code=404,detail='فایل پیدا نشد.')
    if not Path(doc['stored_path']).exists():
        raise HTTPException(status_code=404,detail='فایل ذخیره‌شده پیدا نشد.')
    job_id=_enqueue_reindex_job(document_id,admin['id'])
    audit(admin['id'],'document.reindex_queued',{'document_id':document_id,'job_id':job_id})
    return {'ok':True,'queued':True,'document_id':document_id,'job_id':job_id,'status':'queued'}


@app.post('/api/v1/admin/documents/reindex-all')
def reindex_all_documents(admin: dict[str, Any] = Depends(require_roles('manager','admin'))):
    results=[]
    with get_db() as db:
        documents=db.execute('SELECT id,filename,stored_path FROM documents ORDER BY created_at ASC').fetchall()
    for doc in documents:
        if not Path(doc['stored_path']).exists():
            with get_db() as db:
                db.execute("UPDATE documents SET status='missing' WHERE id=?",(doc['id'],))
            results.append({'id':doc['id'],'filename':doc['filename'],'ok':False,'error':'فایل ذخیره‌شده پیدا نشد.'});continue
        job_id=_enqueue_reindex_job(doc['id'],admin['id'],'در صف بازسازی همه منابع')
        results.append({'id':doc['id'],'filename':doc['filename'],'ok':True,'queued':True,'job_id':job_id})
    queued=sum(1 for x in results if x.get('ok'))
    failed=len(results)-queued
    audit(admin['id'],'documents.reindex_all_queued',{'total':len(results),'queued':queued,'failed':failed})
    return {'ok':failed==0,'total':len(results),'succeeded':queued,'queued':queued,'failed':failed,'results':results}


@app.delete('/api/v1/admin/documents/{document_id}')
def delete_document(document_id: str, admin: dict[str, Any] = Depends(require_roles('admin'))):
    with get_db() as db:
        row = db.execute("SELECT stored_path,is_builtin FROM documents WHERE id = ?", (document_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail='فایل پیدا نشد.')
        if int(row['is_builtin'] or 0):
            raise HTTPException(status_code=409,detail='جزوه پایه قابل حذف نیست؛ از کلید قطع/وصل دسترسی جزوه‌ها استفاده کنید.')
        db.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
        db.execute("DELETE FROM chunks_fts WHERE document_id = ?", (document_id,))
        db.execute("DELETE FROM chunk_semantic_buckets WHERE document_id = ?", (document_id,))
        db.execute("DELETE FROM chunk_semantic WHERE document_id = ?", (document_id,))
        db.execute("DELETE FROM documents WHERE id = ?", (document_id,))
        bump_knowledge_version(db)
    Path(row['stored_path']).unlink(missing_ok=True)
    audit(admin['id'], 'document.deleted', {'document_id': document_id})
    return {'ok': True}


@app.get('/api/v1/admin/audit-logs')
def audit_logs(owner: dict[str,Any]=Depends(require_owner)):
    with get_db() as db:
        rows=db.execute("SELECT a.action,COALESCE(u.username,u.email,'-') actor,a.details,a.created_at FROM audit_logs a LEFT JOIN users u ON u.id=a.user_id ORDER BY a.id DESC LIMIT 1000").fetchall()
    return [dict(r) for r in rows]


@app.get('/api/v1/manager/dashboard')
def manager_dashboard(user: dict[str, Any] = Depends(require_roles('manager','admin'))):
    seven_days_ago = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
    with get_db() as db:
        total_users = db.execute("SELECT COUNT(*) FROM users WHERE is_active = 1").fetchone()[0]
        questions_7d = db.execute("SELECT COUNT(*) FROM messages WHERE role='user' AND created_at >= ?", (seven_days_ago,)).fetchone()[0]
        unanswered_7d = db.execute("SELECT COUNT(*) FROM messages WHERE role='assistant' AND status IN ('knowledge_gap','unanswered','escalated') AND created_at >= ?", (seven_days_ago,)).fetchone()[0]
        ready_files = db.execute("SELECT COUNT(*) FROM documents WHERE status='ready'").fetchone()[0]
        escalations = db.execute("SELECT COUNT(*) FROM messages WHERE role='assistant' AND status='escalated' AND created_at >= ?", (seven_days_ago,)).fetchone()[0]
        useful_answers = db.execute("SELECT COUNT(*) FROM messages WHERE role='assistant' AND status IN ('faq_hit','cache_hit','training_answer','answered','self_analysis') AND created_at >= ?", (seven_days_ago,)).fetchone()[0]
        avg_ms = db.execute("SELECT COALESCE(AVG(response_ms), 0) FROM messages WHERE role='assistant' AND response_ms IS NOT NULL AND created_at >= ?", (seven_days_ago,)).fetchone()[0]
        turns_7d = db.execute("SELECT COUNT(*) FROM messages WHERE created_at >= ?", (seven_days_ago,)).fetchone()[0]
    return {
        'total_users': total_users,
        'questions_7d': questions_7d,
        'unanswered_7d': unanswered_7d,
        'ready_files': ready_files,
        'escalations': escalations,
        'useful_answers': useful_answers,
        'avg_response_ms': int(avg_ms or 0),
        'turns_7d': turns_7d,
    }


@app.get('/api/v1/manager/search-report')
def search_report(q:str|None=None,user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> list[dict[str,Any]]:
    return _question_rows(q,None,None,100,0)


@app.get('/api/v1/manager/report.csv')
def search_report_csv(user:dict[str,Any]=Depends(require_roles('manager','admin'))):
    rows=_question_rows(None,None,None,EXPORT_ROW_LIMIT,0);out=io.StringIO();writer=csv.writer(out)
    writer.writerow(['user','question','answer','status','response_section','api_slot','model_route','confidence_score','prompt_tokens','output_tokens','total_tokens','estimated_cost','response_ms','created_at'])
    for r in rows:writer.writerow([r['requester'],r['question'],r['answer'],r['status'],r.get('response_section'),r.get('api_slot'),r.get('model_route'),r.get('confidence_score'),r['prompt_tokens'],r['output_tokens'],r['total_tokens'],r.get('estimated_cost',0),r['response_ms'],r['asked_at']])
    return StreamingResponse(iter([out.getvalue()]),media_type='text/csv; charset=utf-8',headers={'Content-Disposition':'attachment; filename=search-report.csv'})




# =============================================================================
# R13 RAILWAY BUILD FIX
# =============================================================================

MANAGER_TRAINING_REQUIRES_APPROVAL = os.getenv(
    "MANAGER_TRAINING_REQUIRES_APPROVAL", "true"
).lower() in {"1", "true", "yes", "on"}
KNOWLEDGE_GAP_NOTIFY_THRESHOLD = int(os.getenv("KNOWLEDGE_GAP_NOTIFY_THRESHOLD", "2"))
TRAINING_RELEVANCE_MIN_SCORE = float(os.getenv("TRAINING_RELEVANCE_MIN_SCORE", "0.08"))
DIRECT_TRAINING_MIN_SCORE = float(os.getenv("DIRECT_TRAINING_MIN_SCORE", "0.72"))

_retrieve_document_chunks = retrieve


def _add_column_if_missing(db: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    if column not in column_names(db, table):
        db.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def _safe_username(raw: str | None, user_id: int) -> str:
    base = normalize_text(raw or '').replace(' ', '_')
    base = re.sub(r'[^a-z0-9_.-]', '', base.lower())
    if len(base) < 3:
        base = f'user{user_id}'
    return base[:42]


def _deduplicate_usernames(db: sqlite3.Connection) -> None:
    rows = db.execute("SELECT id,username,email FROM users ORDER BY id").fetchall()
    used: set[str] = set()
    for row in rows:
        raw = row['username']
        if not raw and row['email'] and '@' in row['email']:
            raw = row['email'].split('@',1)[0]
        base = _safe_username(raw, row['id'])
        candidate = base
        suffix = 2
        while candidate.lower() in used:
            tail = f'_{suffix}'
            candidate = base[:50-len(tail)] + tail
            suffix += 1
        used.add(candidate.lower())
        if row['username'] != candidate:
            db.execute("UPDATE users SET username=? WHERE id=?", (candidate,row['id']))
    db.execute("DROP INDEX IF EXISTS uq_users_username_lower")
    db.execute("CREATE UNIQUE INDEX uq_users_username_lower ON users(lower(username)) WHERE username IS NOT NULL")


# R29.7 — exact cargo vehicle catalogue requested by operations.
# Vehicle dimensions are configured only by manager/admin; operators never enter
# the vehicle envelope themselves.
CARGO_VEHICLES = {
    'peykan_flatbed': 'پیکان کفی دار',
    'peykan_no_flatbed': 'پیکان بدون کفی',
    'nissan_flatbed': 'نیسان کفی دار',
    'nissan_no_flatbed': 'نیسان بدون',
    'khavar_covered': 'خاور مسقف',
    'khavar_open': 'خاور رو باز',
}

# R35.2 — manager-taught cargo reference images. This catalogue intentionally
# follows the five choices requested for visual training and is independent
# from the geometric cargo profiles above.
CARGO_IMAGE_TRAINING_VEHICLES = {
    'nissan': 'نیسان',
    'peykan_no_flatbed': 'پیکان بدون کفی',
    'peykan_flatbed': 'پیکان کفی دار',
    'khavar_covered': 'خاور مسقف',
    'khavar_open': 'خاور روباز',
}
CARGO_IMAGE_TRAINING_ALIASES = {
    'nissan':'nissan','نیسان':'nissan','نیسانوانت':'nissan','نیسان وانت':'nissan',
    'peykan_no_flatbed':'peykan_no_flatbed','پیکانبدونکفی':'peykan_no_flatbed','پیکان بدون کفی':'peykan_no_flatbed',
    'peykan_flatbed':'peykan_flatbed','پیکانکفیدار':'peykan_flatbed','پیکان کفی دار':'peykan_flatbed',
    'khavar_covered':'khavar_covered','خاورمسقف':'khavar_covered','خاور مسقف':'khavar_covered',
    'khavar_open':'khavar_open','خاورروباز':'khavar_open','خاور رو باز':'khavar_open','خاور روباز':'khavar_open',
}
CARGO_VEHICLE_ALIASES = {
    'peykan_flatbed':'peykan_flatbed','پیکانکفیدار':'peykan_flatbed','پیکان کفی دار':'peykan_flatbed','پیکان کفی دار':'peykan_flatbed',
    'peykan_no_flatbed':'peykan_no_flatbed','پیکانبدونکفی':'peykan_no_flatbed','پیکان بدون کفی':'peykan_no_flatbed',
    'nissan_flatbed':'nissan_flatbed','نیسانکفیدار':'nissan_flatbed','نیسان کفی دار':'nissan_flatbed','نیسان کفی دار':'nissan_flatbed',
    'nissan_no_flatbed':'nissan_no_flatbed','نیسانبدونکفی':'nissan_no_flatbed','نیسان بدون':'nissan_no_flatbed','نیسان بدون':'nissan_no_flatbed',
    'khavar_covered':'khavar_covered','خاورمسقف':'khavar_covered','خاور مسقف':'khavar_covered','خاور بسته':'khavar_covered',
    'khavar_open':'khavar_open','خاورروباز':'khavar_open','خاور رو باز':'khavar_open','خاور رو باز':'khavar_open',
}
CALCULATION_VEHICLES = {'nissan':'نیسان','peykan':'پیکان','khavar':'خاور'}
CALCULATION_VEHICLE_ALIASES = {
    'nissan':'nissan','نیسان':'nissan',
    'peykan':'peykan','پیکان':'peykan','پیکانوانت':'peykan','پیکان وانت':'peykan',
    'khavar':'khavar','خاور':'khavar',
}


def _ensure_operational_schema(db: sqlite3.Connection) -> None:
    db.executescript("""
        CREATE TABLE IF NOT EXISTS cargo_vehicle_profiles (
            vehicle TEXT PRIMARY KEY,
            vehicle_label TEXT NOT NULL,
            length_cm REAL,
            width_cm REAL,
            height_cm REAL,
            max_weight_kg REAL,
            configured INTEGER NOT NULL DEFAULT 0,
            updated_by INTEGER,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(updated_by) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS cargo_image_training (
            id TEXT PRIMARY KEY,
            vehicle TEXT NOT NULL,
            vehicle_label TEXT NOT NULL,
            cargo_name TEXT NOT NULL,
            normalized_cargo_name TEXT NOT NULL,
            image_path TEXT NOT NULL,
            mime_type TEXT NOT NULL,
            original_filename TEXT NOT NULL,
            file_size_bytes INTEGER NOT NULL DEFAULT 0,
            created_by INTEGER,
            created_at TEXT NOT NULL,
            FOREIGN KEY(created_by) REFERENCES users(id)
        );
        CREATE INDEX IF NOT EXISTS idx_cargo_image_training_vehicle ON cargo_image_training(vehicle,created_at DESC);
        CREATE INDEX IF NOT EXISTS idx_cargo_image_training_name ON cargo_image_training(normalized_cargo_name);
        CREATE TABLE IF NOT EXISTS calculation_settings (
            vehicle TEXT PRIMARY KEY,
            vehicle_label TEXT NOT NULL,
            cancellation_base_toman INTEGER NOT NULL DEFAULT 0,
            waiting_hourly_toman INTEGER NOT NULL DEFAULT 0,
            deviation_per_km_toman INTEGER NOT NULL DEFAULT 0,
            free_wait_minutes INTEGER NOT NULL DEFAULT 0,
            extra_destination_free_minutes INTEGER NOT NULL DEFAULT 15,
            deviation_use_distance INTEGER NOT NULL DEFAULT 1,
            deviation_use_time INTEGER NOT NULL DEFAULT 0,
            updated_by INTEGER,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(updated_by) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS calculation_history (
            id TEXT PRIMARY KEY,
            user_id INTEGER NOT NULL,
            calculation_type TEXT NOT NULL CHECK(calculation_type IN ('cancellation','waiting','deviation')),
            vehicle TEXT NOT NULL,
            input_json TEXT NOT NULL,
            rate_snapshot_json TEXT NOT NULL,
            result_json TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        );
        CREATE INDEX IF NOT EXISTS idx_calculation_history_user_date ON calculation_history(user_id,created_at DESC);
        CREATE TABLE IF NOT EXISTS location_cache (
            query_key TEXT PRIMARY KEY,
            normalized_query TEXT NOT NULL,
            result_json TEXT NOT NULL,
            hit_count INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL,
            expires_at TEXT NOT NULL,
            last_used_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_location_cache_expires ON location_cache(expires_at);
    """)
    _add_column_if_missing(db,'calculation_settings','deviation_time_unit_minutes','INTEGER NOT NULL DEFAULT 0')
    _add_column_if_missing(db,'calculation_settings','deviation_time_unit_toman','INTEGER NOT NULL DEFAULT 0')
    db.execute("""UPDATE calculation_settings SET deviation_time_unit_minutes=60,deviation_time_unit_toman=waiting_hourly_toman
                  WHERE deviation_use_time=1 AND waiting_hourly_toman>0 AND deviation_time_unit_minutes=0 AND deviation_time_unit_toman=0""")
    # Keep the operational catalogue exact: obsolete cargo profiles from older
    # releases must not leak into manager/operator screens.
    allowed_cargo_keys=tuple(CARGO_VEHICLES.keys())
    placeholders=','.join('?' for _ in allowed_cargo_keys)
    db.execute(f"DELETE FROM cargo_vehicle_profiles WHERE vehicle NOT IN ({placeholders})",allowed_cargo_keys)
    ts=now_iso()
    for vehicle,label in CARGO_VEHICLES.items():
        db.execute(
            """INSERT OR IGNORE INTO cargo_vehicle_profiles(vehicle,vehicle_label,configured,updated_at)
               VALUES(?,?,0,?)""",
            (vehicle,label,ts),
        )
    # Optional deployment-provided profiles. These never overwrite a profile that a
    # manager has already configured in the database.
    try:
        configured=json.loads(CARGO_VEHICLE_PROFILES_JSON)
    except json.JSONDecodeError:
        configured={}
    if isinstance(configured,dict):
        for raw_key,profile in configured.items():
            key=_normalize_cargo_vehicle_key(str(raw_key),raise_error=False)
            if not key or not isinstance(profile,dict):
                continue
            try:
                l=float(profile.get('length_cm'));w=float(profile.get('width_cm'));h=float(profile.get('height_cm'))
                max_w=profile.get('max_weight_kg');max_w=float(max_w) if max_w not in (None,'') else None
            except (TypeError,ValueError):
                continue
            if min(l,w,h)<=0:
                continue
            db.execute(
                """UPDATE cargo_vehicle_profiles SET length_cm=?,width_cm=?,height_cm=?,max_weight_kg=?,configured=1,updated_at=?
                   WHERE vehicle=? AND configured=0""",
                (l,w,h,max_w,ts,key),
            )
    defaults=[('nissan','نیسان',60,15),('peykan','پیکان',50,15),('khavar','خاور',120,30)]
    for vehicle,label,free,extra in defaults:
        db.execute(
            """INSERT OR IGNORE INTO calculation_settings(
                   vehicle,vehicle_label,cancellation_base_toman,waiting_hourly_toman,deviation_per_km_toman,
                   free_wait_minutes,extra_destination_free_minutes,deviation_use_distance,deviation_use_time,updated_by,updated_at
               ) VALUES(?,?,0,0,0,?,?,1,0,NULL,?)""",
            (vehicle,label,free,extra,ts),
        )


def _normalize_cargo_vehicle_key(value: str, *, raise_error: bool = True) -> str | None:
    raw=normalize_text(value).replace(' ','')
    key=CARGO_VEHICLE_ALIASES.get(raw) or CARGO_VEHICLE_ALIASES.get(str(value).strip().lower())
    if not key and raise_error:
        raise HTTPException(status_code=400,detail='نوع خودرو باید یکی از گزینه‌های آموزش بررسی بار باشد: پیکان کفی دار، پیکان بدون کفی، نیسان کفی دار، نیسان بدون، خاور مسقف یا خاور رو باز.')
    return key


def _normalize_calculation_vehicle(value: str) -> str:
    raw=normalize_text(value).replace(' ','')
    key=CALCULATION_VEHICLE_ALIASES.get(raw) or CALCULATION_VEHICLE_ALIASES.get(str(value).strip().lower())
    if key not in CALCULATION_VEHICLES:
        raise HTTPException(status_code=400,detail='نوع خودرو باید نیسان، پیکان یا خاور باشد.')
    return key


def ensure_schema() -> None:
    # Preflight old databases before base indexes are created.
    with get_db() as db:
        tables = {r[0] for r in db.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
        if 'users' in tables:
            _add_column_if_missing(db,'users','username','TEXT')
            _add_column_if_missing(db,'users','is_active','INTEGER NOT NULL DEFAULT 1')
            _add_column_if_missing(db,'users','is_owner','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'users','created_by','INTEGER')
            _add_column_if_missing(db,'users','last_login_at','TEXT')
            _add_column_if_missing(db,'users','question_limit','INTEGER')
            _add_column_if_missing(db,'users','questions_used','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'users','daily_question_limit','INTEGER')
            _add_column_if_missing(db,'users','monthly_question_limit','INTEGER')
            _add_column_if_missing(db,'users','daily_questions_used','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'users','monthly_questions_used','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'users','daily_quota_date','TEXT')
            _add_column_if_missing(db,'users','monthly_quota_month','TEXT')
            _add_column_if_missing(db,'users','department','TEXT')
        if 'documents' in tables:
            _add_column_if_missing(db,'documents','status',"TEXT NOT NULL DEFAULT 'ready'")
            _add_column_if_missing(db,'documents','chunk_count','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'documents','version','INTEGER NOT NULL DEFAULT 1')
            _add_column_if_missing(db,'documents','page_count','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'documents','vision_candidate_pages','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'documents','vision_success_pages','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'documents','vision_failed_pages','INTEGER NOT NULL DEFAULT 0')
            _add_column_if_missing(db,'documents','ingestion_quality_pct','REAL NOT NULL DEFAULT 100')
            _add_column_if_missing(db,'documents','page_fidelity_pct','REAL NOT NULL DEFAULT 100')
            _add_column_if_missing(db,'documents','numeric_agreement_pct','REAL NOT NULL DEFAULT 100')
            _add_column_if_missing(db,'documents','quality_gate_reason',"TEXT NOT NULL DEFAULT ''")
            _add_column_if_missing(db,'documents','ingestion_warnings_json',"TEXT NOT NULL DEFAULT '[]'")
            _add_column_if_missing(db,'documents','last_indexed_at','TEXT')
            _add_column_if_missing(db,'documents','ingestion_version','INTEGER NOT NULL DEFAULT 1')
            _add_column_if_missing(db,'documents','content_sha256','TEXT')
            _add_column_if_missing(db,'documents','file_size_bytes','INTEGER NOT NULL DEFAULT 0')
        if 'chunks' in tables:
            _add_column_if_missing(db,'chunks','page_start','INTEGER')
            _add_column_if_missing(db,'chunks','page_end','INTEGER')
            _add_column_if_missing(db,'chunks','section_title','TEXT')
            _add_column_if_missing(db,'chunks','chunk_type',"TEXT NOT NULL DEFAULT 'text'")
            _add_column_if_missing(db,'chunks','search_aliases',"TEXT NOT NULL DEFAULT ''")
        if 'messages' in tables:
            _add_column_if_missing(db,'messages','response_ms','INTEGER')
            _add_column_if_missing(db,'messages','status',"TEXT NOT NULL DEFAULT 'answered'")
    _ensure_base_schema()
    with get_db() as db:
        _add_column_if_missing(db,'documents','is_builtin','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'documents','is_enabled','INTEGER NOT NULL DEFAULT 1')
        _add_column_if_missing(db,'documents','source_key','TEXT')
        _add_column_if_missing(db,'documents','page_fidelity_pct','REAL NOT NULL DEFAULT 100')
        _add_column_if_missing(db,'documents','numeric_agreement_pct','REAL NOT NULL DEFAULT 100')
        _add_column_if_missing(db,'documents','quality_gate_reason',"TEXT NOT NULL DEFAULT ''")
        db.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_documents_source_key ON documents(source_key) WHERE source_key IS NOT NULL")
        _add_column_if_missing(db,'users','is_owner','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'users','last_login_at','TEXT')
        _add_column_if_missing(db,'users','question_limit','INTEGER')
        _add_column_if_missing(db,'users','questions_used','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'users','daily_question_limit','INTEGER')
        _add_column_if_missing(db,'users','monthly_question_limit','INTEGER')
        _add_column_if_missing(db,'users','daily_questions_used','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'users','monthly_questions_used','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'users','daily_quota_date','TEXT')
        _add_column_if_missing(db,'users','monthly_quota_month','TEXT')
        _add_column_if_missing(db,'users','department','TEXT')
        _add_column_if_missing(db,'messages','prompt_tokens','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'messages','output_tokens','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'messages','total_tokens','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'messages','source_count','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'messages','api_slot','INTEGER')
        _add_column_if_missing(db,'messages','model_route','TEXT')
        _add_column_if_missing(db,'messages','estimated_cost','REAL NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'messages','confidence_score','REAL')

        db.executescript(
            """
            CREATE TABLE IF NOT EXISTS schema_migrations (
                revision TEXT PRIMARY KEY,
                release_id TEXT NOT NULL,
                applied_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS usage_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                external_user_id TEXT,
                event_type TEXT NOT NULL,
                prompt_tokens INTEGER NOT NULL DEFAULT 0,
                output_tokens INTEGER NOT NULL DEFAULT 0,
                total_tokens INTEGER NOT NULL DEFAULT 0,
                model TEXT,
                response_ms INTEGER,
                created_at TEXT NOT NULL,
                FOREIGN KEY(user_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS training_rules (
                id TEXT PRIMARY KEY,
                topic TEXT NOT NULL,
                topic_key TEXT NOT NULL,
                canonical_key TEXT,
                supersedes_id TEXT,
                instruction TEXT NOT NULL,
                answer TEXT NOT NULL,
                priority INTEGER NOT NULL DEFAULT 100,
                status TEXT NOT NULL DEFAULT 'pending' CHECK(status IN ('pending','active','rejected','superseded')),
                effective_from TEXT NOT NULL,
                expires_at TEXT,
                created_by INTEGER NOT NULL,
                approved_by INTEGER,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY(created_by) REFERENCES users(id),
                FOREIGN KEY(approved_by) REFERENCES users(id),
                FOREIGN KEY(supersedes_id) REFERENCES training_rules(id)
            );
            CREATE VIRTUAL TABLE IF NOT EXISTS training_fts USING fts5(
                topic,instruction,answer,training_id UNINDEXED,tokenize='unicode61'
            );
            CREATE TABLE IF NOT EXISTS knowledge_gaps (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                normalized_question TEXT NOT NULL UNIQUE,
                sample_question TEXT NOT NULL,
                occurrence_count INTEGER NOT NULL DEFAULT 1,
                status TEXT NOT NULL DEFAULT 'open' CHECK(status IN ('open','in_training','resolved','ignored')),
                assigned_training_id TEXT,
                first_seen TEXT NOT NULL,last_seen TEXT NOT NULL,last_user_id INTEGER,
                FOREIGN KEY(last_user_id) REFERENCES users(id),
                FOREIGN KEY(assigned_training_id) REFERENCES training_rules(id)
            );
            CREATE TABLE IF NOT EXISTS faqs (
                id TEXT PRIMARY KEY,
                question TEXT NOT NULL,
                normalized_question TEXT NOT NULL UNIQUE,
                answer TEXT NOT NULL,
                aliases_json TEXT NOT NULL DEFAULT '[]',
                priority INTEGER NOT NULL DEFAULT 100,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_by INTEGER NOT NULL,
                updated_by INTEGER NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY(created_by) REFERENCES users(id),
                FOREIGN KEY(updated_by) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS answer_cache (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                normalized_question TEXT NOT NULL,
                sample_question TEXT NOT NULL,
                detail_mode INTEGER NOT NULL DEFAULT 0,
                access_scope TEXT NOT NULL,
                knowledge_signature TEXT NOT NULL,
                answer TEXT NOT NULL,
                sources_json TEXT NOT NULL DEFAULT '[]',
                model TEXT,
                hit_count INTEGER NOT NULL DEFAULT 0,
                created_at TEXT NOT NULL,
                last_used_at TEXT NOT NULL,
                UNIQUE(normalized_question, detail_mode, access_scope, knowledge_signature)
            );
            CREATE TABLE IF NOT EXISTS system_settings (
                key TEXT PRIMARY KEY,value TEXT NOT NULL,updated_by INTEGER,updated_at TEXT NOT NULL,
                FOREIGN KEY(updated_by) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS admin_notifications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,kind TEXT NOT NULL,title TEXT NOT NULL,
                details TEXT NOT NULL DEFAULT '{}',severity TEXT NOT NULL DEFAULT 'info',
                is_read INTEGER NOT NULL DEFAULT 0,created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS admin_notification_reads (
                notification_id INTEGER NOT NULL,user_id INTEGER NOT NULL,read_at TEXT NOT NULL,
                PRIMARY KEY(notification_id,user_id),
                FOREIGN KEY(notification_id) REFERENCES admin_notifications(id) ON DELETE CASCADE,
                FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS ai_api_slot_state (
                slot INTEGER PRIMARY KEY,
                blocked_until TEXT,
                last_status TEXT,
                last_error TEXT,
                success_count INTEGER NOT NULL DEFAULT 0,
                failure_count INTEGER NOT NULL DEFAULT 0,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS ai_provider_configs (
                slot INTEGER PRIMARY KEY CHECK(slot BETWEEN 1 AND 20),
                label TEXT NOT NULL, base_url TEXT NOT NULL, api_key TEXT NOT NULL, model TEXT NOT NULL,
                model_economy TEXT, model_standard TEXT, model_advanced TEXT,
                vision_model TEXT, transcription_model TEXT, embedding_model TEXT,
                input_cost_per_1m REAL NOT NULL DEFAULT 0, output_cost_per_1m REAL NOT NULL DEFAULT 0,
                credit_amount REAL NOT NULL DEFAULT 0, credit_currency TEXT NOT NULL DEFAULT 'USD',
                enabled INTEGER NOT NULL DEFAULT 1, created_by INTEGER, created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
                FOREIGN KEY(created_by) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS request_metrics (
                id INTEGER PRIMARY KEY AUTOINCREMENT, request_id TEXT NOT NULL, method TEXT NOT NULL, path TEXT NOT NULL,
                status_code INTEGER NOT NULL, duration_ms INTEGER NOT NULL, created_at TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_request_metrics_created ON request_metrics(created_at);
            CREATE INDEX IF NOT EXISTS idx_request_metrics_path ON request_metrics(path,created_at);
            CREATE TABLE IF NOT EXISTS request_counters (
                identity_key TEXT NOT NULL, window_type TEXT NOT NULL, window_key TEXT NOT NULL,
                count INTEGER NOT NULL DEFAULT 0, updated_at TEXT NOT NULL,
                PRIMARY KEY(identity_key,window_type,window_key)
            );
            CREATE TABLE IF NOT EXISTS chunk_semantic (
                document_id TEXT NOT NULL, chunk_index INTEGER NOT NULL, terms_json TEXT NOT NULL,
                ngrams_json TEXT NOT NULL, updated_at TEXT NOT NULL,
                PRIMARY KEY(document_id,chunk_index),
                FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS chunk_semantic_buckets (
                bucket TEXT NOT NULL, document_id TEXT NOT NULL, chunk_index INTEGER NOT NULL,
                PRIMARY KEY(bucket,document_id,chunk_index),
                FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS chunk_embeddings (
                document_id TEXT NOT NULL, chunk_index INTEGER NOT NULL, model TEXT NOT NULL,
                dimensions INTEGER NOT NULL, vector_blob BLOB NOT NULL, updated_at TEXT NOT NULL,
                PRIMARY KEY(document_id,chunk_index,model),
                FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS knowledge_facts (
                id INTEGER PRIMARY KEY AUTOINCREMENT, document_id TEXT NOT NULL, chunk_index INTEGER NOT NULL,
                subject TEXT NOT NULL DEFAULT '', fact_type TEXT NOT NULL, value_text TEXT NOT NULL DEFAULT '',
                condition_text TEXT NOT NULL DEFAULT '', fact_text TEXT NOT NULL, page_start INTEGER, page_end INTEGER,
                section_title TEXT NOT NULL DEFAULT '', created_at TEXT NOT NULL,
                FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE
            );
            CREATE VIRTUAL TABLE IF NOT EXISTS knowledge_facts_fts USING fts5(
                subject,fact_text,condition_text,section_title,document_id UNINDEXED,chunk_index UNINDEXED,tokenize='unicode61'
            );
            CREATE TABLE IF NOT EXISTS golden_cases (
                id TEXT PRIMARY KEY, question TEXT NOT NULL, expected_answer TEXT NOT NULL,
                expected_source TEXT, is_active INTEGER NOT NULL DEFAULT 1, created_by INTEGER,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
                FOREIGN KEY(created_by) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS golden_runs (
                id TEXT PRIMARY KEY, total INTEGER NOT NULL, passed INTEGER NOT NULL, pass_percent REAL NOT NULL,
                details_json TEXT NOT NULL DEFAULT '[]', created_by INTEGER, created_at TEXT NOT NULL,
                FOREIGN KEY(created_by) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS answer_feedback (
                message_id TEXT NOT NULL, user_id INTEGER, rating TEXT NOT NULL
                    CHECK(rating IN ('correct','incomplete','wrong','wrong_source','more_detail')),
                note TEXT, created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
                PRIMARY KEY(message_id,user_id),
                FOREIGN KEY(message_id) REFERENCES messages(id) ON DELETE CASCADE,
                FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS health_checks (
                id INTEGER PRIMARY KEY AUTOINCREMENT, component TEXT NOT NULL, status TEXT NOT NULL,
                details TEXT NOT NULL DEFAULT '{}', checked_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS backup_history (
                id TEXT PRIMARY KEY, action TEXT NOT NULL, filename TEXT, size_bytes INTEGER NOT NULL DEFAULT 0,
                actor_id INTEGER, status TEXT NOT NULL, details TEXT NOT NULL DEFAULT '{}', created_at TEXT NOT NULL,
                FOREIGN KEY(actor_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS api_call_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT, slot INTEGER, provider_label TEXT, model TEXT,
                model_route TEXT, status TEXT NOT NULL, http_status INTEGER, error_class TEXT,
                error_detail TEXT, response_ms INTEGER NOT NULL DEFAULT 0, attempt_index INTEGER NOT NULL DEFAULT 1,
                was_failover INTEGER NOT NULL DEFAULT 0, created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS document_jobs (
                id TEXT PRIMARY KEY, document_id TEXT NOT NULL, status TEXT NOT NULL,
                progress INTEGER NOT NULL DEFAULT 0, phase TEXT, error TEXT,
                created_by INTEGER, created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
                FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE,
                FOREIGN KEY(created_by) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS document_pages (
                document_id TEXT NOT NULL,page_number INTEGER NOT NULL,base_text TEXT NOT NULL DEFAULT '',
                vision_text TEXT NOT NULL DEFAULT '',combined_text TEXT NOT NULL DEFAULT '',image_count INTEGER NOT NULL DEFAULT 0,
                status TEXT NOT NULL DEFAULT 'text',vision_error TEXT,character_count INTEGER NOT NULL DEFAULT 0,
                PRIMARY KEY(document_id,page_number),FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE
            );
            CREATE INDEX IF NOT EXISTS idx_document_pages_status ON document_pages(document_id,status,page_number);
            CREATE INDEX IF NOT EXISTS idx_document_jobs_updated ON document_jobs(updated_at DESC);
            CREATE INDEX IF NOT EXISTS idx_api_call_events_date ON api_call_events(created_at DESC);
            CREATE INDEX IF NOT EXISTS idx_api_call_events_slot ON api_call_events(slot,created_at DESC);
            CREATE TABLE IF NOT EXISTS login_attempts (
                identifier TEXT NOT NULL,ip_address TEXT NOT NULL,failed_count INTEGER NOT NULL DEFAULT 0,
                first_failed_at TEXT,last_failed_at TEXT,locked_until TEXT,
                PRIMARY KEY(identifier,ip_address)
            );
            CREATE INDEX IF NOT EXISTS idx_usage_user_date ON usage_events(user_id,created_at);
            CREATE INDEX IF NOT EXISTS idx_usage_external_date ON usage_events(external_user_id,created_at);
            CREATE INDEX IF NOT EXISTS idx_training_status_effective ON training_rules(status,effective_from,priority DESC);
            CREATE INDEX IF NOT EXISTS idx_gaps_status_count ON knowledge_gaps(status,occurrence_count DESC);
            CREATE INDEX IF NOT EXISTS idx_answer_cache_lookup ON answer_cache(access_scope,detail_mode,knowledge_signature,last_used_at DESC);
            CREATE INDEX IF NOT EXISTS idx_answer_cache_normalized ON answer_cache(normalized_question);
            CREATE INDEX IF NOT EXISTS idx_faq_active_priority ON faqs(is_active,priority DESC,updated_at DESC);
            CREATE INDEX IF NOT EXISTS idx_faq_normalized ON faqs(normalized_question);
            CREATE INDEX IF NOT EXISTS idx_notification_reads_user ON admin_notification_reads(user_id,read_at DESC);
            CREATE INDEX IF NOT EXISTS idx_request_counters_updated ON request_counters(updated_at);
            CREATE INDEX IF NOT EXISTS idx_semantic_bucket_lookup ON chunk_semantic_buckets(bucket);
            CREATE INDEX IF NOT EXISTS idx_chunk_embeddings_model ON chunk_embeddings(model,document_id,chunk_index);
            CREATE INDEX IF NOT EXISTS idx_knowledge_facts_doc ON knowledge_facts(document_id,chunk_index);
            CREATE INDEX IF NOT EXISTS idx_knowledge_facts_type ON knowledge_facts(fact_type,document_id);
            CREATE INDEX IF NOT EXISTS idx_golden_active ON golden_cases(is_active,updated_at DESC);
            CREATE INDEX IF NOT EXISTS idx_health_checks_date ON health_checks(checked_at DESC);
            CREATE INDEX IF NOT EXISTS idx_feedback_rating ON answer_feedback(rating,updated_at DESC);
            """
        )
        _add_column_if_missing(db,'training_rules','canonical_key','TEXT')
        _add_column_if_missing(db,'training_rules','supersedes_id','TEXT')
        _add_column_if_missing(db,'training_rules','visibility',"TEXT NOT NULL DEFAULT 'public'")
        _add_column_if_missing(db,'training_rules','allowed_roles_json',"TEXT NOT NULL DEFAULT '[]'")
        _add_column_if_missing(db,'training_rules','allowed_user_ids_json',"TEXT NOT NULL DEFAULT '[]'")
        _add_column_if_missing(db,'training_rules','department','TEXT')
        _add_column_if_missing(db,'faqs','visibility',"TEXT NOT NULL DEFAULT 'public'")
        _add_column_if_missing(db,'faqs','allowed_roles_json',"TEXT NOT NULL DEFAULT '[]'")
        _add_column_if_missing(db,'faqs','allowed_user_ids_json',"TEXT NOT NULL DEFAULT '[]'")
        _add_column_if_missing(db,'faqs','department','TEXT')
        _add_column_if_missing(db,'answer_cache','cache_tier',"TEXT NOT NULL DEFAULT 'approved'")
        _add_column_if_missing(db,'answer_cache','expires_at','TEXT')
        _add_column_if_missing(db,'answer_cache','approved_by','INTEGER')
        _add_column_if_missing(db,'usage_events','api_slot','INTEGER')
        _add_column_if_missing(db,'usage_events','provider_label','TEXT')
        _add_column_if_missing(db,'usage_events','model_route','TEXT')
        _add_column_if_missing(db,'usage_events','estimated_cost','REAL NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'documents','page_count','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'documents','vision_candidate_pages','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'documents','vision_success_pages','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'documents','vision_failed_pages','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'documents','ingestion_quality_pct','REAL NOT NULL DEFAULT 100')
        _add_column_if_missing(db,'documents','ingestion_warnings_json',"TEXT NOT NULL DEFAULT '[]'")
        _add_column_if_missing(db,'documents','last_indexed_at','TEXT')
        _add_column_if_missing(db,'documents','ingestion_version','INTEGER NOT NULL DEFAULT 1')
        _add_column_if_missing(db,'documents','content_sha256','TEXT')
        _add_column_if_missing(db,'documents','file_size_bytes','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'documents','reindex_status',"TEXT NOT NULL DEFAULT 'idle'")
        _add_column_if_missing(db,'documents','reindex_error',"TEXT NOT NULL DEFAULT ''")
        _add_column_if_missing(db,'documents','reindex_started_at','TEXT')
        _add_column_if_missing(db,'documents','reindex_completed_at','TEXT')
        _add_column_if_missing(db,'document_jobs','job_type',"TEXT NOT NULL DEFAULT 'ingest'")
        _add_column_if_missing(db,'document_jobs','payload_json',"TEXT NOT NULL DEFAULT '{}'")
        _add_column_if_missing(db,'document_jobs','attempts','INTEGER NOT NULL DEFAULT 0')
        _add_column_if_missing(db,'document_jobs','max_attempts',f'INTEGER NOT NULL DEFAULT {DOCUMENT_JOB_MAX_ATTEMPTS}')
        _add_column_if_missing(db,'document_jobs','next_run_at','TEXT')
        _add_column_if_missing(db,'document_jobs','worker_id','TEXT')
        _add_column_if_missing(db,'document_jobs','priority','INTEGER NOT NULL DEFAULT 50')
        _add_column_if_missing(db,'chunks','page_start','INTEGER')
        _add_column_if_missing(db,'chunks','page_end','INTEGER')
        _add_column_if_missing(db,'chunks','section_title','TEXT')
        _add_column_if_missing(db,'chunks','chunk_type',"TEXT NOT NULL DEFAULT 'text'")
        _add_column_if_missing(db,'chunks','search_aliases',"TEXT NOT NULL DEFAULT ''")
        _deduplicate_usernames(db)
        # Normalize ownership to exactly the configured master account when it exists.
        master = db.execute(
            "SELECT id FROM users WHERE lower(username)=lower(?) OR lower(email)=lower(?) ORDER BY CASE WHEN lower(email)=lower(?) THEN 0 ELSE 1 END,id LIMIT 1",
            (MASTER_ADMIN_USERNAME,MASTER_ADMIN_EMAIL,MASTER_ADMIN_EMAIL),
        ).fetchone()
        if master:
            db.execute("UPDATE users SET is_owner=0 WHERE id != ?", (master['id'],))
            db.execute("UPDATE users SET is_owner=1,role='admin',is_active=1 WHERE id=?", (master['id'],))
        db.execute("DROP INDEX IF EXISTS uq_single_owner")
        db.execute("CREATE UNIQUE INDEX uq_single_owner ON users(is_owner) WHERE is_owner=1")
        defaults={
            'chat_enabled':'true',
            'maintenance_message':'سرویس پاسخ‌گویی موقتاً توسط ادمین اصلی متوقف شده است.',
            'active_api_slot':'1',
            'health_last_state':'unknown',
            'knowledge_version':'1',
            'builtin_sources_enabled':'true' if BUILTIN_SOURCE_GLOBAL_DEFAULT else 'false',
        }
        for key,value in defaults.items():
            db.execute("INSERT OR IGNORE INTO system_settings(key,value,updated_at) VALUES (?,?,?)",(key,value,now_iso()))
        rows=db.execute("SELECT id,topic,instruction,canonical_key FROM training_rules").fetchall()
        for row in rows:
            key=row['canonical_key'] or canonical_training_key(row['topic'],row['instruction'])
            db.execute("UPDATE training_rules SET canonical_key=? WHERE id=?",(key,row['id']))
        # R22: all FAQ and active training entries have equal importance and no expiry.
        db.execute("UPDATE faqs SET priority=100, visibility='public', allowed_roles_json='[]', allowed_user_ids_json='[]', department=NULL")
        db.execute("UPDATE training_rules SET priority=100, expires_at=NULL, visibility='public', allowed_roles_json='[]', allowed_user_ids_json='[]', department=NULL WHERE status IN ('active','pending')")
        db.execute("UPDATE training_rules SET status='active', approved_by=COALESCE(approved_by,created_by), effective_from=COALESCE(created_at,?) WHERE status='pending'",(now_iso(),))
        db.execute("UPDATE training_rules SET priority=100, expires_at=NULL WHERE priority<>100 OR expires_at IS NOT NULL")
        db.execute("UPDATE faqs SET priority=100, visibility='public', allowed_roles_json='[]', allowed_user_ids_json='[]', department=NULL WHERE priority<>100 OR visibility<>'public' OR allowed_roles_json<>'[]' OR allowed_user_ids_json<>'[]' OR department IS NOT NULL")
        rebuild_training_fts(db)
        backfill_semantic_index(db)
        db.execute("INSERT OR IGNORE INTO schema_migrations(revision,release_id,applied_at) VALUES(?,?,?)",(SCHEMA_REVISION,RELEASE_ID,now_iso()))
        db.execute("UPDATE documents SET quality_gate_reason=CASE WHEN status='ready' THEN 'passed' ELSE 'legacy source requires reindex' END WHERE COALESCE(quality_gate_reason,'')='' ")
        _ensure_operational_schema(db)
        # Reset daily/monthly counters when the deployment starts in a new period.
        _refresh_all_quota_periods(db)


def system_setting(key: str, default: str = "") -> str:
    with get_db() as db:
        row = db.execute("SELECT value FROM system_settings WHERE key = ?", (key,)).fetchone()
    return row["value"] if row else default


def system_setting_bool(key: str, default: bool = False) -> bool:
    return system_setting(key, "true" if default else "false").lower() in {
        "1", "true", "yes", "on"
    }


def update_system_setting(key: str, value: str, admin_id: int) -> None:
    with get_db() as db:
        db.execute(
            """
            INSERT INTO system_settings(key, value, updated_by, updated_at)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(key) DO UPDATE SET
                value = excluded.value,
                updated_by = excluded.updated_by,
                updated_at = excluded.updated_at
            """,
            (key, value, admin_id, now_iso()),
        )

def bump_knowledge_version(db: sqlite3.Connection | None = None) -> int:
    """Invalidate answer/query caches in O(1) whenever authoritative knowledge changes."""
    owns = db is None
    conn = db or get_db()
    try:
        row = conn.execute("SELECT value FROM system_settings WHERE key='knowledge_version'").fetchone()
        try:
            current = int(row['value']) if row else 0
        except (TypeError, ValueError):
            current = 0
        new_value = current + 1
        conn.execute(
            """INSERT INTO system_settings(key,value,updated_at) VALUES('knowledge_version',?,?)
               ON CONFLICT(key) DO UPDATE SET value=excluded.value,updated_at=excluded.updated_at""",
            (str(new_value), now_iso()),
        )
        _QUERY_EMBED_CACHE.clear()
        return new_value
    finally:
        if owns:
            conn.commit(); conn.close()


def knowledge_version() -> int:
    try:
        return int(system_setting('knowledge_version','1'))
    except (TypeError, ValueError):
        return 1


class ApiSlotLimitError(Exception):
    def __init__(self, status_code: int, detail: str, retry_after_seconds: int | None = None):
        super().__init__(detail)
        self.status_code = status_code
        self.detail = detail
        self.retry_after_seconds = retry_after_seconds


def _parse_utc(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(value)
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
    except ValueError:
        return None


def _active_api_slot_number() -> int:
    try:
        value = int(system_setting('active_api_slot', '1'))
    except (TypeError, ValueError):
        value = 1
    return max(1, min(MAX_AI_API_SLOTS, value))


def ordered_available_ai_slots() -> list[dict[str, Any]]:
    slots = configured_ai_slots()
    if not slots:
        return []
    active = _active_api_slot_number()
    slots.sort(key=lambda item: ((int(item['slot'])-active) % MAX_AI_API_SLOTS, int(item['slot'])))
    now = datetime.now(timezone.utc)
    with get_db() as db:
        rows = {int(row['slot']): dict(row) for row in db.execute("SELECT * FROM ai_api_slot_state").fetchall()}
    available = []
    for slot in slots:
        state = rows.get(int(slot['slot']))
        blocked_until = _parse_utc(state.get('blocked_until') if state else None)
        if blocked_until and blocked_until > now:
            continue
        available.append(slot)
    return available


def ordered_available_ai_slots_for_route(route: str = 'standard') -> list[dict[str, Any]]:
    """Return healthy providers ordered by measured latency for this workload.

    The original active-slot order is preserved until enough successful samples
    exist, so a fresh deployment stays deterministic. Results are memoized briefly
    to keep provider selection effectively constant-time.
    """
    slots=ordered_available_ai_slots()
    if not PROVIDER_SPEED_ROUTING_ENABLED or len(slots)<2:
        return slots
    key=route if route in {'economy','standard','advanced','vision','voice','embedding'} else 'standard'
    now_mono=time.monotonic()
    with _PROVIDER_SPEED_CACHE_LOCK:
        cached=_PROVIDER_SPEED_CACHE.get(key)
    stats: dict[int,tuple[int,float]]
    if cached and now_mono-cached[0] < PROVIDER_SPEED_CACHE_SECONDS:
        stats=cached[1]
    else:
        cutoff=(datetime.now(timezone.utc)-timedelta(hours=24)).isoformat()
        with get_db() as db:
            rows=db.execute(
                """SELECT slot,COUNT(*) samples,AVG(response_ms) avg_ms
                   FROM api_call_events
                   WHERE status='success' AND slot IS NOT NULL AND slot>0 AND created_at>=? AND model_route=?
                   GROUP BY slot""",(cutoff,key),
            ).fetchall()
        stats={int(row['slot']):(int(row['samples'] or 0),float(row['avg_ms'] or 0.0)) for row in rows}
        with _PROVIDER_SPEED_CACHE_LOCK:
            _PROVIDER_SPEED_CACHE[key]=(now_mono,stats)
    original={int(slot['slot']):idx for idx,slot in enumerate(slots)}
    def sort_key(slot: dict[str,Any]) -> tuple[int,float,int]:
        samples,avg=stats.get(int(slot['slot']),(0,0.0))
        if samples>=PROVIDER_SPEED_MIN_SAMPLES and avg>0:
            return (0,avg,original[int(slot['slot'])])
        return (1,float(original[int(slot['slot'])]),original[int(slot['slot'])])
    return sorted(slots,key=sort_key)


def mark_api_slot_success(slot: int) -> None:
    ts = now_iso()
    with get_db() as db:
        db.execute(
            """INSERT INTO ai_api_slot_state(slot,blocked_until,last_status,last_error,success_count,failure_count,updated_at)
            VALUES (?,NULL,'active',NULL,1,0,?)
            ON CONFLICT(slot) DO UPDATE SET blocked_until=NULL,last_status='active',last_error=NULL,
                success_count=ai_api_slot_state.success_count+1,updated_at=excluded.updated_at""",
            (slot,ts),
        )
        db.execute(
            """INSERT INTO system_settings(key,value,updated_at) VALUES('active_api_slot',?,?)
            ON CONFLICT(key) DO UPDATE SET value=excluded.value,updated_at=excluded.updated_at""",
            (str(slot),ts),
        )


def mark_api_slot_blocked(slot: int, reason: str, seconds: int, status: str='limited') -> None:
    until = (datetime.now(timezone.utc)+timedelta(seconds=max(15,seconds))).isoformat()
    with get_db() as db:
        db.execute(
            """INSERT INTO ai_api_slot_state(slot,blocked_until,last_status,last_error,success_count,failure_count,updated_at)
            VALUES (?,?,?,?,0,1,?)
            ON CONFLICT(slot) DO UPDATE SET blocked_until=excluded.blocked_until,last_status=excluded.last_status,
                last_error=excluded.last_error,failure_count=ai_api_slot_state.failure_count+1,updated_at=excluded.updated_at""",
            (slot,until,status,sanitize_answer_text(reason)[:500],now_iso()),
        )


def record_api_call_event(
    *,
    slot: int | None,
    provider_label: str,
    model: str | None,
    model_route: str,
    status: str,
    response_ms: int,
    attempt_index: int = 1,
    was_failover: bool = False,
    http_status: int | None = None,
    error_class: str | None = None,
    error_detail: str | None = None,
) -> None:
    """Record each provider attempt without storing API keys or request content."""
    with get_db() as db:
        db.execute(
            """INSERT INTO api_call_events(
                slot,provider_label,model,model_route,status,http_status,error_class,error_detail,
                response_ms,attempt_index,was_failover,created_at
            ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                slot, provider_label[:120], (model or '')[:200], model_route[:40], status[:40],
                http_status, (error_class or '')[:80] or None, sanitize_answer_text(error_detail or '')[:500] or None,
                max(0,int(response_ms)), max(1,int(attempt_index)), int(bool(was_failover)), now_iso(),
            ),
        )


def api_call_health_stats(hours: int = 1) -> dict[str, Any]:
    hours=max(1,min(int(hours),168));cutoff=(datetime.now(timezone.utc)-timedelta(hours=hours)).isoformat()
    with get_db() as db:
        row=db.execute(
            """SELECT COUNT(*) attempts,
                      SUM(CASE WHEN status='success' THEN 1 ELSE 0 END) successes,
                      SUM(CASE WHEN status!='success' THEN 1 ELSE 0 END) failures,
                      SUM(CASE WHEN was_failover=1 AND status='success' THEN 1 ELSE 0 END) failover_successes,
                      CAST(COALESCE(AVG(CASE WHEN status='success' THEN response_ms END),0) AS INTEGER) avg_success_ms
               FROM api_call_events WHERE created_at>=?""",
            (cutoff,),
        ).fetchone()
        by_slot=[dict(x) for x in db.execute(
            """SELECT COALESCE(slot,0) slot,provider_label,COUNT(*) attempts,
                      SUM(CASE WHEN status='success' THEN 1 ELSE 0 END) successes,
                      SUM(CASE WHEN status!='success' THEN 1 ELSE 0 END) failures
               FROM api_call_events WHERE created_at>=? GROUP BY COALESCE(slot,0),provider_label ORDER BY slot""",
            (cutoff,),
        ).fetchall()]
    attempts=int(row['attempts'] or 0);failures=int(row['failures'] or 0)
    return {
        'hours':hours,'attempts':attempts,'successes':int(row['successes'] or 0),'failures':failures,
        'failure_rate_percent':round(failures/attempts*100,1) if attempts else 0.0,
        'failover_successes':int(row['failover_successes'] or 0),'avg_success_ms':int(row['avg_success_ms'] or 0),
        'by_slot':by_slot,
    }


def api_pool_public_status() -> dict[str, Any]:
    configured = {int(item['slot']): item for item in configured_ai_slots()}
    active = _active_api_slot_number()
    with get_db() as db:
        states = {int(row['slot']): dict(row) for row in db.execute("SELECT * FROM ai_api_slot_state").fetchall()}
    items=[]
    now=datetime.now(timezone.utc)
    display_slots=sorted(set(configured) | {int(k) for k in states if 1<=int(k)<=MAX_AI_API_SLOTS})
    for slot in display_slots:
        config=configured.get(slot);state=states.get(slot,{})
        blocked=_parse_utc(state.get('blocked_until'))
        items.append({
            'slot':slot,'configured':bool(config),'active':bool(config and slot==active),
            'provider':config.get('label') if config else None,'model':config.get('model') if config else None,
            'base_url':config.get('base_url') if config else None,'managed_by':config.get('managed_by') if config else None,
            'status':'blocked' if blocked and blocked>now else ('active' if config and slot==active else 'ready' if config else 'empty'),
            'blocked_until':state.get('blocked_until'),'last_status':state.get('last_status'),
            'last_error':state.get('last_error'),'success_count':int(state.get('success_count') or 0),
            'failure_count':int(state.get('failure_count') or 0),
        })
    return {'active_slot':active,'slots':items,'configured_count':len(configured),'max_slots':MAX_AI_API_SLOTS}


def notify_admin(kind: str, title: str, details: dict[str, Any] | None = None, severity: str = "info") -> None:
    with get_db() as db:
        db.execute(
            "INSERT INTO admin_notifications(kind, title, details, severity, created_at) VALUES (?, ?, ?, ?, ?)",
            (kind, title, json.dumps(details or {}, ensure_ascii=False), severity, now_iso()),
        )


def notify_admin_once(kind: str, title: str, details: dict[str, Any] | None = None, severity: str = 'info', dedupe_hours: int = 24) -> bool:
    cutoff=(datetime.now(timezone.utc)-timedelta(hours=max(1,dedupe_hours))).isoformat()
    signature=hashlib.sha256(json.dumps(details or {},ensure_ascii=False,sort_keys=True).encode('utf-8')).hexdigest()[:16]
    dedupe_title=f"{title} [{signature}]"
    with get_db() as db:
        exists=db.execute("SELECT 1 FROM admin_notifications WHERE kind=? AND title=? AND created_at>=? LIMIT 1",(kind,dedupe_title,cutoff)).fetchone()
        if exists:
            return False
        db.execute("INSERT INTO admin_notifications(kind,title,details,severity,created_at) VALUES(?,?,?,?,?)",(kind,dedupe_title,json.dumps(details or {},ensure_ascii=False),severity,now_iso()))
    return True


def _backup_health_state() -> dict[str, Any]:
    with get_db() as db:
        row=db.execute("SELECT created_at,filename,details FROM backup_history WHERE action IN ('backup','auto_backup') AND status='success' ORDER BY created_at DESC LIMIT 1").fetchone()
    last_at=row['created_at'] if row else None
    age_hours=None
    if last_at:
        parsed=_parse_utc(last_at)
        if parsed:
            age_hours=round((datetime.now(timezone.utc)-parsed).total_seconds()/3600,1)
    overdue=bool(AUTO_BACKUP_ENABLED and (age_hours is None or age_hours>BACKUP_OVERDUE_WARNING_HOURS))
    return {'enabled':AUTO_BACKUP_ENABLED,'last_success_at':last_at,'last_filename':row['filename'] if row else None,'age_hours':age_hours,'overdue':overdue}


def _health_snapshot() -> dict[str, Any]:
    checked=now_iso();db_ok=False;storage_ok=False;integrity='unknown';error=None
    try:
        with get_db() as db:
            db.execute('SELECT 1').fetchone()
            integrity=str(db.execute('PRAGMA quick_check').fetchone()[0])
        db_ok=integrity.lower()=='ok'
        probe=UPLOAD_DIR/'.health-probe'
        probe.write_text('ok',encoding='utf-8');probe.unlink(missing_ok=True)
        backup_probe=BACKUP_DIR/'.health-probe'
        backup_probe.write_text('ok',encoding='utf-8');backup_probe.unlink(missing_ok=True)
        storage_ok=True
    except Exception as exc:
        error=sanitize_answer_text(str(exc))[:500]
    usage=shutil.disk_usage(DB_PATH.parent)
    disk_percent=round((usage.used/usage.total*100) if usage.total else 0.0,1)
    pool=api_pool_public_status()
    configured=[item for item in pool['slots'] if item['configured']]
    available=[item for item in configured if item['status']!='blocked']
    cutoff=(datetime.now(timezone.utc)-timedelta(hours=1)).isoformat()
    with get_db() as db:
        recent=db.execute("SELECT COUNT(*) total,SUM(CASE WHEN event_type IN ('self_analysis','answered','training_answer') THEN 1 ELSE 0 END) ai_count FROM usage_events WHERE created_at>=?",(cutoff,)).fetchone()
        document_failures=int(db.execute("SELECT COUNT(*) FROM document_jobs WHERE status='error' AND updated_at>=?",(cutoff,)).fetchone()[0])
        pending_jobs=int(db.execute("SELECT COUNT(*) FROM document_jobs WHERE status IN ('queued','processing')").fetchone()[0])
        latest_golden=db.execute("SELECT total,passed,pass_percent,created_at FROM golden_runs ORDER BY created_at DESC LIMIT 1").fetchone()
    api_stats=api_call_health_stats(1)
    backup_state=_backup_health_state()
    status='healthy';reasons=[]
    if not db_ok or not storage_ok:
        status='critical';reasons.append('database_or_storage')
    if configured and not available:
        status='critical';reasons.append('all_api_slots_blocked')
    if api_stats['attempts']>=4 and api_stats['failure_rate_percent']>=HEALTH_ERROR_RATE_WARNING_PERCENT:
        status='warning' if status=='healthy' else status;reasons.append('api_error_rate_high')
    if document_failures>0:
        status='warning' if status=='healthy' else status;reasons.append('document_processing_failed')
    if GOLDEN_EVAL_ENABLED and latest_golden and float(latest_golden['pass_percent'] or 0)<GOLDEN_EVAL_MIN_PASS_PERCENT:
        status='warning' if status=='healthy' else status;reasons.append('golden_eval_below_target')
    if disk_percent>=HEALTH_DISK_WARNING_PERCENT:
        status='warning' if status=='healthy' else status;reasons.append('disk_usage_high')
    if backup_state['overdue']:
        status='warning' if status=='healthy' else status;reasons.append('backup_overdue')
    replica_hint=os.getenv('RAILWAY_REPLICA_ID') or os.getenv('RAILWAY_REPLICA_INDEX')
    return {
        'status':status,'checked_at':checked,'database':db_ok,'database_integrity':integrity,
        'storage':storage_ok,'disk_used_percent':disk_percent,'disk_free_bytes':usage.free,
        'api_configured':len(configured),'api_available':len(available),'api_pool':pool,
        'api_last_hour':api_stats,'recent_requests':int(recent['total'] or 0),
        'document_failures_last_hour':document_failures,'document_jobs_pending':pending_jobs,
        'golden_latest':dict(latest_golden) if latest_golden else None,'backup':backup_state,
        'sqlite_single_replica_mode':SQLITE_SINGLE_REPLICA_MODE,'replica_hint':replica_hint,
        'reasons':reasons,'error':error,
    }


def run_health_check(*, notify: bool=True) -> dict[str, Any]:
    snapshot=_health_snapshot()
    with get_db() as db:
        db.execute("INSERT INTO health_checks(component,status,details,checked_at) VALUES('system',?,?,?)",(snapshot['status'],json.dumps(snapshot,ensure_ascii=False),snapshot['checked_at']))
        cutoff=(datetime.now(timezone.utc)-timedelta(days=14)).isoformat()
        db.execute('DELETE FROM health_checks WHERE checked_at<?',(cutoff,))
        old=db.execute("SELECT value FROM system_settings WHERE key='health_last_state'").fetchone()
        old_state=old['value'] if old else 'unknown'
        db.execute("""INSERT INTO system_settings(key,value,updated_at) VALUES('health_last_state',?,?)
          ON CONFLICT(key) DO UPDATE SET value=excluded.value,updated_at=excluded.updated_at""",(snapshot['status'],snapshot['checked_at']))
    if notify and snapshot['status']!=old_state:
        severity='danger' if snapshot['status']=='critical' else 'warning' if snapshot['status']=='warning' else 'success'
        notify_admin('health_state',f"وضعیت سلامت سیستم: {snapshot['status']}",{'previous':old_state,'reasons':snapshot['reasons'],'disk_used_percent':snapshot['disk_used_percent'],'api_available':snapshot['api_available'],'api_failure_rate_percent':snapshot.get('api_last_hour',{}).get('failure_rate_percent'),'backup_overdue':snapshot.get('backup',{}).get('overdue')},severity)
    return snapshot


async def _health_monitor_loop() -> None:
    while True:
        if _RESTORE_IN_PROGRESS.is_set():
            await asyncio.sleep(min(5.0,HEALTH_MONITOR_INTERVAL_SECONDS))
            continue
        try:
            await asyncio.to_thread(maybe_create_automatic_backup)
        except asyncio.CancelledError:
            raise
        except Exception as exc:
            try:
                notify_admin_once('auto_backup_failed','پشتیبان‌گیری خودکار ناموفق بود',{'error':sanitize_answer_text(str(exc))[:300]},'danger',6)
            except Exception:
                pass
        try:
            await asyncio.to_thread(run_health_check,notify=True)
        except asyncio.CancelledError:
            raise
        except Exception as exc:
            try:
                notify_admin_once('health_monitor_error','خطا در اجرای پایش سلامت',{'error':sanitize_answer_text(str(exc))[:300]},'danger',6)
            except Exception:
                pass
        await asyncio.sleep(HEALTH_MONITOR_INTERVAL_SECONDS)


def _safe_zip_member(name: str) -> bool:
    path=Path(name)
    return bool(name) and not path.is_absolute() and '..' not in path.parts and not name.startswith(('/', '\\'))


def _create_backup_file_unlocked(
    actor_id: int | None = None,
    *,
    action: str = 'backup',
    persist_to_disk: bool = False,
) -> tuple[Path,str,str,int,bool]:
    """Create a ZIP on disk so large backups do not need to live in RAM.

    Returns (path, filename, backup_id, size_bytes, persistent). Caller owns the
    temporary path when persistent is False and should unlink it after streaming.
    """
    if action not in {'backup','auto_backup'}:
        raise ValueError('invalid backup action')
    backup_id=str(uuid.uuid4());stamp=datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
    filename=f'barsan_backup_{stamp}.zip'
    temp_root=Path(tempfile.gettempdir())/f'barsan-backup-file-{backup_id}';temp_root.mkdir(parents=True,exist_ok=True)
    temp_db=temp_root/'barsan.db'
    working_zip=temp_root/filename
    final_path: Path | None=None
    try:
        with get_db() as source, sqlite3.connect(temp_db,factory=ClosingSQLiteConnection) as dest:
            source.backup(dest)
        # The copied database may inherit WAL mode. Reopen in DELETE mode so the
        # sanitized state is checkpointed into the single barsan.db file before ZIP.
        with sqlite3.connect(temp_db,factory=ClosingSQLiteConnection) as clean:
            try:
                clean.execute('PRAGMA journal_mode=DELETE')
                clean.execute('DELETE FROM ai_provider_configs');clean.commit()
                clean.execute('VACUUM');clean.commit()
            except sqlite3.Error:
                clean.rollback();raise
        manifest={
            'format':'barsan-backup-v1','release':RELEASE_ID,'created_at':now_iso(),
            'database':'barsan.db','uploads_prefix':'uploads/','includes_secrets':False,
        }
        with zipfile.ZipFile(working_zip,'w',compression=zipfile.ZIP_DEFLATED,compresslevel=6,allowZip64=True) as archive:
            archive.write(temp_db,'barsan.db')
            archive.writestr('backup_manifest.json',json.dumps(manifest,ensure_ascii=False,indent=2))
            for path in sorted(UPLOAD_DIR.rglob('*')):
                if path.is_file():
                    rel=path.relative_to(UPLOAD_DIR).as_posix()
                    archive.write(path,f'uploads/{rel}')
        size_bytes=working_zip.stat().st_size
        if size_bytes>BACKUP_MAX_MB*1024*1024:
            raise HTTPException(status_code=413,detail=f'حجم فایل پشتیبان از سقف {BACKUP_MAX_MB} مگابایت بیشتر شد.')
        details={'release':RELEASE_ID,'automatic':action=='auto_backup','streamed':True}
        if persist_to_disk:
            BACKUP_DIR.mkdir(parents=True,exist_ok=True)
            final_path=BACKUP_DIR/filename
            shutil.move(str(working_zip),str(final_path));details['stored_path']=str(final_path)
        else:
            # Move outside temp_root so cleanup does not delete the response file.
            final_path=Path(tempfile.gettempdir())/f'barsan-download-{backup_id}.zip'
            shutil.move(str(working_zip),str(final_path))
        with get_db() as db:
            db.execute("INSERT INTO backup_history(id,action,filename,size_bytes,actor_id,status,details,created_at) VALUES(?,?,?,?,?,'success',?,?)",
                (backup_id,action,filename,size_bytes,actor_id,json.dumps(details,ensure_ascii=False),now_iso()))
        return final_path,filename,backup_id,size_bytes,persist_to_disk
    except Exception as exc:
        if final_path and final_path.exists(): final_path.unlink(missing_ok=True)
        try:
            with get_db() as db:
                db.execute("INSERT OR REPLACE INTO backup_history(id,action,filename,size_bytes,actor_id,status,details,created_at) VALUES(?,?,?,?,?,'failed',?,?)",
                    (backup_id,action,filename,0,actor_id,json.dumps({'error':sanitize_answer_text(str(exc))[:500]},ensure_ascii=False),now_iso()))
        except Exception:
            pass
        raise
    finally:
        shutil.rmtree(temp_root,ignore_errors=True)


def create_backup_file(
    actor_id: int | None = None,
    *,
    action: str = 'backup',
    persist_to_disk: bool = False,
) -> tuple[Path,str,str,int,bool]:
    with _BACKUP_LOCK:
        return _create_backup_file_unlocked(actor_id,action=action,persist_to_disk=persist_to_disk)

def _create_backup_bytes_unlocked(
    actor_id: int | None = None,
    *,
    action: str = 'backup',
    persist_to_disk: bool = False,
) -> tuple[bytes,str,str]:
    if action not in {'backup','auto_backup'}:
        raise ValueError('invalid backup action')
    backup_id=str(uuid.uuid4());stamp=datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
    filename=f'barsan_backup_{stamp}.zip'
    temp_root=Path(tempfile.gettempdir())/f'barsan-backup-{backup_id}';temp_root.mkdir(parents=True,exist_ok=True)
    temp_db=temp_root/'barsan.db'
    stored_path: Path | None=None
    try:
        with get_db() as source, sqlite3.connect(temp_db,factory=ClosingSQLiteConnection) as dest:
            source.backup(dest)
        # The copied database may inherit WAL mode. Reopen in DELETE mode so the
        # sanitized state is checkpointed into the single barsan.db file before ZIP.
        with sqlite3.connect(temp_db,factory=ClosingSQLiteConnection) as clean:
            try:
                clean.execute('PRAGMA journal_mode=DELETE')
                clean.execute('DELETE FROM ai_provider_configs');clean.commit()
                clean.execute('VACUUM');clean.commit()
            except sqlite3.Error:
                clean.rollback();raise
        manifest={
            'format':'barsan-backup-v1','release':RELEASE_ID,'created_at':now_iso(),
            'database':'barsan.db','uploads_prefix':'uploads/','includes_secrets':False,
        }
        buffer=io.BytesIO()
        with zipfile.ZipFile(buffer,'w',compression=zipfile.ZIP_DEFLATED,compresslevel=6) as archive:
            archive.write(temp_db,'barsan.db')
            archive.writestr('backup_manifest.json',json.dumps(manifest,ensure_ascii=False,indent=2))
            for path in sorted(UPLOAD_DIR.rglob('*')):
                if path.is_file():
                    rel=path.relative_to(UPLOAD_DIR).as_posix()
                    archive.write(path,f'uploads/{rel}')
        payload=buffer.getvalue()
        if len(payload)>BACKUP_MAX_MB*1024*1024:
            raise HTTPException(status_code=413,detail=f'حجم فایل پشتیبان از سقف {BACKUP_MAX_MB} مگابایت بیشتر شد.')
        details={'release':RELEASE_ID,'automatic':action=='auto_backup'}
        if persist_to_disk:
            BACKUP_DIR.mkdir(parents=True,exist_ok=True)
            stored_path=BACKUP_DIR/filename
            temp_path=stored_path.with_suffix('.zip.tmp')
            temp_path.write_bytes(payload);temp_path.replace(stored_path)
            details['stored_path']=str(stored_path)
        with get_db() as db:
            db.execute("INSERT INTO backup_history(id,action,filename,size_bytes,actor_id,status,details,created_at) VALUES(?,?,?,?,?,'success',?,?)",
                (backup_id,action,filename,len(payload),actor_id,json.dumps(details,ensure_ascii=False),now_iso()))
        return payload,filename,backup_id
    except Exception as exc:
        if stored_path:
            stored_path.unlink(missing_ok=True)
        try:
            with get_db() as db:
                db.execute("INSERT OR REPLACE INTO backup_history(id,action,filename,size_bytes,actor_id,status,details,created_at) VALUES(?,?,?,?,?,'failed',?,?)",
                    (backup_id,action,filename,0,actor_id,json.dumps({'error':sanitize_answer_text(str(exc))[:500]},ensure_ascii=False),now_iso()))
        except Exception:
            pass
        raise
    finally:
        shutil.rmtree(temp_root,ignore_errors=True)


def create_backup_bytes(
    actor_id: int | None = None,
    *,
    action: str = 'backup',
    persist_to_disk: bool = False,
) -> tuple[bytes,str,str]:
    with _BACKUP_LOCK:
        return _create_backup_bytes_unlocked(actor_id, action=action, persist_to_disk=persist_to_disk)


def _prune_automatic_backups() -> None:
    with get_db() as db:
        rows=db.execute("SELECT id,details FROM backup_history WHERE action='auto_backup' AND status='success' ORDER BY created_at DESC").fetchall()
        for row in rows[AUTO_BACKUP_RETENTION:]:
            try:
                details=json.loads(row['details'] or '{}')
            except json.JSONDecodeError:
                details={}
            path=Path(details.get('stored_path') or '')
            if path and path.parent==BACKUP_DIR:
                path.unlink(missing_ok=True)
            db.execute("DELETE FROM backup_history WHERE id=?",(row['id'],))


def maybe_create_automatic_backup() -> dict[str, Any] | None:
    if not AUTO_BACKUP_ENABLED:
        return None
    state=_backup_health_state()
    if state['age_hours'] is not None and state['age_hours']<AUTO_BACKUP_INTERVAL_HOURS:
        return None
    path,filename,backup_id,size_bytes,_=create_backup_file(None,action='auto_backup',persist_to_disk=True)
    _prune_automatic_backups()
    notify_admin('auto_backup','پشتیبان خودکار با موفقیت ایجاد شد',{'backup_id':backup_id,'filename':filename,'size_bytes':size_bytes},'success')
    return {'backup_id':backup_id,'filename':filename,'size_bytes':size_bytes,'stored_path':str(path)}


def _restore_backup_bytes_unlocked(payload: bytes, actor_id: int) -> dict[str,Any]:
    restore_id=str(uuid.uuid4())
    if len(payload)>BACKUP_MAX_MB*1024*1024:
        raise HTTPException(status_code=413,detail=f'حجم فایل بازیابی از سقف {BACKUP_MAX_MB} مگابایت بیشتر است.')
    temp_root=Path(tempfile.gettempdir())/f'barsan-restore-{restore_id}';temp_root.mkdir(parents=True,exist_ok=True)
    try:
        try:
            archive=zipfile.ZipFile(io.BytesIO(payload))
        except zipfile.BadZipFile as exc:
            raise HTTPException(status_code=400,detail='فایل ZIP پشتیبان معتبر نیست.') from exc
        infos=archive.infolist()
        if not infos or any(not _safe_zip_member(info.filename) for info in infos):
            raise HTTPException(status_code=400,detail='مسیر غیرمجاز داخل فایل پشتیبان شناسایی شد.')
        expanded=sum(max(0,info.file_size) for info in infos)
        if expanded>BACKUP_MAX_MB*1024*1024:
            raise HTTPException(status_code=413,detail='حجم بازشده فایل پشتیبان بیش از حد مجاز است.')
        names={info.filename for info in infos}
        if 'barsan.db' not in names or 'backup_manifest.json' not in names:
            raise HTTPException(status_code=400,detail='این فایل، پشتیبان معتبر بارسان نیست.')
        try:
            manifest=json.loads(archive.read('backup_manifest.json'))
        except Exception as exc:
            raise HTTPException(status_code=400,detail='Manifest فایل پشتیبان معتبر نیست.') from exc
        if manifest.get('format')!='barsan-backup-v1':
            raise HTTPException(status_code=400,detail='نسخه ساختار فایل پشتیبان پشتیبانی نمی‌شود.')
        candidate=temp_root/'candidate.db';candidate.write_bytes(archive.read('barsan.db'))
        try:
            with sqlite3.connect(candidate,factory=ClosingSQLiteConnection) as check:
                integrity=str(check.execute('PRAGMA integrity_check').fetchone()[0])
                tables={row[0] for row in check.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
        except sqlite3.DatabaseError as exc:
            raise HTTPException(status_code=400,detail='دیتابیس داخل پشتیبان خراب یا نامعتبر است.') from exc
        if integrity.lower()!='ok' or not {'users','documents','messages'}.issubset(tables):
            raise HTTPException(status_code=400,detail='اعتبارسنجی دیتابیس پشتیبان ناموفق بود.')
        restored_uploads=temp_root/'uploads';restored_uploads.mkdir()
        restored_count=0
        for info in infos:
            if info.filename.startswith('uploads/') and not info.is_dir():
                rel=Path(info.filename).relative_to('uploads')
                if not rel.parts:
                    continue
                target=(restored_uploads/rel).resolve()
                try:
                    target.relative_to(restored_uploads.resolve())
                except ValueError as exc:
                    raise HTTPException(status_code=400,detail='مسیر آپلود نامعتبر در فایل پشتیبان شناسایی شد.') from exc
                target.parent.mkdir(parents=True,exist_ok=True)
                target.write_bytes(archive.read(info.filename));restored_count+=1
        # SQLite backup API replaces the live database atomically at page level.
        with sqlite3.connect(candidate,factory=ClosingSQLiteConnection) as source, sqlite3.connect(DB_PATH,timeout=60,factory=ClosingSQLiteConnection) as destination:
            source.backup(destination)
        UPLOAD_DIR.mkdir(parents=True,exist_ok=True)
        for current in list(UPLOAD_DIR.iterdir()):
            if current.is_dir():
                shutil.rmtree(current,ignore_errors=True)
            else:
                current.unlink(missing_ok=True)
        for restored in restored_uploads.rglob('*'):
            if restored.is_file():
                rel=restored.relative_to(restored_uploads)
                destination=UPLOAD_DIR/rel;destination.parent.mkdir(parents=True,exist_ok=True)
                shutil.copy2(restored,destination)
        ensure_schema();seed_admin()
        with get_db() as db:
            actor_row=db.execute("SELECT id FROM users WHERE is_owner=1 ORDER BY id LIMIT 1").fetchone()
            actor_after=int(actor_row['id']) if actor_row else None
            db.execute("INSERT INTO backup_history(id,action,filename,size_bytes,actor_id,status,details,created_at) VALUES(?,?,?,?,?,'success',?,?)",
                (restore_id,'restore','uploaded-backup.zip',len(payload),actor_after,json.dumps({'source_release':manifest.get('release'),'integrity':integrity},ensure_ascii=False),now_iso()))
        return {'ok':True,'restore_id':restore_id,'source_release':manifest.get('release'),'uploads_restored':restored_count,'actor_id':actor_after}
    finally:
        shutil.rmtree(temp_root,ignore_errors=True)


def restore_backup_bytes(payload: bytes, actor_id: int) -> dict[str,Any]:
    if _RESTORE_IN_PROGRESS.is_set():
        raise HTTPException(status_code=409,detail='یک عملیات بازیابی دیگر در حال انجام است.')
    _RESTORE_IN_PROGRESS.set()
    try:
        # Stop new document work first, wait for an in-flight guarded job to finish,
        # then serialize against backup creation before replacing DB/uploads.
        with _MAINTENANCE_LOCK:
            with _BACKUP_LOCK:
                return _restore_backup_bytes_unlocked(payload, actor_id)
    finally:
        _RESTORE_IN_PROGRESS.clear()


def record_usage(
    *,
    user_id: int | None,
    external_user_id: str | None,
    event_type: str,
    usage: dict[str, Any],
    model: str,
    response_ms: int,
) -> None:
    """Persist actual provider usage plus route/API metadata for cost dashboards."""
    with get_db() as db:
        db.execute(
            """
            INSERT INTO usage_events(
                user_id, external_user_id, event_type, prompt_tokens,
                output_tokens, total_tokens, model, response_ms,
                api_slot, provider_label, model_route, estimated_cost, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                user_id,
                external_user_id,
                event_type,
                int(usage.get("prompt_tokens", 0)),
                int(usage.get("output_tokens", 0)),
                int(usage.get("total_tokens", 0)),
                model,
                response_ms,
                int(usage.get("api_slot")) if usage.get("api_slot") is not None else None,
                str(usage.get("provider_label") or ""),
                str(usage.get("model_route") or "zero_token"),
                float(usage.get("estimated_cost") or 0.0),
                now_iso(),
            ),
        )


def estimate_tokens(text: str) -> int:
    return max(1, math.ceil(len(text) / 4)) if text else 0


def rebuild_training_fts(db: sqlite3.Connection) -> None:
    db.execute("DELETE FROM training_fts")
    rows = db.execute(
        "SELECT id, topic, instruction, answer FROM training_rules WHERE status IN ('active','pending')"
    ).fetchall()
    db.executemany(
        "INSERT INTO training_fts(topic, instruction, answer, training_id) VALUES (?, ?, ?, ?)",
        [(r["topic"], r["instruction"], r["answer"], r["id"]) for r in rows],
    )


_TRAINING_STOPWORDS={'است','هست','شود','می‌شود','می شود','برای','این','آن','از','به','در','با','و','یا','جدید','قبلی','مبلغ','قیمت','مقدار','اعلام','تغییر','هزینه','تومان','هزار','میلیون','جایگزین','چقدر','چقدره','چنده','چیست','چیه','بگو','بگید','میشه','می‌شه','لطفا','لطفاً'}
_TRAINING_TOKEN_ALIASES={'لغو':'کنسلی','لغوی':'کنسلی','جریمه':'کنسلی','خوابار':'خاور','روبار':'روباری'}

_SECTION_NAMES = (
    'بررسی بار','مسیریابی','محاسبات','گزارش عملکرد','گزارش','نرخ یاب','نرخ‌یاب',
    'بخش بار','بخش مسیر','منوی بررسی','منوی مسیریابی'
)
_SECTION_REDIRECT_PATTERNS = (
    r'فقط\s+در\s+بخش', r'در\s+همان\s+بخش', r'وارد\s+بخش', r'برو(?:ید)?\s+(?:به|تو)\s+بخش',
    r'مراجعه\s+کن(?:ید)?', r'از\s+بخش.+استفاده\s+کن', r'این\s+درخواست.+بخش',
    r'مشخصات.+را\s+در\s+همان\s+بخش\s+وارد', r'این\s+قابلیت.+بخش'
)

def is_navigation_only_content(value: str) -> bool:
    """Detect legacy answers that only redirect users to another UI section."""
    if not CHAT_UNIFIED_KNOWLEDGE_MODE or not CHAT_DISABLE_SECTION_REDIRECTS:
        return False
    text=sanitize_answer_text(value)
    normalized=normalize_text(text)
    if not normalized or not any(normalize_text(name) in normalized for name in _SECTION_NAMES):
        return False
    if not any(re.search(pattern,normalized) for pattern in _SECTION_REDIRECT_PATTERNS):
        return False
    # Keep long factual material that happens to mention a UI section.
    factual=bool(re.search(r'\d|[۰-۹]|کیلو|تن|تومان|درصد|حداکثر|حداقل|مجاز|ممنوع|شرط|تبصره',text))
    return (not factual) or len(text) <= 420

def _filter_navigation_only_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [item for item in items if not is_navigation_only_content(str(item.get('answer') or item.get('content') or ''))]

def _significant_tokens(*texts: str) -> set[str]:
    tokens=set()
    for text in texts:
        for raw in normalize_text(text).split():
            token=re.sub(r'^[^0-9A-Za-zآ-ی]+|[^0-9A-Za-zآ-ی]+$','',raw)
            if len(token)>=2 and token not in _TRAINING_STOPWORDS and not token.isdigit():
                tokens.add(_TRAINING_TOKEN_ALIASES.get(token,token))
    return tokens


def canonical_training_key(topic: str, instruction: str='') -> str:
    tokens=sorted(_significant_tokens(topic) or _significant_tokens(instruction))
    return '|'.join(tokens[:16]) or normalize_text(topic or instruction)[:240]


def _token_overlap_score(question: str, *texts: str) -> float:
    return _retrieval_score(question, "\n".join(texts))


def _key_similarity(a: str,b: str) -> float:
    aa={x for x in a.split('|') if x};bb={x for x in b.split('|') if x}
    if not aa or not bb:return 0.0
    return len(aa&bb)/len(aa|bb)


def retrieve_training(question: str, user: dict[str, Any] | None = None, integration: bool = False) -> list[dict[str, Any]]:
    """Authoritative training retrieval.

    Training is the first knowledge layer. We use FTS for speed, then a bounded full
    rescue scan so colloquial Persian wording, typos, and paraphrases do not make an
    existing manager training silently fall through to documents.
    """
    now=now_iso();match_expr=build_fts_query(question)
    candidate_rows=[]
    with get_db() as db:
        try:
            candidate_rows=db.execute(
                """SELECT tr.*,bm25(training_fts) AS fts_rank
                   FROM training_fts JOIN training_rules tr ON tr.id=training_fts.training_id
                   WHERE training_fts MATCH ? AND tr.status='active' AND tr.effective_from<=?
                     AND (tr.expires_at IS NULL OR tr.expires_at>?)
                   ORDER BY fts_rank ASC,tr.updated_at DESC LIMIT 320""",
                (match_expr,now,now),
            ).fetchall()
        except sqlite3.OperationalError:
            candidate_rows=[]

        # Always keep a rescue pool available. This is intentionally bounded and
        # ordered by recency so training remains reliable without making chat slow.
        rescue_rows=db.execute(
            """SELECT * FROM training_rules WHERE status='active' AND effective_from<=?
               AND (expires_at IS NULL OR expires_at>?)
               ORDER BY updated_at DESC LIMIT ?""",
            (now,now,TRAINING_RESCUE_SCAN_LIMIT),
        ).fetchall()

    seen=set();results=[]
    def consider(row: sqlite3.Row) -> None:
        rid=str(row['id'])
        if rid in seen:return
        seen.add(rid)
        if not content_accessible(row,user,integration): return
        if is_navigation_only_content(row['answer']) or is_navigation_only_content(row['instruction']): return
        intent_text=f"{row['topic']}\n{row['instruction']}"
        answer_text=str(row['answer'] or '')
        intent_relevance=_retrieval_score(question,intent_text,row['topic'])
        answer_relevance=_retrieval_score(question,answer_text,row['topic'])
        q_key=canonical_training_key(question,'')
        row_key=row['canonical_key'] or canonical_training_key(row['topic'],row['instruction'])
        key_similarity=_key_similarity(q_key,row_key)
        # Intent/topic must actually resemble the question. This prevents a generic
        # word inside an answer (for example «تمدید») from hijacking an unrelated
        # training entry, while still allowing short colloquial questions.
        credible_intent=(intent_relevance>=0.30 or key_similarity>=0.34 or (intent_relevance>=0.22 and answer_relevance>=0.68))
        if not credible_intent:return
        relevance=min(1.5,intent_relevance*0.68+answer_relevance*0.32+0.10*key_similarity)
        if relevance<TRAINING_RELEVANCE_MIN_SCORE:return
        results.append({'source_type':'training','training_id':rid,'document_id':f"training:{rid}",
            'file_name':f"آموزش مدیریتی — {row['topic']}",'chunk_index':0,'content':row['answer'],'answer':row['answer'],
            'priority':100,'score':round(relevance,4),'excerpt':row['answer'][:500],'effective_from':row['effective_from']})

    for row in candidate_rows:consider(row)
    # Rescue scan runs when FTS did not find a sufficiently strong training hit.
    if not results or max(float(x.get('score') or 0) for x in results)<TRAINING_STAGE_MIN_SCORE:
        for row in rescue_rows:consider(row)
    results.sort(key=lambda item:item['score'],reverse=True)
    return results[:max(8,min(20,RETRIEVAL_TOP_K))]

def normalize_faq_aliases(values: list[str] | str | None) -> list[str]:
    if values is None:
        return []
    raw = values if isinstance(values,list) else re.split(r'[|؛;\n]+',str(values))
    result=[];seen=set()
    for value in raw:
        clean=sanitize_answer_text(str(value)).strip()
        key=canonical_question_for_cache(clean)
        if clean and key and key not in seen:
            result.append(clean[:1000]);seen.add(key)
    return result[:50]


def _faq_match_score(question: str, candidate: str) -> float:
    q=canonical_question_for_cache(question);c=canonical_question_for_cache(candidate)
    if not q or not c:return 0.0
    if q==c:return 1.0
    qa={x for x in search_tokens(q) if x not in _PERSIAN_STOPWORDS};ca={x for x in search_tokens(c) if x not in _PERSIAN_STOPWORDS}
    jac=len(qa&ca)/len(qa|ca) if qa and ca else 0.0
    containment=len(qa&ca)/len(qa) if qa else 0.0
    seq=SequenceMatcher(None,q,c).ratio()
    score=min(1.0,0.42*jac+0.33*containment+0.25*seq)
    # A short FAQ alias fully contained in the operator's wording is a strong match.
    if len(ca)>=2 and ca.issubset(qa):score=max(score,0.93)
    if len(qa)>=2 and qa.issubset(ca):score=max(score,0.90)
    return score


def find_faq_answer(question: str, user: dict[str, Any] | None = None, integration: bool = False) -> dict[str, Any] | None:
    normalized=canonical_question_for_cache(question)
    if len(normalized)<2:return None
    with get_db() as db:
        exact=db.execute("SELECT * FROM faqs WHERE is_active=1 AND normalized_question=? ORDER BY updated_at DESC LIMIT 1",(normalized,)).fetchone()
        if exact and content_accessible(exact,user,integration):
            row=exact;score=1.0
        else:
            tokens=[x for x in search_tokens(question) if len(x)>=2 and x not in _PERSIAN_STOPWORDS][:4]
            if tokens:
                clauses=' OR '.join('(question LIKE ? OR aliases_json LIKE ?)' for _ in tokens)
                params=[value for token in tokens for value in (f'%{token}%',f'%{token}%')]
                rows=db.execute(f"SELECT * FROM faqs WHERE is_active=1 AND ({clauses}) ORDER BY updated_at DESC LIMIT 500",params).fetchall()
            else:
                rows=[]
            if not rows:
                rows=db.execute("SELECT * FROM faqs WHERE is_active=1 ORDER BY updated_at DESC LIMIT 500").fetchall()
            row=None;score=0.0
            for candidate in rows:
                if not content_accessible(candidate,user,integration):
                    continue
                try:aliases=json.loads(candidate['aliases_json'] or '[]')
                except json.JSONDecodeError:aliases=[]
                candidate_score=max([_faq_match_score(question,candidate['question'])]+[_faq_match_score(question,a) for a in aliases])
                if candidate_score>score:
                    row=candidate;score=candidate_score
        if row is None or score<FAQ_FUZZY_THRESHOLD:return None
    answer_text=sanitize_answer_text(row['answer'])
    if is_navigation_only_content(answer_text):
        return None
    return {
        'faq_id':row['id'],'question':row['question'],'answer':answer_text,
        'score':round(score,4),'priority':int(row['priority']),
        'source':{'source_type':'faq','document_id':f"faq:{row['id']}",'file_name':'سؤالات متداول','chunk_index':0,'score':round(score,4),'excerpt':answer_text[:500]},
    }


def _faq_row_value(row: dict[str, Any], *keys: str) -> Any:
    normalized={normalize_text(str(k)).replace(' ',''):v for k,v in row.items() if k is not None}
    for key in keys:
        token=normalize_text(key).replace(' ','')
        if token in normalized:return normalized[token]
    return None


def parse_faq_import(filename: str, data: bytes) -> list[dict[str, Any]]:
    ext=Path(filename).suffix.lower();rows=[]
    if ext=='.xlsx':
        wb=load_workbook(io.BytesIO(data),read_only=True,data_only=True);ws=wb.active
        values=list(ws.iter_rows(values_only=True))
        if not values:return []
        headers=[str(x or '').strip() for x in values[0]]
        rows=[dict(zip(headers,line)) for line in values[1:]]
    elif ext=='.csv':
        text=data.decode('utf-8-sig',errors='replace');rows=list(csv.DictReader(io.StringIO(text)))
    elif ext=='.json':
        parsed=json.loads(data.decode('utf-8-sig',errors='strict'))
        if isinstance(parsed,dict):parsed=parsed.get('faqs') or parsed.get('items') or []
        if not isinstance(parsed,list):raise HTTPException(status_code=400,detail='ساختار JSON باید آرایه‌ای از سؤال و پاسخ باشد.')
        rows=[item for item in parsed if isinstance(item,dict)]
    elif ext in {'.txt','.md'}:
        text=data.decode('utf-8-sig',errors='replace');rows=[]
        for line in text.splitlines():
            line=line.strip()
            if not line:continue
            parts=re.split(r'\t|\s*\|\s*',line,maxsplit=2)
            if len(parts)>=2:rows.append({'question':parts[0],'answer':parts[1],'aliases':parts[2] if len(parts)>2 else ''})
    else:
        raise HTTPException(status_code=400,detail='برای سؤالات متداول فقط XLSX، CSV، JSON، TXT یا MD مجاز است.')
    result=[]
    for row in rows:
        q=_faq_row_value(row,'question','سوال','سؤال','پرسش')
        a=_faq_row_value(row,'answer','جواب','پاسخ')
        if not q or not a:continue
        aliases=_faq_row_value(row,'aliases','alias','کلیدواژه','عبارت های مشابه','سوالات مشابه') or ''
        active_raw=_faq_row_value(row,'active','is_active','فعال')
        is_active=str(active_raw).strip().lower() not in {'0','false','no','off','غیرفعال'}
        result.append({'question':sanitize_answer_text(str(q)),'answer':sanitize_answer_text(str(a)),'aliases':normalize_faq_aliases(aliases),'is_active':is_active})
    return result


def upsert_faq_rows(items: list[dict[str, Any]], actor_id: int) -> dict[str, int]:
    created=updated=skipped=0;ts=now_iso()
    with get_db() as db:
        for item in items[:FAQ_MAX_ROWS]:
            question=sanitize_answer_text(item.get('question',''));answer=sanitize_answer_text(item.get('answer',''))
            normalized=canonical_question_for_cache(question)
            if len(normalized)<2 or not answer:skipped+=1;continue
            existing=db.execute("SELECT id FROM faqs WHERE normalized_question=?",(normalized,)).fetchone()
            aliases=json.dumps(normalize_faq_aliases(item.get('aliases')),ensure_ascii=False)
            if existing:
                db.execute("""UPDATE faqs SET question=?,answer=?,aliases_json=?,priority=?,is_active=?,visibility=?,allowed_roles_json=?,allowed_user_ids_json=?,department=?,updated_by=?,updated_at=? WHERE id=?""",
                    (question,answer,aliases,100,int(bool(item.get('is_active',True))),'public',json.dumps([],ensure_ascii=False),json.dumps([]),None,actor_id,ts,existing['id']));updated+=1
            else:
                db.execute("""INSERT INTO faqs(id,question,normalized_question,answer,aliases_json,priority,is_active,visibility,allowed_roles_json,allowed_user_ids_json,department,created_by,updated_by,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (str(uuid.uuid4()),question,normalized,answer,aliases,100,int(bool(item.get('is_active',True))),'public',json.dumps([],ensure_ascii=False),json.dumps([]),None,actor_id,actor_id,ts,ts));created+=1
        if created or updated:
            bump_knowledge_version(db)
    return {'created':created,'updated':updated,'skipped':skipped,'total':created+updated}


def retrieve_combined(question: str, user: dict[str, Any] | None, integration: bool = False) -> list[dict[str, Any]]:
    training = retrieve_training(question,user,integration)
    documents = _retrieve_document_chunks(question, user, integration=integration)
    for item in documents:
        item.setdefault('source_type', 'document')
        item.setdefault('priority', 0)
    return (training + documents[:RETRIEVAL_TOP_K])[:RETRIEVAL_TOTAL_ITEMS]


def _numeric_claims(text: str) -> set[str]:
    normalized=normalize_text(text)
    return set(re.findall(r"\b\d+(?:[.,]\d+)?(?:\s*(?:درصد|تومان|ریال|روز|ماه|سال|ساعت|دقیقه))?\b",normalized))


def detect_source_conflict(training: list[dict[str, Any]], documents: list[dict[str, Any]]) -> bool:
    """Conservative conflict signal: strong training/document matches with incompatible numeric claims."""
    if not training or not documents:
        return False
    top_t=training[0];top_d=documents[0]
    if float(top_t.get('score') or 0)<TRAINING_STAGE_MIN_SCORE or float(top_d.get('score') or 0)<RETRIEVAL_MEDIUM_CONFIDENCE:
        return False
    t_claims=_numeric_claims(str(top_t.get('answer') or top_t.get('content') or ''))
    d_claims=_numeric_claims(str(top_d.get('content') or ''))
    return bool(t_claims and d_claims and t_claims.isdisjoint(d_claims))


_CONTRADICTORY_INTENT_PAIRS=(
    ({'تمدید','فعال'}, {'لغو','کنسلی','سلب','غیرفعال'}),
    ({'افزایش','بیشتر'}, {'کاهش','کمتر'}),
    ({'ورود','ثبت'}, {'حذف','پاک'}),
    ({'بارگیری'}, {'تخلیه'}),
)

def _has_contradictory_intent(question: str, content: str) -> bool:
    q=_significant_tokens(question);c=_significant_tokens(content)
    for left,right in _CONTRADICTORY_INTENT_PAIRS:
        if (q & left and c & right) or (q & right and c & left):
            return True
    return False

def _credible_document_items(question: str, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    q_tokens=_significant_tokens(question)
    result=[]
    for item in items:
        content=f"{item.get('file_name','')}\n{item.get('content','')}"
        if _has_contradictory_intent(question,content):
            continue
        c_tokens=_significant_tokens(content)
        shared=len(q_tokens & c_tokens)
        score=float(item.get('score') or 0)
        # Two shared meaningful concepts is strong evidence. For very short questions,
        # one concept plus a medium/high retrieval score is allowed. Very high hybrid
        # similarity can also pass when wording differs substantially.
        credible=(shared>=2 or score>=0.72 or (len(q_tokens)<=2 and shared>=1 and score>=RETRIEVAL_MEDIUM_CONFIDENCE))
        if credible:
            result.append(item)
    return result


def _retrieval_stage_cache_key(question: str, user: dict[str,Any] | None, integration: bool) -> str:
    scope=cache_access_scope(user,integration)
    normalized=canonical_question_for_cache(question)
    return f"{knowledge_version()}|{scope}|{normalized}"


def _retrieval_stage_cache_get(question: str, user: dict[str,Any] | None, integration: bool) -> tuple[str,list[dict[str,Any]]] | None:
    key=_retrieval_stage_cache_key(question,user,integration)
    now_mono=time.monotonic()
    with _RETRIEVAL_STAGE_CACHE_LOCK:
        cached=_RETRIEVAL_STAGE_CACHE.get(key)
        if not cached:
            return None
        if now_mono-cached[0] > RETRIEVAL_STAGE_CACHE_TTL_SECONDS:
            _RETRIEVAL_STAGE_CACHE.pop(key,None);return None
        stage=cached[1];items=[dict(item) for item in cached[2]]
    return stage,items


def _retrieval_stage_cache_put(question: str, user: dict[str,Any] | None, integration: bool, stage: str, items: list[dict[str,Any]]) -> None:
    key=_retrieval_stage_cache_key(question,user,integration)
    with _RETRIEVAL_STAGE_CACHE_LOCK:
        if len(_RETRIEVAL_STAGE_CACHE)>=RETRIEVAL_STAGE_CACHE_MAX_ENTRIES:
            oldest=min(_RETRIEVAL_STAGE_CACHE.items(),key=lambda item:item[1][0])[0]
            _RETRIEVAL_STAGE_CACHE.pop(oldest,None)
        _RETRIEVAL_STAGE_CACHE[key]=(time.monotonic(),stage,[dict(item) for item in items])


def retrieve_priority_stage(question: str, user: dict[str, Any] | None, integration: bool = False) -> tuple[str, list[dict[str, Any]]]:
    """Strict authoritative order: manager training first, documents second.

    A relevant training hit is never mixed with documents, because the user's own
    training must remain the authoritative answer. Documents are searched only when
    no sufficiently relevant training exists.
    """
    memo=_retrieval_stage_cache_get(question,user,integration)
    if memo is not None:
        return memo
    training=_filter_navigation_only_items(retrieve_training(question,user,integration))
    # retrieve_training already applies intent/relevance guards. Any surviving
    # manager training is authoritative and must outrank every organizational source.
    if training:
        result=training[:RETRIEVAL_TOTAL_ITEMS];_retrieval_stage_cache_put(question,user,integration,'training',result)
        return 'training',result

    documents=_retrieve_document_chunks(question,user,integration=integration)
    for item in documents:
        item.setdefault('source_type','document');item.setdefault('priority',0)
    documents=_credible_document_items(question,documents)
    if documents:
        best=max(float(item.get('score') or 0) for item in documents)
        if best >= max(RETRIEVAL_MIN_SCORE, RETRIEVAL_MEDIUM_CONFIDENCE*0.55):
            result=documents[:RETRIEVAL_TOTAL_ITEMS];_retrieval_stage_cache_put(question,user,integration,'document',result)
            return 'document',result
    _retrieval_stage_cache_put(question,user,integration,'none',[])
    return 'none',[]


def retrieve_deep_priority_stage(question: str, user: dict[str, Any] | None, integration: bool = False) -> tuple[str, list[dict[str, Any]], dict[str, Any]]:
    """R34 query planning + multi-retrieval + semantic rerank.

    Manager training remains the authoritative layer. If a decomposed query exposes a
    relevant training rule that the raw wording missed, it wins only after the deep
    reranker confirms meaningful alignment with the original question.
    """
    plan=analyze_query(question,max_subqueries=DEEP_MULTI_RETRIEVAL_MAX_QUERIES)
    if not DEEP_QUERY_ANALYSIS_ENABLED or not DEEP_MULTI_RETRIEVAL_ENABLED:
        stage,items=retrieve_priority_stage(question,user,integration=integration)
        return stage,items,{'query_plan':plan.as_dict(),'deep_retrieval':False}

    queries=plan.subqueries[:DEEP_MULTI_RETRIEVAL_MAX_QUERIES]
    training_sets=[]
    for q in queries:
        rows=_filter_navigation_only_items(retrieve_training(q,user,integration))
        if rows:
            training_sets.append((q,rows))
    if training_sets:
        merged=merge_multiretrieval(plan,training_sets,top_n=RETRIEVAL_TOTAL_ITEMS*2)
        if DEEP_SEMANTIC_RERANK_ENABLED:
            merged=semantic_rerank_candidates(plan,merged,top_n=RETRIEVAL_TOTAL_ITEMS)
        else:
            merged=merged[:RETRIEVAL_TOTAL_ITEMS]
        best=max([float(x.get('deep_semantic_score') or x.get('score') or 0) for x in merged] or [0.0])
        if best>=max(0.20,TRAINING_STAGE_MIN_SCORE*0.85):
            return 'training',merged,{'query_plan':plan.as_dict(),'deep_retrieval':True,'retrieval_queries':queries,'candidate_count':sum(len(rows) for _,rows in training_sets)}

    document_sets=[]
    for q in queries:
        rows=_retrieve_document_chunks(q,user,integration=integration)
        for item in rows:
            item.setdefault('source_type','document');item.setdefault('priority',0)
        if rows:
            document_sets.append((q,rows))
    if document_sets:
        merged=merge_multiretrieval(plan,document_sets,top_n=max(RETRIEVAL_TOTAL_ITEMS*3,HYBRID_RERANK_TOP_N))
        if DEEP_SEMANTIC_RERANK_ENABLED:
            merged=semantic_rerank_candidates(plan,merged,top_n=RETRIEVAL_TOTAL_ITEMS)
        else:
            merged=merged[:RETRIEVAL_TOTAL_ITEMS]
        best=max([float(x.get('deep_semantic_score') or x.get('score') or 0) for x in merged] or [0.0])
        if best>=max(RETRIEVAL_MIN_SCORE,RETRIEVAL_MEDIUM_CONFIDENCE*0.50):
            return 'document',merged,{'query_plan':plan.as_dict(),'deep_retrieval':True,'retrieval_queries':queries,'candidate_count':sum(len(rows) for _,rows in document_sets)}

    return 'none',[],{'query_plan':plan.as_dict(),'deep_retrieval':True,'retrieval_queries':queries,'candidate_count':0}


async def retrieve_priority_stage_async(question: str, user: dict[str, Any] | None, integration: bool = False) -> tuple[str, list[dict[str, Any]]]:
    """Non-blocking facade for request handlers.

    SQLite ranking, local semantic scans and any provider embedding fallback execute
    in a worker thread, so an expensive retrieval cannot stall the FastAPI event loop.
    """
    return await asyncio.to_thread(retrieve_priority_stage, question, user, integration)


async def retrieve_deep_priority_stage_async(question: str, user: dict[str, Any] | None, integration: bool = False) -> tuple[str, list[dict[str, Any]], dict[str, Any]]:
    return await asyncio.to_thread(retrieve_deep_priority_stage, question, user, integration)


_DETAIL_REQUEST_PATTERNS = (
    r"(?:کامل|دقیق|مفصل)\s+(?:توضیح|شرح)\s*(?:بده|دهید|بدین)?",
    r"(?:توضیح|شرح)\s+(?:کامل|دقیق|مفصل)",
    r"با\s+جزئیات(?:\s+(?:توضیح|شرح|بگو|بده))?",
    r"(?:مرحله|قدم)\s+به\s+(?:مرحله|قدم)",
    r"همه\s+(?:جزئیات|مراحل|موارد)\s+را\s+(?:بگو|توضیح بده)",
    r"\bin\s+detail\b|\bdetailed\s+explanation\b|\bexplain\s+fully\b",
)
_CACHE_NOISE_TOKENS = {
    'لطفا','لطفاً','خواهشاً','میشه','میشود','می‌شود','ممکنه','ممکن','است','هست','بگو','بگید',
    'بفرمایید','پاسخ','پاسخ بده','جواب','جواب بده','کن','کنید','توضیح','شرح','کامل','دقیق','مفصل',
}


def is_detailed_request(question: str) -> bool:
    normalized = normalize_text(question)
    return any(re.search(pattern, normalized, flags=re.IGNORECASE) for pattern in _DETAIL_REQUEST_PATTERNS)


def canonical_question_for_cache(question: str) -> str:
    normalized = normalize_text(question)
    for pattern in _DETAIL_REQUEST_PATTERNS:
        normalized = re.sub(pattern, ' ', normalized, flags=re.IGNORECASE)
    normalized = re.sub(r'[؟?!،؛:,.]+', ' ', normalized)
    tokens = [token for token in normalized.split() if token not in _CACHE_NOISE_TOKENS]
    return ' '.join(tokens).strip() or normalized


def sanitize_answer_text(value: str) -> str:
    text = unicodedata.normalize('NFKC', str(value or ''))
    text = text.translate(str.maketrans({'ي':'ی','ى':'ی','ك':'ک','ة':'ه','ۀ':'ه'}))
    text = re.sub(r'[\u200b-\u200f\u202a-\u202e\u2066-\u2069\ufeff\ufffd]', '', text)
    text = text.replace('\r\n','\n').replace('\r','\n')
    text = text.replace('**','').replace('__','').replace('`','')
    text = re.sub(r'(?m)^\s*[•*]+\s*', '– ', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\s+([،؛:؟!,.])', r'\1', text)
    text = re.sub(r'([،؛:؟!,.])(?=[^\s\n])', r'\1 ', text)
    text = re.sub(r'\n[ \t]+', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()



getcontext().prec = 40


def _normalize_math_expression(value: str) -> str:
    trans=str.maketrans('۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩','01234567890123456789')
    expr=str(value or '').translate(trans)
    expr=expr.replace('×','*').replace('÷','/').replace('−','-').replace('—','-').replace('–','-')
    expr=expr.replace('٫','.').replace('٬','').replace(',','').replace('٪','%').replace('^','**')
    expr=re.sub(r'(?<=\d)\s+(?=\d)','',expr)
    return expr.strip()


def _pure_math_expression(question: str) -> str | None:
    """Return an expression only when the complete request is arithmetic."""
    raw=str(question or '').strip()
    if not raw:
        return None
    raw=re.sub(r'^(?:محاسبه|حساب|حساب\s*کن|نتیجه)\s*[:：]?\s*','',raw,flags=re.IGNORECASE)
    raw=re.sub(r'(?:چقدر\s*(?:می\s*شود|می‌شود)|چنده|برابر\s+چند\s*(?:است|می\s*شود|می‌شود)?)\s*[؟?]*\s*$','',raw,flags=re.IGNORECASE).strip()
    raw=re.sub(r'[؟?=]+\s*$','',raw).strip()
    expr=_normalize_math_expression(raw)
    if not re.fullmatch(r'[0-9+\-*/().%\s]+',expr):
        return None
    if len(re.findall(r'\d+(?:\.\d+)?',expr))<2:
        return None
    if not re.search(r'[+\-*/%]',expr):
        return None
    return expr


def _requires_context_synthesis(question: str) -> bool:
    normalized=normalize_text(question)
    markers=(
        'محاسبه','حساب','جمع','تفریق','ضرب','تقسیم','درصد','مقایسه','بررسی',
        'مجاز','غیرمجاز','اگر','در صورتی','با توجه','ظرفیت نهایی','چقدر می شود',
        'چقدر می‌شود','مسیر','مبدا','مبدأ','مقصد','وزن بار','نوع بار'
    )
    return any(normalize_text(x) in normalized for x in markers)


def _decimal_eval_node(node: ast.AST) -> Decimal:
    if isinstance(node,ast.Expression): return _decimal_eval_node(node.body)
    if isinstance(node,ast.Constant) and isinstance(node.value,(int,float)):
        return Decimal(str(node.value))
    if isinstance(node,ast.UnaryOp) and isinstance(node.op,(ast.UAdd,ast.USub)):
        value=_decimal_eval_node(node.operand)
        return value if isinstance(node.op,ast.UAdd) else -value
    if isinstance(node,ast.BinOp):
        left=_decimal_eval_node(node.left);right=_decimal_eval_node(node.right)
        if isinstance(node.op,ast.Add): return left+right
        if isinstance(node.op,ast.Sub): return left-right
        if isinstance(node.op,ast.Mult): return left*right
        if isinstance(node.op,ast.Div):
            if right==0: raise ZeroDivisionError('تقسیم بر صفر مجاز نیست.')
            return left/right
        if isinstance(node.op,ast.Mod):
            if right==0: raise ZeroDivisionError('باقی‌مانده تقسیم بر صفر تعریف نشده است.')
            return left%right
        if isinstance(node.op,ast.Pow):
            if right!=right.to_integral_value() or abs(int(right))>12:
                raise ValueError('توان باید عدد صحیح بین منفی ۱۲ و ۱۲ باشد.')
            return left**int(right)
    raise ValueError('عبارت ریاضی نامعتبر است.')


def safe_decimal_calculate(expression: str) -> Decimal:
    expr=_normalize_math_expression(expression)
    if len(expr)>180 or not re.fullmatch(r'[0-9+\-*/().%\s]+',expr):
        raise ValueError('عبارت ریاضی فقط باید شامل عدد و عملگرهای استاندارد باشد.')
    parsed=ast.parse(expr,mode='eval')
    return _decimal_eval_node(parsed)


def _format_decimal_result(value: Decimal) -> str:
    if value==value.to_integral_value(): return f'{int(value):,}'
    normalized=format(value.normalize(),'f')
    if '.' in normalized: normalized=normalized.rstrip('0').rstrip('.')
    return normalized


def apply_verified_calculations(answer: str) -> str:
    """Replace model calculation markers with backend-verified Decimal results."""
    def replace(match: re.Match[str]) -> str:
        expr=match.group(1).strip()
        try:return _format_decimal_result(safe_decimal_calculate(expr))
        except Exception:return 'محاسبه نامعتبر'
    return re.sub(r'\[\[CALC:(.*?)\]\]',replace,answer,flags=re.S)


def direct_math_answer(question: str) -> str | None:
    expr=_pure_math_expression(question)
    if not expr:
        return None
    try:result=safe_decimal_calculate(expr)
    except Exception:return None
    display=expr.replace('**','^')
    return f'نتیجه محاسبه {display.strip()} برابر با {_format_decimal_result(result)} است.'


def _context_has_rules_and_exceptions(items: list[dict[str, Any]]) -> bool:
    text=normalize_text(' '.join(str(x.get('content') or x.get('answer') or '') for x in items))
    markers=('تبصره','استثنا','اما','ولی','در صورتی','در صورت','به شرط','مگر','به جز','حداکثر','حداقل','بالاتر','پایین تر','منطقه','شرایط خاص')
    return any(normalize_text(marker) in text for marker in markers)


def _question_requests_self_analysis(question: str) -> bool:
    normalized=normalize_text(question)
    return any(x in normalized for x in ('تحلیل خودت','نظر خودت','بدون منبع','دانش عمومی','خودت بررسی کن'))


def _sentence_chunks(text: str) -> list[str]:
    lines = [line.strip(' -–•\t') for line in text.splitlines() if line.strip()]
    chunks: list[str] = []
    for line in lines:
        parts = re.split(r'(?<=[.!؟!…])\s+', line)
        chunks.extend(part.strip() for part in parts if part.strip())
    return chunks


def format_answer_for_mode(answer: str, detailed: bool) -> str:
    clean = sanitize_answer_text(answer)
    if detailed:
        return clean[:AI_DETAILED_MAX_ANSWER_CHARS].rstrip()
    chunks = _sentence_chunks(clean)
    if not chunks:
        return clean[:AI_DEFAULT_MAX_ANSWER_CHARS].rstrip()
    selected: list[str] = []
    total = 0
    for chunk in chunks:
        candidate = (' '.join(selected + [chunk])).strip()
        if selected and (len(selected) >= 2 or len(candidate) > AI_DEFAULT_MAX_ANSWER_CHARS):
            break
        selected.append(chunk)
        total = len(candidate)
        if total >= AI_DEFAULT_MAX_ANSWER_CHARS:
            break
    result = ' '.join(selected).strip() or clean
    if len(result) > AI_DEFAULT_MAX_ANSWER_CHARS:
        cut = max(result.rfind('؟',0,AI_DEFAULT_MAX_ANSWER_CHARS), result.rfind('.',0,AI_DEFAULT_MAX_ANSWER_CHARS), result.rfind('!',0,AI_DEFAULT_MAX_ANSWER_CHARS))
        if cut >= 80:
            result = result[:cut+1]
    return result.strip()


def cache_access_scope(user: dict[str, Any] | None, integration: bool = False) -> str:
    vis = ','.join(allowed_visibilities(user, integration))
    if user:
        return f"vis={vis};role={user.get('role','user')};user={int(user.get('id') or 0)};dept={normalize_text(user.get('department') or '')}"
    return f"vis={vis};role={'integration' if integration else 'guest'};user=0;dept="


def current_knowledge_signature(access_scope: str) -> str:
    # R29: signature generation must stay constant-time even with tens of thousands
    # of chunks. Every training/document/FAQ mutation bumps knowledge_version.
    payload=f"{ANSWER_CACHE_NAMESPACE}|v={knowledge_version()}|{access_scope}"
    return hashlib.sha256(payload.encode('utf-8')).hexdigest()[:24]


def _cache_similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    seq = SequenceMatcher(None,a,b).ratio()
    aa = {token for token in a.split() if token not in _PERSIAN_STOPWORDS}
    bb = {token for token in b.split() if token not in _PERSIAN_STOPWORDS}
    jac = len(aa & bb) / len(aa | bb) if aa and bb else 0.0
    return min(seq, jac) if aa != bb else max(seq, jac)


def _cached_sources_eligible(sources: list[dict[str,Any]], user: dict[str,Any] | None, integration: bool=False) -> bool:
    """Reject cached answers whose document source is no longer answer-eligible.

    This closes the gap where a document can be disabled/blocked after an answer was
    cached without necessarily changing the cache row itself.
    """
    ids=document_source_ids(sources)
    if not ids:
        return True
    placeholders=','.join('?' for _ in ids)
    with get_db() as db:
        rows=db.execute(f"SELECT id,visibility,status,is_enabled FROM documents WHERE id IN ({placeholders})",ids).fetchall()
    if len(rows)!=len(ids):
        return False
    by_id={str(row['id']):row for row in rows}
    allowed=set(allowed_visibilities(user,integration))
    for document_id in ids:
        row=by_id.get(document_id)
        if not row or str(row['status'] or '')!='ready' or not bool(row['is_enabled']):
            return False
        if str(row['visibility'] or 'public') not in allowed:
            return False
    return True


def find_exact_cached_answer(question: str, detailed: bool, user: dict[str,Any] | None, integration: bool = False) -> dict[str,Any] | None:
    """O(1)-style cache fast path safe across training/source priority changes.

    knowledge_signature includes knowledge_version, so any manager training, source,
    or FAQ mutation makes an old answer ineligible immediately. No fuzzy matching is
    used here; fuzzy cache remains a later fallback.
    """
    if not ANSWER_CACHE_ENABLED or not FAST_EXACT_CACHE_ENABLED:
        return None
    normalized=canonical_question_for_cache(question)
    if len(normalized)<3:
        return None
    scope=cache_access_scope(user,integration);signature=current_knowledge_signature(scope);now_value=now_iso()
    with get_db() as db:
        row=db.execute(
            """SELECT * FROM answer_cache WHERE normalized_question=? AND detail_mode=? AND access_scope=? AND knowledge_signature=?
               AND (cache_tier='approved' OR (cache_tier='temporary' AND expires_at IS NOT NULL AND expires_at>?))
               ORDER BY CASE cache_tier WHEN 'approved' THEN 0 ELSE 1 END,id DESC LIMIT 1""",
            (normalized,int(detailed),scope,signature,now_value),
        ).fetchone()
        if row is None or is_navigation_only_content(row['answer']):
            return None
    try:sources=json.loads(row['sources_json'] or '[]')
    except json.JSONDecodeError:sources=[]
    if not _cached_sources_eligible(sources,user,integration):
        return None
    with get_db() as db:
        db.execute("UPDATE answer_cache SET hit_count=hit_count+1,last_used_at=? WHERE id=?",(now_value,row['id']))
    return {'answer':sanitize_answer_text(row['answer']),'sources':sources,'model':row['model'] or 'barsan-answer-cache','cache_id':row['id'],'cache_tier':row['cache_tier'] or 'approved','expires_at':row['expires_at']}


def find_cached_answer(question: str, detailed: bool, user: dict[str, Any] | None, integration: bool = False) -> dict[str, Any] | None:
    if not ANSWER_CACHE_ENABLED:
        return None
    normalized = canonical_question_for_cache(question)
    if len(normalized) < 3:
        return None
    scope = cache_access_scope(user, integration)
    signature = current_knowledge_signature(scope)
    now_value = now_iso()
    validity = "(cache_tier='approved' OR (cache_tier='temporary' AND expires_at IS NOT NULL AND expires_at>?))"
    with get_db() as db:
        row = db.execute(
            f"""SELECT * FROM answer_cache WHERE normalized_question=? AND detail_mode=? AND access_scope=?
                AND knowledge_signature=? AND {validity}
                ORDER BY CASE cache_tier WHEN 'approved' THEN 0 ELSE 1 END,id DESC LIMIT 1""",
            (normalized, int(detailed), scope, signature, now_value),
        ).fetchone()
        if row is None and len(normalized) >= 10:
            candidates = db.execute(
                f"""SELECT * FROM answer_cache WHERE detail_mode=? AND access_scope=? AND knowledge_signature=?
                    AND {validity} ORDER BY CASE cache_tier WHEN 'approved' THEN 0 ELSE 1 END,last_used_at DESC LIMIT ?""",
                (int(detailed), scope, signature, now_value, ANSWER_CACHE_FUZZY_CANDIDATES),
            ).fetchall()
            best = None; best_score = 0.0
            for candidate in candidates:
                score = _cache_similarity(normalized, candidate['normalized_question'])
                if score >= ANSWER_CACHE_FUZZY_THRESHOLD and score > best_score:
                    best,best_score=candidate,score
            row=best
        if row is None:
            return None
        if is_navigation_only_content(row['answer']):
            return None
        try:
            sources=json.loads(row['sources_json'] or '[]')
        except json.JSONDecodeError:
            sources=[]
        if not _cached_sources_eligible(sources,user,integration):
            return None
        db.execute("UPDATE answer_cache SET hit_count=hit_count+1,last_used_at=? WHERE id=?", (now_value,row['id']))
        return {
            'answer':sanitize_answer_text(row['answer']),'sources':sources,
            'model':row['model'] or 'barsan-answer-cache','cache_id':row['id'],
            'cache_tier':row['cache_tier'] or 'approved','expires_at':row['expires_at'],
        }


def cached_answer_matches_stage(cached: dict[str, Any] | None, stage: str) -> bool:
    if not cached or stage not in {'training','document'}:
        return False
    source_types={str(item.get('source_type') or '') for item in (cached.get('sources') or [])}
    if stage=='training':
        return bool(source_types) and source_types.issubset({'training'})
    return bool(source_types & {'document','resource','pdf','google_doc'}) and 'training' not in source_types


def store_cached_answer(
    question: str,
    detailed: bool,
    user: dict[str, Any] | None,
    integration: bool,
    answer: str,
    sources: list[dict[str, Any]],
    model: str,
    *,
    cache_tier: str = 'approved',
    approved_by: int | None = None,
) -> None:
    if not ANSWER_CACHE_ENABLED or not answer.strip() or is_navigation_only_content(answer):
        return
    if cache_tier not in {'approved','temporary'}:
        raise ValueError('cache_tier must be approved or temporary')
    normalized=canonical_question_for_cache(question)
    if len(normalized)<3:return
    scope=cache_access_scope(user,integration);signature=current_knowledge_signature(scope);ts=now_iso()
    expires_at=(datetime.now(timezone.utc)+timedelta(hours=TEMPORARY_CACHE_HOURS)).isoformat() if cache_tier=='temporary' else None
    with get_db() as db:
        db.execute(
            """INSERT INTO answer_cache(normalized_question,sample_question,detail_mode,access_scope,knowledge_signature,
                      answer,sources_json,model,cache_tier,expires_at,approved_by,created_at,last_used_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
               ON CONFLICT(normalized_question,detail_mode,access_scope,knowledge_signature) DO UPDATE SET
               sample_question=excluded.sample_question,answer=excluded.answer,sources_json=excluded.sources_json,
               model=excluded.model,cache_tier=excluded.cache_tier,expires_at=excluded.expires_at,
               approved_by=excluded.approved_by,created_at=excluded.created_at,last_used_at=excluded.last_used_at""",
            (normalized,question,int(detailed),scope,signature,answer,json.dumps(sources,ensure_ascii=False),model,cache_tier,expires_at,approved_by,ts,ts),
        )
        db.execute("DELETE FROM answer_cache WHERE cache_tier='temporary' AND expires_at IS NOT NULL AND expires_at<=?", (ts,))
        for tier,limit in (('approved',APPROVED_CACHE_MAX_ROWS),('temporary',TEMPORARY_CACHE_MAX_ROWS)):
            count=int(db.execute('SELECT COUNT(*) FROM answer_cache WHERE cache_tier=?',(tier,)).fetchone()[0])
            if count>limit:
                db.execute("DELETE FROM answer_cache WHERE id IN (SELECT id FROM answer_cache WHERE cache_tier=? ORDER BY last_used_at ASC LIMIT ?)",(tier,count-limit))


def record_knowledge_gap(question: str, user_id: int | None) -> None:
    normalized = normalize_text(question)
    if not normalized:
        return
    with get_db() as db:
        db.execute(
            """
            INSERT INTO knowledge_gaps(
                normalized_question, sample_question, occurrence_count,
                status, first_seen, last_seen, last_user_id
            ) VALUES (?, ?, 1, 'open', ?, ?, ?)
            ON CONFLICT(normalized_question) DO UPDATE SET
                occurrence_count = knowledge_gaps.occurrence_count + 1,
                sample_question = excluded.sample_question,
                last_seen = excluded.last_seen,
                last_user_id = excluded.last_user_id
            """,
            (normalized, question, now_iso(), now_iso(), user_id),
        )
        row = db.execute(
            "SELECT id, occurrence_count FROM knowledge_gaps WHERE normalized_question = ?",
            (normalized,),
        ).fetchone()
    if row and row["occurrence_count"] in {KNOWLEDGE_GAP_NOTIFY_THRESHOLD, 5, 10, 20}:
        notify_admin(
            "knowledge_gap",
            "نیاز آموزشی جدید شناسایی شد",
            {"gap_id": row["id"], "question": question, "count": row["occurrence_count"]},
            "warning",
        )


async def _request_json_with_retries(
    *,
    url: str,
    payload: dict[str, Any],
    params: dict[str, str] | None,
    headers: dict[str, str] | None,
    timeout_seconds: float,
    max_retries: int,
    provider_label: str,
    rotate_on_quota: bool = False,
) -> dict[str, Any]:
    timeout = httpx.Timeout(
        timeout_seconds,
        connect=min(20.0, timeout_seconds),
        read=timeout_seconds,
        write=min(30.0, timeout_seconds),
        pool=min(20.0, timeout_seconds),
    )
    retryable_statuses = {408, 409, 425, 429, 500, 502, 503, 504}
    last_error: Exception | None = None
    client=_HTTP_CLIENT
    owns_client=False
    if client is None:
        client=httpx.AsyncClient(timeout=timeout,follow_redirects=True)
        owns_client=True
    try:
        for attempt in range(max_retries + 1):
            try:
                response = await client.post(url, params=params, headers=headers, json=payload, timeout=timeout)
                detail = None
                if response.status_code >= 400:
                    try:
                        body = response.json()
                        error_obj = body.get('error') if isinstance(body,dict) else None
                        detail = (error_obj.get('message') if isinstance(error_obj,dict) else None) or (body.get('message') if isinstance(body,dict) else None) or (body.get('detail') if isinstance(body,dict) else None)
                    except Exception:
                        detail = response.text[:500]
                    detail_text = str(detail or '')
                    limit_markers = ('quota','rate limit','rate_limit','credit','balance','insufficient','exceeded','limit reached','سقف','اعتبار','موجودی','محدودیت')
                    is_key_or_quota_error = response.status_code in {401,402,403,429} or any(marker in detail_text.lower() for marker in limit_markers)
                    if rotate_on_quota and is_key_or_quota_error:
                        retry_after_raw = response.headers.get('retry-after')
                        retry_after_seconds = None
                        try:
                            retry_after_seconds = int(float(retry_after_raw)) if retry_after_raw else None
                        except (TypeError,ValueError):
                            retry_after_seconds = None
                        raise ApiSlotLimitError(response.status_code,detail_text or f'HTTP {response.status_code}',retry_after_seconds)
                if response.status_code in retryable_statuses and attempt < max_retries:
                    retry_after = response.headers.get('retry-after')
                    try:
                        delay = min(12.0, max(0.5, float(retry_after))) if retry_after else min(8.0, 0.75 * (2 ** attempt))
                    except ValueError:
                        delay = min(8.0, 0.75 * (2 ** attempt))
                    await asyncio.sleep(delay)
                    continue
                if response.status_code >= 400:
                    raise HTTPException(
                        status_code=502,
                        detail=f"خطای {provider_label}: HTTP {response.status_code}" + (f" — {detail}" if detail else ""),
                    )
                return response.json()
            except HTTPException:
                raise
            except (httpx.TimeoutException, httpx.NetworkError, httpx.RemoteProtocolError) as exc:
                last_error = exc
                if attempt >= max_retries:
                    break
                await asyncio.sleep(min(8.0, 0.75 * (2 ** attempt)))
    finally:
        if owns_client:
            await client.aclose()
    raise HTTPException(status_code=502, detail=f"ارتباط با {provider_label} پس از چند تلاش برقرار نشد.") from last_error


def _empty_usage() -> dict[str, Any]:
    return {"prompt_tokens":0,"output_tokens":0,"total_tokens":0,"estimated_cost":0.0}


def _add_usage(total: dict[str, Any], current: dict[str, Any]) -> None:
    for key in ("prompt_tokens","output_tokens","total_tokens"):
        total[key]=int(total.get(key,0))+int(current.get(key,0))
    total['estimated_cost']=float(total.get('estimated_cost',0.0))+float(current.get('estimated_cost',0.0) or 0.0)
    for key in ('api_slot','provider_label','model','model_route','reasoning_enabled','reasoning_mode','reasoning_effort','reasoning_tokens'):
        if current.get(key) is not None:
            total[key]=current.get(key)


def select_model_route(question: str, stage: str, sources: list[dict[str, Any]], detailed: bool) -> str:
    # R34 routing also considers deterministic query complexity and reasoning flags.
    if stage in {'faq','cache'}:
        return 'zero_token'
    if stage=='training' and sources and float(sources[0].get('score') or 0)>=DIRECT_TRAINING_MIN_SCORE and not _requires_context_synthesis(question):
        return 'direct_training'
    normalized=normalize_text(question)
    plan=analyze_query(question,max_subqueries=DEEP_MULTI_RETRIEVAL_MAX_QUERIES)
    distinct_docs={str(item.get('document_id')) for item in sources}
    complex_request=(
        detailed or plan.complexity>=0.58 or 'reasoning' in plan.flags or 'exception_sensitive' in plan.flags
        or len(distinct_docs)>=MODEL_ROUTE_MULTI_SOURCE_THRESHOLD
        or any(normalize_text(k) in normalized for k in MODEL_ROUTE_COMPLEX_KEYWORDS)
    )
    if complex_request:
        return 'advanced'
    top_score=max([float(x.get('deep_semantic_score') or x.get('score') or 0) for x in sources] or [0])
    if len(question.split())<=MODEL_ROUTE_SIMPLE_MAX_WORDS and len(distinct_docs)<=1 and top_score>=RETRIEVAL_HIGH_CONFIDENCE:
        return 'economy'
    return 'standard'


def _merge_continuation(base: str, continuation: str) -> str:
    base = base.rstrip()
    continuation = continuation.lstrip()
    if not base:
        return continuation
    if not continuation:
        return base
    max_overlap = min(500, len(base), len(continuation))
    for overlap in range(max_overlap, 19, -1):
        if base[-overlap:].strip() == continuation[:overlap].strip():
            return base + continuation[overlap:]
    separator = "" if base.endswith((" ", "\n", "-", "—", "/")) else "\n"
    return base + separator + continuation


async def _generate_ai_text(
    messages: list[dict[str, str]],
    *,
    max_tokens: int,
    temperature: float,
    route: str = 'standard',
) -> tuple[str, dict[str, Any], str]:
    if AI_PROVIDER == "local":
        return "", _empty_usage(), "stop"

    if AI_PROVIDER in OPENAI_COMPATIBLE_PROVIDERS:
        path = AI_CHAT_COMPLETIONS_PATH if AI_CHAT_COMPLETIONS_PATH.startswith('/') else f'/{AI_CHAT_COMPLETIONS_PATH}'
        slots = ordered_available_ai_slots_for_route(route)
        if not slots:
            raise HTTPException(status_code=503,detail='همه APIهای فعال در محدودیت یا توقف موقت هستند. یک API جدید از پنل ادمین اضافه کنید یا محدودیت‌ها را آزاد کنید.')
        last_errors: list[str] = []
        for attempt_index, slot in enumerate(slots,1):
            attempt_started=time.perf_counter()
            headers = {"Content-Type": "application/json"}
            headers[AI_AUTH_HEADER] = f"{AI_AUTH_SCHEME} {slot['api_key']}".strip()
            try:
                headers.update({str(k): str(v) for k, v in json.loads(AI_EXTRA_HEADERS_JSON).items()})
            except (json.JSONDecodeError, AttributeError):
                pass

            selected_model = slot.get('models',{}).get(route) or slot['model']

            reasoning_effort=(AI_REASONING_EFFORT_ADVANCED if route=='advanced' else AI_REASONING_EFFORT_STANDARD)
            capabilities=provider_capabilities(slot,selected_model)
            capability_mode=str(capabilities.get('reasoning_mode') or 'none')
            reasoning_parameter=(capability_mode if AI_REASONING_PARAMETER=='auto' else AI_REASONING_PARAMETER)
            if reasoning_parameter=='intrinsic': reasoning_parameter='none'
            reasoning_requested=bool(AI_REASONING_ENABLED and route in {'standard','advanced'} and capability_mode in {'reasoning','reasoning_effort'} and reasoning_parameter!='none')
            intrinsic_reasoning=bool(AI_REASONING_ENABLED and route in {'standard','advanced'} and capability_mode=='intrinsic')

            async def compatible_request(token_parameter: str, model_name: str = selected_model, *, use_reasoning: bool = False) -> dict[str, Any]:
                payload: dict[str, Any] = {
                    "model": model_name,
                    "messages": messages,
                    token_parameter: max_tokens,
                    "stream": False,
                }
                if use_reasoning:
                    if reasoning_parameter=='reasoning':
                        payload['reasoning']={'effort':reasoning_effort}
                    else:
                        payload['reasoning_effort']=reasoning_effort
                if AI_SEND_TEMPERATURE and (not use_reasoning or AI_REASONING_SEND_TEMPERATURE):
                    payload["temperature"] = temperature
                return await _request_json_with_retries(
                    url=f"{slot['base_url']}{path}",
                    payload=payload,
                    params=None,
                    headers=headers,
                    timeout_seconds=AI_TIMEOUT_SECONDS,
                    max_retries=AI_MAX_RETRIES,
                    provider_label=f"{slot['label']} (API {slot['slot']})",
                    rotate_on_quota=True,
                )

            token_parameters = [AI_TOKEN_PARAMETER] if AI_TOKEN_PARAMETER != 'auto' else ['max_tokens','max_completion_tokens']
            model_candidates=[selected_model]+([slot['model']] if selected_model!=slot['model'] else [])
            reasoning_modes=[True,False] if reasoning_requested else [False]
            used_reasoning=intrinsic_reasoning
            used_reasoning_mode='intrinsic' if intrinsic_reasoning else 'none'
            try:
                data=None;compatibility_error: HTTPException | None=None
                for model_candidate in model_candidates:
                    for use_reasoning in reasoning_modes:
                        for token_parameter in token_parameters:
                            try:
                                data=await compatible_request(token_parameter,model_candidate,use_reasoning=use_reasoning)
                                selected_model=model_candidate
                                used_reasoning=bool(use_reasoning or intrinsic_reasoning)
                                used_reasoning_mode=capability_mode if used_reasoning else 'none'
                                break
                            except HTTPException as exc:
                                detail_text=str(exc.detail)
                                if any(code in detail_text for code in ('HTTP 400','HTTP 404','HTTP 422')):
                                    compatibility_error=exc
                                    continue
                                raise
                        if data is not None:
                            break
                    if data is not None:
                        break
                if data is None:
                    raise compatibility_error or HTTPException(status_code=502,detail='پارامتر یا مدل سازگار با سرویس پیدا نشد.')
            except ApiSlotLimitError as exc:
                elapsed_ms=int((time.perf_counter()-attempt_started)*1000)
                record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=selected_model,model_route=route,
                    status='limited',http_status=exc.status_code,error_class='quota_or_rate_limit',error_detail=exc.detail,
                    response_ms=elapsed_ms,attempt_index=attempt_index,was_failover=attempt_index>1)
                cooldown = exc.retry_after_seconds or (AI_KEY_RATE_LIMIT_COOLDOWN_SECONDS if exc.status_code == 429 else AI_KEY_QUOTA_COOLDOWN_SECONDS)
                mark_api_slot_blocked(int(slot['slot']),exc.detail,cooldown,'limited')
                last_errors.append(f"API {slot['slot']}: {exc.detail}")
                continue
            except HTTPException as exc:
                elapsed_ms=int((time.perf_counter()-attempt_started)*1000)
                transient=AI_ROTATE_ON_TRANSIENT_ERRORS and any(code in str(exc.detail) for code in ('HTTP 400','HTTP 404','HTTP 408','HTTP 409','HTTP 422','HTTP 425','HTTP 500','HTTP 502','HTTP 503','HTTP 504','ارتباط با'))
                record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=selected_model,model_route=route,
                    status='transient_error' if transient else 'error',http_status=None,error_class='provider_http_error',error_detail=str(exc.detail),
                    response_ms=elapsed_ms,attempt_index=attempt_index,was_failover=attempt_index>1)
                if transient:
                    mark_api_slot_blocked(int(slot['slot']),str(exc.detail),AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'transient_error')
                    last_errors.append(f"API {slot['slot']}: {exc.detail}")
                    continue
                raise

            choices = data.get("choices") or []
            if not choices:
                elapsed_ms=int((time.perf_counter()-attempt_started)*1000)
                record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=selected_model,model_route=route,
                    status='invalid_response',error_class='missing_choices',error_detail='پاسخ بدون choices',response_ms=elapsed_ms,
                    attempt_index=attempt_index,was_failover=attempt_index>1)
                mark_api_slot_blocked(int(slot['slot']),'پاسخ بدون choices',AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'invalid_response')
                last_errors.append(f"API {slot['slot']}: پاسخ خالی")
                continue
            choice = choices[0]
            message = choice.get("message") or {}
            raw_content = message.get("content") if isinstance(message, dict) else None
            if isinstance(raw_content, list):
                answer = "\n".join(
                    str(part.get("text") or part.get("content") or "")
                    for part in raw_content
                    if isinstance(part, dict) and (part.get("text") or part.get("content"))
                ).strip()
            else:
                answer = str(raw_content or choice.get("text") or "").strip()
            if not answer:
                elapsed_ms=int((time.perf_counter()-attempt_started)*1000)
                record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=selected_model,model_route=route,
                    status='invalid_response',error_class='empty_content',error_detail='متن پاسخ خالی بود',response_ms=elapsed_ms,
                    attempt_index=attempt_index,was_failover=attempt_index>1)
                mark_api_slot_blocked(int(slot['slot']),'متن پاسخ خالی بود',AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'invalid_response')
                last_errors.append(f"API {slot['slot']}: متن خالی")
                continue
            elapsed_ms=int((time.perf_counter()-attempt_started)*1000)
            record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=selected_model,model_route=route,
                status='success',response_ms=elapsed_ms,attempt_index=attempt_index,was_failover=attempt_index>1)
            mark_api_slot_success(int(slot['slot']))
            raw_usage = data.get("usage") or {}
            prompt_tokens = int(raw_usage.get("prompt_tokens") or estimate_tokens(json.dumps(messages, ensure_ascii=False)))
            output_tokens = int(raw_usage.get("completion_tokens") or raw_usage.get("output_tokens") or estimate_tokens(answer))
            completion_details=raw_usage.get('completion_tokens_details') or raw_usage.get('output_tokens_details') or {}
            reasoning_tokens=int(completion_details.get('reasoning_tokens') or 0) if isinstance(completion_details,dict) else 0
            estimated_cost=(prompt_tokens/1_000_000.0)*float(slot.get('input_cost_per_1m') or 0)+(output_tokens/1_000_000.0)*float(slot.get('output_cost_per_1m') or 0)
            usage = {
                "prompt_tokens":prompt_tokens,"output_tokens":output_tokens,
                "total_tokens":int(raw_usage.get("total_tokens") or prompt_tokens+output_tokens),
                "api_slot":int(slot['slot']),"provider_label":slot['label'],"model":selected_model,
                "model_route":route,"estimated_cost":round(estimated_cost,8),
                "reasoning_enabled":bool(used_reasoning),"reasoning_mode":used_reasoning_mode,"reasoning_effort":reasoning_effort if used_reasoning and used_reasoning_mode not in {'none','intrinsic'} else None,
                "reasoning_tokens":reasoning_tokens,
            }
            finish_reason = str(choice.get("finish_reason") or "stop").lower()
            if finish_reason in {"max_tokens", "max_output_tokens", "max_completion_tokens"}:
                finish_reason = "length"
            return answer, usage, finish_reason
        notify_admin('api_pool_exhausted','تمام APIهای هوش مصنوعی در محدودیت یا خطا قرار گرفتند',{'errors':last_errors},'danger')
        raise HTTPException(status_code=503,detail='تمام APIهای هوش مصنوعی فعلاً محدود یا نامتاح هستند؛ سیستم به‌صورت خودکار همه کلیدهای تنظیم‌شده را بررسی کرد.')


    prompt = "\n\n".join(
        f"{('دستور سیستم' if item.get('role') == 'system' else 'دستیار' if item.get('role') == 'assistant' else 'کاربر')}:\n{item.get('content', '')}"
        for item in messages
    )
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "maxOutputTokens": max_tokens,
            "temperature": temperature,
        },
    }
    data = await _request_json_with_retries(
        url=f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent",
        payload=payload,
        params={"key": GEMINI_API_KEY},
        headers={"Content-Type": "application/json"},
        timeout_seconds=GEMINI_TIMEOUT_SECONDS,
        max_retries=GEMINI_MAX_RETRIES,
        provider_label="Gemini",
    )
    candidates = data.get("candidates") or []
    if not candidates:
        raise HTTPException(status_code=502, detail="Gemini پاسخی برنگرداند.")
    candidate = candidates[0]
    parts = candidate.get("content", {}).get("parts", [])
    answer = "\n".join(str(part.get("text") or "") for part in parts if part.get("text")).strip()
    if not answer:
        raise HTTPException(status_code=502, detail="متن پاسخ Gemini خالی است.")
    meta = data.get("usageMetadata") or {}
    prompt_tokens = int(meta.get("promptTokenCount") or estimate_tokens(prompt))
    output_tokens = int(meta.get("candidatesTokenCount") or estimate_tokens(answer))
    usage = {
        "prompt_tokens":prompt_tokens,"output_tokens":output_tokens,
        "total_tokens":int(meta.get("totalTokenCount") or prompt_tokens+output_tokens),
        "api_slot":0,"provider_label":"Gemini","model":GEMINI_MODEL,"model_route":route,"estimated_cost":0.0,
    }
    finish_reason = str(candidate.get("finishReason") or "STOP").lower()
    if finish_reason in {"max_tokens", "max_output_tokens"}:
        finish_reason = "length"
    return answer, usage, finish_reason


def _looks_incomplete(answer: str) -> bool:
    value = answer.rstrip()
    if not value:
        return True
    if value.endswith((".", "؟", "!", "…", ")", "]", "}", "»")):
        return False
    if len(value) < 25:
        return True
    tail = normalize_text(value[-90:]).split()
    unfinished_words = {"و", "یا", "اما", "ولی", "زیرا", "چون", "که", "همچنین", "سپس", "بنابراین", "شامل"}
    return bool(tail and tail[-1] in unfinished_words) or value.endswith((":", "،", ",", "-", "—"))


def _should_expand_answer(question: str, answer: str, context_items: list[dict[str, Any]]) -> bool:
    if not AI_ENFORCE_DETAILED_ANSWER:
        return False
    context_chars = sum(len(str(item.get('content') or '')) for item in context_items)
    return (
        _looks_incomplete(answer)
        or (len(answer.strip()) < AI_MIN_RESPONSE_CHARS and context_chars >= 1400 and len(search_tokens(question)) >= 2)
    )


def _remote_verification_required(question: str, selected: list[dict[str,Any]], answer: str) -> bool:
    """Use a second model call only when deterministic evidence checks are insufficient."""
    if not SOURCE_VERIFICATION_FAST_GUARD:
        return True
    source_blob='\n'.join(str(item.get('content') or item.get('answer') or '') for item in selected)
    answer_numbers=_numeric_claims(answer);source_numbers=_numeric_claims(source_blob)
    if answer_numbers and not answer_numbers.issubset(source_numbers):
        return True
    conditional=_context_has_rules_and_exceptions(selected)
    if conditional:
        # When the compact answer preserves all small numeric claims and explicitly
        # signals a condition/exception, a second network round-trip adds little value.
        important_numbers=set(list(source_numbers)[:6])
        normalized_answer=normalize_text(answer)
        has_condition_word=any(normalize_text(x) in normalized_answer for x in ('اما','شرط','استثنا','در شرایط','در صورت','مگر','حداکثر','حداقل'))
        if important_numbers and not important_numbers.issubset(answer_numbers):
            return True
        if not has_condition_word:
            return True
    distinct_docs={str(item.get('document_id') or '') for item in selected}
    top_score=max([float(item.get('score') or 0) for item in selected] or [0.0])
    if len(distinct_docs)>1 and top_score<RETRIEVAL_HIGH_CONFIDENCE:
        return True
    return False


def calibrated_confidence_threshold(base_threshold: float) -> tuple[float,dict[str,Any]]:
    """Calibrate acceptance threshold from the latest labelled Golden run.

    The shift is intentionally bounded so a small/eccentric eval set cannot make
    production dramatically more permissive or restrictive.
    """
    global _CONFIDENCE_CALIBRATION_CACHE
    if not CONFIDENCE_CALIBRATION_ENABLED:
        return base_threshold,{'active':False,'reason':'disabled'}
    now=time.monotonic()
    with _CONFIDENCE_CALIBRATION_LOCK:
        if now-_CONFIDENCE_CALIBRATION_CACHE[0]<60 and _CONFIDENCE_CALIBRATION_CACHE[1] is not None:
            return float(_CONFIDENCE_CALIBRATION_CACHE[1]),dict(_CONFIDENCE_CALIBRATION_CACHE[2])
    try:
        with get_db() as db:
            row=db.execute('SELECT details_json,created_at FROM golden_runs ORDER BY created_at DESC LIMIT 1').fetchone()
        details=json.loads(row['details_json'] or '[]') if row else []
    except (sqlite3.Error,json.JSONDecodeError,TypeError):
        details=[];row=None
    points=[]
    for item in details if isinstance(details,list) else []:
        try: score=float(item.get('evidence_confidence'))
        except (TypeError,ValueError): continue
        points.append((score,bool(item.get('passed'))))
    positives=sum(1 for _,label in points if label); negatives=len(points)-positives
    if len(points)<CONFIDENCE_CALIBRATION_MIN_CASES or not positives or not negatives:
        # Fresh installs still use an explicit conservative baseline instead of an
        # 'inactive' calibrator. As soon as labelled Golden data is sufficient the
        # adaptive calibration below automatically takes over.
        safe_threshold=max(base_threshold,DEEP_CONFIDENCE_MIN)
        meta={'active':True,'mode':'safe_default','reason':'awaiting_labelled_golden_cases','cases':len(points),'positive':positives,'negative':negatives,'threshold':round(safe_threshold,4)}
        with _CONFIDENCE_CALIBRATION_LOCK:_CONFIDENCE_CALIBRATION_CACHE=(now,safe_threshold,meta)
        return safe_threshold,meta
    best=(base_threshold,-1.0)
    candidates=sorted(set([base_threshold]+[round(0.30+i*0.01,2) for i in range(61)]))
    for threshold in candidates:
        tp=sum(1 for score,label in points if label and score>=threshold); fn=positives-tp
        tn=sum(1 for score,label in points if (not label) and score<threshold); fp=negatives-tn
        tpr=tp/max(1,tp+fn);tnr=tn/max(1,tn+fp);balanced=(tpr+tnr)/2
        distance=abs(threshold-base_threshold)
        best_distance=abs(best[0]-base_threshold)
        if balanced>best[1]+1e-9 or (abs(balanced-best[1])<1e-9 and distance<best_distance):
            best=(threshold,balanced)
    low=max(0.20,base_threshold-CONFIDENCE_CALIBRATION_MAX_SHIFT);high=min(0.95,base_threshold+CONFIDENCE_CALIBRATION_MAX_SHIFT)
    threshold=max(low,min(high,best[0]))
    meta={'active':True,'cases':len(points),'positive':positives,'negative':negatives,'balanced_accuracy':round(best[1],4),'raw_threshold':best[0],'threshold':round(threshold,4),'run_created_at':row['created_at'] if row else None}
    with _CONFIDENCE_CALIBRATION_LOCK:_CONFIDENCE_CALIBRATION_CACHE=(now,threshold,meta)
    return threshold,meta


async def _model_semantic_rerank(question: str, items: list[dict[str,Any]], *, route: str='standard') -> tuple[list[dict[str,Any]],dict[str,Any]]:
    """Optional LLM reranker over a small pre-ranked candidate set.

    It can only reorder supplied evidence; it cannot add facts or remove all
    candidates. Any provider/JSON failure returns the deterministic order.
    """
    if not MODEL_SEMANTIC_RERANK_ENABLED or len(items)<3 or AI_PROVIDER=='local':
        return items,_empty_usage()
    candidates=items[:MODEL_SEMANTIC_RERANK_MAX_ITEMS]
    lines=[]
    for idx,item in enumerate(candidates):
        content=sanitize_answer_text(str(item.get('content') or item.get('answer') or ''))[:850]
        lines.append(f"[{idx}] {content}")
    messages=[
        {'role':'system','content':'فقط ارتباط شواهد با سؤال را رتبه‌بندی کن. هیچ پاسخ یا توضیحی نساز. خروجی فقط JSON معتبر به شکل {"order":[0,2,1]} باشد و همه شماره‌های ورودی را دقیقاً یک بار بیاور.'},
        {'role':'user','content':f"سؤال: {question}\n\nشواهد:\n"+'\n'.join(lines)},
    ]
    try:
        raw,usage,finish=await _generate_ai_text(messages,max_tokens=MODEL_SEMANTIC_RERANK_MAX_TOKENS,temperature=0.0,route=('standard' if route=='economy' else route))
        if finish=='length': return items,usage
        match=re.search(r'\{.*\}',raw,flags=re.S)
        data=json.loads(match.group(0) if match else raw)
        order=data.get('order') if isinstance(data,dict) else None
        if not isinstance(order,list): return items,usage
        normalized=[]
        for value in order:
            try: idx=int(value)
            except (TypeError,ValueError): continue
            if 0<=idx<len(candidates) and idx not in normalized: normalized.append(idx)
        if len(normalized)!=len(candidates): return items,usage
        reranked=[]
        for rank,idx in enumerate(normalized):
            clone=dict(candidates[idx]);clone['model_rerank_rank']=rank;clone['model_reranked']=True;reranked.append(clone)
        reranked.extend(items[len(candidates):])
        return reranked,usage
    except Exception:
        return items,_empty_usage()


async def ask_ai(question: str, context_items: list[dict[str, Any]], detailed: bool = False, route: str = 'standard', memory: str = '') -> tuple[str, dict[str, Any]]:
    if not context_items:
        return "پاسخ این پرسش در منابع و آموزش‌های فعال موجود نیست.", _empty_usage()

    context_items = _filter_navigation_only_items(context_items)
    if not context_items:
        return "پاسخ این پرسش در منابع موجود پیدا نشد.", _empty_usage()
    query_plan=analyze_query(question,max_subqueries=DEEP_MULTI_RETRIEVAL_MAX_QUERIES)

    training_items = sorted(
        [item for item in context_items if item.get("source_type") == "training"],
        key=lambda item: (float(item.get('score',0)),int(item.get('priority',0))), reverse=True,
    )
    document_items = sorted(
        [item for item in context_items if item.get("source_type") != "training"],
        key=lambda item: float(item.get('score',0)), reverse=True,
    )

    rerank_usage=_empty_usage()
    if not training_items and query_plan.complexity>=MODEL_SEMANTIC_RERANK_MIN_COMPLEXITY:
        document_items,rerank_usage=await _model_semantic_rerank(question,document_items,route=route)

    if DIRECT_TRAINING_ANSWER_ENABLED and training_items and training_items[0]["score"] >= DIRECT_TRAINING_MIN_SCORE and not _requires_context_synthesis(question):
        canonical_answer = str(training_items[0]["answer"]).strip()
        if detailed or len(canonical_answer) <= AI_DEFAULT_MAX_ANSWER_CHARS * 2:
            answer = format_answer_for_mode(canonical_answer, detailed)
            usage=_empty_usage();usage.update({'model_route':'direct_training','provider_label':'Barsan knowledge','model':'barsan-training-direct-zero-token','verification_status':'deterministic_verified','evidence_confidence':0.99,'deep_query_plan':query_plan.as_dict()})
            return answer, usage

    # R30 adaptive context: high-confidence single-document answers use a compact
    # prompt; ambiguous/multi-source questions automatically keep the wider context.
    # This cuts model latency without dropping rule/exception context when it matters.
    top_score=max([float(item.get('score') or 0.0) for item in context_items] or [0.0])
    distinct_docs={str(item.get('document_id') or '') for item in context_items if item.get('source_type')!='training'}
    conditional_query=_context_has_rules_and_exceptions(context_items)
    if detailed:
        item_limit=min(RETRIEVAL_TOTAL_ITEMS,12);char_limit=min(RETRIEVAL_MAX_CONTEXT_CHARS,28000)
    elif top_score>=RETRIEVAL_HIGH_CONFIDENCE and len(distinct_docs)<=1:
        item_limit=min(RETRIEVAL_TOTAL_ITEMS,5 if conditional_query else 4);char_limit=min(RETRIEVAL_MAX_CONTEXT_CHARS,14000 if conditional_query else 10000)
    elif top_score>=RETRIEVAL_MEDIUM_CONFIDENCE:
        item_limit=min(RETRIEVAL_TOTAL_ITEMS,7);char_limit=min(RETRIEVAL_MAX_CONTEXT_CHARS,18000)
    else:
        item_limit=min(RETRIEVAL_TOTAL_ITEMS,10);char_limit=min(RETRIEVAL_MAX_CONTEXT_CHARS,24000)
    selected: list[dict[str, Any]] = []
    used_chars = 0
    # Only one authoritative layer is normally present: training OR documents.
    for item in (training_items + document_items):
        if len(selected) >= item_limit or used_chars >= char_limit:
            break
        content = sanitize_answer_text(str(item.get('content') or ''))
        if not content:
            continue
        remaining = char_limit-used_chars
        if len(content) > remaining:
            content = content[:remaining].rsplit(' ',1)[0].rstrip() or content[:remaining]
        clone = dict(item)
        clone['content'] = content
        selected.append(clone)
        used_chars += len(content)

    training_context = '\n'.join(f"[آموزش {i}] {item['content']}" for i,item in enumerate([x for x in selected if x.get('source_type')=='training'],1))
    document_context = '\n'.join(f"[منبع {i}] {item['content']}" for i,item in enumerate([x for x in selected if x.get('source_type')!='training'],1))
    rule_map=build_rule_exception_map(selected) if DEEP_RULE_ENGINE_ENABLED else {}
    rule_context=format_rule_map_for_prompt(rule_map) if rule_map else ''
    response_rule = (
        "کاربر صریحاً توضیح کامل خواسته است. همه قواعد، شروط، تبصره‌ها و استثناهای مرتبط را ساختاریافته و بدون رونویسی مطالب نامرتبط توضیح بده."
        if detailed else
        "پاسخ را فشرده و معمولاً در یک یا دو جمله بده؛ اما برای کوتاه‌کردن، شرط یا تبصره مرتبط را حذف نکن. ابتدا قاعده عمومی را بگو و سپس استثنا یا شرایط خاص را با «اما» یا «در شرایط...» اضافه کن."
    )
    system_prompt = f"""شما دستیار سازمانی بارسان هستید و پاسخ باید فقط از شواهد ارائه‌شده ساخته شود.
قاعده مرجع: اگر «آموزش مدیر» در ورودی وجود دارد، فقط همان آموزش مرجع قطعی پاسخ است. فایل سازمانی فقط زمانی به ورودی می‌آید که آموزش مرتبط وجود نداشته باشد.
{response_rule}
پیش از پاسخ، سؤال و شواهد را درونی و مرحله‌ای تحلیل کن؛ زنجیره استدلال داخلی را نمایش نده و فقط نتیجه مستند را بنویس.
اگر «نقشه قواعد و استثناها» ارائه شده، برای تعیین قاعده عمومی، شرط و تبصره از آن استفاده کن اما متن منبع همچنان مرجع نهایی است.
قوانین قطعی:
1) مقدار یا حکم استثنایی را به‌جای قاعده عمومی معرفی نکن.
2) اگر منبع یک مقدار اصلی و یک مقدار مشروط دارد، هر دو را همراه شرط دقیق اعلام کن.
3) اعداد، واحدها، نام مناطق، محدودیت‌ها، «حداقل/حداکثر»، تبصره و استثنا را عیناً حفظ کن.
4) برای هر جمع، تفریق، ضرب یا تقسیم، عبارت را با اعداد انگلیسی داخل [[CALC:...]] بنویس تا سامانه نتیجه را قطعی محاسبه کند؛ خودت نتیجه را حدس نزن.
5) هیچ ادعایی خارج از منابع اضافه نکن و متن را کامل تمام کن.
6) هر ماژول مستقل است: پاسخ را در همان ماژول و به‌صورت مستقیم بده و کاربر را به منوی دیگری ارجاع نده. اگر اطلاعات در آموزش یا منابع وجود دارد، همان‌جا پاسخ کامل را ارائه کن.
7) دستورهای قدیمی که فقط می‌گویند «به فلان بخش برو» دانش پاسخ‌گویی نیستند و باید نادیده گرفته شوند.
اگر پاسخ در منابع نیست، فقط بگو: «پاسخ این پرسش در منابع موجود نیست.»""".strip()
    sections = [f"سؤال: {question}"]
    if memory:
        sections.append(f"زمینه کوتاه مکالمه قبلی (فقط برای رفع ارجاع‌های مبهم):\n{memory}")
    if training_context:
        sections.append(f"آموزش مدیر:\n{training_context}")
    if document_context:
        sections.append(f"بخش‌های مرتبط منابع:\n{document_context}")
    if rule_context:
        sections.append(f"نقشه قواعد و استثناها (استخراج ساختاریافته از همین منابع):\n{rule_context}")
    sections.append(f"طرح تحلیل سؤال: intent={query_plan.intent}; flags={','.join(query_plan.flags) or 'none'}; entities={','.join(query_plan.entities) or 'none'}; numbers={','.join(query_plan.numbers) or 'none'}")
    user_prompt = '\n\n'.join(sections)

    if AI_PROVIDER == "local":
        raw = training_items[0]["answer"] if training_items else document_items[0]["content"]
        answer = format_answer_for_mode(str(raw), detailed)
        confidence,confidence_parts=evidence_confidence(query_plan,selected,rule_map,verification_status='deterministic_verified')
        usage = {"prompt_tokens":0,"output_tokens":0,"total_tokens":0,"verification_status":"deterministic_verified","evidence_confidence":confidence,"confidence_parts":confidence_parts,"deep_query_plan":query_plan.as_dict()}
        return answer, usage

    messages = [{"role":"system","content":system_prompt},{"role":"user","content":user_prompt}]
    max_tokens = active_max_completion_tokens(detailed)
    answer, first_usage, finish_reason = await _generate_ai_text(messages,max_tokens=max_tokens,temperature=active_temperature(),route=route)
    total_usage = _empty_usage(); _add_usage(total_usage,rerank_usage); _add_usage(total_usage,first_usage) 

    continuation_rounds = 0
    while (finish_reason == "length" or _looks_incomplete(answer)) and continuation_rounds < AI_CONTINUATION_ROUNDS:
        continuation_rounds += 1
        continuation_messages = [
            {"role":"system","content":"فقط ادامه کوتاه و دقیق جمله نیمه‌تمام را تولید کن؛ متن قبلی یا منابع را تکرار نکن."},
            {"role":"user","content":f"سؤال اصلی: {question}\nانتهای پاسخ قبلی: {answer[-1200:]}"},
        ]
        extra,extra_usage,finish_reason = await _generate_ai_text(
            continuation_messages,max_tokens=min(max_tokens,700),temperature=active_temperature(),route=route
        )
        answer = _merge_continuation(answer,extra); _add_usage(total_usage,extra_usage)

    if finish_reason == "length" or _looks_incomplete(answer):
        raise HTTPException(status_code=502,detail="پاسخ مدل کامل نشد؛ دوباره تلاش کنید.")

    # A second evidence check is used when the source contains numeric claims or
    # conditional rules. It prevents an exception from replacing the base rule.
    source_blob='\n'.join(str(item.get('content') or '') for item in selected)
    verification_status='not_required'
    needs_verification=(
        _context_has_rules_and_exceptions(selected) or bool(re.search(r'\d|[۰-۹]',source_blob))
        or route=='advanced' or query_plan.complexity>=0.58 or bool(rule_map.get('potential_numeric_conflicts'))
    )
    remote_verification=(
        DEEP_ANSWER_VERIFICATION_ENABLED and needs_verification
        and (route=='advanced' or query_plan.complexity>=0.58 or _remote_verification_required(question,selected,answer))
    )
    if SOURCE_ANSWER_VERIFICATION and remote_verification:
        verify_messages=[
            {"role":"system","content":"پاسخ پیشنهادی را فقط با متن منابع تطبیق بده. قاعده عمومی را اول، سپس همه شرط‌ها و استثناهای مرتبط را کوتاه بیاور. عدد، واحد یا شرطی را حذف یا جابه‌جا نکن. اگر محاسبه لازم است از [[CALC:expression]] استفاده کن. فقط نسخه اصلاح‌شده پاسخ را بده."},
            {"role":"user","content":f"سؤال: {question}\n\nمنابع:\n{source_blob[:RETRIEVAL_MAX_CONTEXT_CHARS]}\n\nنقشه قواعد/استثناها:\n{rule_context[:7000]}\n\nپاسخ پیشنهادی:\n{answer}"},
        ]
        verification_status='verified'
        try:
            checked,checked_usage,checked_finish=await _generate_ai_text(verify_messages,max_tokens=SOURCE_VERIFICATION_MAX_TOKENS,temperature=0.0,route=route)
            _add_usage(total_usage,checked_usage)
            if checked_finish!='length' and checked.strip() and not _looks_incomplete(checked):
                answer=checked.strip()
            else:
                verification_status='incomplete'
        except Exception as exc:
            # _generate_ai_text already attempts provider failover. If every verifier is
            # unavailable, never silently trust new numeric claims: keep the candidate
            # only when every number is present in the authoritative evidence.
            verification_status='provider_unavailable'
            answer_numbers=_numeric_claims(answer)
            source_numbers=_numeric_claims(source_blob)
            if not answer_numbers.issubset(source_numbers):
                factual_lines=[]
                for item in selected:
                    text=sanitize_answer_text(str(item.get('content') or item.get('answer') or ''))
                    for line in text.splitlines():
                        line=line.strip()
                        if line and (set(search_tokens(question)) & set(search_tokens(line)) or _numeric_claims(line)):
                            factual_lines.append(line)
                    if len(factual_lines)>=5: break
                if factual_lines:
                    answer=' '.join(factual_lines[:5])
            total_usage['verification_error']=sanitize_answer_text(str(exc))[:160]
    elif needs_verification:
        answer_numbers=_numeric_claims(answer);source_numbers=_numeric_claims(source_blob)
        verification_status='deterministic_verified' if answer_numbers.issubset(source_numbers) else 'failed'
    total_usage['verification_status']=verification_status

    answer = apply_verified_calculations(format_answer_for_mode(answer,detailed))
    if is_navigation_only_content(answer):
        factual_items=[item for item in selected if not is_navigation_only_content(str(item.get('content') or item.get('answer') or ''))]
        factual_blob='\n\n'.join(str(item.get('content') or item.get('answer') or '') for item in factual_items)[:RETRIEVAL_MAX_CONTEXT_CHARS]
        if factual_blob:
            rescue_messages=[
                {"role":"system","content":"فقط از متن منبع پاسخ بده. هیچ اشاره‌ای به منو، بخش بررسی بار، مسیریابی یا رفتن به صفحه دیگر نکن. پاسخ دانشی را مستقیم و دقیق بده؛ عدد، شرط، قاعده و استثنا را حفظ کن."},
                {"role":"user","content":f"سؤال: {question}\n\nمتن منبع:\n{factual_blob}"},
            ]
            rescued,rescue_usage,rescue_finish=await _generate_ai_text(rescue_messages,max_tokens=CHAT_REDIRECT_RETRY_MAX_TOKENS,temperature=0.0,route=route)
            _add_usage(total_usage,rescue_usage)
            if rescue_finish!='length' and rescued.strip() and not is_navigation_only_content(rescued):
                answer=apply_verified_calculations(format_answer_for_mode(rescued,detailed))
        if is_navigation_only_content(answer):
            answer='پاسخ این پرسش در منابع موجود پیدا نشد.'
    if detailed and AI_ENFORCE_DETAILED_ANSWER and _should_expand_answer(question,answer,selected):
        revision_messages = messages + [
            {"role":"assistant","content":answer},
            {"role":"user","content":"فقط نکات مرتبطی را که از منابع جا افتاده اضافه کن؛ از تکرار و رونویسی کامل منبع خودداری کن."},
        ]
        revised,revised_usage,revised_finish = await _generate_ai_text(
            revision_messages,max_tokens=max_tokens,temperature=active_temperature(),route=route
        )
        _add_usage(total_usage,revised_usage)
        if revised_finish != 'length' and not _looks_incomplete(revised):
            answer = apply_verified_calculations(format_answer_for_mode(revised,True))
    confidence,confidence_parts=evidence_confidence(query_plan,selected,rule_map,verification_status=verification_status)
    total_usage['evidence_confidence']=confidence
    total_usage['confidence_parts']=confidence_parts
    total_usage['deep_query_plan']=query_plan.as_dict()
    total_usage['rule_map_summary']={key:len(rule_map.get(key) or []) for key in ('base_rules','conditions','exceptions','limits','numeric_facts')}
    base_confidence_threshold=DEEP_CONFIDENCE_COMPLEX_MIN if query_plan.complexity>=0.58 else DEEP_CONFIDENCE_MIN
    confidence_threshold,calibration_meta=calibrated_confidence_threshold(base_confidence_threshold)
    total_usage['confidence_threshold']=confidence_threshold
    total_usage['confidence_calibration']=calibration_meta
    training_authoritative=bool(training_items)
    if DEEP_CONFIDENCE_GATE_ENABLED and not training_authoritative and confidence < confidence_threshold:
        total_usage['confidence_gated']=True
        total_usage['confidence_gate_reason']='insufficient_evidence'
        return 'برای این پرسش شواهد کافی و قابل‌اتکایی در منابع فعال پیدا نشد؛ برای جلوگیری از پاسخ حدسی، پاسخ قطعی ارائه نمی‌کنم.',total_usage
    total_usage['confidence_gated']=False
    return answer.strip(),total_usage


async def ask_ai_without_sources(question: str, detailed: bool = False, route: str = 'advanced', memory: str = '') -> tuple[str, dict[str, Any]]:
    if not SELF_ANALYSIS_ENABLED:
        return "پاسخ این پرسش در سؤالات متداول، کش، آموزش‌ها و منابع موجود نیست.", _empty_usage()
    response_rule = (
        "کاربر توضیح کامل خواسته است؛ تحلیل را مرحله‌بندی‌شده، روشن و بدون ادعای قطعیت بی‌دلیل بنویس."
        if detailed else
        "تحلیل را در یک یا حداکثر دو جمله مستقیم و کاربردی ارائه کن."
    )
    system_prompt = f"""شما دستیار سازمانی بارسان هستید. هیچ پاسخ مرتبطی در داده‌های سازمان پیدا نشده است.
{response_rule}
از دانش عمومی خود تحلیل کن، مرز اطمینان و فرض‌ها را رعایت کن، فارسی روان بنویس و متن را کامل تمام کن.
عبارت «تحلیل خودم /» را ننویس؛ سامانه آن را اضافه می‌کند.""".strip()
    user_content=question + (f"\n\nزمینه کوتاه مکالمه قبلی:\n{memory}" if memory else '')
    messages=[{"role":"system","content":system_prompt},{"role":"user","content":user_content}]
    max_tokens=active_max_completion_tokens(detailed)
    answer,usage,finish=await _generate_ai_text(messages,max_tokens=max_tokens,temperature=active_temperature(),route=route)
    rounds=0
    while (finish=='length' or _looks_incomplete(answer)) and rounds<AI_CONTINUATION_ROUNDS:
        rounds+=1
        continuation_messages=[{"role":"system","content":"فقط جمله نیمه‌تمام را ادامه بده؛ متن قبلی را تکرار نکن."},{"role":"user","content":f"سؤال: {question}\nانتهای پاسخ: {answer[-1200:]}"}]
        extra,extra_usage,finish=await _generate_ai_text(continuation_messages,max_tokens=min(max_tokens,700),temperature=active_temperature(),route=route)
        answer=_merge_continuation(answer,extra);_add_usage(usage,extra_usage)
    if finish=='length' or _looks_incomplete(answer):
        raise HTTPException(status_code=502,detail='تحلیل مدل کامل نشد؛ دوباره تلاش کنید.')
    body=format_answer_for_mode(answer,detailed)
    return f"تحلیل خودم / {body}",usage



RESPONSE_SECTION_META: dict[str, dict[str, Any]] = {
    'faq': {
        'label': 'سؤالات متداول',
        'description': 'پاسخ قطعی از بانک FAQ؛ بدون فراخوانی هوش مصنوعی و بدون مصرف توکن.',
        'zero_token_expected': True,
    },
    'cache': {
        'label': 'کش پاسخ‌های قبلی',
        'description': 'پاسخ ذخیره‌شده سؤال تکراری؛ بدون مراجعه دوباره به مدل و منابع.',
        'zero_token_expected': True,
    },
    'training': {
        'label': 'آموزش مدیر',
        'description': 'پاسخ با اولویت آموزش‌های فعال مدیر یا ادمین تولید شده است.',
        'zero_token_expected': False,
    },
    'resources': {
        'label': 'منابع بارگذاری‌شده',
        'description': 'پاسخ بر اساس قطعات مرتبط فایل‌ها و جزوات سازمان تولید شده است.',
        'zero_token_expected': False,
    },
    'self_analysis': {
        'label': 'تحلیل خود مدل',
        'description': 'در FAQ، کش، آموزش و منابع پاسخ معتبری پیدا نشده و مدل تحلیل مستقل ارائه کرده است.',
        'zero_token_expected': False,
    },
    'knowledge_gap': {
        'label': 'بدون پاسخ سازمانی',
        'description': 'پاسخ معتبر سازمانی پیدا نشده یا تحلیل مستقل غیرفعال بوده است.',
        'zero_token_expected': True,
    },
    'other': {
        'label': 'سایر',
        'description': 'وضعیت قدیمی یا خارج از مسیرهای استاندارد پاسخ‌گویی.',
        'zero_token_expected': False,
    },
}


def _safe_sources(value: Any) -> list[dict[str, Any]]:
    if isinstance(value, list):
        return [item for item in value if isinstance(item, dict)]
    if isinstance(value, str):
        try:
            parsed=json.loads(value or '[]')
            return [item for item in parsed if isinstance(item, dict)] if isinstance(parsed,list) else []
        except (TypeError,json.JSONDecodeError):
            return []
    return []


def response_section_key(status: str | None, model: str | None = None, sources: Any = None) -> str:
    normalized=(status or '').strip().lower()
    model_name=(model or '').strip().lower()
    source_types={str(item.get('source_type') or '').strip().lower() for item in _safe_sources(sources)}
    if normalized=='faq_hit' or 'faq-zero-token' in model_name or 'faq' in source_types:
        return 'faq'
    if normalized=='cache_hit' or 'answer-cache' in model_name:
        return 'cache'
    if normalized=='training_answer' or 'training' in model_name or 'training' in source_types:
        return 'training'
    if normalized in {'self_analysis'} or 'analysis' in source_types:
        return 'self_analysis'
    if normalized in {'knowledge_gap','unanswered','escalated'}:
        return 'knowledge_gap'
    if normalized=='answered' or 'document' in source_types:
        return 'resources'
    return 'other'


def response_source_names(sources: Any, limit: int = 4) -> list[str]:
    names=[]
    for item in _safe_sources(sources):
        raw_name=str(item.get('file_name') or '').strip()
        name=re.sub(r'[\u200e\u200f\u202a-\u202e\u2066-\u2069]', '', raw_name)
        name=re.sub(r'\s+', ' ', name).strip()
        page=item.get('page_start')
        if name and page:
            name=f"{name} — صفحه {int(page)}"
        if name and name not in names:
            names.append(name)
        if len(names)>=limit:
            break
    return names


def response_section_details(status: str | None, model: str | None = None, sources: Any = None, total_tokens: int = 0) -> dict[str, Any]:
    key=response_section_key(status,model,sources)
    meta=RESPONSE_SECTION_META[key]
    source_names=response_source_names(sources)
    return {
        'response_section':key,
        'response_section_label':meta['label'],
        'response_section_description':meta['description'],
        'response_source_names':source_names,
        'zero_token':int(total_tokens or 0)==0,
    }


def response_source_stats(days: int = 30) -> dict[str, Any]:
    days=max(1,min(int(days),3650))
    start_iso=(datetime.now(timezone.utc)-timedelta(days=days-1)).isoformat()
    with get_db() as db:
        rows=db.execute(
            """SELECT status,model,sources_json,total_tokens,response_ms
               FROM messages WHERE role='assistant' AND created_at>=?""",
            (start_iso,),
        ).fetchall()
    aggregates={key:{
        'key':key,'label':meta['label'],'description':meta['description'],
        'value':0,'percentage':0.0,'total_tokens':0,'zero_token_count':0,
        'zero_token_percentage':0.0,'avg_response_ms':0,
    } for key,meta in RESPONSE_SECTION_META.items()}
    response_ms_sum={key:0 for key in aggregates}
    response_ms_count={key:0 for key in aggregates}
    for row in rows:
        key=response_section_key(row['status'],row['model'],row['sources_json'])
        bucket=aggregates[key]
        bucket['value']+=1
        tokens=int(row['total_tokens'] or 0)
        bucket['total_tokens']+=tokens
        if tokens==0:
            bucket['zero_token_count']+=1
        if row['response_ms'] is not None:
            response_ms_sum[key]+=int(row['response_ms'] or 0)
            response_ms_count[key]+=1
    total=sum(item['value'] for item in aggregates.values())
    zero_token_total=sum(item['zero_token_count'] for item in aggregates.values())
    ordered=[]
    for key in ('faq','cache','training','resources','self_analysis','knowledge_gap','other'):
        item=aggregates[key]
        item['percentage']=round((item['value']/total*100) if total else 0.0,1)
        item['zero_token_percentage']=round((item['zero_token_count']/item['value']*100) if item['value'] else 0.0,1)
        item['avg_response_ms']=round(response_ms_sum[key]/response_ms_count[key]) if response_ms_count[key] else 0
        if item['value']:
            ordered.append(item)
    return {
        'days':days,
        'total_answers':total,
        'zero_token_answers':zero_token_total,
        'zero_token_percentage':round((zero_token_total/total*100) if total else 0.0,1),
        'ai_assisted_answers':max(0,total-zero_token_total),
        'ai_assisted_percentage':round(((total-zero_token_total)/total*100) if total else 0.0,1),
        'items':ordered,
    }

def provider_usage_dashboard(days: int = 30) -> dict[str, Any]:
    days=max(1,min(int(days),3650))
    start=(datetime.now(timezone.utc)-timedelta(days=days-1)).isoformat()
    with get_db() as db:
        slot_rows=db.execute("""SELECT COALESCE(api_slot,0) api_slot,COALESCE(provider_label,'') provider_label,
            SUM(prompt_tokens) prompt_tokens,SUM(output_tokens) output_tokens,SUM(total_tokens) total_tokens,
            SUM(estimated_cost) estimated_cost,COUNT(*) request_count,
            CAST(COALESCE(AVG(response_ms),0) AS INTEGER) avg_response_ms
            FROM usage_events WHERE created_at>=? GROUP BY COALESCE(api_slot,0),COALESCE(provider_label,'')""",(start,)).fetchall()
        route_rows=db.execute("""SELECT COALESCE(model_route,'unknown') model_route,COUNT(*) request_count,
            SUM(total_tokens) total_tokens,SUM(estimated_cost) estimated_cost
            FROM usage_events WHERE created_at>=? GROUP BY COALESCE(model_route,'unknown') ORDER BY request_count DESC""",(start,)).fetchall()
        daily=db.execute("""SELECT substr(created_at,1,10) day,SUM(total_tokens) total_tokens,SUM(estimated_cost) estimated_cost,COUNT(*) request_count
            FROM usage_events WHERE created_at>=? GROUP BY substr(created_at,1,10) ORDER BY day""",(start,)).fetchall()
        attempt_rows=db.execute("""SELECT COALESCE(slot,0) slot,COUNT(*) attempts,
            SUM(CASE WHEN status='success' THEN 1 ELSE 0 END) successful_attempts,
            SUM(CASE WHEN status!='success' THEN 1 ELSE 0 END) failed_attempts,
            SUM(CASE WHEN was_failover=1 AND status='success' THEN 1 ELSE 0 END) failover_successes,
            SUM(CASE WHEN status='limited' THEN 1 ELSE 0 END) limit_hits
            FROM api_call_events WHERE created_at>=? GROUP BY COALESCE(slot,0)""",(start,)).fetchall()
    by_slot={int(row['api_slot']):dict(row) for row in slot_rows}
    attempts_by_slot={int(row['slot']):dict(row) for row in attempt_rows}
    configs={int(item['slot']):item for item in configured_ai_slots()}
    pool=api_pool_public_status();states={int(item['slot']):item for item in pool['slots']}
    items=[]
    relevant_slots=sorted(set(configs) | {k for k in by_slot if k>0} | {k for k in attempts_by_slot if k>0})
    for slot in relevant_slots:
        cfg=configs.get(slot);row=by_slot.get(slot,{});attempt=attempts_by_slot.get(slot,{})
        spent=round(float(row.get('estimated_cost') or 0.0),8)
        credit=float(cfg.get('credit_amount') or 0.0) if cfg else 0.0
        items.append({
            'slot':slot,'configured':bool(cfg),'provider':cfg.get('label') if cfg else None,
            'model':cfg.get('model') if cfg else None,'status':states.get(slot,{}).get('status','empty'),
            'prompt_tokens':int(row.get('prompt_tokens') or 0),'output_tokens':int(row.get('output_tokens') or 0),
            'total_tokens':int(row.get('total_tokens') or 0),'request_count':int(row.get('request_count') or 0),
            'avg_response_ms':int(row.get('avg_response_ms') or 0),'estimated_cost':spent,
            'configured_credit':credit,'estimated_remaining_credit':round(max(0.0,credit-spent),8) if credit>0 else None,
            'currency':cfg.get('credit_currency') if cfg else None,
            'success_count':int(states.get(slot,{}).get('success_count') or 0),
            'failure_count':int(states.get(slot,{}).get('failure_count') or 0),
            'attempts':int(attempt.get('attempts') or 0),'successful_attempts':int(attempt.get('successful_attempts') or 0),
            'failed_attempts':int(attempt.get('failed_attempts') or 0),'failover_successes':int(attempt.get('failover_successes') or 0),
            'limit_hits':int(attempt.get('limit_hits') or 0),
            'blocked_until':states.get(slot,{}).get('blocked_until'),
        })
    zero=by_slot.get(0,{})
    return {
        'days':days,'items':items,
        'zero_token_events':int(zero.get('request_count') or 0),
        'total_tokens':sum(item['total_tokens'] for item in items),
        'estimated_cost':round(sum(item['estimated_cost'] for item in items),8),
        'provider_attempts':sum(item['attempts'] for item in items),
        'provider_failures':sum(item['failed_attempts'] for item in items),
        'successful_failovers':sum(item['failover_successes'] for item in items),
        'limit_hits':sum(item['limit_hits'] for item in items),
        'routes':[dict(row) for row in route_rows],
        'daily':[dict(row) for row in daily],
        'credit_note':'اعتبار نمایش‌داده‌شده بر اساس مقدار تنظیم‌شده در Variables و هزینه تخمینی هر یک میلیون توکن است؛ موجودی زنده فقط در صورت ارائه API استاندارد توسط سرویس‌دهنده قابل دریافت است.',
    }


def feedback_quality_dashboard(days: int = 30) -> dict[str, Any]:
    days=max(1,min(int(days),3650));start=(datetime.now(timezone.utc)-timedelta(days=days-1)).isoformat()
    with get_db() as db:
        rows=db.execute("SELECT rating,COUNT(*) value FROM answer_feedback WHERE created_at>=? GROUP BY rating",(start,)).fetchall()
    counts={row['rating']:int(row['value']) for row in rows};total=sum(counts.values())
    positive=counts.get('correct',0)
    return {'days':days,'total':total,'correct_percentage':round(positive/total*100,1) if total else 0.0,'items':counts}


def public_source_items(sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    result=[]
    for source in sources:
        result.append({
            'source_type':source.get('source_type','document'),
            'document_id':source.get('document_id') or source.get('training_id') or source.get('faq_id') or 'unknown',
            'file_name':source.get('file_name') or 'منبع',
            'chunk_index':int(source.get('chunk_index') or 0),
            'score':float(source.get('score') or 0),
            'confidence':source.get('confidence'),
            'lexical_score':float(source.get('lexical_score') or 0),
            'semantic_score':float(source.get('semantic_score') or 0),
            'embedding_score':float(source.get('embedding_score') or 0),
            'rerank_score':float(source.get('rerank_score') or source.get('score') or 0),
            'page_start':source.get('page_start'),
            'page_end':source.get('page_end'),
            'section_title':sanitize_answer_text(str(source.get('section_title') or ''))[:240] or None,
            'chunk_type':source.get('chunk_type'),
            'conflict_detected':bool(source.get('conflict_detected')),
            'excerpt':sanitize_answer_text(source.get('excerpt') or source.get('content') or source.get('answer') or '')[:500],
        })
    return result


def needs_conversation_memory(question: str) -> bool:
    normalized=normalize_text(question)
    markers=('این','آن','اون','مورد قبل','مورد دوم','ادامه','بیشتر توضیح','همون','همین','قبلی','منظورت')
    return CONVERSATION_MEMORY_ENABLED and (len(question.split())<=8 or any(normalize_text(x) in normalized for x in markers))


def conversation_memory(conversation_id: str | None, question: str) -> str:
    if not conversation_id or not needs_conversation_memory(question) or CONVERSATION_MEMORY_MESSAGES<=0:
        return ''
    with get_db() as db:
        rows=db.execute(
            "SELECT role,content FROM messages WHERE conversation_id=? ORDER BY created_at DESC LIMIT ?",
            (conversation_id,CONVERSATION_MEMORY_MESSAGES),
        ).fetchall()
    parts=[];used=0
    for row in reversed(rows):
        content=sanitize_answer_text(row['content'])
        remaining=CONVERSATION_MEMORY_MAX_CHARS-used
        if remaining<=0:break
        content=content[:remaining]
        parts.append(('کاربر' if row['role']=='user' else 'دستیار')+': '+content)
        used+=len(content)
    return '\n'.join(parts)


def save_message(
    conversation_id: str,
    role: str,
    content: str,
    sources: list[dict[str, Any]] | None = None,
    model: str | None = None,
    response_ms: int | None = None,
    status: str = "answered",
    usage: dict[str, Any] | None = None,
) -> str:
    message_id = str(uuid.uuid4())
    usage = usage or {}
    source_rows = sources or []
    confidence = float(usage.get('evidence_confidence')) if usage.get('evidence_confidence') is not None else max((float(item.get("score") or 0.0) for item in source_rows), default=0.0)
    with get_db() as db:
        db.execute(
            """
            INSERT INTO messages(
                id, conversation_id, role, content, sources_json, model,
                response_ms, status, prompt_tokens, output_tokens,
                total_tokens, source_count, api_slot, model_route,
                estimated_cost, confidence_score, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                message_id,
                conversation_id,
                role,
                content,
                json.dumps(source_rows, ensure_ascii=False),
                model,
                response_ms,
                status,
                int(usage.get("prompt_tokens", 0)),
                int(usage.get("output_tokens", 0)),
                int(usage.get("total_tokens", 0)),
                len(source_rows),
                int(usage.get("api_slot")) if usage.get("api_slot") is not None else None,
                str(usage.get("model_route") or ("zero_token" if role == "assistant" and not usage.get("total_tokens") else "")),
                float(usage.get("estimated_cost") or 0.0),
                confidence,
                now_iso(),
            ),
        )
        db.execute("UPDATE conversations SET updated_at = ? WHERE id = ?", (now_iso(), conversation_id))
    return message_id


async def process_operational_module(*, module: str, question: str, user: dict[str, Any]) -> dict[str, Any]:
    """Independent operational engine with shared authoritative knowledge.

    Knowledge order is immutable: Training -> Sources -> FAQ -> Cache. Operational
    modules do not use or create chat conversations, so their memory and state never
    leak into each other or into the main chat.
    """
    if module not in {'cargo','route','calc'}:
        raise HTTPException(status_code=404,detail='ماژول عملیاتی نامعتبر است.')
    if not system_setting_bool('chat_enabled', True):
        raise HTTPException(status_code=503,detail=system_setting('maintenance_message','سرویس پاسخ‌گویی موقتاً متوقف است.'))
    enforce_rate_limit(f"tool:{module}:user:{user['id']}")
    quota_reserved=False
    started=time.perf_counter()
    detailed=False
    try:
        quota_reserved=reserve_question_quota(user)
        if module=='calc':
            math_answer=direct_math_answer(question)
            if math_answer:
                usage=_empty_usage();duration_ms=int((time.perf_counter()-started)*1000)
                model_name='barsan-deterministic-calculator-zero-token'
                record_usage(user_id=user['id'],external_user_id=None,event_type='tool_calc',usage=usage,model=model_name,response_ms=duration_ms)
                return {'module':module,'answer':math_answer,'sources':[],'model':model_name,'response_ms':duration_ms,'usage':usage,'status':'calculator','response_section':'other','response_section_label':'محاسبه قطعی','response_source_names':[],'confidence_score':1.0,'cached':False,'model_route':'deterministic_math','zero_token':True,'quota':question_quota_snapshot(user['id'])}

        authoritative_training=_filter_navigation_only_items(retrieve_training(question,user,False))
        exact_cached=None if authoritative_training else find_exact_cached_answer(question,detailed,user,False)
        if exact_cached:
            duration_ms=int((time.perf_counter()-started)*1000);usage=_empty_usage()
            answer=format_answer_for_mode(exact_cached['answer'],detailed);public_sources=exact_cached['sources']
            status='cache_hit';model_name='barsan-tool-exact-cache-fast-path-zero-token';route='exact_cache_fast_path';cached_flag=True
            record_usage(user_id=user['id'],external_user_id=None,event_type=f'tool_{module}_{status}',usage=usage,model=model_name,response_ms=duration_ms)
            details=response_section_details(status,model_name,public_sources,0)
            return {'module':module,'answer':answer,'sources':public_sources,'model':model_name,'response_ms':duration_ms,'usage':usage,'status':status,'response_section':response_section_key(status,model_name,public_sources),'response_section_label':details['response_section_label'],'response_source_names':response_source_names(public_sources),'confidence_score':max([float(x.get('score') or 0) for x in public_sources] or [1.0]),'cached':True,'model_route':route,'zero_token':True,'quota':question_quota_snapshot(user['id'])}

        stage,sources,deep_meta=await retrieve_deep_priority_stage_async(question,user,integration=False)
        if sources:
            cached=find_cached_answer(question,detailed,user,False)
            if cached_answer_matches_stage(cached,stage):
                answer=format_answer_for_mode(cached['answer'],detailed);public_sources=cached['sources'];usage=_empty_usage()
                status='cache_hit';model_name='barsan-tool-stage-cache-zero-token';route='zero_token_stage_safe';cached_flag=True
            else:
                route=select_model_route(question,stage,sources,detailed)
                answer,usage=await ask_ai(question,sources,detailed=detailed,route=route,memory='')
                status='knowledge_gap' if usage.get('confidence_gated') else ('training_answer' if stage=='training' else 'answered')
                if usage.get('confidence_gated'):
                    record_knowledge_gap(question,user['id'])
                public_sources=public_source_items(sources)
                model_name='barsan-training-direct-zero-token' if stage=='training' and usage['total_tokens']==0 else str(usage.get('model') or active_model_name())
                cached_flag=False
                if not usage.get('confidence_gated'):
                    store_cached_answer(question,detailed,user,False,answer,public_sources,model_name,cache_tier='approved',approved_by=user['id'] if user.get('role') in {'manager','admin'} else None)
        else:
            faq=find_faq_answer(question,user,False)
            if faq:
                answer=format_answer_for_mode(faq['answer'],detailed);public_sources=[faq['source']];usage=_empty_usage()
                status='faq_hit';model_name='barsan-tool-faq-fallback-zero-token';route='zero_token_fallback';cached_flag=False
            else:
                cached=find_cached_answer(question,detailed,user,False)
                if cached:
                    answer=format_answer_for_mode(cached['answer'],detailed);public_sources=cached['sources'];usage=_empty_usage()
                    status='cache_hit';model_name='barsan-tool-cache-fallback-zero-token';route='zero_token_fallback';cached_flag=True
                else:
                    record_knowledge_gap(question,user['id'])
                    answer='پاسخ دقیق این مورد در آموزش‌ها، منابع یا سؤالات متداول موجود پیدا نشد.'
                    public_sources=[];usage=_empty_usage();status='knowledge_gap';model_name='barsan-source-first-no-match';route='source_first_no_match';cached_flag=False

        duration_ms=int((time.perf_counter()-started)*1000)
        record_usage(user_id=user['id'],external_user_id=None,event_type=f'tool_{module}_{status}',usage=usage,model=model_name,response_ms=duration_ms)
        details=response_section_details(status,model_name,public_sources,usage.get('total_tokens',0))
        return {'module':module,'answer':answer,'sources':public_sources,'model':model_name,'response_ms':duration_ms,'usage':usage,'status':status,'response_section':response_section_key(status,model_name,public_sources),'response_section_label':details['response_section_label'],'response_source_names':response_source_names(public_sources),'confidence_score':float(usage.get('evidence_confidence')) if usage.get('evidence_confidence') is not None else max([float(x.get('score') or 0) for x in public_sources] or [0.0]),'cached':cached_flag,'model_route':usage.get('model_route') or route,'api_slot':usage.get('api_slot'),'estimated_cost':usage.get('estimated_cost',0.0),'zero_token':usage.get('total_tokens',0)==0,'quota':question_quota_snapshot(user['id'])}
    except Exception:
        refund_question_quota(user,quota_reserved)
        raise


async def process_chat(
    *,
    message: str,
    conversation_id: str | None,
    user: dict[str, Any] | None,
    external_user_id: str | None,
    integration: bool,
    rate_identity: str | None = None,
    rate_per_minute: int | None = None,
    rate_daily: int = 0,
    rate_monthly: int = 0,
) -> dict[str, Any]:
    if not system_setting_bool("chat_enabled", True):
        raise HTTPException(status_code=503,detail=system_setting("maintenance_message", "سرویس پاسخ‌گویی موقتاً متوقف است."))
    identity = rate_identity or (f"user:{user['id']}" if user else f"external:{external_user_id}")
    enforce_rate_limit(identity, per_minute=rate_per_minute, daily_limit=rate_daily, monthly_limit=rate_monthly)
    quota_reserved=False
    try:
        quota_reserved=reserve_question_quota(user)
        if conversation_id:
            ensure_conversation_access(conversation_id,user,external_user_id)
        else:
            conversation_id = create_conversation(user["id"] if user else None,external_user_id,message[:80])
        memory=conversation_memory(conversation_id,message)
        save_message(conversation_id,"user",message,status="asked")
        started = time.perf_counter()
        detailed = is_detailed_request(message)

        # Deterministic arithmetic route. Pure calculations never rely on a language
        # model, so addition/subtraction/multiplication/division are reproducible.
        math_answer=direct_math_answer(message)
        if math_answer:
            usage=_empty_usage();duration_ms=int((time.perf_counter()-started)*1000)
            model_name='barsan-deterministic-calculator-zero-token'
            assistant_message_id=save_message(conversation_id,'assistant',math_answer,[],model_name,duration_ms,'calculator',usage)
            record_usage(user_id=user['id'] if user else None,external_user_id=external_user_id,event_type='calculator',usage=usage,model=model_name,response_ms=duration_ms)
            return {'conversation_id':conversation_id,'assistant_message_id':assistant_message_id,'answer':math_answer,'sources':[],'model':model_name,'response_ms':duration_ms,'usage':usage,'status':'calculator','response_section':'other','response_section_label':'محاسبه قطعی','response_source_names':[],'confidence_score':1.0,'cached':False,'cache_tier':'approved','model_route':'deterministic_math','zero_token':True,'detailed':detailed,'quota':question_quota_snapshot(user['id'] if user else None)}

        # R30 exact-cache fast path. It is safe before retrieval because the cache
        # signature contains knowledge_version; a new/changed manager training or
        # document automatically invalidates every older answer.
        authoritative_training=_filter_navigation_only_items(retrieve_training(message,user,integration))
        exact_cached=None if authoritative_training else find_exact_cached_answer(message,detailed,user,integration)
        if exact_cached:
            duration_ms=int((time.perf_counter()-started)*1000);usage=_empty_usage()
            answer=format_answer_for_mode(exact_cached['answer'],detailed);public_sources=exact_cached['sources']
            model_name='barsan-exact-cache-fast-path-zero-token'
            assistant_message_id=save_message(conversation_id,'assistant',answer,public_sources,model_name,duration_ms,'cache_hit',usage)
            record_usage(user_id=user['id'] if user else None,external_user_id=external_user_id,event_type='cache_hit',usage=usage,model=model_name,response_ms=duration_ms)
            return {'conversation_id':conversation_id,'assistant_message_id':assistant_message_id,'answer':answer,'sources':public_sources,'model':model_name,'response_ms':duration_ms,'usage':usage,'status':'cache_hit','response_section':'cache','response_section_label':RESPONSE_SECTION_META['cache']['label'],'response_source_names':response_source_names(public_sources),'confidence_score':max([float(x.get('score') or 0) for x in public_sources] or [1.0]),'cached':True,'cache_tier':exact_cached.get('cache_tier','approved'),'model_route':'exact_cache_fast_path','zero_token':True,'detailed':detailed,'quota':question_quota_snapshot(user['id'] if user else None)}

        # Knowledge priority is strict: manager training first, then organizational sources.
        stage,sources,deep_meta=await retrieve_deep_priority_stage_async(message,user,integration=integration)
        if sources:
            # Cache may accelerate a repeated question only AFTER the authoritative
            # layer has been resolved. It therefore cannot outrank training or sources.
            cached=find_cached_answer(message,detailed,user,integration)
            if cached_answer_matches_stage(cached,stage):
                duration_ms=int((time.perf_counter()-started)*1000);usage=_empty_usage()
                answer=format_answer_for_mode(cached['answer'],detailed);public_sources=cached['sources']
                model_name='barsan-answer-cache-stage-safe-zero-token'
                assistant_message_id=save_message(conversation_id,'assistant',answer,public_sources,model_name,duration_ms,'cache_hit',usage)
                record_usage(user_id=user['id'] if user else None,external_user_id=external_user_id,event_type='cache_hit',usage=usage,model=model_name,response_ms=duration_ms)
                return {'conversation_id':conversation_id,'assistant_message_id':assistant_message_id,'answer':answer,'sources':public_sources,'model':model_name,'response_ms':duration_ms,'usage':usage,'status':'cache_hit','response_section':'cache','response_section_label':RESPONSE_SECTION_META['cache']['label'],'response_source_names':response_source_names(public_sources),'confidence_score':max([float(x.get('score') or 0) for x in public_sources] or [1.0]),'cached':True,'cache_tier':cached.get('cache_tier','approved'),'model_route':'zero_token_stage_safe','zero_token':True,'detailed':detailed,'quota':question_quota_snapshot(user['id'] if user else None)}
            model_route=select_model_route(message,stage,sources,detailed)
            answer,usage=await ask_ai(message,sources,detailed=detailed,route=model_route,memory=memory)
            answer_status='knowledge_gap' if usage.get('confidence_gated') else ('training_answer' if stage=='training' else 'answered')
            if usage.get('confidence_gated'):
                record_knowledge_gap(message,user['id'] if user else None)
            public_sources=public_source_items(sources)
            if stage=='training' and usage['total_tokens']==0:
                model_name='barsan-training-direct-zero-token'
            else:
                model_name=str(usage.get('model') or active_model_name())
        else:
            # Only after training AND organizational sources fail do we consider FAQ
            # and cache. They can no longer override manager training or a document.
            faq=find_faq_answer(message,user,integration)
            if faq:
                answer=format_answer_for_mode(faq['answer'],detailed)
                usage=_empty_usage();duration_ms=int((time.perf_counter()-started)*1000)
                public_sources=[faq['source']];model_name='barsan-faq-fallback-zero-token'
                assistant_message_id=save_message(conversation_id,'assistant',answer,public_sources,model_name,duration_ms,'faq_hit',usage)
                record_usage(user_id=user['id'] if user else None,external_user_id=external_user_id,event_type='faq_hit',usage=usage,model=model_name,response_ms=duration_ms)
                return {'conversation_id':conversation_id,'assistant_message_id':assistant_message_id,'answer':answer,'sources':public_sources,'model':model_name,'response_ms':duration_ms,'usage':usage,'status':'faq_hit','response_section':'faq','response_section_label':RESPONSE_SECTION_META['faq']['label'],'response_source_names':response_source_names(public_sources),'confidence_score':float(faq.get('score') or 1.0),'cached':False,'cache_tier':'approved','model_route':'zero_token_fallback','zero_token':True,'detailed':detailed,'quota':question_quota_snapshot(user['id'] if user else None)}

            cached=find_cached_answer(message,detailed,user,integration)
            if cached:
                duration_ms=int((time.perf_counter()-started)*1000);usage=_empty_usage()
                answer=format_answer_for_mode(cached['answer'],detailed);public_sources=cached['sources']
                model_name='barsan-answer-cache-fallback-zero-token'
                assistant_message_id=save_message(conversation_id,'assistant',answer,public_sources,model_name,duration_ms,'cache_hit',usage)
                record_usage(user_id=user['id'] if user else None,external_user_id=external_user_id,event_type='cache_hit',usage=usage,model=model_name,response_ms=duration_ms)
                return {'conversation_id':conversation_id,'assistant_message_id':assistant_message_id,'answer':answer,'sources':public_sources,'model':model_name,'response_ms':duration_ms,'usage':usage,'status':'cache_hit','response_section':'cache','response_section_label':RESPONSE_SECTION_META['cache']['label'],'response_source_names':response_source_names(public_sources),'confidence_score':max([float(x.get('score') or 0) for x in public_sources] or [1.0]),'cached':True,'cache_tier':cached.get('cache_tier','approved'),'model_route':'zero_token_fallback','zero_token':True,'detailed':detailed,'quota':question_quota_snapshot(user['id'] if user else None)}

            # Strict source-first mode avoids confident guesses for organizational
            # questions. Independent analysis is available only when explicitly asked.
            record_knowledge_gap(message,user['id'] if user else None)
            model_route='advanced'
            if SOURCE_FIRST_STRICT and not _question_requests_self_analysis(message):
                answer='پاسخ دقیق این پرسش در منابع، تصاویر PDF، آموزش‌ها یا سؤالات متداول موجود پیدا نشد.'
                usage=_empty_usage();answer_status='knowledge_gap'
                public_sources=[];model_name='barsan-source-first-no-match'
            else:
                answer,usage=await ask_ai_without_sources(message,detailed=detailed,route=model_route,memory=memory)
                answer_status='self_analysis' if SELF_ANALYSIS_ENABLED else 'knowledge_gap'
                public_sources=[{'source_type':'analysis','document_id':'analysis:self','file_name':'تحلیل مدل بدون منبع سازمانی','chunk_index':0,'score':0.0,'excerpt':sanitize_answer_text(answer)[:500]}]
                model_name=str(usage.get('model') or active_model_name()) if usage['total_tokens'] else 'barsan-no-source'

        duration_ms=int((time.perf_counter()-started)*1000)
        assistant_message_id=save_message(conversation_id,'assistant',answer,public_sources,model_name,duration_ms,answer_status,usage)
        record_usage(user_id=user['id'] if user else None,external_user_id=external_user_id,event_type=answer_status,usage=usage,model=model_name,response_ms=duration_ms)
        if answer_status in {'answered','training_answer'}:
            store_cached_answer(message,detailed,user,integration,answer,public_sources,model_name,cache_tier='approved',approved_by=user['id'] if user and user.get('role') in {'manager','admin'} else None)
        elif answer_status=='self_analysis':
            store_cached_answer(message,detailed,user,integration,answer,public_sources,model_name,cache_tier='temporary')
        return {'conversation_id':conversation_id,'assistant_message_id':assistant_message_id,'answer':answer,'sources':public_sources,'model':model_name,'response_ms':duration_ms,'usage':usage,'status':answer_status,'response_section':response_section_key(answer_status,model_name,public_sources),'response_section_label':response_section_details(answer_status,model_name,public_sources,usage['total_tokens'])['response_section_label'],'response_source_names':response_source_names(public_sources),'confidence_score':float(usage.get('evidence_confidence')) if usage.get('evidence_confidence') is not None else max([float(x.get('score') or 0) for x in public_sources] or [0.0]),'cached':False,'cache_tier':'temporary' if answer_status=='self_analysis' else 'approved','model_route':usage.get('model_route') or model_route,'api_slot':usage.get('api_slot'),'estimated_cost':usage.get('estimated_cost',0.0),'zero_token':usage['total_tokens']==0,'detailed':detailed,'quota':question_quota_snapshot(user['id'] if user else None)}
    except Exception:
        refund_question_quota(user,quota_reserved)
        raise


@app.get('/api/v1/faqs')
def list_faqs(q: str | None = None, user: dict[str,Any] = Depends(require_roles('manager','admin'))) -> list[dict[str,Any]]:
    sql="""SELECT f.*,cu.username creator_username,uu.username updater_username
    FROM faqs f LEFT JOIN users cu ON cu.id=f.created_by LEFT JOIN users uu ON uu.id=f.updated_by"""
    params=[]
    if q:
        sql += " WHERE f.question LIKE ? OR f.answer LIKE ? OR f.aliases_json LIKE ?"
        like=f"%{q.strip()}%";params=[like,like,like]
    sql += " ORDER BY f.is_active DESC,f.priority DESC,f.updated_at DESC LIMIT ?"
    params.append(FAQ_MAX_ROWS)
    with get_db() as db:rows=db.execute(sql,params).fetchall()
    result=[]
    for row in rows:
        item=dict(row)
        try:item['aliases']=json.loads(item.pop('aliases_json') or '[]')
        except json.JSONDecodeError:item['aliases']=[]
        item['allowed_roles']=_json_list(item.pop('allowed_roles_json',None))
        item['allowed_user_ids']=_json_list(item.pop('allowed_user_ids_json',None))
        item['is_active']=bool(item['is_active']);result.append(item)
    return result


@app.post('/api/v1/faqs')
def create_faq(data: FaqInput, user: dict[str,Any] = Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    result=upsert_faq_rows([data.model_dump()],user['id'])
    audit(user['id'],'faq.upserted',{'question':data.question,'result':result})
    return result


@app.patch('/api/v1/faqs/{faq_id}')
def update_faq(faq_id: str, data: FaqUpdateInput, user: dict[str,Any] = Depends(require_roles('manager','admin'))) -> dict[str,bool]:
    with get_db() as db:
        row=db.execute("SELECT * FROM faqs WHERE id=?",(faq_id,)).fetchone()
        if not row:raise HTTPException(status_code=404,detail='سؤال متداول پیدا نشد.')
        current=dict(row)
        question=sanitize_answer_text(data.question) if data.question is not None else current['question']
        answer=sanitize_answer_text(data.answer) if data.answer is not None else current['answer']
        aliases=normalize_faq_aliases(data.aliases) if data.aliases is not None else json.loads(current['aliases_json'] or '[]')
        priority=100
        active=int(data.is_active) if data.is_active is not None else int(current['is_active'])
        visibility='public'
        allowed_roles=[]
        allowed_user_ids=[]
        department=None
        normalized=canonical_question_for_cache(question)
        try:
            db.execute("""UPDATE faqs SET question=?,normalized_question=?,answer=?,aliases_json=?,priority=?,is_active=?,visibility=?,allowed_roles_json=?,allowed_user_ids_json=?,department=?,updated_by=?,updated_at=? WHERE id=?""",
                (question,normalized,answer,json.dumps(aliases,ensure_ascii=False),priority,active,visibility,json.dumps(allowed_roles,ensure_ascii=False),json.dumps(allowed_user_ids),department,user['id'],now_iso(),faq_id))
        except sqlite3.IntegrityError as exc:
            raise HTTPException(status_code=409,detail='سؤال مشابه دیگری در سؤالات متداول وجود دارد.') from exc
        bump_knowledge_version(db)
    audit(user['id'],'faq.updated',{'faq_id':faq_id})
    return {'ok':True}


@app.delete('/api/v1/faqs/{faq_id}')
def delete_faq(faq_id: str, user: dict[str,Any] = Depends(require_roles('manager','admin'))) -> dict[str,bool]:
    with get_db() as db:
        cur=db.execute("DELETE FROM faqs WHERE id=?",(faq_id,))
        if cur.rowcount!=1:raise HTTPException(status_code=404,detail='سؤال متداول پیدا نشد.')
        bump_knowledge_version(db)
    audit(user['id'],'faq.deleted',{'faq_id':faq_id})
    return {'ok':True}


@app.post('/api/v1/faqs/import')
async def import_faqs(file: UploadFile = File(...), user: dict[str,Any] = Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    payload=await file.read(10*1024*1024+1)
    if len(payload)>10*1024*1024:raise HTTPException(status_code=413,detail='حجم فایل سؤالات متداول باید حداکثر ۱۰ مگابایت باشد.')
    rows=parse_faq_import(file.filename or 'faqs.csv',payload)
    if not rows:raise HTTPException(status_code=400,detail='هیچ ردیف معتبر سؤال و پاسخ در فایل پیدا نشد.')
    result=upsert_faq_rows(rows,user['id'])
    audit(user['id'],'faq.imported',{'filename':file.filename,'rows':len(rows),**result})
    return {'filename':file.filename,'rows_read':len(rows),**result}


@app.get('/api/v1/faqs/template.csv')
def faq_template(user: dict[str,Any] = Depends(require_roles('manager','admin'))) -> StreamingResponse:
    out=io.StringIO();writer=csv.writer(out);writer.writerow(['question','answer','aliases','active']);writer.writerow(['هزینه لغو سرویس چقدر است؟','هزینه لغو سرویس ۵۰۰ هزار تومان است.','قیمت کنسلی|مبلغ لغو',1])
    payload=('\ufeff'+out.getvalue()).encode('utf-8')
    return StreamingResponse(io.BytesIO(payload),media_type='text/csv; charset=utf-8',headers={'Content-Disposition':'attachment; filename=barsan_faq_template.csv'})



def _transcription_model_for_slot(slot: int) -> str:
    config=_provider_config_for_slot(slot)
    return str(config.get('transcription_model') or '') if config else '' 


def _clean_transcript_text(value: str) -> str:
    text=sanitize_answer_text(value)
    words=text.split()
    compact=[]
    for word in words:
        token=re.sub(r'[،,.؟!؛:]','',word)
        previous=re.sub(r'[،,.؟!؛:]','',compact[-1]) if compact else ''
        if token and token==previous: continue
        compact.append(word)
    for size in range(6,1,-1):
        changed=True
        while changed:
            changed=False
            for index in range(size,len(compact)-size+1):
                if compact[index-size:index]==compact[index:index+size]:
                    del compact[index:index+size];changed=True;break
    return ' '.join(compact).strip()


async def transcribe_audio_bytes(filename: str, content_type: str, payload: bytes) -> tuple[str, dict[str, Any]]:
    errors=[]
    timeout=httpx.Timeout(AI_TIMEOUT_SECONDS,connect=min(20.0,AI_TIMEOUT_SECONDS),read=AI_TIMEOUT_SECONDS,write=AI_TIMEOUT_SECONDS,pool=20.0)
    client=_HTTP_CLIENT
    owns_client=False
    if client is None:
        client=httpx.AsyncClient(timeout=timeout,follow_redirects=True);owns_client=True
    try:
        slots=ordered_available_ai_slots()
        for attempt_index,slot in enumerate(slots,1):
            model=_transcription_model_for_slot(int(slot['slot']))
            if not model: continue
            headers={AI_AUTH_HEADER:f"{AI_AUTH_SCHEME} {slot['api_key']}".strip()}
            started=time.perf_counter()
            try:
                response=await client.post(
                    f"{slot['base_url']}/audio/transcriptions",
                    headers=headers,
                    data={'model':model,'language':'fa','response_format':'json','temperature':'0',
                          'prompt':'متن فارسی سازمانی را دقیق، بدون تکرار کلمات، با حفظ اعداد، نام‌ها و علائم نگارشی بنویس.'},
                    files={'file':(Path(filename or 'voice.webm').name,payload,content_type or 'application/octet-stream')},
                    timeout=timeout,
                )
                elapsed=int((time.perf_counter()-started)*1000)
                if response.status_code>=400:
                    detail=sanitize_answer_text(response.text)[:300]
                    record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=model,model_route='voice',status='error',response_ms=elapsed,attempt_index=attempt_index,was_failover=attempt_index>1,http_status=response.status_code,error_class='HTTPError',error_detail=detail)
                    if response.status_code in {401,402,403,429}:
                        mark_api_slot_blocked(int(slot['slot']),detail,AI_KEY_RATE_LIMIT_COOLDOWN_SECONDS if response.status_code==429 else AI_KEY_QUOTA_COOLDOWN_SECONDS,'limited')
                    elif response.status_code>=500:
                        mark_api_slot_blocked(int(slot['slot']),detail,AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'transient_error')
                    errors.append(f"API {slot['slot']}: HTTP {response.status_code}")
                    continue
                data=response.json()
                text=_clean_transcript_text(data.get('text') or data.get('transcript') or '')
                if text:
                    mark_api_slot_success(int(slot['slot']))
                    record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=model,model_route='voice',status='success',response_ms=elapsed,attempt_index=attempt_index,was_failover=attempt_index>1,http_status=response.status_code)
                    return text,{'api_slot':slot['slot'],'provider_label':slot['label'],'model':model}
                errors.append(f"API {slot['slot']}: متن خالی")
                mark_api_slot_blocked(int(slot['slot']),'پاسخ تبدیل گفتار خالی بود',AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'invalid_response')
            except Exception as exc:
                elapsed=int((time.perf_counter()-started)*1000)
                detail=sanitize_answer_text(str(exc))[:300]
                record_api_call_event(slot=int(slot['slot']),provider_label=slot['label'],model=model,model_route='voice',status='exception',response_ms=elapsed,attempt_index=attempt_index,was_failover=attempt_index>1,error_class=type(exc).__name__,error_detail=detail)
                mark_api_slot_blocked(int(slot['slot']),detail,AI_KEY_TRANSIENT_COOLDOWN_SECONDS,'transient_error')
                errors.append(f"API {slot['slot']}: {detail[:180]}")
    finally:
        if owns_client: await client.aclose()
    raise HTTPException(status_code=503,detail='تبدیل ویس به متن انجام نشد. مدل تبدیل گفتار در Variables تنظیم یا در دسترس نیست. '+(' | '.join(errors[-3:])))


@app.post('/api/v1/speech/transcribe')
async def transcribe_voice(file: UploadFile=File(...),user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    payload=await file.read(TRANSCRIPTION_MAX_MB*1024*1024+1)
    if len(payload)>TRANSCRIPTION_MAX_MB*1024*1024:
        raise HTTPException(status_code=413,detail=f'حجم فایل صوتی باید حداکثر {TRANSCRIPTION_MAX_MB} مگابایت باشد.')
    if not payload:
        raise HTTPException(status_code=400,detail='فایل صوتی خالی است.')
    text,meta=await transcribe_audio_bytes(file.filename or 'voice.webm',file.content_type or 'audio/webm',payload)
    audit(user['id'],'speech.transcribed',{'filename':file.filename,'size_bytes':len(payload),'api_slot':meta.get('api_slot'),'model':meta.get('model')})
    return {'text':text,**meta}


class TrainingChatInput(BaseModel):
    # Manager training is authoritative, immediately active, public to chatbot users,
    # never expires, and all training items have equal intrinsic priority.
    message: str = Field(min_length=3, max_length=6000)


class TrainingReviewInput(BaseModel):
    status: str


class KnowledgeGapStatusInput(BaseModel):
    status: str


class SystemControlInput(BaseModel):
    chat_enabled: bool
    maintenance_message: str | None = Field(default=None, max_length=500)


def _persian_number_to_int(value: str) -> int | None:
    mapping = {"یک": 1, "دو": 2, "سه": 3, "چهار": 4, "پنج": 5, "شش": 6, "هفت": 7}
    if value in mapping:
        return mapping[value]
    trans = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")
    value = value.translate(trans)
    return int(value) if value.isdigit() else None


def _normalize_user_datetime(value: str | None) -> str | None:
    if not value: return None
    try:
        dt=datetime.fromisoformat(value.replace('Z','+00:00'))
    except ValueError as exc:
        raise HTTPException(status_code=400,detail='قالب تاریخ و زمان معتبر نیست.') from exc
    if dt.tzinfo is None:
        dt=dt.replace(tzinfo=timezone(timedelta(minutes=LOCAL_TIMEZONE_OFFSET_MINUTES)))
    return dt.astimezone(timezone.utc).isoformat()


def _infer_effective_date(message: str, supplied: str | None) -> str:
    normalized=_normalize_user_datetime(supplied)
    if normalized: return normalized
    match=re.search(r'(یک|دو|سه|چهار|پنج|شش|هفت|[۰-۹0-9]+)\s*روز\s*(?:دیگر|دیگه|بعد)',message)
    if match:
        days=_persian_number_to_int(match.group(1)) or 0
        return (datetime.now(timezone.utc)+timedelta(days=days)).isoformat()
    return now_iso()


async def interpret_training(message: str, effective_from: str | None) -> tuple[dict[str, Any], dict[str, int]]:
    topic_match = re.search(r"(مبلغ\s+کنسلی\s+[\wآ-ی]+)", message)
    fallback_topic = (
        topic_match.group(1)
        if topic_match
        else re.split(r"\s+(?:از|به|تغییر|جایگزین|است|می‌شود)\s+", message, maxsplit=1)[0]
    ).strip()
    fallback = {
        "topic": (fallback_topic or message[:100].strip())[:160],
        "answer": message.strip(),
        "effective_from": _infer_effective_date(message, effective_from),
    }
    if AI_PROVIDER == "local":
        usage = {
            "prompt_tokens": estimate_tokens(message),
            "output_tokens": estimate_tokens(fallback["answer"]),
            "total_tokens": estimate_tokens(message) + estimate_tokens(fallback["answer"]),
        }
        return fallback, usage

    prompt = f"""
پیام زیر یک آموزش سازمانی برای دستیار بارسان است.
آن را بدون تغییر معنا به یک شیء JSON معتبر تبدیل کن و فقط JSON بده:
{{"topic":"موضوع کوتاه","answer":"پاسخ قطعی و کامل برای کارشناس"}}
تاریخ‌ها، اعداد، نام‌ها و استثناها باید دقیقاً حفظ شوند. هیچ داده‌ای اضافه نکن.
پیام: {message}
""".strip()
    try:
        text, usage, _ = await _generate_ai_text(
            [
                {"role": "system", "content": "فقط JSON معتبر و بدون Markdown تولید کن."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=700,
            temperature=0.0,
        )
        json_match = re.search(r"\{.*\}", text, flags=re.S)
        parsed = json.loads(json_match.group(0)) if json_match else {}
        fallback["topic"] = str(parsed.get("topic") or fallback["topic"])[:160]
        fallback["answer"] = str(parsed.get("answer") or fallback["answer"]).strip()
        return fallback, usage
    except Exception:
        usage = {
            "prompt_tokens": estimate_tokens(prompt),
            "output_tokens": estimate_tokens(fallback["answer"]),
            "total_tokens": estimate_tokens(prompt) + estimate_tokens(fallback["answer"]),
        }
        return fallback, usage


def activate_training(db: sqlite3.Connection,training_id: str,approver_id: int) -> None:
    row=db.execute("SELECT id,topic,instruction,canonical_key FROM training_rules WHERE id=?",(training_id,)).fetchone()
    if not row: raise HTTPException(status_code=404,detail='آموزش پیدا نشد.')
    new_key=row['canonical_key'] or canonical_training_key(row['topic'],row['instruction'])
    active=db.execute("SELECT id,canonical_key,topic,instruction FROM training_rules WHERE status='active' AND id!=?",(training_id,)).fetchall()
    superseded=[]
    for old in active:
        old_key=old['canonical_key'] or canonical_training_key(old['topic'],old['instruction'])
        if old_key==new_key or _key_similarity(old_key,new_key)>=0.60:
            superseded.append(old['id'])
    for old_id in superseded:
        db.execute("UPDATE training_rules SET status='superseded',updated_at=? WHERE id=?",(now_iso(),old_id))
    db.execute("UPDATE training_rules SET status='active',approved_by=?,canonical_key=?,supersedes_id=?,updated_at=? WHERE id=?",(approver_id,new_key,superseded[0] if superseded else None,now_iso(),training_id))
    rebuild_training_fts(db)
    bump_knowledge_version(db)


@app.post('/api/v1/training/chat')
async def training_chat(data:TrainingChatInput,user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    started=time.perf_counter();parsed,usage=await interpret_training(data.message,None)
    training_id=str(uuid.uuid4())
    status='active'
    topic_key=normalize_text(parsed['topic'])[:240] or normalize_text(data.message)[:240]
    canonical_key=canonical_training_key(parsed['topic'],data.message)
    effective_from=now_iso();expires_at=None
    with get_db() as db:
        visibility='public'
        allowed_roles=[]
        allowed_user_ids=[]
        department=None
        db.execute("""INSERT INTO training_rules(id,topic,topic_key,canonical_key,instruction,answer,priority,status,effective_from,expires_at,visibility,allowed_roles_json,allowed_user_ids_json,department,created_by,approved_by,created_at,updated_at)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (training_id,parsed['topic'],topic_key,canonical_key,data.message.strip(),parsed['answer'],100,status,effective_from,expires_at,visibility,json.dumps(allowed_roles,ensure_ascii=False),json.dumps(allowed_user_ids),department,user['id'],user['id'],now_iso(),now_iso()))
        activate_training(db,training_id,user['id'])
    duration_ms=int((time.perf_counter()-started)*1000)
    record_usage(user_id=user['id'],external_user_id=None,event_type='training',usage=usage,model=active_model_name() if AI_PROVIDER != 'local' else 'local-training-parser',response_ms=duration_ms)
    audit(user['id'],'training.created',{'training_id':training_id,'status':status,'topic':parsed['topic']})
    if status=='pending':notify_admin('training_pending','آموزش جدید مدیر نیازمند تأیید است',{'training_id':training_id,'manager':user['username'],'topic':parsed['topic']},'warning')
    else:notify_admin('training_active','آموزش اولویت‌دار فعال شد',{'training_id':training_id,'topic':parsed['topic']},'success')
    return {'training_id':training_id,'topic':parsed['topic'],'answer':parsed['answer'],'effective_from':effective_from,'expires_at':expires_at,'status':status,'assistant_message':'آموزش ثبت و فعال شد و از این پس برای همه کاربران قابل استفاده است.','usage':usage}


@app.get("/api/v1/training/rules")
def training_rules(
    status: str | None = None,
    user: dict[str, Any] = Depends(require_roles("manager", "admin")),
) -> list[dict[str, Any]]:
    sql = """
    SELECT tr.*, cu.username AS creator_username, au.username AS approver_username
    FROM training_rules tr
    LEFT JOIN users cu ON cu.id = tr.created_by
    LEFT JOIN users au ON au.id = tr.approved_by
    """
    params: list[Any] = []
    if status:
        sql += " WHERE tr.status = ?"
        params.append(status)
    sql += " ORDER BY tr.updated_at DESC LIMIT 500"
    with get_db() as db:
        rows = db.execute(sql, params).fetchall()
    return [dict(r) for r in rows]


@app.patch('/api/v1/admin/training/{training_id}')
def review_training(training_id:str,data:TrainingReviewInput,owner:dict[str,Any]=Depends(require_owner)) -> dict[str,Any]:
    if data.status not in {'active','rejected'}:raise HTTPException(status_code=400,detail='وضعیت فقط active یا rejected است.')
    with get_db() as db:
        row=db.execute("SELECT id,topic FROM training_rules WHERE id=?",(training_id,)).fetchone()
        if not row:raise HTTPException(status_code=404,detail='آموزش پیدا نشد.')
        if data.status=='active':activate_training(db,training_id,owner['id'])
        else:
            db.execute("UPDATE training_rules SET status='rejected',approved_by=?,updated_at=? WHERE id=?",(owner['id'],now_iso(),training_id));rebuild_training_fts(db);bump_knowledge_version(db)
    audit(owner['id'],'training.reviewed',{'training_id':training_id,'status':data.status});notify_admin('training_review','وضعیت آموزش تغییر کرد',{'training_id':training_id,'status':data.status})
    return {'ok':True,'status':data.status}


@app.get('/api/v1/manager/analytics/users')
def user_analytics(user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> list[dict[str,Any]]:
    with get_db() as db:
        registered=[dict(r) for r in db.execute("""WITH qs AS(SELECT c.user_id,COUNT(m.id) question_count,MAX(m.created_at) last_question FROM conversations c JOIN messages m ON m.conversation_id=c.id AND m.role='user' WHERE c.user_id IS NOT NULL GROUP BY c.user_id),us AS(SELECT user_id,SUM(prompt_tokens) prompt_tokens,SUM(output_tokens) output_tokens,SUM(total_tokens) total_tokens,CAST(COALESCE(AVG(response_ms),0) AS INTEGER) avg_response_ms,MAX(created_at) last_usage FROM usage_events WHERE user_id IS NOT NULL GROUP BY user_id) SELECT u.id,u.username,u.email,u.name,u.role,u.is_active,u.is_owner,'registered' entity_type,COALESCE(qs.question_count,0) question_count,COALESCE(us.prompt_tokens,0) prompt_tokens,COALESCE(us.output_tokens,0) output_tokens,COALESCE(us.total_tokens,0) total_tokens,COALESCE(us.avg_response_ms,0) avg_response_ms,MAX(COALESCE(us.last_usage,''),COALESCE(qs.last_question,''),u.created_at) last_activity FROM users u LEFT JOIN qs ON qs.user_id=u.id LEFT JOIN us ON us.user_id=u.id ORDER BY total_tokens DESC,question_count DESC""").fetchall()]
        external=[dict(r) for r in db.execute("""WITH qs AS(SELECT c.external_user_id,COUNT(m.id) question_count,MAX(m.created_at) last_question FROM conversations c JOIN messages m ON m.conversation_id=c.id AND m.role='user' WHERE c.user_id IS NULL AND c.external_user_id IS NOT NULL GROUP BY c.external_user_id),us AS(SELECT external_user_id,SUM(prompt_tokens) prompt_tokens,SUM(output_tokens) output_tokens,SUM(total_tokens) total_tokens,CAST(COALESCE(AVG(response_ms),0) AS INTEGER) avg_response_ms,MAX(created_at) last_usage FROM usage_events WHERE user_id IS NULL AND external_user_id IS NOT NULL GROUP BY external_user_id) SELECT NULL id,COALESCE(qs.external_user_id,us.external_user_id) username,NULL email,COALESCE(qs.external_user_id,us.external_user_id) name,'external' role,1 is_active,0 is_owner,'external' entity_type,COALESCE(qs.question_count,0) question_count,COALESCE(us.prompt_tokens,0) prompt_tokens,COALESCE(us.output_tokens,0) output_tokens,COALESCE(us.total_tokens,0) total_tokens,COALESCE(us.avg_response_ms,0) avg_response_ms,MAX(COALESCE(us.last_usage,''),COALESCE(qs.last_question,'')) last_activity FROM qs LEFT JOIN us ON us.external_user_id=qs.external_user_id UNION SELECT NULL,us.external_user_id,NULL,us.external_user_id,'external',1,0,'external',0,us.prompt_tokens,us.output_tokens,us.total_tokens,us.avg_response_ms,us.last_usage FROM us LEFT JOIN qs ON qs.external_user_id=us.external_user_id WHERE qs.external_user_id IS NULL""").fetchall()]
    return sorted(registered+external,key=lambda r:(r.get('total_tokens') or 0,r.get('question_count') or 0),reverse=True)


def _question_rows(search:str|None=None,user_id:int|None=None,external_user_id:str|None=None,limit:int=2000,offset:int=0) -> list[dict[str,Any]]:
    sql="""
    WITH timeline AS (
      SELECT m.*,
        LEAD(role) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_role,
        LEAD(content) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_content,
        LEAD(status) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_status,
        LEAD(prompt_tokens) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_prompt_tokens,
        LEAD(output_tokens) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_output_tokens,
        LEAD(total_tokens) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_total_tokens,
        LEAD(response_ms) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_response_ms,
        LEAD(model) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_model,
        LEAD(sources_json) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_sources_json,
        LEAD(api_slot) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_api_slot,
        LEAD(model_route) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_model_route,
        LEAD(estimated_cost) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_estimated_cost,
        LEAD(confidence_score) OVER(PARTITION BY conversation_id ORDER BY created_at,id) next_confidence_score
      FROM messages m
    )
    SELECT q.id question_id,q.content question,q.created_at asked_at,c.id conversation_id,c.user_id,c.external_user_id,
      COALESCE(u.username,u.email,c.external_user_id,'-') requester,u.name requester_name,
      CASE WHEN q.next_role='assistant' THEN q.next_content END answer,
      CASE WHEN q.next_role='assistant' THEN COALESCE(q.next_status,'answered') ELSE 'unanswered' END status,
      CASE WHEN q.next_role='assistant' THEN COALESCE(q.next_prompt_tokens,0) ELSE 0 END prompt_tokens,
      CASE WHEN q.next_role='assistant' THEN COALESCE(q.next_output_tokens,0) ELSE 0 END output_tokens,
      CASE WHEN q.next_role='assistant' THEN COALESCE(q.next_total_tokens,0) ELSE 0 END total_tokens,
      CASE WHEN q.next_role='assistant' THEN COALESCE(q.next_response_ms,0) ELSE 0 END response_ms,
      CASE WHEN q.next_role='assistant' THEN q.next_model END model,
      CASE WHEN q.next_role='assistant' THEN COALESCE(q.next_sources_json,'[]') ELSE '[]' END sources_json,
      CASE WHEN q.next_role='assistant' THEN q.next_api_slot END api_slot,
      CASE WHEN q.next_role='assistant' THEN q.next_model_route END model_route,
      CASE WHEN q.next_role='assistant' THEN COALESCE(q.next_estimated_cost,0) ELSE 0 END estimated_cost,
      CASE WHEN q.next_role='assistant' THEN q.next_confidence_score END confidence_score
    FROM timeline q JOIN conversations c ON c.id=q.conversation_id LEFT JOIN users u ON u.id=c.user_id
    WHERE q.role='user'
    """
    params=[]
    if search:
        sql+=" AND (q.content LIKE ? OR COALESCE(u.username,u.email,u.name,c.external_user_id,'') LIKE ?)";like=f'%{search}%';params.extend([like,like])
    if user_id is not None:sql+=' AND c.user_id=?';params.append(user_id)
    if external_user_id is not None:sql+=' AND c.external_user_id=?';params.append(external_user_id)
    sql+=' ORDER BY q.created_at DESC LIMIT ? OFFSET ?';params.extend([limit,offset])
    with get_db() as db:rows=db.execute(sql,params).fetchall()
    result=[]
    for row in rows:
        item=dict(row)
        item.update(response_section_details(item.get('status'),item.get('model'),item.get('sources_json'),item.get('total_tokens') or 0))
        result.append(item)
    return result


def _question_count(search:str|None=None,user_id:int|None=None,external_user_id:str|None=None) -> int:
    sql="SELECT COUNT(*) FROM messages q JOIN conversations c ON c.id=q.conversation_id LEFT JOIN users u ON u.id=c.user_id WHERE q.role='user'";params=[]
    if search:sql+=" AND (q.content LIKE ? OR COALESCE(u.username,u.email,u.name,c.external_user_id,'') LIKE ?)";like=f'%{search}%';params.extend([like,like])
    if user_id is not None:sql+=' AND c.user_id=?';params.append(user_id)
    if external_user_id is not None:sql+=' AND c.external_user_id=?';params.append(external_user_id)
    with get_db() as db:return int(db.execute(sql,params).fetchone()[0])


@app.get('/api/v1/manager/analytics/questions')
def question_analytics(q:str|None=None,user_id:int|None=None,external_user_id:str|None=None,page:int=1,page_size:int=100,user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    page=max(1,page);page_size=max(10,min(page_size,500));total=_question_count(q,user_id,external_user_id)
    return {'items':_question_rows(q,user_id,external_user_id,page_size,(page-1)*page_size),'total':total,'page':page,'page_size':page_size,'pages':max(1,(total+page_size-1)//page_size)}


@app.get("/api/v1/manager/knowledge-gaps")
def knowledge_gaps(
    status: str | None = None,
    user: dict[str, Any] = Depends(require_roles("manager", "admin")),
) -> list[dict[str, Any]]:
    sql = "SELECT * FROM knowledge_gaps"
    params: list[Any] = []
    if status:
        sql += " WHERE status = ?"
        params.append(status)
    sql += " ORDER BY occurrence_count DESC, last_seen DESC LIMIT 500"
    with get_db() as db:
        rows = db.execute(sql, params).fetchall()
    return [dict(r) for r in rows]


@app.patch("/api/v1/manager/knowledge-gaps/{gap_id}")
def update_gap_status(
    gap_id: int,
    data: KnowledgeGapStatusInput,
    user: dict[str, Any] = Depends(require_roles("manager", "admin")),
) -> dict[str, Any]:
    if data.status not in {"open", "in_training", "resolved", "ignored"}:
        raise HTTPException(status_code=400, detail="وضعیت شکاف دانشی نامعتبر است.")
    with get_db() as db:
        cursor = db.execute("UPDATE knowledge_gaps SET status = ? WHERE id = ?", (data.status, gap_id))
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="مورد پیدا نشد.")
    audit(user["id"], "knowledge_gap.updated", {"gap_id": gap_id, "status": data.status})
    return {"ok": True}



def _validate_dynamic_provider_url(value: str) -> str:
    value=str(value or '').strip().rstrip('/')
    if not value.startswith('https://'):
        raise HTTPException(status_code=400,detail='Base URL باید با https:// شروع شود.')
    return value

def _mask_api_key(value: str) -> str:
    value=str(value or '')
    if len(value)<=10:return '••••••••'
    return value[:4]+'…'+value[-4:]

def _dynamic_provider_public_rows() -> list[dict[str,Any]]:
    with get_db() as db:
        rows=db.execute('SELECT * FROM ai_provider_configs ORDER BY slot').fetchall()
    result=[]
    for row in rows:
        item=dict(row);stored=item.pop('api_key',None);
        try: plain=_unprotect_api_key(stored)
        except RuntimeError: plain=''
        item['api_key_masked']=_mask_api_key(plain);item['managed_by']='admin'
        result.append(item)
    return result

@app.get('/api/v1/admin/api-providers')
def list_dynamic_api_providers(admin:dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    dynamic=_dynamic_provider_public_rows()
    static=[{'slot':x['slot'],'label':x['label'],'base_url':x['base_url'],'model':x['model'],'enabled':True,'managed_by':'variables','api_key_masked':_mask_api_key(x['api_key'])} for x in _static_ai_slots()]
    return {'items':static+dynamic,'configured_count':len(configured_ai_slots()),'max_slots':MAX_AI_API_SLOTS,'dynamic_slots_available':max(0,MAX_AI_API_SLOTS-len(configured_ai_slots()))}

@app.post('/api/v1/admin/api-providers')
def create_dynamic_api_provider(data:DynamicApiProviderInput,admin:dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    base=_validate_dynamic_provider_url(data.base_url)
    with get_db() as db:
        used={int(r[0]) for r in db.execute('SELECT slot FROM ai_provider_configs').fetchall()}
        used.update(int(item['slot']) for item in _static_ai_slots())
        slot=next((i for i in range(1,MAX_AI_API_SLOTS+1) if i not in used),None)
        if slot is None:raise HTTPException(status_code=409,detail=f'حداکثر {MAX_AI_API_SLOTS} API قابل تنظیم است.')
        ts=now_iso();db.execute('''INSERT INTO ai_provider_configs(slot,label,base_url,api_key,model,model_economy,model_standard,model_advanced,vision_model,transcription_model,embedding_model,input_cost_per_1m,output_cost_per_1m,credit_amount,credit_currency,enabled,created_by,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',(slot,data.label.strip(),base,_protect_api_key(data.api_key),data.model.strip(),(data.model_economy or '').strip(),(data.model_standard or '').strip(),(data.model_advanced or '').strip(),(data.vision_model or '').strip(),(data.transcription_model or '').strip(),(data.embedding_model or '').strip(),data.input_cost_per_1m,data.output_cost_per_1m,data.credit_amount,data.credit_currency.strip().upper(),1,admin['id'],ts,ts))
    with _PROVIDER_SPEED_CACHE_LOCK:_PROVIDER_SPEED_CACHE.clear()
    audit(admin['id'],'api_provider.created',{'slot':slot,'label':data.label})
    return {'ok':True,'slot':slot,'configured_count':len(configured_ai_slots())}

@app.patch('/api/v1/admin/api-providers/{slot}')
def update_dynamic_api_provider(slot:int,data:DynamicApiProviderUpdate,admin:dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    if slot<1 or slot>MAX_AI_API_SLOTS:raise HTTPException(status_code=400,detail='شماره API نامعتبر است.')
    if any(int(x['slot'])==slot for x in _static_ai_slots()):raise HTTPException(status_code=400,detail='این اسلات از Railway Variables مدیریت می‌شود.')
    changes=data.model_dump(exclude_unset=True)
    if 'base_url' in changes:changes['base_url']=_validate_dynamic_provider_url(changes['base_url'])
    if 'credit_currency' in changes and changes['credit_currency']:changes['credit_currency']=changes['credit_currency'].strip().upper()
    if 'api_key' in changes and changes['api_key'] is not None:changes['api_key']=_protect_api_key(str(changes['api_key']))
    allowed={'label','base_url','api_key','model','model_economy','model_standard','model_advanced','vision_model','transcription_model','embedding_model','input_cost_per_1m','output_cost_per_1m','credit_amount','credit_currency','enabled'}
    changes={k:v for k,v in changes.items() if k in allowed}
    if not changes:return {'ok':True,'slot':slot}
    changes['updated_at']=now_iso();cols=','.join(f'{k}=?' for k in changes);vals=list(changes.values())+[slot]
    with get_db() as db:
        cur=db.execute(f'UPDATE ai_provider_configs SET {cols} WHERE slot=?',vals)
        if cur.rowcount==0:raise HTTPException(status_code=404,detail='API پیدا نشد.')
        if changes.get('enabled'):
            db.execute("UPDATE ai_api_slot_state SET blocked_until=NULL,last_status='ready',last_error=NULL,updated_at=? WHERE slot=?",(now_iso(),slot))
    with _PROVIDER_SPEED_CACHE_LOCK:_PROVIDER_SPEED_CACHE.clear()
    audit(admin['id'],'api_provider.updated',{'slot':slot,'fields':sorted(k for k in changes if k!='updated_at')})
    return {'ok':True,'slot':slot}

@app.delete('/api/v1/admin/api-providers/{slot}')
def delete_dynamic_api_provider(slot:int,admin:dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    if slot<1 or slot>MAX_AI_API_SLOTS:raise HTTPException(status_code=400,detail='شماره API نامعتبر است.')
    if any(int(x['slot'])==slot for x in _static_ai_slots()):raise HTTPException(status_code=400,detail='این اسلات فقط از Railway Variables حذف می‌شود.')
    with get_db() as db:
        cur=db.execute('DELETE FROM ai_provider_configs WHERE slot=?',(slot,));db.execute('DELETE FROM ai_api_slot_state WHERE slot=?',(slot,))
        if cur.rowcount==0:raise HTTPException(status_code=404,detail='API پیدا نشد.')
    with _PROVIDER_SPEED_CACHE_LOCK:_PROVIDER_SPEED_CACHE.clear()
    audit(admin['id'],'api_provider.deleted',{'slot':slot})
    return {'ok':True}

@app.post('/api/v1/admin/api-providers/{slot}/unblock')
def unblock_dynamic_api_provider(slot:int,admin:dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    if slot<1 or slot>MAX_AI_API_SLOTS:raise HTTPException(status_code=400,detail='شماره API نامعتبر است.')
    with get_db() as db:db.execute("UPDATE ai_api_slot_state SET blocked_until=NULL,last_status='ready',last_error=NULL,updated_at=? WHERE slot=?",(now_iso(),slot))
    audit(admin['id'],'api_provider.unblocked',{'slot':slot});return {'ok':True}

@app.get('/api/v1/admin/api-pool')
def get_api_pool(admin: dict[str,Any] = Depends(require_roles('admin'))) -> dict[str,Any]:
    return api_pool_public_status()


@app.post('/api/v1/admin/api-pool/reset')
def reset_api_pool(admin: dict[str,Any] = Depends(require_roles('admin'))) -> dict[str,Any]:
    configured=configured_ai_slots();first_slot=int(configured[0]['slot']) if configured else 1
    with get_db() as db:
        db.execute("UPDATE ai_api_slot_state SET blocked_until=NULL,last_status='ready',last_error=NULL,updated_at=?",(now_iso(),))
        db.execute("""INSERT INTO system_settings(key,value,updated_by,updated_at) VALUES('active_api_slot',?,?,?)
        ON CONFLICT(key) DO UPDATE SET value=excluded.value,updated_by=excluded.updated_by,updated_at=excluded.updated_at""",(str(first_slot),admin['id'],now_iso()))
    audit(admin['id'],'api_pool.reset',{'active_slot':first_slot})
    return api_pool_public_status()

@app.get('/api/v1/admin/analytics/provider-usage')
def admin_provider_usage(days: int=30, admin: dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    return provider_usage_dashboard(days)


@app.get('/api/v1/admin/analytics/request-health')
def admin_request_health(hours: int=24, admin: dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    hours=max(1,min(24*30,int(hours)))
    cutoff=(datetime.now(timezone.utc)-timedelta(hours=hours)).isoformat()
    with get_db() as db:
        summary=db.execute("""SELECT COUNT(*) total,COALESCE(AVG(duration_ms),0) avg_ms,
            SUM(CASE WHEN status_code>=500 THEN 1 ELSE 0 END) server_errors,
            SUM(CASE WHEN status_code>=400 THEN 1 ELSE 0 END) all_errors
            FROM request_metrics WHERE created_at>=?""",(cutoff,)).fetchone()
        rows=db.execute("""SELECT path,COUNT(*) requests,ROUND(AVG(duration_ms),1) avg_ms,MAX(duration_ms) max_ms,
            SUM(CASE WHEN status_code>=400 THEN 1 ELSE 0 END) errors
            FROM request_metrics WHERE created_at>=? GROUP BY path ORDER BY requests DESC,avg_ms DESC LIMIT 30""",(cutoff,)).fetchall()
    total=int(summary['total'] or 0)
    return {'hours':hours,'total_requests':total,'average_ms':round(float(summary['avg_ms'] or 0),1),
        'server_error_rate':round(int(summary['server_errors'] or 0)/max(1,total),4),'error_rate':round(int(summary['all_errors'] or 0)/max(1,total),4),
        'paths':[dict(row) for row in rows]}


@app.get('/api/v1/manager/analytics/quality')
def manager_quality(days: int=30, user: dict[str,Any]=Depends(require_roles('manager','admin'))) -> dict[str,Any]:
    return feedback_quality_dashboard(days)


@app.get('/api/v1/admin/health')
def admin_health(run_now: bool=False, admin: dict[str,Any]=Depends(require_roles('admin'))) -> dict[str,Any]:
    if run_now:
        return run_health_check(notify=True)
    with get_db() as db:
        row=db.execute("SELECT details FROM health_checks WHERE component='system' ORDER BY checked_at DESC LIMIT 1").fetchone()
    return json.loads(row['details']) if row else run_health_check(notify=False)


@app.get('/api/v1/admin/backups')
def backup_history(admin: dict[str,Any]=Depends(require_owner)) -> list[dict[str,Any]]:
    with get_db() as db:
        rows=db.execute("SELECT id,action,filename,size_bytes,status,details,created_at FROM backup_history ORDER BY created_at DESC LIMIT 100").fetchall()
    result=[]
    for row in rows:
        item=dict(row)
        try: details=json.loads(item.get('details') or '{}')
        except json.JSONDecodeError: details={}
        stored=Path(details.get('stored_path') or '')
        item['automatic']=bool(details.get('automatic') or item.get('action')=='auto_backup')
        item['stored_available']=bool(stored and stored.parent==BACKUP_DIR and stored.is_file())
        item['details']=details
        result.append(item)
    return result


@app.get('/api/v1/admin/backups/saved/{backup_id}')
def download_saved_backup(backup_id: str, owner: dict[str,Any]=Depends(require_owner)) -> StreamingResponse:
    with get_db() as db:
        row=db.execute("SELECT filename,details FROM backup_history WHERE id=? AND status='success'",(backup_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404,detail='فایل پشتیبان پیدا نشد.')
    try: details=json.loads(row['details'] or '{}')
    except json.JSONDecodeError: details={}
    path=Path(details.get('stored_path') or '')
    if not path or path.parent!=BACKUP_DIR or not path.is_file():
        raise HTTPException(status_code=404,detail='نسخه ذخیره‌شده این پشتیبان موجود نیست.')
    return StreamingResponse(path.open('rb'),media_type='application/zip',headers={'Content-Disposition':f'attachment; filename="{Path(row["filename"]).name}"'})


@app.get('/api/v1/admin/backups/download')
def download_backup(owner: dict[str,Any]=Depends(require_owner)) -> StreamingResponse:
    path,filename,backup_id,size_bytes,_=create_backup_file(owner['id'])
    audit(owner['id'],'backup.created',{'backup_id':backup_id,'filename':filename,'size_bytes':size_bytes,'streamed':True})
    headers={'Content-Disposition':f'attachment; filename="{filename}"','X-Barsan-Backup-Id':backup_id}
    def iterator():
        try:
            with path.open('rb') as handle:
                while True:
                    block=handle.read(1024*1024)
                    if not block: break
                    yield block
        finally:
            path.unlink(missing_ok=True)
    return StreamingResponse(iterator(),media_type='application/zip',headers=headers)


@app.post('/api/v1/admin/backups/restore')
async def restore_backup(file: UploadFile=File(...), owner: dict[str,Any]=Depends(require_owner)) -> dict[str,Any]:
    payload=await file.read(BACKUP_MAX_MB*1024*1024+1)
    result=await asyncio.to_thread(restore_backup_bytes,payload,owner['id'])
    audit(result.get('actor_id'),'backup.restored',{'filename':Path(file.filename or 'backup.zip').name,**result})
    notify_admin('backup_restored','بازیابی فایل پشتیبان انجام شد',{'actor':owner['username'],'source_release':result.get('source_release')},'warning')
    return result


@app.get('/api/v1/admin/system-control')
def get_system_control(owner:dict[str,Any]=Depends(require_owner)) -> dict[str,Any]:
    with get_db() as db:
        unread=db.execute("SELECT COUNT(*) FROM admin_notifications n LEFT JOIN admin_notification_reads r ON r.notification_id=n.id AND r.user_id=? WHERE r.notification_id IS NULL",(owner['id'],)).fetchone()[0]
        last=db.execute("SELECT MAX(updated_at) FROM system_settings").fetchone()[0]
    return {'chat_enabled':system_setting_bool('chat_enabled',True),'maintenance_message':system_setting('maintenance_message'),'unread_notifications':unread,'last_update':last,'note':'این کلید سرویس پاسخ‌گویی را کنترل می‌کند؛ پنل مدیریت روشن می‌ماند.'}


@app.patch('/api/v1/admin/system-control')
def set_system_control(data:SystemControlInput,owner:dict[str,Any]=Depends(require_owner)) -> dict[str,Any]:
    update_system_setting('chat_enabled','true' if data.chat_enabled else 'false',owner['id'])
    if data.maintenance_message is not None:update_system_setting('maintenance_message',data.maintenance_message.strip(),owner['id'])
    audit(owner['id'],'system.chat_toggled',{'chat_enabled':data.chat_enabled});notify_admin('system_control','وضعیت سرویس پاسخ‌گویی تغییر کرد',{'chat_enabled':data.chat_enabled,'admin':owner['username']},'warning' if not data.chat_enabled else 'success')
    return get_system_control(owner)


@app.get('/api/v1/admin/notifications')
def admin_notifications(owner:dict[str,Any]=Depends(require_owner)) -> list[dict[str,Any]]:
    with get_db() as db:
        rows=db.execute("SELECT n.*,CASE WHEN r.notification_id IS NULL THEN 0 ELSE 1 END is_read_by_me FROM admin_notifications n LEFT JOIN admin_notification_reads r ON r.notification_id=n.id AND r.user_id=? ORDER BY n.created_at DESC LIMIT 300",(owner['id'],)).fetchall()
    return [dict(r) for r in rows]


@app.post('/api/v1/admin/notifications/read-all')
def read_all_notifications(owner:dict[str,Any]=Depends(require_owner)) -> dict[str,bool]:
    with get_db() as db:
        db.execute("INSERT OR IGNORE INTO admin_notification_reads(notification_id,user_id,read_at) SELECT id,?,? FROM admin_notifications",(owner['id'],now_iso()))
    return {'ok':True}








def _excel_style_sheet(ws) -> None:
    dark='0B1A30';gold='B6913A';white='FFFFFF';thin=Side(style='thin',color='E5E7EB')
    for cell in ws[1]:
        cell.fill=PatternFill('solid',fgColor=dark);cell.font=Font(color=white,bold=True);cell.alignment=Alignment(horizontal='center',vertical='center',wrap_text=True);cell.border=Border(bottom=thin)
    ws.freeze_panes='A2';ws.auto_filter.ref=ws.dimensions
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment=Alignment(vertical='top',wrap_text=True);cell.border=Border(bottom=thin)
    for col in ws.columns:
        letter=col[0].column_letter
        max_len=max((len(str(c.value)) if c.value is not None else 0) for c in col[:250])
        ws.column_dimensions[letter].width=min(max(max_len+2,11),48)


def build_excel_report(include_audit: bool=False) -> bytes:
    users=user_analytics({'role':'admin'})
    total_questions=_question_count()
    questions=_question_rows(None,None,None,EXPORT_ROW_LIMIT,0)
    gaps=knowledge_gaps(None,{'role':'admin'})
    trainings=training_rules(None,{'role':'admin'})
    wb=Workbook();summary=wb.active;summary.title='Summary'
    summary.append(['گزارش سازمانی بارسان','مقدار'])
    summary.append(['زمان تولید',now_iso()])
    summary.append(['تعداد کاربران و کانال‌های خارجی',0])
    summary.append(['کل سؤال‌های دیتابیس',total_questions])
    summary.append(['سؤال‌های صادرشده',0])
    summary.append(['کل توکن ورودی',0])
    summary.append(['کل توکن خروجی',0])
    summary.append(['کل مصرف توکن',0])
    summary.append(['میانگین زمان پاسخ (ms)',0])
    summary.append(['شکاف‌های دانشی باز',0])
    summary.append(['آموزش‌های فعال',0])
    summary.append(['وضعیت محدودیت خروجی','کامل' if total_questions<=EXPORT_ROW_LIMIT else f'محدود به {EXPORT_ROW_LIMIT} ردیف'])

    ws=wb.create_sheet('Users')
    ws.append(['شناسه','شناسه نمایشی','نام','ایمیل','نوع/نقش','تعداد سؤال','توکن ورودی','توکن خروجی','کل توکن','میانگین پاسخ ms','آخرین فعالیت','وضعیت'])
    for r in users:
        ws.append([r.get('id'),r.get('username'),r.get('name'),r.get('email'),r.get('role'),r.get('question_count',0),r.get('prompt_tokens',0),r.get('output_tokens',0),r.get('total_tokens',0),r.get('avg_response_ms',0),r.get('last_activity'),'فعال' if r.get('is_active') else 'غیرفعال'])
    user_end=max(2,ws.max_row);total_row=ws.max_row+1
    ws.append(['جمع','','','','',f'=SUM(F2:F{user_end})',f'=SUM(G2:G{user_end})',f'=SUM(H2:H{user_end})',f'=SUM(I2:I{user_end})',f'=IFERROR(AVERAGEIF(J2:J{user_end},">0"),0)','',''])
    for c in ws[total_row]:c.font=Font(bold=True,color='FFFFFF');c.fill=PatternFill('solid',fgColor='B6913A')

    qws=wb.create_sheet('Questions')
    qws.append(['شناسه سؤال','کاربر','نام کاربر','سؤال','پاسخ','بخش پاسخ‌دهی','منابع استفاده‌شده','بدون مصرف توکن','وضعیت','توکن ورودی','توکن خروجی','کل توکن','زمان پاسخ ms','زمان سؤال','شناسه گفتگو'])
    for r in questions:qws.append([r['question_id'],r['requester'],r.get('requester_name'),r['question'],r.get('answer'),r.get('response_section_label'), ' | '.join(r.get('response_source_names') or []), 'بله' if r.get('zero_token') else 'خیر',r['status'],r['prompt_tokens'],r['output_tokens'],r['total_tokens'],r['response_ms'],r['asked_at'],r['conversation_id']])
    q_end=max(2,qws.max_row)

    source_summary=response_source_stats(3650)
    sws=wb.create_sheet('Answer_Sources')
    sws.append(['بخش پاسخ‌دهی','تعداد پاسخ','درصد کارکرد','پاسخ بدون توکن','درصد بدون توکن در بخش','کل توکن','میانگین زمان پاسخ ms','توضیح'])
    for item in source_summary['items']:
        sws.append([item['label'],item['value'],item['percentage']/100,item['zero_token_count'],item['zero_token_percentage']/100,item['total_tokens'],item['avg_response_ms'],item['description']])
    for row in range(2,sws.max_row+1):
        sws.cell(row=row,column=3).number_format='0.0%'
        sws.cell(row=row,column=5).number_format='0.0%'

    gws=wb.create_sheet('Knowledge_Gaps');gws.append(['شناسه','نمونه سؤال','تعداد تکرار','اولین مشاهده','وضعیت','آخرین مشاهده','آموزش مرتبط'])
    for r in gaps:gws.append([r['id'],r['sample_question'],r['occurrence_count'],r['first_seen'],r['status'],r['last_seen'],r['assigned_training_id']])
    tws=wb.create_sheet('Training');tws.append(['شناسه','موضوع','آموزش','پاسخ قطعی','اولویت','وضعیت','زمان اجرا UTC','انقضا UTC','سازنده','تأییدکننده','آخرین تغییر','جایگزین آموزش'])
    for r in trainings:tws.append([r['id'],r['topic'],r['instruction'],r['answer'],r['priority'],r['status'],r['effective_from'],r['expires_at'],r['creator_username'],r['approver_username'],r['updated_at'],r.get('supersedes_id')])

    if include_audit:
        aws=wb.create_sheet('Audit');aws.append(['عملیات','عامل','جزئیات','زمان'])
        with get_db() as db:audits=db.execute("SELECT a.action,COALESCE(u.username,u.email,'-') actor,a.details,a.created_at FROM audit_logs a LEFT JOIN users u ON u.id=a.user_id ORDER BY a.id DESC LIMIT ?",(EXPORT_ROW_LIMIT,)).fetchall()
        for r in audits:aws.append([r['action'],r['actor'],r['details'],r['created_at']])

    summary['B3']=f'=COUNTA(Users!B2:B{user_end})'
    summary['B5']=f'=COUNTA(Questions!A2:A{q_end})'
    summary['B6']=f'=Users!G{total_row}'
    summary['B7']=f'=Users!H{total_row}'
    summary['B8']=f'=Users!I{total_row}'
    summary['B9']=f'=Users!J{total_row}'
    summary['B10']='=COUNTIF(Knowledge_Gaps!E:E,"open")'
    summary['B11']='=COUNTIF(Training!F:F,"active")'
    summary.column_dimensions['A'].width=34;summary.column_dimensions['B'].width=28
    for sheet in wb.worksheets:_excel_style_sheet(sheet)
    wb.calculation.fullCalcOnLoad=True;wb.calculation.forceFullCalc=True;wb.calculation.calcMode='auto'
    buffer=io.BytesIO();wb.save(buffer);return buffer.getvalue()


@app.get('/api/v1/manager/reports/organization.xlsx')
def organization_excel_report(user:dict[str,Any]=Depends(require_roles('manager','admin'))) -> StreamingResponse:
    payload=build_excel_report(include_audit=False);filename=f"barsan-manager-report-{datetime.now(timezone.utc).date().isoformat()}.xlsx"
    return StreamingResponse(iter([payload]),media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',headers={'Content-Disposition':f'attachment; filename="{filename}"'})


@app.get('/api/v1/admin/reports/full.xlsx')
def admin_full_excel_report(owner:dict[str,Any]=Depends(require_owner)) -> StreamingResponse:
    payload=build_excel_report(include_audit=True);filename=f"barsan-owner-full-report-{datetime.now(timezone.utc).date().isoformat()}.xlsx"
    return StreamingResponse(iter([payload]),media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',headers={'Content-Disposition':f'attachment; filename="{filename}"'})




@app.get("/api/v1/manager/analytics/response-sources")
def response_source_analytics(
    days: int = 30,
    user: dict[str, Any] = Depends(require_roles("manager", "admin")),
) -> dict[str, Any]:
    return response_source_stats(days)



@app.get("/api/v1/manager/analytics/charts")
def analytics_charts(
    days: int = 30,
    user: dict[str, Any] = Depends(require_roles("manager", "admin")),
) -> dict[str, Any]:
    days = max(7, min(days, 180))
    start_dt = datetime.now(timezone.utc) - timedelta(days=days - 1)
    start_iso = start_dt.isoformat()
    date_keys = [
        (start_dt + timedelta(days=offset)).date().isoformat()
        for offset in range(days)
    ]

    with get_db() as db:
        question_rows = db.execute(
            """
            SELECT substr(created_at, 1, 10) AS day, COUNT(*) AS questions
            FROM messages
            WHERE role = 'user' AND created_at >= ?
            GROUP BY day
            """,
            (start_iso,),
        ).fetchall()
        usage_rows = db.execute(
            """
            SELECT substr(created_at, 1, 10) AS day,
                   COALESCE(SUM(total_tokens), 0) AS tokens,
                   CAST(COALESCE(AVG(response_ms), 0) AS INTEGER) AS avg_response_ms
            FROM usage_events
            WHERE created_at >= ?
            GROUP BY day
            """,
            (start_iso,),
        ).fetchall()
        status_rows = db.execute(
            """
            SELECT COALESCE(status, 'unknown') AS label, COUNT(*) AS value
            FROM messages
            WHERE role = 'assistant' AND created_at >= ?
            GROUP BY status ORDER BY value DESC
            """,
            (start_iso,),
        ).fetchall()
        role_rows = db.execute(
            """
            SELECT role AS label, COUNT(*) AS value
            FROM users WHERE is_active = 1
            GROUP BY role ORDER BY value DESC
            """
        ).fetchall()
        training_rows = db.execute(
            """
            SELECT status AS label, COUNT(*) AS value
            FROM training_rules GROUP BY status ORDER BY value DESC
            """
        ).fetchall()
        gap_rows = db.execute(
            """
            SELECT status AS label, COALESCE(SUM(occurrence_count), 0) AS value
            FROM knowledge_gaps GROUP BY status ORDER BY value DESC
            """
        ).fetchall()
        audit_rows = db.execute(
            """
            SELECT substr(created_at, 1, 10) AS day, COUNT(*) AS events
            FROM audit_logs WHERE created_at >= ?
            GROUP BY day
            """,
            (start_iso,),
        ).fetchall()

    q_map = {row['day']: int(row['questions']) for row in question_rows}
    u_map = {
        row['day']: {
            'tokens': int(row['tokens'] or 0),
            'avg_response_ms': int(row['avg_response_ms'] or 0),
        }
        for row in usage_rows
    }
    a_map = {row['day']: int(row['events']) for row in audit_rows}
    daily = [
        {
            'date': day,
            'label': day[5:].replace('-', '/'),
            'questions': q_map.get(day, 0),
            'tokens': u_map.get(day, {}).get('tokens', 0),
            'avg_response_ms': u_map.get(day, {}).get('avg_response_ms', 0),
            'admin_events': a_map.get(day, 0),
        }
        for day in date_keys
    ]

    top_users = [
        {
            'label': row.get('username') or row.get('email') or row.get('name') or '-',
            'tokens': int(row.get('total_tokens') or 0),
            'questions': int(row.get('question_count') or 0),
        }
        for row in user_analytics({'role': 'admin'})[:8]
    ]
    source_summary=response_source_stats(days)
    return {
        'days': days,
        'daily': daily,
        'top_users': top_users,
        'answer_status': [dict(row) for row in status_rows],
        'response_sources': source_summary['items'],
        'response_source_summary': {k:v for k,v in source_summary.items() if k!='items'},
        'roles': [dict(row) for row in role_rows],
        'training_status': [dict(row) for row in training_rows],
        'knowledge_gap_status': [dict(row) for row in gap_rows],
    }

from ui_templates import MAIN_HTML, WIDGET_HTML

THINKING_LOADER_PATH = Path(__file__).resolve().with_name('thinking_loader.mp4')

@app.get('/thinking-loader.mp4', include_in_schema=False)
def thinking_loader_video() -> FileResponse:
    if not THINKING_LOADER_PATH.exists():
        raise HTTPException(status_code=404, detail='Thinking loader asset not found')
    return FileResponse(
        THINKING_LOADER_PATH,
        media_type='video/mp4',
        headers={'Cache-Control': 'no-store, no-cache, must-revalidate, max-age=0', 'Pragma': 'no-cache'},
    )


@app.get('/', response_class=HTMLResponse)
def index() -> str:
    return MAIN_HTML


@app.get('/widget', response_class=HTMLResponse)
def widget() -> str:
    return WIDGET_HTML


@app.get('/widget.js', response_class=PlainTextResponse)
def widget_js() -> str:
    return f"""
(function(){{
  var b=document.createElement('button');
  b.textContent='پشتیبانی';
  b.style='position:fixed;bottom:20px;left:20px;z-index:2147483647;padding:12px 18px;border:0;border-radius:999px;background:#b99745;font-weight:bold;cursor:pointer';
  var f=document.createElement('iframe');
  f.src='{PUBLIC_BASE_URL}/widget';
  f.style='display:none;position:fixed;bottom:75px;left:20px;width:360px;max-width:calc(100vw - 30px);height:460px;border:0;border-radius:16px;z-index:2147483647;box-shadow:0 15px 50px #0008';
  b.onclick=function(){{f.style.display=f.style.display==='none'?'block':'none'}};
  document.body.appendChild(f);document.body.appendChild(b);
}})();
""".strip()


@app.exception_handler(Exception)
async def unhandled(_: Request, exc: Exception):
    if isinstance(exc, HTTPException):
        raise exc
    return JSONResponse({'detail': 'خطای داخلی سرور رخ داد.', 'type': exc.__class__.__name__}, status_code=500)
